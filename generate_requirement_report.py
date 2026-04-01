import os
import subprocess
import sys

try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document

def create_report():
    doc = Document()
    doc.add_heading('BÁO CÁO PHÂN TÍCH YÊU CẦU - PR.02 CÔNG THỨC TÍNH PHÍ', 0)
    
    doc.add_heading('PHẦN A: TÓM TẮT NGHIỆP VỤ CHUYÊN SÂU (Requirements Breakdown)', 1)
    
    doc.add_heading('1. Thông điệp cốt lõi (Core Business Value)', 2)
    doc.add_paragraph('Cung cấp công cụ cho người quản trị (Maker) linh hoạt thiết lập cơ chế/công thức tính phí và điều kiện áp dụng cho từng nhóm khách hàng. Phục vụ hệ thống tính toán tự động chi phí giao dịch trên core.')
    
    doc.add_heading('2. Luồng chính (Happy Path Workflow)', 2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Người dùng đăng nhập với vai trò Maker.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Mở màn hình "Thiết lập mã phí" -> Cập nhật thông tin mã phí chung.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Cấu hình "Điều kiện tính phí theo tài khoản" để siết chặt đối tượng áp dụng (VD: Số dư bình quân < 2.000.000 VND).')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Tại danh sách "Quy tắc tính phí", cấu hình các biểu thức/công thức tĩnh hoặc động theo từng "Nhóm khách hàng".')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Lưu thay đổi, hệ thống khởi tạo tác vụ hiển thị ở mục "Tác vụ chờ duyệt".')

    doc.add_heading('3. Các luồng rẽ nhánh / Ngoại lệ (Alternative & Exception Flows)', 2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Nhập thiếu dữ liệu bắt buộc: Hệ thống chặn lưu và popup cảnh báo (BR_01).')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Huỷ thao tác trung gian: Nhấn nút "Đóng" làm sạch toàn bộ dữ liệu đang nhập (BR_02).')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Trùng lặp cấu hình: Trùng Mã/Tên -> Hệ thống báo lỗi tồn tại (BR_03).')

    doc.add_heading('4. Bảng Điều Kiện Tiên Quyết / Dữ Liệu Nền (Pre-conditions & Master Data)', 2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Session đăng nhập khả dụng, thuộc nhãn phân quyền Maker.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Các dictionary Master (Nhóm KH, Tiền tệ, Các Operator >, <, =, Danh sách Trường biến số / Hàm date_diff) phải được nạp đầy đủ vào hệ thống trước.')

    doc.add_heading('PHẦN B: DANH SÁCH CẢNH BÁO LỖ HỔNG & Q&A DÀNH CHO BA', 1)
    doc.add_paragraph('Tiến hành tham chiếu chéo giữa Mockup (Wireframe) và Tài liệu text URD. Phát hiện sự lệch pha cực kỳ nghiêm trọng ảnh hưởng đến rủi ro dự án.')
    
    table = doc.add_table(rows=1, cols=5)
    table.style = 'TableGrid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ID Câu Hỏi'
    hdr_cells[1].text = 'Tham chiếu'
    hdr_cells[2].text = 'Nội dung Lỗ hổng / Khuyến nghị'
    hdr_cells[3].text = 'Phân loại'
    hdr_cells[4].text = 'Đề xuất xử lý từ QA'
    
    data = [
        ("Q&A-01", "URD BR_04, BR_05 vs Hình ảnh Popup UI", "MÂU THUẪN MÔ HÌNH TOÁN HỌC:\n- Theo URD: Hệ thống yêu cầu công thức CỰC ĐƠN GIẢN. Tối đa 3 component nối bằng phép [+]. Component chỉ chọn 1 trong 3 template mẫu.\n- Theo Wireframe: Form thiết kế đang cấp quyền kiểu 'Expression Builder' TỰ DO CỰC PHỨC TẠP. Cho phép (+) (-) (x) (/) lồng ghép ngoặc tròn () và các hàm nâng cao (DATE_DIFF, MONTHS_BETWEEN). Sự chênh lệch này là khổng lồ về mặt xử lý Core backend.", "Nghiệp vụ cốt lõi", "BA BẮT BUỘC CHỐT LẠI YÊU CẦU Scope: Nếu làm theo hình học UI, phải đập toàn bộ BR_04 & BR_05. Phải bổ sung tài liệu về Validator Parser cho chuỗi công thức động (Nỗ lực Dev/Test sẽ bị X10 lần)."),
        ("Q&A-02", "Luồng Màn hình URD vs UI hiển thị", "SAI LỆCH LUỒNG MÀN HÌNH NỀN:\n- URD PR.02 mô tả việc 'Thêm mới Công Thức Tính phí' là một màn hình danh mục riêng độc lập.\n- Khung cảnh tổng thể của Wireframe lại là luồng 'Thiết lập MÃ PHÍ' (Gán thuộc tính theo từng nhóm KH, set điều kiện SD). Dường như Wireframe UI này thuộc về tài liệu PR.03 Danh mục Mã Phí thay vì mẫu Công thức?", "Flow Nghiệp Vụ", "Đề nghị BA làm rõ: Bối cảnh của Wireframe này đang trỏ vào chức năng Thêm Cấu Hình Chuyên Sâu, hay tạo Master Rule Công thức?"),
        ("Q&A-03", "UI Thiết lập công thức", "RÀO CẢN VALIDATION / SYNTAX ERROR:\nBox nhập công thức trên UI cho phép Input từ bàn phím hay buộc User phải click chèn từ button? Nếu User có quyền gõ text, rủi ro người dùng gõ sai ngoặc '(Ax B', chia cho 0 (:0), hoặc gọi biến số không tồn tại thì hệ thống văng lỗi Exception hay có Validator Check ngay lập tức?", "UI / Security Validation", "Yêu cầu bổ sung ràng buộc (Constraint & Validation Rules): Button Xác nhận phải có luồng Validate Parser Công thức. Không cho phép lưu khi Cú pháp Expression sai."),
        ("Q&A-04", "Wireframe Component Data", "XỬ LÝ NULL / DEFECT KHI CALL VALUE:\nDữ liệu công thức chứa các trường #MarginValue#, #LC_Amount#. Trên runtime hệ thống thực tế chạy, nếu giao dịch đổ vào không truy xuất ra được các giá trị này (Bị Null, hoặc Missing), phương pháp Fallback data là gì? Hệ thống sẽ văng Failed Giao Dịch hay auto map giá trị bằng số 0?", "Logic Hệ Thống", "Bổ sung BR dự phòng Default Fallback Variables. VD: Null value của thuộc tính tính giá ưu tiên parse mặc định = 0.")
    ]
    
    for row_data in data:
        row = table.add_row().cells
        row[0].text = row_data[0]
        row[1].text = row_data[1]
        row[2].text = row_data[2]
        row[3].text = row_data[3]
        row[4].text = row_data[4]
        
    doc.save('/Users/mac/antigravity-testing-kit/Phan_Tich_URD_PR02_Cong_Thuc_Phi.docx')
    print("Document created successfully at /Users/mac/antigravity-testing-kit/Phan_Tich_URD_PR02_Cong_Thuc_Phi.docx")

create_report()
