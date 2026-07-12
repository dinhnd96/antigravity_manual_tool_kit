# -*- coding: utf-8 -*-
import docx
import os
import re

def verify_updates():
    original_path = "/Users/mac/antigravity-testing-kit/Tài liệu toàn hệ thống Profix/US1-10/US06_Part_C_TestCase_Coverage.docx"
    updated_path = "/Users/mac/antigravity-testing-kit/Tài liệu toàn hệ thống Profix/US1-10/US06_Part_C_TestCase_Coverage_Updated.docx"
    
    if not os.path.exists(original_path) or not os.path.exists(updated_path):
        print("Error: Missing one of the document files for comparison.")
        return False
        
    doc_orig = docx.Document(original_path)
    doc_upd = docx.Document(updated_path)
    
    errors = 0
    warnings = 0
    passed = 0
    
    print("=== VERIFYING TEST CASE REFERENCES UPDATES ===")
    
    for t_idx in range(min(len(doc_orig.tables), len(doc_upd.tables))):
        table_orig = doc_orig.tables[t_idx]
        table_upd = doc_upd.tables[t_idx]
        
        if len(table_orig.rows) == 0:
            continue
            
        # Find column indices
        headers = [cell.text.strip() for cell in table_orig.rows[0].cells]
        sc_col_idx = -1
        ref_col_idx = -1
        
        for idx, h in enumerate(headers):
            if "Mã SC" in h or h == "Mã":
                sc_col_idx = idx
            elif "Trích dẫn tài liệu" in h or "Reference" in h:
                ref_col_idx = idx
                
        if sc_col_idx == -1 or ref_col_idx == -1:
            continue
            
        print(f"\nAuditing Table {t_idx} (Headers: {headers[sc_col_idx]}, {headers[ref_col_idx]})")
        
        for r_idx in range(1, len(table_orig.rows)):
            sc_id = table_orig.rows[r_idx].cells[sc_col_idx].text.strip()
            old_ref = table_orig.rows[r_idx].cells[ref_col_idx].text.strip()
            new_ref = table_upd.rows[r_idx].cells[ref_col_idx].text.strip()
            
            if not sc_id:
                continue
                
            print(f"  {sc_id}:")
            print(f"    - Before: {old_ref}")
            print(f"    - After : {new_ref}")
            
            # Constraint 1: Length check <= 200
            ref_len = len(new_ref)
            if ref_len > 200:
                print(f"    ❌ ERROR: Length is {ref_len} (> 200 chars limit)")
                errors += 1
            else:
                print(f"    ✓ Length: {ref_len} chars")
                
            # Constraint 2: Verbatim structure containing "-" or "–"
            if " – " not in new_ref and " - " not in new_ref:
                print("    ❌ ERROR: Missing standard separator ' – ' (Location - Quote)")
                errors += 1
                
            # Constraint 3: Verbatim quotation indicator *" or *“
            if '*"' not in new_ref and '*“' not in new_ref and '*”' not in new_ref:
                print("    ⚠️ WARNING: Missing verbatim quote indicator *\" or *“")
                warnings += 1
                
            # Constraint 4: No line numbers
            line_match = re.search(r'\b(dòng|line)\b\s*\d+', new_ref, re.IGNORECASE)
            if line_match:
                print(f"    ❌ ERROR: Contains banned line number reference: '{line_match.group()}'")
                errors += 1
                
            passed += 1
            
    print("\n=== VERIFICATION SUMMARY ===")
    print(f"Total test cases verified: {passed}")
    print(f"Errors found: {errors}")
    print(f"Warnings found: {warnings}")
    
    if errors > 0:
        print("❌ Verification FAILED due to errors. Please fix mapping values.")
        return False
    else:
        print("✅ Verification PASSED successfully!")
        return True

if __name__ == "__main__":
    verify_updates()
