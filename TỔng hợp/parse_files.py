import pandas as pd
import docx
import json
import sys

def parse_docx(file_path):
    doc = docx.Document(file_path)
    text = []
    for para in doc.paragraphs:
        text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            text.append(" | ".join(row_data))
    return "\n".join(text)

def parse_excel(file_path):
    xl = pd.ExcelFile(file_path)
    result = {}
    for sheet in xl.sheet_names:
        df = xl.parse(sheet)
        if not df.empty:
            result[sheet] = df.to_string()
    return result

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python parse_files.py <docx_path> <excel_path>")
        sys.exit(1)
        
    docx_text = parse_docx(sys.argv[1])
    try:
        excel_data = parse_excel(sys.argv[2])
    except Exception as e:
        excel_data = str(e)
        
    # write to output files
    with open('output_docx.txt', 'w', encoding='utf-8') as f:
        f.write(docx_text)
        
    with open('output_excel.txt', 'w', encoding='utf-8') as f:
        for k, v in excel_data.items():
            f.write(f"Sheet: {k}\n")
            f.write(v)
            f.write("\n\n")
    print("Done")
