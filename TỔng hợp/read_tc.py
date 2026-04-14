import pandas as pd
import docx
import sys
import json

def read_data(docx_path, excel_path):
    print("--- DOCX CONTENT ---")
    doc = docx.Document(docx_path)
    for p in doc.paragraphs:
        if p.text.strip():
            print(p.text)
    for table in doc.tables:
        for row in table.rows:
            print(" | ".join([c.text.replace("\n", " ") for c in row.cells]))
            
    print("\n--- EXCEL TEST CASES ---")
    xl = pd.ExcelFile(excel_path)
    df = xl.parse("TestCases")
    
    # We only want to print rows that have meaningful content
    # I'll print them as JSON for easy reading
    records = df.to_dict(orient="records")
    for row in records:
        if pd.notna(row['TC_ID']) or pd.notna(row['Title']):
            print(json.dumps({k: str(v) for k, v in row.items() if pd.notna(v)}, ensure_ascii=False, indent=2))

if __name__ == "__main__":
    read_data(sys.argv[1], sys.argv[2])
