from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('BÁO CÁO TIẾN ĐỘ & CHẤT LƯỢNG KIỂM THỬ', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Summary Info
doc.add_paragraph('Dự án: ProfiX Phase 1 - Feature 02 (SA Tham Số Hệ Thống)')
doc.add_paragraph('Ngày báo cáo: 10/04/2026')
doc.add_paragraph('')

# 1. Executive Summary
h1 = doc.add_heading('1. Executive Summary (Tóm lược điều hành)', level=1)
p1 = doc.add_paragraph()
run1 = p1.add_run('Tình trạng dự án: HIGH RISK - NGUY CƠ TRỄ TIẾN ĐỘ (DELAY)\n')
run1.bold = True
run1.font.color.rgb = RGBColor(255, 165, 0) # Orange
p1.add_run('Tiến độ thực thi đang bị nghẽn nghiêm trọng ở mức 82.9%, tỉ lệ Blocked rất cao (30 cases) do liên tục bị tắc nghẽn bởi các feature thuộc nhóm PR (PR01, PR05) và một số module nội tại (SA03, SA07, SA12) chưa được team Dev bàn giao.')

# 2. Dashboard
doc.add_heading('2. Dashboard Chỉ số (Key Metrics)', level=1)
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Hạng mục'
hdr_cells[1].text = 'Số lượng'
hdr_cells[2].text = 'Tỉ lệ %'
hdr_cells[3].text = 'Ghi chú'

metrics = [
    ('Tổng số kịch bản (Test Cases)', '164', '100%', 'Toàn bộ Feature 02 (SA)'),
    ('Đã thực thi (Executed)', '136', '82.9%', 'Tốc độ đang bị chững lại'),
    ('  - Pass (Đạt)', '85', '51.8%', ''),
    ('  - Failed (Lỗi)', '21', '12.8%', 'Đang vướng lỗi tại SA09, SA10, SA11'),
    ('  - Blocked (Bị chặn)', '30', '18.3%', 'Hầu hết do chờ tính năng khác'),
    ('Còn lại (Remaining/Waiting)', '28', '17.1%', 'Pending hoàn toàn ở SA07 và SA12')
]

for item in metrics:
    row_cells = table.add_row().cells
    row_cells[0].text = item[0]
    row_cells[1].text = item[1]
    row_cells[2].text = item[2]
    row_cells[3].text = item[3]

doc.add_paragraph('')

# 3. Risk & Impact
doc.add_heading('3. Phân tích Rủi ro & Tác động (Risk & Impact Analysis)', level=1)
doc.add_paragraph('Dự án hiện đang tồn đọng 3 nút thắt chính yếu (Bottlenecks) ảnh hưởng trực tiếp tới mốc Go-live:')

doc.add_paragraph('Rủi ro cao từ Sự Phụ Thuộc (Dependency Risks): Có một số lượng lớn các kịch bản đang bị đánh Blocked (chủ yếu ở SA06 và các tính năng tính phí) do team Dev chưa bàn giao PR01 và PR05. Việc thiếu tham số đầu vào khiến kiểm thử Liên kết Phí - Điều kiện Phí hoàn toàn bị tê liệt, chặn đứng luồng End-to-End.', style='List Bullet')

doc.add_paragraph('Chưa Bàn Giao Module (Non-Delivery): Phân hệ SA07 (18 cases) và SA12 (10 cases) hiện tại vẫn là con số 0 (chưa thể test) do chưa được bàn giao mã nguồn. Tương đương với gần 20% scope công việc dự án đang phải bảo lưu.', style='List Bullet')

doc.add_paragraph('Thiếu hụt tính năng ngay trong module đang test: Ngay tại phân hệ SA03, có một số tính năng con Dev chưa làm xong hoặc chưa Merge code, khiến tỷ lệ Blocked của module này tăng lên (10/19 cases Blocked), QA phải xé lẻ test case để test "lom dom" phần có sẵn.', style='List Bullet')

doc.add_paragraph('Tình trạng Lỗi (Failed Cases): Có 21 kịch bản đang Failed diện rộng (đặc biệt là lỗi thao tác UI, validation và API ở SA09, SA10, SA11). Việc dồn lỗi cũ chưa fix và tính năng hệ trọng chưa bàn giao tạo ra rủi ro nợ kỹ thuật (Technical Debt) quá lớn về cuối Release.', style='List Bullet')

# 4. Recommendations
doc.add_heading('4. Hành động & Đề xuất (Recommendations)', level=1)
doc.add_paragraph('Để tháo gỡ điểm nghẽn, QA Team đề nghị thực hiện khẩn cấp các hành động sau:')

doc.add_paragraph('Yêu cầu Team Kỹ thuật (Dev/PM):', style='List Number')
doc.add_paragraph('Chốt cho QA Deadline chính xác (ETA) bàn giao các nhóm chức năng PR01, PR05 cũng như SA07 & SA12.\nBổ sung hoàn thiện ngay các tính năng đang thiếu tại SA03 để QA đóng module này lại.\nTập trung xử lý dứt điểm các lỗi làm Failed 21 cases hiện có.')

doc.add_paragraph('Kế hoạch Dự phòng rủi ro:', style='List Number')
doc.add_paragraph('Nếu Dev không thể giao PR01, PR05 + SA07, SA12 trước thời hạn, đề xuất với Ban Giám Đốc/PM cho lùi ngày UAT/Go-live ít nhất 3-5 ngày để Team QA có đủ lượng "Lead Time" test full nghiệp vụ, đảm bảo hệ thống không gãy khi tích hợp.')

doc.save('Reports/QA_Management_Report_20260410.docx')
print('Tạo file báo cáo docx thành công: Reports/QA_Management_Report_20260410.docx')
