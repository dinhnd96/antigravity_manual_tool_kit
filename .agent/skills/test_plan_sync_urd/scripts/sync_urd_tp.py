import docx
import pandas as pd
import re
import numpy as np
import os

def sync_urd_to_testplan(urd_path, tp_path, output_path):
    if not os.path.exists(urd_path) or not os.path.exists(tp_path):
        return "File not found."

    # 1. Read URD Features (Basic logic)
    doc = docx.Document(urd_path)
    urd_features = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if re.search(r'([A-Z]{2}\.\d{2})', t):
            urd_features.append(t)
    
    # Read API table
    api_list = []
    for table in doc.tables:
        if "API" in "".join([c.text for c in table.rows[0].cells]):
            for row in table.rows[1:]:
                api_list.append(" - ".join([c.text.strip().replace("\n", " ") for c in row.cells[:2]]))
    
    # 2. Read Test Plan
    df = pd.read_excel(tp_path, sheet_name="Test Plan")
    
    # Logic: Đây là nơi bạn sẽ thực hiện so sánh (diff) và merge.
    # Code template này chỉ ví dụ cách thao tác file.
    
    # 3. Ghi file mới
    # (Tại đây bạn sẽ code logic thay đổi như tôi đã làm trong phiên làm việc)
    
    print(f"Logic đồng bộ hóa đã sẵn sàng. URD: {len(urd_features)} features, {len(api_list)} APIs.")
    return True

# Sử dụng: python sync_urd_tp.py <urd_path> <tp_path>
if __name__ == "__main__":
    import sys
    if len(sys.argv) > 2:
        sync_urd_to_testplan(sys.argv[1], sys.argv[2], "Updated_TP.xlsx")
