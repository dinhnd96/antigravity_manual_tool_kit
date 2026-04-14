from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

OUTPUT_PATH = "/Users/mac/antigravity-testing-kit/Feature_02_SA_Tham_So_He_Thong/review/US01_QA_Requirements_Review.docx"

doc = Document()

# ─── Page Margins ───
section = doc.sections[0]
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2)

# ─── Helpers ───
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color="1F3864"):
    p = doc.add_heading(text, level=level)
    run = p.runs[0]
    run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_paragraph(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

# ─── STYLES ───
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)

# ─── TITLE ───
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("BÁO CÁO RÀ SOÁT TÀI LIỆU YÊU CẦU")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
run.font.name = 'Times New Roman'

sub_title = doc.add_paragraph()
sub_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub_title.add_run("US01 – Khai báo Danh mục Sản phẩm Dịch vụ (SPDV)")
run2.bold = True
run2.font.size = Pt(13)
run2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
run2.font.name = 'Times New Roman'

doc.add_paragraph()
meta = doc.add_paragraph()
meta.add_run(f"Ngày tạo: {datetime.date.today().strftime('%d/%m/%Y')}    |    Người thực hiện: QA Team    |    Trạng thái: Bản nháp (Draft)")
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.runs[0].font.size = Pt(10)
meta.runs[0].italic = True

doc.add_paragraph()

# ─── SECTION 1: QA UNDERSTANDING ───
add_heading(doc, "PHẦN 1 — TÓM TẮT NGHIỆP VỤ (QA Understanding Summary)", level=1)
p = doc.add_paragraph()
p.add_run("Mục đích: ").bold = True
p.add_run("Phần này trình bày cách QA hiểu tài liệu US01. BA vui lòng xác nhận đúng/sai trước khi chuyển sang Phần 2.")

summaries = [
    {
        "id": "US01-A",
        "title": "Thêm mới Nghiệp vụ (Cấp 1 trong cây SPDV)",
        "actor": "Maker",
        "trigger": "Click '+ Thêm mới' tại màn hình Danh sách SPDV",
        "flow": "Maker nhập Tên, Mô tả, Ngày hiệu lực, Ngày hết hiệu lực → Xác nhận → Hệ thống tự sinh Mã (2 chữ số, tự tăng 01-99) → Lưu trạng thái Chờ duyệt → Hiển thị tại 'Tác vụ Pending' của Maker và 'Tác vụ chờ duyệt' của Checker.",
        "post": "Sau khi Checker duyệt, bản ghi hiển thị chính thức trên lưới Danh mục SPDV."
    },
    {
        "id": "US01-B",
        "title": "Chỉnh sửa Nghiệp vụ (Cấp 1 — chỉ sửa được Ngày)",
        "actor": "Maker",
        "trigger": "Click 'Chỉnh sửa' tại bản ghi Nghiệp vụ đã được duyệt",
        "flow": "Maker chỉ được thay đổi Ngày hiệu lực và/hoặc Ngày hết hiệu lực → Xác nhận → Hệ thống tạo bản ghi mới trạng thái Chờ duyệt.",
        "post": "Thông tin mới chỉ hiển thị trên lưới sau khi Checker phê duyệt theo US25."
    },
    {
        "id": "US01-C",
        "title": "Thêm mới SPDV chi tiết (Cấp con trong cây)",
        "actor": "Maker",
        "trigger": "Click '+' tại bản ghi SPDV cấp cha liền trước trên cây",
        "flow": "Tương tự US01-A. Mã được sinh theo công thức: Mã cấp cha + 2 chữ số tự tăng từ 01-99. Hệ thống tự hiển thị Mã-Tên SPDV cấp cha liền trước (Read-only).",
        "post": "Cây SPDV được mở rộng thêm 1 nhánh sau khi Checker duyệt."
    },
    {
        "id": "US01-D",
        "title": "Batch Job — Tự động cập nhật Trạng thái theo ngày T",
        "actor": "Hệ thống (Batch Job chạy đầu ngày)",
        "trigger": "Hàng ngày tại thời điểm đầu ngày (giờ cụ thể chưa được xác định)",
        "flow": "Quét toàn bộ bản ghi:\n• Nếu ngày T = Ngày hiệu lực và trạng thái = Không hoạt động → cập nhật thành Hoạt động.\n• Nếu ngày T > Ngày hết hiệu lực và trạng thái = Hoạt động → cập nhật thành Không hoạt động.",
        "post": "Trạng thái hoạt động của tất cả Nghiệp vụ/SPDV được đồng bộ mỗi ngày."
    },
]

for s in summaries:
    add_heading(doc, f"{s['id']}: {s['title']}", level=2, color="2E75B6")
    tbl = doc.add_table(rows=4, cols=2)
    tbl.style = 'Table Grid'
    tbl.columns[0].width = Cm(4)
    tbl.columns[1].width = Cm(12)
    labels = ["Actor", "Trigger", "Happy Flow", "Post-condition"]
    values = [s["actor"], s["trigger"], s["flow"], s["post"]]
    for i, (label, val) in enumerate(zip(labels, values)):
        row = tbl.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_bg(row.cells[0], "D9E1F2")
        row.cells[1].text = val
    doc.add_paragraph()

# ─── SECTION 2: Q&A TABLE ───
add_heading(doc, "PHẦN 2 — DANH SÁCH CÂU HỎI Q&A GỬI BA", level=1)

notice = doc.add_paragraph()
notice.add_run("🔴 Cao").bold = True
notice.add_run(": Không có câu trả lời thì không thể viết Test Case.   ")
notice.add_run("🟡 Trung bình").bold = True
notice.add_run(": Cần xác nhận trước Sprint TC.   ")
notice.add_run("🟠 Thấp").bold = True
notice.add_run(": Edge case, hỏi khi có thời gian.")

doc.add_paragraph()

BA_REPLY_PLACEHOLDER = "(BA điền câu trả lời vào đây)"

qa_items = [
    ("1","US01-A","Validation","Khi Mã Nghiệp vụ đạt ngưỡng >99, thông báo lỗi chính xác là gì? Hệ thống có gợi ý hành động tiếp theo không?","Gợi ý: Toast lỗi 'Mã Nghiệp vụ đã đạt giới hạn tối đa (99). Vui lòng liên hệ Quản trị viên.' và chặn lưu.","🔴 Cao",BA_REPLY_PLACEHOLDER),
    ("2","US01-A","Validation","Trường Tên Nghiệp vụ và Mô tả có giới hạn số ký tự tối đa không? Nhập vượt quá thì xử lý thế nào?","Gợi ý: Tên ≤ 100 ký tự, Mô tả ≤ 500 ký tự. Hiển thị counter ký tự còn lại, chặn nhập thêm khi đạt giới hạn.","🔴 Cao",BA_REPLY_PLACEHOLDER),
    ("3","US01-A","Validation","Ngày hiệu lực không được chọn ngày quá khứ — vậy ngày hôm nay (T) có được chọn không?","Gợi ý: Nên cho phép chọn ngày hôm nay vì theo logic: ngày T = Ngày hiệu lực → trạng thái Hoạt động.","🔴 Cao",BA_REPLY_PLACEHOLDER),
    ("4","US01-A","State Constraint","Khi bản ghi Nghiệp vụ đang ở trạng thái Chờ duyệt, Maker có được phép thêm SPDV chi tiết con vào Nghiệp vụ đó không?","Gợi ý: Không cho phép. Nghiệp vụ chưa được duyệt nên chưa tồn tại chính thức. Nên ẩn nút '+' hoặc báo lỗi rõ ràng.","🔴 Cao",BA_REPLY_PLACEHOLDER),
    ("5","US01-B","State Constraint","Khi Maker sửa và tạo bản Chờ duyệt mới, Checker thấy giao diện như thế nào? Có hiển thị so sánh giá trị cũ và mới không?","Gợi ý: Nên có màn hình diff hiển thị before/after (Ngày hiệu lực cũ → mới; Ngày hết hiệu lực cũ → mới) để Checker duyệt có căn cứ.","🔴 Cao",BA_REPLY_PLACEHOLDER),
    ("6","US01-B","State Constraint","Nếu bản ghi đã có 1 bản Chờ duyệt chưa được xử lý, Maker có được tạo thêm bản Chờ duyệt thứ 2 không?","Gợi ý: Chặn. Thông báo: 'Bản ghi đang có tác vụ chờ duyệt. Vui lòng chờ xử lý xong trước khi thực hiện thay đổi mới.'","🔴 Cao",BA_REPLY_PLACEHOLDER),
    ("7","US01-B","Validation","Khi sửa Ngày hiệu lực của Nghiệp vụ cấp cha mà vi phạm ràng buộc với cấp con, thông báo lỗi có liệt kê tên SPDV con vi phạm không?","Gợi ý: 'Ngày hiệu lực không hợp lệ. Các SPDV chi tiết [Tên 1, Tên 2] có Ngày hiệu lực sớm hơn. Vui lòng kiểm tra lại.'","🔴 Cao",BA_REPLY_PLACEHOLDER),
    ("8","US01-C","Validation","Mã SPDV chi tiết tự sinh theo Mã cha + 2 chữ số. Nếu số tự tăng >99, thông báo lỗi là gì?","Gợi ý: Báo rõ tên SPDV cấp cha bị đầy để BA/Admin biết cụm nào cần tái cấu trúc.","🔴 Cao",BA_REPLY_PLACEHOLDER),
    ("9","US01-C","Ambiguity","Cây SPDV có giới hạn số cấp con tối đa không? (VD: tối đa 3 cấp: Nghiệp vụ → SPDV cấp 2 → SPDV cấp 3...)","Gợi ý: Cần quy định rõ số cấp tối đa để QA viết TC boundary và Dev thiết kế DB đệ quy đúng.","🔴 Cao",BA_REPLY_PLACEHOLDER),
    ("10","US01-D","State Constraint","Khi Batch Job tự động chuyển Nghiệp vụ sang Không hoạt động, các SPDV chi tiết cấp con có bị cascade đổi trạng thái không?","Gợi ý: Nên cascade — nếu cấp cha Không hoạt động thì toàn bộ cây con cũng Không hoạt động, đảm bảo tính nhất quán.","🔴 Cao",BA_REPLY_PLACEHOLDER),
    ("11","US01-D","Integration","Batch Job chạy vào đầu ngày — cụ thể là mấy giờ? Nếu hệ thống downtime đúng lúc chạy, cơ chế retry như thế nào?","Gợi ý: Nên retry tự động sau 30 phút và ghi log cảnh báo gửi Admin nếu thất bại sau N lần retry.","🟡 Trung bình",BA_REPLY_PLACEHOLDER),
    ("12","US01-A","Authorization","Màn hình Danh mục SPDV có phân quyền theo đơn vị/chi nhánh không? Đơn vị khác có nhìn thấy SPDV của đơn vị khác không?","Gợi ý: Nếu là danh mục chung toàn hệ thống thì mọi đơn vị đều thấy (view-only đối với SPDV không thuộc đơn vị mình).","🟡 Trung bình",BA_REPLY_PLACEHOLDER),
    ("13","US01-A","UI Clarity","Lưới Danh mục SPDV hiển thị dạng cây (tree view) hay bảng phẳng? Có hỗ trợ collapse/expand từng nhánh không?","Gợi ý: Tree view có collapse/expand theo cấp. Cần mockup để QA viết case cho thao tác UI mở rộng/thu gọn cây.","🟡 Trung bình",BA_REPLY_PLACEHOLDER),
    ("14","US01-A","UI Clarity","Màn hình Danh mục SPDV có chức năng Tìm kiếm/Filter không? Tìm theo tiêu chí nào (Tên, Mã, Trạng thái)?","Gợi ý: Nên có filter theo Tên và Trạng thái ít nhất. Không có filter sẽ khó dùng khi cây có hàng trăm nút.","🟡 Trung bình",BA_REPLY_PLACEHOLDER),
    ("15","US01-B","UI Clarity","Sau khi Maker submit sửa, lưới hiển thị giá trị cũ hay mới của Ngày hiệu lực trong khi bản ghi đang Chờ duyệt?","Gợi ý: Hiển thị giá trị cũ (đang hiệu lực), kèm badge/cờ 'Đang chờ duyệt' để phân biệt trạng thái.","🟡 Trung bình",BA_REPLY_PLACEHOLDER),
    ("16","US01-A/B","UI Clarity","Nút Đóng trên form có popup xác nhận 'Bạn có chắc muốn đóng? Dữ liệu chưa lưu sẽ bị mất.' không hay đóng ngay lập tức?","Gợi ý: Nếu form đã có dữ liệu nhập (dirty state), nên bật popup xác nhận trước khi đóng để tránh mất dữ liệu.","🟡 Trung bình",BA_REPLY_PLACEHOLDER),
]

# Table header
tbl = doc.add_table(rows=1, cols=7)
tbl.style = 'Table Grid'
headers = ["STT", "US_ID", "Chiều phân tích", "Câu hỏi", "Đề xuất hướng xử lý", "Ưu tiên", "Phản hồi của BA"]
widths  = [Cm(0.9), Cm(1.6), Cm(2.5), Cm(5.5), Cm(4.8), Cm(1.4), Cm(4.3)]

hdr_row = tbl.rows[0]
for i, (h, w) in enumerate(zip(headers, widths)):
    cell = hdr_row.cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(cell, "1F3864")
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    cell.width = w

priority_colors = {"🔴 Cao": "FFE0E0", "🟡 Trung bình": "FFF8DC", "🟠 Thấp": "FFF0E0"}

for item in qa_items:
    row = tbl.add_row()
    values = list(item)
    for i, (val, w) in enumerate(zip(values, widths)):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(9.5)
        # Center STT and Priority columns
        if i == 0 or i == 5:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # BA response column: make text italic and light gray
        if i == 6:
            run.italic = True
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            set_cell_bg(cell, "FFFDE7")  # Light yellow background for BA to fill
        cell.width = w
    # Color priority cell
    priority = item[5]
    bg = priority_colors.get(priority, "FFFFFF")
    set_cell_bg(row.cells[5], bg)
    # Alternate row shading (only for non-BA-response cols)
    if qa_items.index(item) % 2 == 1:
        for j in range(5):
            set_cell_bg(row.cells[j], "F2F2F2")

doc.add_paragraph()

# ─── FOOTER NOTE ───
note = doc.add_paragraph()
note.add_run("📌 Ghi chú: ").bold = True
note.add_run("Cột 'Đề xuất hướng xử lý' mang tính tham khảo từ góc độ QA, không phải quyết định cuối cùng. "
             "BA/Dev hoàn toàn có quyền điều chỉnh theo thiết kế thực tế của hệ thống.")
note.runs[1].italic = True
note.runs[1].font.size = Pt(10)

doc.save(OUTPUT_PATH)
print(f"Saved: {OUTPUT_PATH}")
