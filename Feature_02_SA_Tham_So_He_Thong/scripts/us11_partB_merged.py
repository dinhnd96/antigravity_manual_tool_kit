"""Logic file — sinh US11_PartB_QA_Merged.docx từ data files."""
import sys, os
sys.path.insert(0, os.path.dirname(__file__))
from us11_partB_merged_data import QA_DATA
from us11_partB_merged_data2 import QA_DATA_2

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT

ALL_DATA = QA_DATA + QA_DATA_2

doc = Document()
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(9)

doc.add_heading("PHẦN B TỔNG HỢP: DANH SÁCH CẢNH BÁO & Q&A — US11", level=1)
doc.add_paragraph("Chương trình ưu đãi có đánh giá định kỳ KH — Merged (AI + VA)")

p = doc.add_paragraph()
p.add_run("Ghi chú cột 'Nguồn': ").bold = True
p.add_run("AI = Phát hiện bởi AI | VA = Phát hiện bởi Validation Analyst | BOTH = Cả 2 nguồn phát hiện")

HEADERS = ["ID", "Trích xuất (Reference)", "Câu hỏi / Sự cố", "Phân loại", "Đề xuất từ QA", "Nguồn", "Trả lời của BA"]
COL_WIDTHS = [Cm(2.0), Cm(3.8), Cm(7.5), Cm(1.8), Cm(6.5), Cm(1.2), Cm(4.4)]

# Build categories
categories = {
    "Nghiệp vụ": ("HM1 — 🔶 Vấn đề Nghiệp vụ / Luồng xử lý", []),
    "Giới hạn": ("HM2 — 🔴 Giới hạn hệ thống & Exception Handling", []),
    "Toàn vẹn dữ liệu": ("HM3 — 🟠 Toàn vẹn dữ liệu & Ràng buộc", []),
    "UI-UX": ("HM4 — 🔵 UI/UX & Giao diện", []),
}
for item in ALL_DATA:
    cat = item[3]
    categories[cat][1].append(item)

total = 0
for cat_key in ["Nghiệp vụ", "Giới hạn", "Toàn vẹn dữ liệu", "UI-UX"]:
    title, items = categories[cat_key]
    if not items:
        continue
    doc.add_heading(title, level=2)

    table = doc.add_table(rows=1 + len(items), cols=7, style='Light Grid Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for i, h in enumerate(HEADERS):
        cell = table.cell(0, i)
        cell.text = h
        cell.width = COL_WIDTHS[i]
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(8)

    for row_idx, (qid, ref, question, cat, suggestion, source) in enumerate(items, 1):
        vals = [qid, ref, question, cat, suggestion, source, ""]
        for col_idx, val in enumerate(vals):
            cell = table.cell(row_idx, col_idx)
            cell.text = val
            cell.width = COL_WIDTHS[col_idx]
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
    total += len(items)
    doc.add_paragraph("")

# Summary
doc.add_heading("Tổng kết", level=2)
summary = doc.add_paragraph()
summary.add_run(f"Tổng số câu hỏi: {total}\n").bold = True
summary.add_run(f"  • Phát hiện bởi AI: {sum(1 for d in ALL_DATA if d[5]=='AI')}\n")
summary.add_run(f"  • Phát hiện bởi VA: {sum(1 for d in ALL_DATA if d[5]=='VA')}\n")
summary.add_run(f"  • Phát hiện bởi cả 2: {sum(1 for d in ALL_DATA if d[5]=='BOTH')}\n")

output_path = "/Users/mac/antigravity-testing-kit/Feature_02_SA_Tham_So_He_Thong/output/US11_PartB_QA_Merged.docx"
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f"✅ Saved: {output_path}")
print(f"📊 Total Q&A: {total} (AI: {sum(1 for d in ALL_DATA if d[5]=='AI')}, VA: {sum(1 for d in ALL_DATA if d[5]=='VA')}, BOTH: {sum(1 for d in ALL_DATA if d[5]=='BOTH')})")
