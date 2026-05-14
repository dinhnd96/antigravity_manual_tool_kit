"""US38 Part A - Tóm tắt Nghiệp vụ Chuyên Sâu"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# --- Page setup: Landscape, narrow margins ---
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)

def add_heading_styled(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
    return h

# ========== TITLE ==========
title = doc.add_heading('US38 – Phân Tích Nghiệp Vụ (Part A)', level=1)
for run in title.runs:
    run.font.name = 'Times New Roman'

sub = doc.add_paragraph()
sub.add_run('Tính năng: ').bold = True
sub.add_run('Tự động thu phí bảo lãnh định kỳ đối với các khoản bảo lãnh chưa xác định thời hạn')
sub2 = doc.add_paragraph()
sub2.add_run('Feature: ').bold = True
sub2.add_run('SA – Tham Số Hệ Thống / Phí Bảo Lãnh Định Kỳ')
sub3 = doc.add_paragraph()
sub3.add_run('Phiên bản: ').bold = True
sub3.add_run('v1.0 – Phase 1 Analysis')

doc.add_paragraph('')

# ========== A.1 Core Business Value ==========
add_heading_styled(doc, 'A.1. Thông Điệp Cốt Lõi (Core Business Value)', level=2)

p = doc.add_paragraph()
p.add_run('Mục đích: ').bold = True
p.add_run('US38 quy định cơ chế thu phí bảo lãnh định kỳ cho các khoản bảo lãnh CHƯA XÁC ĐỊNH THỜI HẠN. '
           'Đây là loại cam kết bảo lãnh không có ngày chấm dứt cụ thể, hiệu lực kéo dài cho đến khi nghĩa vụ '
           'hoàn thành hoặc các bên chấm dứt thỏa thuận.')

p2 = doc.add_paragraph()
p2.add_run('Người dùng cuối: ').bold = True
p2.add_run('Hệ thống ProfiX (xử lý batch tự động) + User Maker/Checker trên T24 (thu thủ công).')

p3 = doc.add_paragraph()
p3.add_run('Đặc thù quan trọng: ').bold = True
p3.add_run('Khác với bảo lãnh có thời hạn (đã biết ngày kết thúc), bảo lãnh chưa xác định thời hạn '
            'thu phí 1 tháng/lần, lần đầu thu ngay khi phát hành. Khoảng thời gian tính phí kỳ đầu = '
            'số ngày của tháng hiện tại (VD: tháng 5 = 31 ngày).')

# ========== A.2 Flow Structure & Module Mapping ==========
add_heading_styled(doc, 'A.2. Cấu Trúc Luồng Nghiệp Vụ & Phân Bổ Module', level=2)

# --- Module 1 ---
add_heading_styled(doc, 'Module 1: Thu phí bảo lãnh lần đầu khi phát hành (Kỳ đầu tiên)', level=3)

p = doc.add_paragraph()
p.add_run('Luồng chính (Happy Path):').bold = True
doc.add_paragraph('1. User nhập thông tin khoản bảo lãnh chưa xác định thời hạn trên T24.', style='List Bullet')
doc.add_paragraph('2. T24 gọi API đến ProfiX yêu cầu tính phí cho kỳ thu phí đầu tiên.', style='List Bullet')
doc.add_paragraph('3. ProfiX tính phí theo cơ chế US34 (tự động tính phí trên kênh quầy/nội bộ).', style='List Bullet')
doc.add_paragraph('4. Khoảng thời gian tính phí = số ngày của tháng hiện tại (VD: phát hành ngày 07/05 → 31 ngày).', style='List Bullet')
doc.add_paragraph('5. Phí được hạch toán kèm theo giao dịch gốc phát hành bảo lãnh.', style='List Bullet')

p = doc.add_paragraph()
p.add_run('Luồng rẽ nhánh / Ngoại lệ:').bold = True
doc.add_paragraph('• Tham chiếu toàn bộ luồng lỗi/ngoại lệ của US34 (API timeout, tính phí thất bại, v.v.).', style='List Bullet')
doc.add_paragraph('• Tài liệu KHÔNG mô tả riêng luồng lỗi cho kỳ đầu tiên → mặc định áp dụng US34.', style='List Bullet')

# --- Module 2 ---
add_heading_styled(doc, 'Module 2: Đồng bộ dữ liệu bảo lãnh vào ProfiX (ETL)', level=3)

p = doc.add_paragraph()
p.add_run('Luồng chính (Happy Path):').bold = True
doc.add_paragraph('1. Sau khi ngân hàng khóa sổ cuối ngày, dữ liệu T-1 được đồng bộ về DWH.', style='List Bullet')
doc.add_paragraph('2. DWH thực hiện ETL dữ liệu bảo lãnh chưa xác định thời hạn vào bảng Tài khoản ProfiX.', style='List Bullet')
doc.add_paragraph('3. Dữ liệu bao gồm 17 trường (xem bảng chi tiết bên dưới).', style='List Bullet')
doc.add_paragraph('4. Loại tài khoản mặc định = "MD" (Tài khoản bảo lãnh), phân biệt với ACC (Casa) và LD (Tiền vay).', style='List Bullet')
doc.add_paragraph('5. Chỉ đồng bộ tài khoản còn hoạt động.', style='List Bullet')

p = doc.add_paragraph()
p.add_run('Luồng rẽ nhánh / Ngoại lệ:').bold = True
doc.add_paragraph('• ETL dữ liệu T-1 → Tham chiếu cơ chế đồng bộ tại US33.', style='List Bullet')
doc.add_paragraph('• Tài liệu ghi "Chi tiết bảng ETL sẽ được thống nhất trong quá trình triển khai" → chưa chốt.', style='List Bullet')

# --- Data Table ETL ---
add_heading_styled(doc, 'Bảng dữ liệu ETL Tối thiểu cần đồng bộ (17 trường):', level=4)

etl_table = doc.add_table(rows=18, cols=3)
etl_table.style = 'Table Grid'
headers = ['STT', 'Trường dữ liệu', 'Giá trị / Mô tả']
for i, h in enumerate(headers):
    cell = etl_table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(10)

data = [
    ('1', 'Loại tài khoản', 'Mặc định = MD (Tài khoản bảo lãnh)'),
    ('2', 'CIF', 'Mã khách hàng'),
    ('3', 'Số MD', 'Số tài khoản bảo lãnh'),
    ('4', 'Loại tiền', 'Currency code'),
    ('5', 'Sản phẩm bảo lãnh', 'Loại sản phẩm bảo lãnh'),
    ('6', 'Trạng thái tài khoản', 'Chỉ lấy tài khoản còn hoạt động'),
    ('7', 'Xác định thời hạn?', 'Có / Không'),
    ('8', 'Giá trị bảo lãnh', 'Số tiền bảo lãnh'),
    ('9', 'Có ký quỹ?', 'Có / Không'),
    ('10', 'Tỷ lệ ký quỹ', '% ký quỹ'),
    ('11', 'Giá trị ký quỹ', 'Số tiền ký quỹ'),
    ('12', 'Có TSBĐ?', 'Có / Không'),
    ('13', 'Loại TSBĐ', 'Tiền gửi PVCB / Tiền gửi khác / Tài sản khác'),
    ('14', 'Giá trị bảo đảm bằng TSBĐ', 'Số tiền bảo đảm'),
    ('15', 'Giá trị không có TSBĐ', 'Phần không có tài sản bảo đảm'),
    ('16', 'Tần suất thu phí', 'Tháng / Quý'),
    ('17', 'Thu phí tự động', 'Có / Không'),
]
for idx, (stt, field, val) in enumerate(data, 1):
    row = etl_table.rows[idx]
    row.cells[0].text = stt
    row.cells[1].text = field
    row.cells[2].text = val
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)

# --- Module 3 ---
add_heading_styled(doc, 'Module 3: Thu phí bảo lãnh định kỳ tự động (Batch Job)', level=3)

p = doc.add_paragraph()
p.add_run('Luồng chính (Happy Path):').bold = True
doc.add_paragraph('1. Hằng ngày, đầu ngày, ProfiX tự động kiểm tra các khoản bảo lãnh chưa xác định thời hạn đến hạn thu phí.', style='List Bullet')
doc.add_paragraph('2. Sinh dữ liệu phí bảo lãnh định kỳ đến hạn → tham chiếu US35 (thu phí định kỳ theo lịch).', style='List Bullet')
doc.add_paragraph('3. Kiểm tra tham số "Thu phí tự động":', style='List Bullet')
doc.add_paragraph('   → Nếu Thu phí tự động = Có: tiếp tục xử lý thu phí theo US35.', style='List Bullet')
doc.add_paragraph('   → Nếu Thu phí tự động = Không: hệ thống KHÔNG xử lý gì tiếp theo.', style='List Bullet')

p = doc.add_paragraph()
p.add_run('Luồng rẽ nhánh / Ngoại lệ:').bold = True
doc.add_paragraph('• Toàn bộ luồng lỗi khi thu tự động → tham chiếu US35.', style='List Bullet')
doc.add_paragraph('• Phí chưa thu thành công → chuyển sang luồng thu thủ công (Module 4).', style='List Bullet')

# --- Module 4 ---
add_heading_styled(doc, 'Module 4: Thu phí bảo lãnh thủ công (Manual Collection on T24)', level=3)

p = doc.add_paragraph()
p.add_run('Luồng chính (Happy Path) — Flowchart 2:').bold = True
doc.add_paragraph('1. User Maker khởi tạo giao dịch thu phí bảo lãnh định kỳ thủ công trên T24.', style='List Bullet')
doc.add_paragraph('2. T24 call API đến ProfiX → lấy danh sách các khoản phí định kỳ đến hạn.', style='List Bullet')
doc.add_paragraph('3. ProfiX trả ra danh sách các khoản định kỳ đến hạn.', style='List Bullet')
doc.add_paragraph('4. T24 hiển thị danh sách các kỳ thu phí bảo lãnh định kỳ đến hạn.', style='List Bullet')
doc.add_paragraph('5. User Maker lựa chọn các kỳ cần thu.', style='List Bullet')
doc.add_paragraph('6. Khi Maker commit giao dịch, T24 call API lại ProfiX → lấy LẠI danh sách khoản phí đến hạn.', style='List Bullet')
doc.add_paragraph('7. ProfiX trả lại danh sách.', style='List Bullet')
doc.add_paragraph('8. T24 kiểm tra lại trạng thái thanh toán của các kỳ đã chọn:', style='List Bullet')
doc.add_paragraph('   → Hợp lệ: chuyển bước 9.', style='List Bullet')
doc.add_paragraph('   → Không hợp lệ: T24 hiển thị thông báo lỗi (bước 8.1).', style='List Bullet')
doc.add_paragraph('9. User Checker phê duyệt giao dịch.', style='List Bullet')
doc.add_paragraph('10. T24 xử lý hạch toán, call API ProfiX ghi nhận thông tin thu phí thành công.', style='List Bullet')
doc.add_paragraph('11. ProfiX lưu giao dịch thu phí thành công.', style='List Bullet')

p = doc.add_paragraph()
p.add_run('Luồng rẽ nhánh / Ngoại lệ:').bold = True
doc.add_paragraph('• Bước 8 → Trạng thái không hợp lệ → T24 hiển thị thông báo lỗi (bước 8.1).', style='List Bullet')
doc.add_paragraph('• Flowchart thiếu nhánh: Checker từ chối phê duyệt → chưa được mô tả.', style='List Bullet')
doc.add_paragraph('• Flowchart thiếu: ProfiX trả lỗi API ở bước 2, 6, 10 → chưa được mô tả.', style='List Bullet')

# ========== A.3 Pre-conditions ==========
add_heading_styled(doc, 'A.3. Điều Kiện Tiên Quyết & Cấu Hình', level=2)

t = doc.add_table(rows=7, cols=2)
t.style = 'Table Grid'
t.rows[0].cells[0].text = 'Điều kiện'
t.rows[0].cells[1].text = 'Chi tiết'
for p in t.rows[0].cells[0].paragraphs:
    for r in p.runs: r.bold = True
for p in t.rows[0].cells[1].paragraphs:
    for r in p.runs: r.bold = True

precond = [
    ('Dữ liệu ETL T-1', 'Bảng dữ liệu Tài khoản trên ProfiX đã được đồng bộ từ DWH (khóa sổ cuối ngày xong).'),
    ('Code phí bảo lãnh', 'Đã khai báo Code phí phù hợp với sản phẩm bảo lãnh (tham chiếu US34).'),
    ('Biểu phí bảo lãnh', 'Đã cấu hình Biểu phí với điều kiện tính phí mapping đúng dữ liệu ETL (tham chiếu US34).'),
    ('Tham số "Thu phí tự động"', 'Cờ Có/Không trên từng khoản bảo lãnh, quyết định hệ thống có tự động thu hay chờ thu thủ công.'),
    ('Tham số "Tần suất thu phí"', 'Tháng hoặc Quý – quyết định chu kỳ thu phí định kỳ.'),
    ('Kết nối API T24 ↔ ProfiX', 'API kết nối giữa T24 và ProfiX phải hoạt động (cho cả thu tự động lần đầu và thu thủ công).'),
]
for i, (cond, detail) in enumerate(precond, 1):
    t.rows[i].cells[0].text = cond
    t.rows[i].cells[1].text = detail

# ========== A.4 QTC Rules ==========
add_heading_styled(doc, 'A.4. Quy Tắc Chung Áp Dụng', level=2)

qtc_items = [
    'QTC-11: Xử lý lỗi FE-First – FE chặn lỗi đầu tiên, BE trả mã lỗi kỹ thuật nếu FE chưa chặn.',
    'QTC-12: Luồng Maker-Checker – Áp dụng cho thu phí thủ công trên T24 (Maker commit → Checker phê duyệt).',
    'QTC-14.1: No-Change Guard – Nếu Maker commit nhưng không có thay đổi → FE chặn.',
    'QTC-01.4: Định dạng Number – Giá trị bảo lãnh, phí hiển thị theo format số ProfiX.',
    'QTC-01.5: Định dạng Date – Ngày thu phí, kỳ thu phí hiển thị dd/mm/yyyy.',
]

for item in qtc_items:
    doc.add_paragraph(item, style='List Bullet')

p_note = doc.add_paragraph()
p_note.add_run('Lưu ý: ').bold = True
p_note.add_run('US38 có giao diện = N/A và mô tả chi tiết các trường = N/A. Toàn bộ giao diện thu thủ công '
               'nằm trên T24 (hệ thống Core Banking), ProfiX chỉ đóng vai trò cung cấp API và xử lý backend. '
               'Không có màn hình ProfiX nào cần test UI trực tiếp trong US38.')

# ========== A.5 Dependency Map ==========
add_heading_styled(doc, 'A.5. Bản Đồ Phụ Thuộc (Dependency Map)', level=2)

dep_table = doc.add_table(rows=4, cols=3)
dep_table.style = 'Table Grid'
dep_table.rows[0].cells[0].text = 'US tham chiếu'
dep_table.rows[0].cells[1].text = 'Mối liên hệ'
dep_table.rows[0].cells[2].text = 'Ảnh hưởng'
for cell in dep_table.rows[0].cells:
    for p in cell.paragraphs:
        for r in p.runs: r.bold = True

deps = [
    ('US34', 'Tính phí trên kênh quầy/nội bộ', 'Thu phí lần đầu khi phát hành bảo lãnh áp dụng toàn bộ logic US34. Khoảng thời gian tính phí = số ngày tháng hiện tại.'),
    ('US33', 'Đồng bộ dữ liệu nguồn tính phí', 'Dữ liệu bảo lãnh chưa xác định thời hạn được đồng bộ qua ETL vào bảng Tài khoản ProfiX theo cơ chế US33.'),
    ('US35', 'Thu phí định kỳ theo lịch', 'Các kỳ thu phí tiếp theo (sau lần đầu) áp dụng toàn bộ quy trình US35.'),
]
for i, (us, rel, impact) in enumerate(deps, 1):
    dep_table.rows[i].cells[0].text = us
    dep_table.rows[i].cells[1].text = rel
    dep_table.rows[i].cells[2].text = impact

# ========== SAVE ==========
output_dir = '/Users/mac/antigravity-testing-kit/Feature_02_SA_Tham_So_He_Thong/output'
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, 'US38_PartA_Summary.docx')
doc.save(output_path)
print(f'✅ Part A saved: {output_path}')
