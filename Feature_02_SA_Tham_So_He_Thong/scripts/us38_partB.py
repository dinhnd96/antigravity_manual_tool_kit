"""US38 Part B - Danh Sách Cảnh Báo & Q&A (Loopholes Discovery)"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# --- Page setup: Landscape, narrow margins ---
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(10)

def set_cell_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = tc.makeelement(qn('w:tcW'), {})
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(int(width_cm * 567)))
    tcW.set(qn('w:type'), 'dxa')

def add_heading_styled(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
    return h

# ========== TITLE ==========
title = doc.add_heading('US38 – Danh Sách Cảnh Báo & Q&A (Part B)', level=1)
for run in title.runs:
    run.font.name = 'Times New Roman'

sub = doc.add_paragraph()
sub.add_run('Tính năng: ').bold = True
sub.add_run('Tự động thu phí bảo lãnh định kỳ – Bảo lãnh chưa xác định thời hạn')
sub2 = doc.add_paragraph()
sub2.add_run('Mục đích: ').bold = True
sub2.add_run('Khai quật các điểm thiếu sót, luồng rẽ nhánh chưa rõ, hoặc mâu thuẫn trong tài liệu US38 để BA giải đáp trước khi viết Test Case.')

doc.add_paragraph('')

# ========== Q&A TABLE ==========
# Columns: ID | Trích xuất | Câu hỏi / Sự cố | Phân loại | Đề xuất từ QA | Trả lời của BA
# Widths (cm): 2.5 | 4.0 | 6.5 | 2.0 | 6.5 | 5.5 = 27 cm total
col_widths = [2.5, 4.0, 6.5, 2.0, 6.5, 5.5]

qa_data = [
    # ===== HẠNG MỤC 1: Nghiệp vụ / Luồng xử lý =====
    ('US38-QA-01.1',
     'Mục "Yêu cầu nghiệp vụ", đoạn mô tả khoảng thời gian tính phí kỳ đầu',
     'Tài liệu ghi "khoảng thời gian tính phí gửi cho ProfiX là số ngày của tháng hiện tại" (VD: ngày 07/05 → 31 ngày). Điều này có nghĩa phí kỳ đầu luôn tính cho NGUYÊN THÁNG bất kể ngày phát hành nằm giữa tháng? Hay tính theo số ngày thực tế từ ngày phát hành đến cuối tháng (VD: 07/05 → 31/05 = 25 ngày)?',
     'Nghiệp vụ',
     'Đề xuất: BA cần làm rõ "số ngày của tháng hiện tại" là (a) tổng số ngày lịch của tháng (31 ngày) hay (b) số ngày còn lại trong tháng từ ngày phát hành. Nếu là (a), khách hàng sẽ bị tính phí nhiều hơn thực tế sử dụng trong kỳ đầu.',
     ''),

    ('US38-QA-01.2',
     'Mục "Yêu cầu nghiệp vụ", đoạn mô tả tham số "Tần suất thu phí"',
     'Tài liệu ghi "Tần suất thu phí = Tháng/Quý", nhưng toàn bộ phần yêu cầu nghiệp vụ chỉ mô tả logic thu phí 1 tháng/lần. Nếu Tần suất = Quý thì khoảng thời gian tính phí kỳ đầu là bao nhiêu? Và các kỳ thu định kỳ tiếp theo tính theo quý (3 tháng) hay vẫn theo tháng?',
     'Nghiệp vụ',
     'Đề xuất: BA cần làm rõ logic khi Tần suất = Quý. Cụ thể: (1) Kỳ đầu tính phí bao nhiêu ngày? (2) Các kỳ tiếp theo thu vào ngày nào? (3) Nếu phát hành giữa quý, kỳ đầu có tính pro-rata không?',
     ''),

    ('US38-QA-01.3',
     'Flowchart 2 "Thu phí bảo lãnh thủ công", bước 6-7',
     'Khi Maker commit giao dịch, T24 call API lại ProfiX lấy LẠI danh sách khoản phí đến hạn (bước 6-7). Sau đó bước 8 kiểm tra trạng thái. Câu hỏi: Trạng thái "không hợp lệ" cụ thể là gì? Ví dụ: kỳ thu phí đã được thanh toán bởi batch job tự động (race condition giữa thu tự động và thu thủ công)?',
     'Nghiệp vụ',
     'Đề xuất: BA cần liệt kê các trạng thái thanh toán hợp lệ/không hợp lệ của kỳ thu phí. Đặc biệt cần xác định: nếu giữa lúc Maker chọn kỳ (bước 5) và commit (bước 6), batch job tự động đã thu thành công kỳ đó → T24 xử lý thế nào?',
     ''),

    ('US38-QA-01.4',
     'Flowchart 2 "Thu phí bảo lãnh thủ công", sau bước 8.1',
     'Khi T24 hiển thị thông báo lỗi (bước 8.1 – trạng thái không hợp lệ), Flowchart KHÔNG vẽ nhánh tiếp theo. Maker có được quay lại chọn kỳ khác không? Hay giao dịch bị hủy hoàn toàn?',
     'Nghiệp vụ',
     'Đề xuất: Flowchart cần bổ sung nhánh sau bước 8.1: (a) Maker có thể chọn lại các kỳ hợp lệ và commit lại, HOẶC (b) Giao dịch bị rollback, Maker phải tạo lại từ đầu.',
     ''),

    ('US38-QA-01.5',
     'Flowchart 2 "Thu phí bảo lãnh thủ công", bước 9',
     'Tài liệu ghi "User Checker phê duyệt giao dịch" nhưng không mô tả luồng khi Checker TỪ CHỐI phê duyệt. Khi từ chối: (a) Giao dịch về trạng thái nào? (b) Maker có thể chỉnh sửa và submit lại không? (c) Các kỳ thu phí đã chọn có bị lock trong thời gian chờ duyệt không?',
     'Nghiệp vụ',
     'Đề xuất: BA bổ sung luồng Checker từ chối vào Flowchart. Cần xác định rõ: trạng thái giao dịch sau từ chối, quyền Maker chỉnh sửa lại, và cơ chế lock/unlock kỳ thu phí.',
     ''),

    ('US38-QA-01.6',
     'Mục "Yêu cầu nghiệp vụ", đoạn thu phí kỳ đầu tiên',
     'Tài liệu ghi thu phí kỳ đầu "tương tự US34". US34 mô tả tính phí trên kênh quầy/nội bộ. Khoản bảo lãnh chưa xác định thời hạn có được phép miễn/giảm phí theo Chương trình ưu đãi (CTƯĐ) không? Nếu có, CTƯĐ áp dụng cho kỳ đầu có áp dụng tiếp cho các kỳ định kỳ sau không?',
     'Nghiệp vụ',
     'Đề xuất: BA xác nhận bảo lãnh chưa xác định thời hạn có nằm trong phạm vi áp dụng CTƯĐ hay không, và nếu có thì CTƯĐ có hiệu lực xuyên suốt các kỳ thu phí hay chỉ kỳ đầu.',
     ''),

    ('US38-QA-01.7',
     'Mục "Yêu cầu nghiệp vụ", đoạn "Maker lựa chọn các kỳ cần thu"',
     'Khi thu thủ công, Maker có thể chọn nhiều kỳ cùng lúc. Câu hỏi: (a) Nếu chọn 3 kỳ nhưng bước 8 phát hiện 1 kỳ không hợp lệ → cả 3 kỳ đều bị rollback hay chỉ reject kỳ không hợp lệ? (b) Phí từng kỳ có thể khác nhau (do thay đổi biểu phí giữa các kỳ) – T24 hiển thị chi tiết phí từng kỳ không?',
     'Nghiệp vụ',
     'Đề xuất: BA xác nhận cơ chế xử lý: (a) All-or-nothing hay partial success khi chọn nhiều kỳ? (b) Maker có thấy breakdown phí từng kỳ trước khi commit không?',
     ''),

    # ===== HẠNG MỤC 2: Giới hạn hệ thống & Exception Handling =====
    ('US38-QA-02.1',
     'Flowchart 1 "Đồng bộ dữ liệu", bước 1',
     'ETL dữ liệu T-1 chạy "sau khi ngân hàng hoàn tất khóa sổ cuối ngày". Nếu ETL thất bại hoặc bị timeout giữa chừng, ProfiX sẽ dùng dữ liệu cũ (T-2) để tính phí? Có cơ chế retry ETL không? Có cảnh báo khi dữ liệu ETL lỗi thời không?',
     'Giới hạn',
     'Đề xuất: BA/SA cần xác định cơ chế fallback khi ETL fail: (a) ProfiX tạm dừng batch thu phí cho đến khi ETL thành công, HOẶC (b) Dùng dữ liệu T-2 kèm cảnh báo. Cần có monitoring/alert cho DWH team.',
     ''),

    ('US38-QA-02.2',
     'Flowchart 2 "Thu phí bảo lãnh thủ công", bước 2 và 6',
     'T24 call API ProfiX 2 lần: lần 1 để lấy danh sách (bước 2), lần 2 khi Maker commit (bước 6). Nếu API ProfiX timeout/down ở bước 2 hoặc bước 6, T24 xử lý thế nào? Có cơ chế retry hay báo lỗi và rollback giao dịch?',
     'Giới hạn',
     'Đề xuất: BA/SA cần định nghĩa timeout và cơ chế xử lý cho từng lần call API. Đặc biệt bước 10 (hạch toán + call API ghi nhận thành công): nếu T24 hạch toán thành công nhưng API ProfiX ghi nhận thất bại → dữ liệu bất đồng bộ.',
     ''),

    ('US38-QA-02.3',
     'Flowchart 2 "Thu phí bảo lãnh thủ công", bước 10',
     'Bước 10: "T24 xử lý hạch toán, call API đến ProfiX để ghi nhận thông tin thu phí thành công". Đây là 2 action trong 1 bước. Nếu hạch toán trên T24 thành công NHƯNG call API ProfiX thất bại → T24 đã trừ tiền KH nhưng ProfiX chưa ghi nhận → nguy cơ thu phí trùng ở kỳ sau. Cơ chế rollback/compensation thế nào?',
     'Giới hạn',
     'Đề xuất: BA/SA cần xác định cơ chế đảm bảo tính nhất quán (consistency): (a) T24 chỉ hạch toán SAU KHI ProfiX confirm, HOẶC (b) Có luồng compensation (hoàn tiền/đánh dấu) nếu API fail sau hạch toán.',
     ''),

    ('US38-QA-02.4',
     'Mục "Yêu cầu nghiệp vụ", đoạn batch job đầu ngày',
     'Batch job kiểm tra "đến hạn thu phí bảo lãnh định kỳ trong ngày". Nếu batch job chạy lâu hoặc fail giữa chừng (đã xử lý 50/100 khoản bảo lãnh), cơ chế resume/retry thế nào? Có xử lý idempotent (tránh thu trùng khi retry)?',
     'Giới hạn',
     'Đề xuất: Tham chiếu US35 về cơ chế retry batch. BA xác nhận batch US38 cũng áp dụng cùng cơ chế, bao gồm idempotency check.',
     ''),

    # ===== HẠNG MỤC 3: Toàn vẹn dữ liệu & Ràng buộc =====
    ('US38-QA-03.1',
     'Mục "Yêu cầu nghiệp vụ", danh sách trường ETL, trường "Xác định thời hạn? = Có/Không"',
     'Bảng ETL có trường "Xác định thời hạn? = Có/Không". US38 chỉ xử lý bảo lãnh "chưa xác định thời hạn" (= Không). Câu hỏi: Bảng ETL có đồng bộ CẢ HAI loại (Có và Không)? Nếu có, ProfiX dùng trường này để filter, hay DWH chỉ gửi loại "Không"?',
     'Toàn vẹn dữ liệu',
     'Đề xuất: BA xác nhận scope dữ liệu ETL: (a) DWH chỉ gửi bảo lãnh "Không xác định thời hạn" → ProfiX không cần filter, HOẶC (b) DWH gửi tất cả → ProfiX cần filter theo trường "Xác định thời hạn = Không". Nếu (b), cần test case kiểm tra ProfiX không thu phí nhầm bảo lãnh có thời hạn.',
     ''),

    ('US38-QA-03.2',
     'Mục "Yêu cầu nghiệp vụ", danh sách trường ETL, trường "Trạng thái tài khoản"',
     'Tài liệu ghi "mặc định chỉ lấy các tài khoản còn hoạt động". Nếu khoản bảo lãnh bị tất toán/đóng giữa kỳ (sau khi ETL T-1 đã đồng bộ), batch job đầu ngày sẽ vẫn thu phí cho khoản đã đóng? Có cơ chế kiểm tra trạng thái realtime trước khi thu phí không?',
     'Toàn vẹn dữ liệu',
     'Đề xuất: BA xác nhận: (a) ProfiX chỉ dựa vào dữ liệu ETL T-1 (có thể trễ 1 ngày) → chấp nhận rủi ro thu phí khoản đã tất toán trong ngày, HOẶC (b) ProfiX call API T24 kiểm tra trạng thái realtime trước khi thu.',
     ''),

    ('US38-QA-03.3',
     'Mục "Yêu cầu nghiệp vụ", danh sách trường ETL, trường "Giá trị bảo lãnh"',
     'Giá trị bảo lãnh là cơ sở tính phí. Nếu giá trị bảo lãnh thay đổi giữa các kỳ (tăng/giảm bảo lãnh), ProfiX tính phí dựa trên giá trị nào? Giá trị tại thời điểm ETL T-1 (snapshot cuối ngày hôm trước) hay giá trị tại thời điểm thu phí?',
     'Toàn vẹn dữ liệu',
     'Đề xuất: BA xác nhận cơ sở tính phí: (a) Giá trị bảo lãnh tại snapshot ETL T-1 (đơn giản, nhất quán), HOẶC (b) Giá trị realtime tại thời điểm tính phí (phức tạp hơn nhưng chính xác hơn).',
     ''),

    ('US38-QA-03.4',
     'Mục "Yêu cầu nghiệp vụ", đoạn thu phí thủ công và thu phí tự động',
     'Khi thu thủ công thành công (Flowchart 2, bước 11), ProfiX lưu giao dịch thu phí. Câu hỏi: Cơ chế đánh dấu kỳ đã thu thành công ra sao? Batch job chạy sau đó có "nhìn thấy" kỳ này đã thu và bỏ qua không? Ngược lại, nếu batch tự động thu trước rồi, thu thủ công có filter ra kỳ đó không?',
     'Toàn vẹn dữ liệu',
     'Đề xuất: BA xác nhận cơ chế quản lý trạng thái kỳ thu phí để tránh thu trùng: ProfiX phải có flag/status trên mỗi kỳ (VD: Chưa thu / Đang thu / Đã thu thành công / Thu thất bại) và cả thu tự động lẫn thu thủ công đều check flag này.',
     ''),

    ('US38-QA-03.5',
     'Mục "Yêu cầu nghiệp vụ", danh sách trường ETL, trường "Thu phí tự động = Có/Không"',
     'Tham số "Thu phí tự động" nằm trên dữ liệu ETL. Nếu tham số này bị thay đổi từ "Có" sang "Không" (hoặc ngược lại) giữa các kỳ, hệ thống áp dụng giá trị nào? Giá trị tại thời điểm ETL T-1 hay giá trị tại thời điểm batch job chạy?',
     'Toàn vẹn dữ liệu',
     'Đề xuất: BA xác nhận nguồn dữ liệu quyết định "Thu phí tự động": dữ liệu ETL (snapshot T-1) hay ProfiX call realtime về T24. Nếu dùng ETL, cần chấp nhận trễ 1 ngày khi thay đổi tham số.',
     ''),

    # ===== HẠNG MỤC 4: UI/UX & Giao diện =====
    ('US38-QA-04.1',
     'Mục "Giao diện" và "Mô tả chi tiết các trường"',
     'Cả 2 mục đều ghi "N/A". US38 không có màn hình ProfiX. Tuy nhiên, luồng thu thủ công (Flowchart 2) diễn ra trên T24 với các tương tác UI: hiển thị danh sách kỳ (bước 4), chọn kỳ (bước 5), hiển thị lỗi (bước 8.1). BA xác nhận: (a) ProfiX team không cần test giao diện T24? (b) Specification giao diện T24 cho thu phí bảo lãnh thủ công nằm ở tài liệu nào?',
     'UI-UX',
     'Đề xuất: BA cần xác nhận rõ phạm vi test của ProfiX team: chỉ test API response hay bao gồm cả kiểm tra giao diện T24? Nếu chỉ test API, cần có API specification chi tiết (request/response schema).',
     ''),

    ('US38-QA-04.2',
     'Flowchart 2 "Thu phí bảo lãnh thủ công", bước 4',
     'T24 hiển thị danh sách các kỳ thu phí bảo lãnh định kỳ đến hạn. Câu hỏi: Danh sách này bao gồm những cột/thông tin gì? (VD: Kỳ, Ngày đến hạn, Số tiền phí, Trạng thái, Loại tiền...) Tài liệu không mô tả chi tiết cấu trúc dữ liệu API response.',
     'UI-UX',
     'Đề xuất: BA/SA cần bổ sung API specification cho endpoint "lấy danh sách khoản phí định kỳ đến hạn", bao gồm: request parameters (CIF, Số MD, Loại bảo lãnh...) và response fields (danh sách kỳ với chi tiết từng kỳ).',
     ''),
]

# --- Build table ---
add_heading_styled(doc, 'Bảng Q&A – US38', level=2)

table = doc.add_table(rows=1 + len(qa_data), cols=6)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['ID', 'Trích xuất (Reference)', 'Câu hỏi / Sự cố', 'Phân loại', 'Đề xuất từ QA', 'Trả lời của BA']

# Header row
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    set_cell_width(cell, col_widths[i])
    for p in cell.paragraphs:
        p.alignment = 1  # CENTER
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9)
            r.font.name = 'Times New Roman'

# Data rows
for row_idx, (qid, ref, question, cat, proposal, ba_answer) in enumerate(qa_data, 1):
    row = table.rows[row_idx]
    values = [qid, ref, question, cat, proposal, ba_answer]
    for col_idx, val in enumerate(values):
        cell = row.cells[col_idx]
        cell.text = val
        set_cell_width(cell, col_widths[col_idx])
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
                r.font.name = 'Times New Roman'

# ========== Summary stats ==========
doc.add_paragraph('')
summary = doc.add_paragraph()
summary.add_run('Thống kê Q&A:').bold = True
doc.add_paragraph(f'• Tổng số câu hỏi: {len(qa_data)}', style='List Bullet')
cat_count = {}
for item in qa_data:
    c = item[3]
    cat_count[c] = cat_count.get(c, 0) + 1
labels = {'Nghiệp vụ': '🔶 Hạng mục 1 – Nghiệp vụ/Luồng xử lý',
          'Giới hạn': '🔴 Hạng mục 2 – Giới hạn hệ thống & Exception',
          'Toàn vẹn dữ liệu': '🟠 Hạng mục 3 – Toàn vẹn dữ liệu & Ràng buộc',
          'UI-UX': '🔵 Hạng mục 4 – UI/UX & Giao diện'}
for cat, label in labels.items():
    count = cat_count.get(cat, 0)
    doc.add_paragraph(f'• {label}: {count} câu hỏi', style='List Bullet')

doc.add_paragraph('')
note = doc.add_paragraph()
note.add_run('⚠️ Lưu ý: ').bold = True
note.add_run('Cột "Trả lời của BA" hiện đang để trống. Vui lòng chuyển file này cho BA điền đáp án. '
             'Sau khi nhận câu trả lời, AI sẽ tiến hành Phase 2 – Tổng hợp Kịch bản Test Case.')

# ========== SAVE ==========
output_dir = '/Users/mac/antigravity-testing-kit/Feature_02_SA_Tham_So_He_Thong/output'
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, 'US38_PartB_QA.docx')
doc.save(output_path)
print(f'✅ Part B saved: {output_path}')
