# US39 Part B - Q&A Generator
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "output")
os.makedirs(OUTPUT_DIR, exist_ok=True)

doc = Document()

# Page setup: Landscape, narrow margins
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)

title = doc.add_heading('US39 — Phần B: Danh Sách Cảnh Báo & Q&A', level=1)
doc.add_paragraph('Feature: SA — Tham Số Hệ Thống')
doc.add_paragraph('User Story: US39 — Tính phí trả nợ trước hạn trên kênh offline (T24)')
doc.add_paragraph('Dành cho: BA — Giải đáp các điểm mù, mâu thuẫn trong tài liệu.')

# ===== Q&A DATA =====
# Format: [ID, Reference, Issue, Category, Proposal, BA_Answer]
qa_data = [
    # === HẠNG MỤC 1: NGHIỆP VỤ / LUỒNG XỬ LÝ ===
    [
        "US39-QA-01.1",
        "Mục \"Yêu cầu nghiệp vụ\", đoạn \"T24 sẽ gọi API để yêu cầu ProfiX tính phí tự động cho giao dịch\"",
        "Tài liệu tham chiếu US33 cho logic tính phí, nhưng US33 mô tả kênh online. Với kênh offline (T24), request/response API tính phí có khác gì so với kênh online không? Cụ thể: (1) Các tham số đầu vào API có giống nhau? (2) Kênh offline có cần truyền thêm tham số nào đặc thù (VD: mã chi nhánh, mã user T24)?",
        "Nghiệp vụ",
        "Đề xuất: BA làm rõ API spec cho kênh offline T24. Nếu hoàn toàn giống US33, cần ghi rõ \"API input/output tương tự US33, không có tham số đặc thù kênh offline\".",
        ""
    ],
    [
        "US39-QA-01.2",
        "Lưu đồ, Bước 6 (Decision): \"User thực hiện?\"",
        "Flowchart chỉ vẽ 2 nhánh: (1) Sửa số tiền phí → commit, (2) Nhấn commit trực tiếp. Thiếu nhánh: User KHÔNG muốn tiếp tục → Hủy giao dịch. Khi user hủy giao dịch tại bước này, ProfiX có cần nhận thông báo gì không? Hay chỉ T24 tự xử lý hủy?",
        "Nghiệp vụ",
        "Đề xuất: Bổ sung nhánh \"Hủy giao dịch\" tại bước 6 trên Flowchart. Nếu hủy chỉ ở phía T24 (không call API ProfiX), cần ghi rõ.",
        ""
    ],
    [
        "US39-QA-01.3",
        "Mục \"Yêu cầu nghiệp vụ\", đoạn \"T24 kiểm tra số tiền phí người dùng nhập vào có nằm ngoài khoảng Số tiền phí tối thiểu/tối đa (nếu có) theo Code phí hay không\"",
        "Tài liệu ghi \"nếu có\" cho min/max. Trường hợp Code phí KHÔNG khai báo min/max: (1) User được sửa số tiền phí thành BẤT KỲ giá trị nào (kể cả 0 hoặc số âm)? (2) Hay T24 chặn giá trị <= 0?",
        "Nghiệp vụ",
        "Đề xuất: Khi Code phí không có min/max, T24 vẫn phải validate: số tiền phí > 0 và không vượt quá số tiền trả nợ trước hạn.",
        ""
    ],
    [
        "US39-QA-01.4",
        "Lưu đồ, Bước 7→8→9",
        "Flowchart chỉ vẽ luồng Checker phê duyệt thành công (happy path). Thiếu hoàn toàn nhánh Checker TỪ CHỐI giao dịch. Khi Checker từ chối: (1) Giao dịch quay về Maker để sửa lại? (2) Giao dịch bị hủy hoàn toàn? (3) ProfiX có cần nhận thông báo từ chối không?",
        "Nghiệp vụ",
        "Đề xuất: Bổ sung nhánh từ chối tại bước 7 trên Flowchart. Giao dịch bị từ chối → Maker có thể sửa và commit lại → ProfiX không nhận thông báo cho đến khi giao dịch được duyệt thành công.",
        ""
    ],
    [
        "US39-QA-01.5",
        "Mục \"Yêu cầu nghiệp vụ\", đoạn \"Phí trả nợ trước hạn theo Chương trình ưu đãi cho vay\"",
        "Tài liệu mô tả: Code phí gom nhóm các bậc thang có cùng THKV + tỷ lệ phí, với điều kiện Mã hạch toán tương ứng. Trường hợp 1 khoản vay (LD) có Mã hạch toán match với NHIỀU Code phí trong cùng Biểu phí CTƯĐ (VD: do khai báo chồng chéo điều kiện THKV), ProfiX xử lý thế nào? Lấy Code phí nào?",
        "Nghiệp vụ",
        "Đề xuất: BA làm rõ quy tắc ưu tiên khi match nhiều Code phí. Nếu không cho phép chồng chéo, cần có validate khi khai báo Code phí.",
        ""
    ],
    [
        "US39-QA-01.6",
        "Mục \"Yêu cầu nghiệp vụ\", đoạn \"Sau khi người dùng nhấn commit giao dịch, T24 sẽ chuyển giao dịch sang bước phê duyệt\"",
        "Khoảng thời gian giữa Maker commit (bước 6) và Checker duyệt (bước 7) có thể kéo dài. Trong khoảng thời gian này, nếu dữ liệu ETL được cập nhật (T-1 mới) dẫn đến THKV thay đổi → phí tính lại có thể khác. ProfiX có tính lại phí tại thời điểm Checker duyệt không? Hay dùng phí đã tính ban đầu?",
        "Nghiệp vụ",
        "Đề xuất: ProfiX KHÔNG tính lại phí khi Checker duyệt. Phí đã chốt tại thời điểm Maker commit. Tuy nhiên cần BA xác nhận.",
        ""
    ],

    # === HẠNG MỤC 2: GIỚI HẠN HỆ THỐNG & EXCEPTION ===
    [
        "US39-QA-02.1",
        "Lưu đồ, Bước 3→4 (T24 call API → ProfiX tính phí)",
        "Tài liệu không mô tả xử lý khi API call thất bại (timeout, lỗi mạng, ProfiX down). T24 phản hồi user như thế nào? Có cơ chế retry không? Giao dịch bị block hay cho phép user commit mà không có phí?",
        "Giới hạn",
        "Đề xuất: Khi API call thất bại, T24 hiển thị lỗi và KHÔNG cho phép user commit giao dịch. Cần retry cơ chế hoặc user phải thử lại thủ công.",
        ""
    ],
    [
        "US39-QA-02.2",
        "Lưu đồ, Bước 8 (T24 call API → ProfiX ghi nhận thu phí thành công)",
        "Sau khi Checker duyệt, T24 call API ghi nhận thu phí thành công lên ProfiX. Nếu API call này thất bại (bước 8): (1) Giao dịch T24 đã hoàn tất hạch toán nhưng ProfiX chưa ghi nhận → dữ liệu mất đồng bộ. (2) Có cơ chế reconciliation / retry / compensating transaction không?",
        "Giới hạn",
        "Đề xuất: Cần có cơ chế retry tự động hoặc job đối soát (reconciliation) giữa T24 và ProfiX để xử lý trường hợp API bước 8 thất bại. BA cần làm rõ.",
        ""
    ],
    [
        "US39-QA-02.3",
        "Mục \"Yêu cầu nghiệp vụ\", đoạn \"hệ thống ProfiX cần được đồng bộ dữ liệu LD vào bảng ETL Tài khoản\"",
        "Trường hợp T24 gọi API tính phí cho 1 số LD mà LD đó CHƯA được ETL vào ProfiX (VD: LD mới mở trong ngày → chưa có trong dữ liệu T-1). ProfiX response trả về gì? T24 xử lý ra sao?",
        "Giới hạn",
        "Đề xuất: ProfiX trả về lỗi \"Không tìm thấy tài khoản LD\" → T24 hiển thị cảnh báo, cho phép user nhập phí thủ công hoặc chờ ETL ngày hôm sau.",
        ""
    ],
    [
        "US39-QA-02.4",
        "Mục \"Yêu cầu nghiệp vụ\", đoạn \"Mỗi Chương trình ưu đãi sẽ có Mã hạch toán riêng\"",
        "Trường hợp LD có Mã hạch toán nhưng Mã hạch toán đó KHÔNG match với bất kỳ Code phí nào trong Biểu phí CTƯĐ (VD: CTƯĐ đã hết hiệu lực, Code phí bị vô hiệu hóa). ProfiX xử lý thế nào? Trả phí = 0? Hay fallback sang Biểu phí dịch vụ chung?",
        "Giới hạn",
        "Đề xuất: BA làm rõ logic fallback: (1) Không match CTƯĐ → áp dụng Biểu phí chung, hoặc (2) Trả lỗi để user xử lý thủ công.",
        ""
    ],

    # === HẠNG MỤC 3: TOÀN VẸN DỮ LIỆU ===
    [
        "US39-QA-03.1",
        "Mục \"Yêu cầu nghiệp vụ\", danh sách dữ liệu ETL tối thiểu",
        "Tài liệu liệt kê \"dự kiến một số trường dữ liệu LD tối thiểu cần được đồng bộ\" — dùng từ \"dự kiến\" và \"tối thiểu\". Danh sách này đã chốt chưa? Có thêm trường nào khác cần ETL (VD: Lãi suất, Số tiền giải ngân, Ngày giải ngân)?",
        "Toàn vẹn dữ liệu",
        "Đề xuất: BA chốt danh sách trường ETL chính thức. Trường \"Ngày giải ngân\" quan trọng vì ảnh hưởng đến tính THKV.",
        ""
    ],
    [
        "US39-QA-03.2",
        "Mục \"Yêu cầu nghiệp vụ\", đoạn \"Số tháng tồn tại của khoản vay tính đến ngày ETL (dùng để khai báo vào điều kiện THKV của Code phí)\"",
        "THKV được tính đến \"ngày ETL\" (tức T-1). Nhưng giao dịch trả nợ trước hạn xảy ra tại ngày T. Vậy có sai lệch 1 ngày giữa THKV (tính đến T-1) và thời điểm thực tế trả nợ (ngày T). Sai lệch này có chấp nhận được không? Hay cần tính THKV đến ngày T (ngày giao dịch)?",
        "Toàn vẹn dữ liệu",
        "Đề xuất: Sai lệch 1 ngày có thể ảnh hưởng khi LD rơi vào đúng ranh giới bậc thang (VD: THKV = 12 tháng tại T-1 nhưng = 12 tháng + 1 ngày tại T). BA cần xác nhận chấp nhận sai lệch T-1 hay cần tính real-time.",
        ""
    ],
    [
        "US39-QA-03.3",
        "Mục \"Yêu cầu nghiệp vụ\", đoạn \"Trạng thái: mặc định chỉ lấy các tài khoản còn hoạt động\"",
        "ETL chỉ lấy tài khoản LD \"còn hoạt động\". Nếu 1 LD được đóng (trạng thái không hoạt động) sau ETL nhưng trước khi user khởi tạo giao dịch trả nợ trước hạn → LD vẫn tồn tại trên ProfiX (do ETL T-1). ProfiX có validate trạng thái LD tại thời điểm tính phí không? Hay dựa hoàn toàn vào dữ liệu ETL snapshot?",
        "Toàn vẹn dữ liệu",
        "Đề xuất: ProfiX nên validate trạng thái LD real-time hoặc T24 phải chặn giao dịch trên LD đã đóng trước khi call API ProfiX.",
        ""
    ],
    [
        "US39-QA-03.4",
        "Bảng diễn giải, Bước 0 (Đồng bộ dữ liệu T-1)",
        "Tài liệu ghi \"Mô tả chi tiết về các bảng dữ liệu ETL sẽ được các bên thống nhất trong quá trình triển khai\". Điều này có nghĩa mapping bảng ETL chưa chốt. Khi nào sẽ có tài liệu ETL mapping chính thức? Test team có cần chờ tài liệu này để viết test case ETL không?",
        "Toàn vẹn dữ liệu",
        "Đề xuất: BA cung cấp timeline dự kiến chốt tài liệu ETL mapping. Test team cần tài liệu này để thiết kế test data và test case cho Module 0.",
        ""
    ],

    # === HẠNG MỤC 4: UI/UX ===
    [
        "US39-QA-04.1",
        "Mục \"Giao diện\", \"Màn hình\" và \"Mô tả chi tiết các trường\" đều ghi N/A",
        "US39 không cung cấp Mockup UI cũng như bảng mô tả trường. Toàn bộ UI nằm ở phía T24 (hệ thống giao dịch nội bộ). Vậy phạm vi test của QA ProfiX chỉ bao gồm API (request/response) và logic tính phí? Không test UI giao dịch T24?",
        "UI-UX",
        "Đề xuất: BA xác nhận phạm vi test QA ProfiX: (1) Chỉ test API + logic tính phí ProfiX, hoặc (2) Cần phối hợp test E2E trên T24. Nếu (1), cần API spec chi tiết (endpoint, request/response schema).",
        ""
    ],
    [
        "US39-QA-04.2",
        "Mục \"Yêu cầu nghiệp vụ\", đoạn \"hệ thống T24 hiển thị thông tin phí\"",
        "Thông tin phí mà T24 hiển thị cho user bao gồm những gì? Chỉ có số tiền phí? Hay bao gồm: Code phí, Tên Code phí, Tỷ lệ phí, Biểu phí áp dụng, Chương trình ưu đãi (nếu có)?",
        "UI-UX",
        "Đề xuất: ProfiX API response nên trả đủ thông tin: Số tiền phí, Code phí, Tên Code phí, Biểu phí áp dụng, CTƯĐ (nếu có). T24 hiển thị tùy theo design riêng.",
        ""
    ],
]

# ===== BUILD TABLE =====
doc.add_heading('Bảng Q&A — Phân loại theo 4 Hạng mục', level=2)

HEADERS = ['ID', 'Trích xuất (Reference)', 'Câu hỏi / Sự cố', 'Phân loại', 'Đề xuất từ QA', 'Trả lời của BA']
COL_WIDTHS = [Cm(2.2), Cm(4.5), Cm(6.5), Cm(2), Cm(5.5), Cm(4.5)]

table = doc.add_table(rows=1 + len(qa_data), cols=6, style='Table Grid')
table.autofit = False

# Header row
for j, header in enumerate(HEADERS):
    cell = table.cell(0, j)
    cell.text = header
    cell.width = COL_WIDTHS[j]
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(9)
    # Gray background
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): 'D9E2F3',
        qn('w:val'): 'clear',
    })
    shading.append(shading_elm)

# Data rows
for i, row_data in enumerate(qa_data):
    for j, cell_text in enumerate(row_data):
        cell = table.cell(i + 1, j)
        cell.text = cell_text
        cell.width = COL_WIDTHS[j]
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(9)

# Category separator headings
doc.add_paragraph('')
doc.add_paragraph('Ghi chú: Các câu hỏi đã được đối chiếu với QTC-01 đến QTC-15. Không có câu hỏi nào trùng với quy tắc chung đã định nghĩa.')

# ===== SAVE =====
output_path = os.path.join(OUTPUT_DIR, "US39_PartB_QA.docx")
doc.save(output_path)
print(f"✅ Part B saved: {output_path}")
