"""US38 Part B Merged - Logic file (import data from 2 data files)"""
import sys, os
sys.path.insert(0, os.path.dirname(__file__))
from us38_partB_merged_data import QA_DATA
from us38_partB_merged_data2 import QA_DATA_2

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

ALL_DATA = QA_DATA + QA_DATA_2

doc = Document()
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10)

def set_w(cell, w_cm):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None: tcW = tc.makeelement(qn('w:tcW'), {}); tcPr.append(tcW)
    tcW.set(qn('w:w'), str(int(w_cm * 567))); tcW.set(qn('w:type'), 'dxa')

title = doc.add_heading('US38 – Danh Sách Q&A Tổng Hợp (AI + VA Merged)', level=1)
for r in title.runs: r.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.add_run('Tính năng: ').bold = True
p.add_run('Tự động thu phí bảo lãnh định kỳ – Bảo lãnh chưa xác định thời hạn')
p2 = doc.add_paragraph()
p2.add_run('Nguồn: ').bold = True
p2.add_run('Merge từ AI-generated Part B (18 câu) + VA-reviewed Part B (20 câu) → 28 câu unique')

doc.add_paragraph('')

# Column widths: ID(2.2) | Ref(3.5) | Question(6.0) | Cat(1.8) | Proposal(5.5) | Source(1.5) | BA(5.5) = ~26cm
col_w = [2.2, 3.5, 6.0, 1.8, 5.5, 1.5, 5.5]
headers = ['ID', 'Trích xuất', 'Câu hỏi / Sự cố', 'Phân loại', 'Đề xuất từ QA', 'Nguồn', 'Trả lời của BA']

# Section headers for grouping
section_rows = {
    0: '🔶 Hạng mục 1: Vấn đề Nghiệp vụ / Luồng xử lý (11 câu)',
    11: '🔴 Hạng mục 2: Giới hạn hệ thống & Exception Handling (7 câu)',
    18: '🟠 Hạng mục 3: Toàn vẹn dữ liệu & Ràng buộc (7 câu)',
    25: '🔵 Hạng mục 4: UI/UX & Giao diện (3 câu)',
}

# Build table
total_rows = 1 + len(ALL_DATA) + len(section_rows)  # header + data + section headers
table = doc.add_table(rows=total_rows, cols=7)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    set_w(cell, col_w[i])
    for p in cell.paragraphs:
        p.alignment = 1
        for r in p.runs:
            r.bold = True; r.font.size = Pt(9); r.font.name = 'Times New Roman'

# Data rows with section headers
row_idx = 1
data_idx = 0
for data_idx_iter, item in enumerate(ALL_DATA):
    # Check if need section header
    if data_idx_iter in section_rows:
        row = table.rows[row_idx]
        for ci in range(7):
            row.cells[ci].text = section_rows[data_idx_iter]
            for p in row.cells[ci].paragraphs:
                for r in p.runs:
                    r.bold = True; r.font.size = Pt(9); r.font.name = 'Times New Roman'
        row_idx += 1

    qid, ref, question, cat, proposal, ba, source = item
    row = table.rows[row_idx]
    values = [qid, ref, question, cat, proposal, source, ba]
    for ci, val in enumerate(values):
        cell = row.cells[ci]
        cell.text = val
        set_w(cell, col_w[ci])
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9); r.font.name = 'Times New Roman'
    row_idx += 1

# Summary
doc.add_paragraph('')
s = doc.add_paragraph()
s.add_run('Thống kê:').bold = True
doc.add_paragraph(f'• Tổng: {len(ALL_DATA)} câu hỏi', style='List Bullet')
cats = {}
for item in ALL_DATA:
    cats[item[3]] = cats.get(item[3], 0) + 1
for c, n in cats.items():
    doc.add_paragraph(f'• {c}: {n} câu', style='List Bullet')

sources = {}
for item in ALL_DATA:
    sources[item[6]] = sources.get(item[6], 0) + 1
doc.add_paragraph('')
s2 = doc.add_paragraph()
s2.add_run('Phân bổ nguồn:').bold = True
for src, n in sources.items():
    doc.add_paragraph(f'• {src}: {n} câu', style='List Bullet')

doc.add_paragraph('')
note = doc.add_paragraph()
note.add_run('⚠️ ').bold = True
note.add_run('Cột "Trả lời của BA" để trống. Vui lòng chuyển file cho BA điền đáp án.')

output_dir = '/Users/mac/antigravity-testing-kit/Feature_02_SA_Tham_So_He_Thong/output'
os.makedirs(output_dir, exist_ok=True)
path = os.path.join(output_dir, 'US38_PartB_QA_Merged.docx')
doc.save(path)
print(f'✅ Merged Part B saved: {path}')
print(f'   Total questions: {len(ALL_DATA)}')
