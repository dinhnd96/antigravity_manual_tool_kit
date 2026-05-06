#!/usr/bin/env python3
"""US06 Part A: Requirements Breakdown (based on US06-v2 + BA answers)."""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.section import WD_ORIENT
import os

doc = Document()
for s in doc.sections:
    s.orientation = WD_ORIENT.LANDSCAPE
    s.page_width, s.page_height = Cm(29.7), Cm(21.0)
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Cm(1.27)

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10)

# ============================================================
# PART A
# ============================================================
doc.add_heading("US06 – Phần A: Tóm Tắt Nghiệp Vụ Chuyên Sâu (Dành cho Tester)", level=1)
doc.add_paragraph("Tài liệu gốc: US06-v2 – Khai báo Biểu phí\nPhiên bản phân tích: v2 (đã tích hợp BA feedback)")

# A.1 Core Business Value
doc.add_heading("A.1. Thông Điệp Cốt Lõi (Core Business Value)", level=2)
doc.add_paragraph(
    "US06 cho phép người dùng (Maker) khai báo và quản lý Biểu phí – là tập hợp các Code phí áp dụng "
    "cho sản phẩm/dịch vụ của Ngân hàng trong từng thời kỳ. Mỗi Biểu phí quy định chi tiết về mức phí, "
    "loại khách hàng, nhóm khách hàng áp dụng, phương thức thu, tần suất thu. Biểu phí có vòng đời hiệu lực "
    "được xác định bởi Ngày hiệu lực và Ngày hết hiệu lực.\n\n"
    "Người dùng cuối: Maker (khai báo/chỉnh sửa) và Checker (phê duyệt) thuộc bộ phận Tham số hệ thống.\n"
    "Navigation: Tham số >> Danh mục Biểu phí >> Thêm mới / Chỉnh sửa."
)

# A.2 Flow Structure
doc.add_heading("A.2. Cấu Trúc Luồng Nghiệp Vụ & Phân Bổ Module", level=2)

# Module 1
doc.add_heading("Module 1: Thêm Mới Biểu Phí", level=3)
doc.add_paragraph("Luồng chính (Happy Path):", style='List Bullet')
steps = [
    "User nhấn 'Thêm mới' tại Danh mục Biểu phí → FE hiển thị màn hình Thêm mới.",
    "User nhập Thông tin chung: Mã (★, unique), Tên Biểu phí (★), Ngày ban hành (★, <= Ngày HT), "
    "Ngày hiệu lực (★, >= Ngày HT), Ngày hết hiệu lực (★, >= Ngày hiệu lực), Số văn bản (★), "
    "Tên văn bản (★), Link iDoc (tùy chọn, validate URL http/https, max 500 ký tự).",
    "User gán Code phí vào Biểu phí bằng 1 trong 2 cách:\n"
    "  • Cách 1: Tích chọn tại trường SPDV/Code phí → hiển thị cây SPDV (trạng thái Hoạt động). "
    "Lưu ý BA xác nhận: Code phí Ngừng hoạt động VẪN được phép chọn.\n"
    "  • Cách 2: Upload file Excel (tham chiếu US07).\n"
    "  ⚠️ Quy tắc: Upload SAU sẽ THAY THẾ toàn bộ Code phí đã tích chọn trước đó.",
    "Code phí đã gán hiển thị tại Lưới Thông tin chi tiết (readonly). User có thể:\n"
    "  • Chỉnh sửa: Sửa Điều kiện tính phí (Mô tả, Giá trị), Quy tắc tính phí (giá trị số, Tối thiểu, Tối đa).\n"
    "  • Xem chi tiết: Giống màn hình Xem code phí tại Danh mục SPDV.\n"
    "  • Bỏ gắn (Xóa): Bỏ Code phí khỏi Biểu phí.",
    "User nhấn 'Xác nhận' → FE ghi nhận → BE validate → Lưu thành công.\n"
    "  ⚠️ BẮT BUỘC có ít nhất 1 Code phí (BA đã xác nhận QA-01.8).",
    "Bản ghi chờ duyệt với hành động:\n"
    "  • 'Thêm mới': nếu KHÔNG sửa Code phí.\n"
    "  • 'Thêm mới – Sửa Code phí': nếu CÓ sửa Code phí.",
    "Biểu phí chỉ hiển thị trên lưới chính thức sau khi Checker cấp cuối phê duyệt (US25).",
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

doc.add_paragraph("\nCác Luồng Rẽ Nhánh / Ngoại lệ (Module 1):", style='List Bullet')
exceptions1 = [
    "Nhấn 'Đóng': Discard thay đổi, quay về Danh mục Biểu phí (QTC-15).",
    "Validate thất bại (BE): Hiển thị thông báo lỗi, giữ nguyên màn hình.",
    "Mã Biểu phí trùng: BE trả lỗi unique constraint.",
    "Upload file lỗi: Tham chiếu US07 (QTC-07).",
    "Bỏ gắn tất cả Code phí → FE disable nút Xác nhận (BA xác nhận).",
]
for e in exceptions1:
    doc.add_paragraph(e, style='List Bullet 2')

# Module 2
doc.add_heading("Module 2: Chỉnh Sửa Biểu Phí", level=3)
doc.add_paragraph("Điều kiện tiên quyết:", style='List Bullet')
doc.add_paragraph(
    "Biểu phí phải đã được duyệt Thêm mới thành công (hiển thị trên lưới). "
    "KHÔNG tồn tại tác vụ Chỉnh sửa/Chỉnh sửa-Sửa Code phí ở trạng thái Chờ duyệt "
    "(BA xác nhận QA-01.7: case Thêm mới chờ duyệt không tồn tại vì Biểu phí chưa hiển thị trên lưới).",
    style='List Bullet 2'
)

doc.add_paragraph("\nLuồng chính phân nhánh theo trạng thái:", style='List Bullet')

# 2a
doc.add_paragraph("Trạng thái = Chưa hiệu lực:", style='List Bullet 2')
items_2a = [
    "Cho phép sửa toàn bộ trừ Mã Biểu phí. Validate tương tự thêm mới.",
    "Cho phép sửa Code phí (Điều kiện tính phí, Quy tắc tính phí).",
    "Cho phép bỏ gắn Code phí.",
    "⚠️ KHÔNG cho phép Thêm mới (gắn thêm) Code phí lẻ. Muốn thêm → Upload Excel qua US07 (BA xác nhận QA-01.5).",
]
for it in items_2a:
    doc.add_paragraph(it, style='List Bullet 3')

# 2b
doc.add_paragraph("Trạng thái = Đang hiệu lực:", style='List Bullet 2')
items_2b = [
    "CHỈ cho phép sửa Ngày hết hiệu lực.",
    "Ràng buộc: Ngày hết hiệu lực mới > Ngày hệ thống VÀ > Ngày hiệu lực.",
    "KHÔNG tác động đến Code phí (BA xác nhận QA-01.6).",
]
for it in items_2b:
    doc.add_paragraph(it, style='List Bullet 3')

# 2c
doc.add_paragraph("Trạng thái = Hết hiệu lực:", style='List Bullet 2')
doc.add_paragraph(
    "Chỉ cho phép chuyển đổi Code phí sang Biểu phí mới (trạng thái = Đang hiệu lực). "
    "Tham chiếu US09.", style='List Bullet 3'
)

doc.add_paragraph("\nSau khi sửa → Bản ghi chờ duyệt:", style='List Bullet')
doc.add_paragraph("'Chỉnh sửa' hoặc 'Chỉnh sửa – Sửa Code phí'.", style='List Bullet 2')

doc.add_paragraph("\nCác Luồng Rẽ Nhánh / Ngoại lệ (Module 2):", style='List Bullet')
exceptions2 = [
    "Có tác vụ Chỉnh sửa đang Chờ duyệt → Hệ thống chặn, không cho sửa.",
    "Nhấn 'Đóng': Discard, quay về Danh mục (QTC-15).",
    "Validate thất bại: Hiển thị lỗi.",
]
for e in exceptions2:
    doc.add_paragraph(e, style='List Bullet 2')

# Module 3
doc.add_heading("Module 3: Trạng Thái Vòng Đời Biểu Phí", level=3)
doc.add_paragraph(
    "Trạng thái được xác định REALTIME (computed) dựa trên Ngày hệ thống (BA xác nhận QA-03.3):\n\n"
    "  • Chưa hiệu lực: Ngày HT < Ngày hiệu lực\n"
    "  • Đang hiệu lực: Ngày hiệu lực <= Ngày HT <= Ngày hết hiệu lực\n"
    "  • Hết hiệu lực: Ngày HT > Ngày hết hiệu lực"
)

# A.3 Pre-conditions
doc.add_heading("A.3. Điều Kiện Tiên Quyết & Dữ Liệu Mồi", level=2)
precond = [
    "Đã có SPDV và Code phí được khai báo (US01-US05) ở trạng thái Hoạt động hoặc Ngừng hoạt động.",
    "User có quyền Maker tại chức năng Danh mục Biểu phí.",
    "Checker có quyền phê duyệt theo Ma trận phê duyệt (US25).",
    "File Excel template cho upload Code phí (US07).",
]
for p in precond:
    doc.add_paragraph(p, style='List Bullet')

# A.4 QTC Applied
doc.add_heading("A.4. Quy Tắc Chung Áp Dụng", level=2)
qtc_items = [
    "QTC-01.5: Date format dd/mm/yyyy.",
    "QTC-01.6: Mã, Tên max 50 ký tự. Link iDoc max 500 ký tự (BA xác nhận QA-02.2).",
    "QTC-02: Tìm kiếm nhanh cây SPDV – gần đúng, case-insensitive, bỏ dấu, auto-trim.",
    "QTC-07: Upload file Excel (.xlsx). Validate format/dung lượng.",
    "QTC-08: Lịch sử tác động – Ngày cập nhật, Tác động, Người cập nhật, Người phê duyệt.",
    "QTC-11: FE-First Error Handling.",
    "QTC-12: Luồng Maker-Checker. Mã Biểu phí sinh sau khi Checker cấp cuối duyệt (nếu sinh tự động).",
    "QTC-14.1: No-Change Guard – FE disable nút Xác nhận nếu không thay đổi.",
    "QTC-15: Nút Đóng → discard, không popup xác nhận.",
]
for q in qtc_items:
    doc.add_paragraph(q, style='List Bullet')

# A.5 Data constraints from BA
doc.add_heading("A.5. Ràng Buộc Dữ Liệu Đặc Biệt (Từ BA Feedback)", level=2)
ba_items = [
    "Code phí Ngừng hoạt động VẪN được chọn vào Biểu phí (QA-01.1).",
    "1 Code phí có thể gán cho NHIỀU Biểu phí, nhưng khoảng thời gian hiệu lực các Biểu phí "
    "KHÔNG được trùng nhau (QA-03.1).",
    "Khi Chỉnh sửa Biểu phí Chưa hiệu lực: KHÔNG cho phép gắn thêm Code phí lẻ. "
    "Phải dùng Upload Excel (US07) nếu muốn thêm (QA-01.5).",
    "Khi Duyệt: BE validate lại toàn bộ ràng buộc như khi tạo tác vụ (QA-01.4 → sẽ thành QTC mới).",
    "Biểu phí phải có ít nhất 1 Code phí trước khi Xác nhận (QA-01.8).",
    "Không highlight Code phí đã sửa tại màn hình Checker – tính năng cân nhắc ở phase sau (QA-01.3).",
]
for b in ba_items:
    doc.add_paragraph(b, style='List Bullet')

# Save
out = os.path.join(os.path.dirname(os.path.dirname(__file__)), "US06_Part_A_v2.docx")
doc.save(out)
print(f"✅ Part A saved: {out}")
