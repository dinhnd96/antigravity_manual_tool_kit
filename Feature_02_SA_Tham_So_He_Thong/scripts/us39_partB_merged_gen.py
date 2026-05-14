# US39 Part B Merged - Docx generator
# Imports data from us39_partB_merged_data.py and us39_partB_merged_data2.py
import os, sys
sys.path.insert(0, os.path.dirname(__file__))
from us39_partB_merged_data import QA_DATA
from us39_partB_merged_data2 import QA_DATA_2
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ALL_DATA = QA_DATA + QA_DATA_2
OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "output")
os.makedirs(OUTPUT_DIR, exist_ok=True)

doc = Document()
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)

doc.add_heading('US39 — Phần B Tổng Hợp: Danh Sách Cảnh Báo & Q&A (AI + VA Merged)', level=1)
doc.add_paragraph('Feature: SA — Tham Số Hệ Thống | US39 — Tính phí trả nợ trước hạn kênh offline (T24)')
doc.add_paragraph(f'Tổng số câu hỏi: {len(ALL_DATA)} | Nguồn: AI (ban đầu) + VA (review bổ sung)')

# Summary stats
ai_only = sum(1 for d in ALL_DATA if d[6] == "AI")
va_only = sum(1 for d in ALL_DATA if d[6] == "VA")
both = sum(1 for d in ALL_DATA if d[6] == "AI+VA")
p = doc.add_paragraph()
p.add_run(f'Phân bổ nguồn: AI gốc = {ai_only} | VA mới = {va_only} | Trùng/Merge = {both}').bold = True

HEADERS = ['ID', 'Trích xuất', 'Câu hỏi / Sự cố', 'Phân loại', 'Đề xuất từ QA', 'Trả lời BA', 'Nguồn']
COL_WIDTHS = [Cm(2), Cm(3.8), Cm(6), Cm(1.8), Cm(5), Cm(4), Cm(1.5)]

# Category separators
HM_LABELS = {
    "US39-QA-01": "🔶 Hạng mục 1: Vấn đề Nghiệp vụ / Luồng xử lý",
    "US39-QA-02": "🔴 Hạng mục 2: Giới hạn hệ thống & Exception Handling",
    "US39-QA-03": "🟠 Hạng mục 3: Toàn vẹn dữ liệu & Ràng buộc",
    "US39-QA-04": "🔵 Hạng mục 4: UI/UX & Giao diện",
}

# Build table with category separators
rows_with_sep = []
last_hm = None
for d in ALL_DATA:
    hm_key = d[0][:12]  # e.g. "US39-QA-01"
    if hm_key != last_hm and hm_key in HM_LABELS:
        rows_with_sep.append(("SEP", HM_LABELS[hm_key]))
        last_hm = hm_key
    rows_with_sep.append(("DATA", d))

table = doc.add_table(rows=1 + len(rows_with_sep), cols=7, style='Table Grid')
table.autofit = False

# Header
for j, header in enumerate(HEADERS):
    cell = table.cell(0, j)
    cell.text = header
    cell.width = COL_WIDTHS[j]
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(8)
    shading = cell._element.get_or_add_tcPr()
    elm = shading.makeelement(qn('w:shd'), {qn('w:fill'): 'D9E2F3', qn('w:val'): 'clear'})
    shading.append(elm)

# Data rows
for i, item in enumerate(rows_with_sep):
    row_idx = i + 1
    if item[0] == "SEP":
        # Merge all cells for separator
        cell = table.cell(row_idx, 0)
        for k in range(1, 7):
            cell = cell.merge(table.cell(row_idx, k))
        cell.text = item[1]
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(9)
        shading = cell._element.get_or_add_tcPr()
        elm = shading.makeelement(qn('w:shd'), {qn('w:fill'): 'FFF2CC', qn('w:val'): 'clear'})
        shading.append(elm)
    else:
        d = item[1]
        for j in range(7):
            cell = table.cell(row_idx, j)
            cell.text = d[j]
            cell.width = COL_WIDTHS[j]
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(8)

doc.add_paragraph('')
doc.add_paragraph('Ghi chú: Các câu hỏi đã được đối chiếu với QTC-01 đến QTC-15. Không có câu hỏi nào trùng với quy tắc chung.')

output_path = os.path.join(OUTPUT_DIR, "US39_PartB_QA_Merged.docx")
doc.save(output_path)
print(f"✅ Merged Part B saved: {output_path}")
print(f"   Total Q&A: {len(ALL_DATA)} (AI={ai_only}, VA={va_only}, AI+VA={both})")
