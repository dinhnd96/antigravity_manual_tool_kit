"""US32 Part A - Business Summary Generator"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Page setup - Landscape, narrow margins
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    return h

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

# ===== TITLE =====
title = doc.add_heading('US32 — PHẦN A: TÓM TẮT NGHIỆP VỤ CHUYÊN SÂU', level=1)
doc.add_paragraph('Feature: Dashboard — Báo cáo tổng quan hoạt động thu phí')
doc.add_paragraph('Phiên bản: v1.0 | Ngày: 14/05/2026 | Người phân tích: AI QA Lead')
doc.add_paragraph('─' * 80)

# ===== A.1 CORE BUSINESS VALUE =====
add_heading('A.1. Thông Điệp Cốt Lõi (Core Business Value)', 2)
add_bullet('Tính năng này cung cấp cho người dùng (Management, Head of Department) một trang báo cáo tổng quan (Dashboard) về hoạt động thu phí dịch vụ của toàn ngân hàng.')
add_bullet('Người dùng cuối: Quản lý cấp trung/cao, nhân viên có quyền truy cập module Dashboard.')
add_bullet('Giá trị: Theo dõi real-time (MTD) các chỉ số doanh thu phí, số giao dịch, số khách hàng phát sinh phí; so sánh tháng hiện tại với tháng liền trước; phân tích theo 3 chiều: Chi nhánh, Sản phẩm, Khối.')
add_bullet('Navigation: Dashboard >> Dashboard theo Chi nhánh / Dashboard theo Sản phẩm / Dashboard theo Khối.')

# ===== A.2 MODULE MAPPING =====
add_heading('A.2. Cấu Trúc Luồng Nghiệp Vụ & Phân Bổ Module', 2)

# Module 0: Common KPI
add_heading('Module 0: Phân Vùng Chỉ Số Tổng Quan (Common KPI Cards)', 3)
add_bullet('Xuất hiện trên CẢ 3 màn hình Dashboard, layout/label/icon KHÔNG thay đổi khi chuyển tab.')
add_bullet('Đơn vị: Triệu VND cho tất cả số tiền trên dashboard.')
p = doc.add_paragraph()
p.add_run('3 Card chỉ số:').bold = True

items = [
    ('Card 1 — Tổng doanh thu:', 'Tổng doanh thu phí tháng hiện tại | Tháng trước | % thay đổi = (Tháng này − Tháng trước) / Tháng trước × 100(%). Xanh nếu ≥100%, Đỏ nếu <100%.'),
    ('Card 2 — Số giao dịch thu phí:', 'Tổng số giao dịch thu phí tháng hiện tại (phí thực thu VND > 0, không tính reverse) | Tháng trước | % thay đổi. Màu tương tự.'),
    ('Card 3 — Tổng số KH phát sinh phí:', 'Tổng số KH có phát sinh giao dịch thu phí tháng này (phí thực thu > 0, không tính reverse) | Tháng trước | % thay đổi. Màu tương tự.'),
]
for prefix, desc in items:
    add_bullet(desc, prefix + ' ')

# Module 1: Dashboard theo Chi nhánh
add_heading('Module 1: Dashboard Theo Chi Nhánh (Mặc định)', 3)
doc.add_paragraph('Là màn hình mặc định khi truy cập Dashboard. Gồm 3 biểu đồ + 1 phần chọn chi nhánh:')

add_heading('1.1 Số giao dịch thu phí theo chi nhánh (Bar Chart Horizontal)', 4)
add_bullet('Loại biểu đồ: Bar chart horizontal. Sắp xếp: chi nhánh có số giao dịch cao nhất lên top.')
add_bullet('Tooltip hover: Tên chi nhánh + Số giao dịch thu phí.')
add_bullet('Tùy chỉnh: Mặc định top 5, người dùng chọn n = 5/10/15.')
add_bullet('Dữ liệu: Tổng giao dịch thu phí theo chi nhánh trong tháng hiện tại (phí thực thu VND > 0, không tính reverse). Tiêu đề biểu đồ format: "Số giao dịch thu phí theo chi nhánh tháng MM/YYYY".')

add_heading('1.2 Doanh thu phí dịch vụ theo chi nhánh (Donut Chart)', 4)
add_bullet('Loại biểu đồ: Donut chart. Phần trung tâm: Tổng doanh thu phí toàn ngân hàng trong tháng.')
add_bullet('Các phần donut: Doanh thu từng chi nhánh top, phần còn lại gộp thành "Còn lại".')
add_bullet('Tùy chỉnh: Mặc định top 5, chọn n = 5/10/15 → chia thành (n+1) phần.')
add_bullet('Tooltip hover: Tên chi nhánh / "Còn lại" + Doanh thu phí trong tháng.')
add_bullet('Dữ liệu: Tổng doanh thu phí giao dịch (phí thực thu VND > 0, không tính reverse) theo từng chi nhánh. Tiêu đề: "Doanh thu phí dịch vụ theo chi nhánh MM/YYYY".')

add_heading('1.3 Nợ phí và truy thu theo chi nhánh (Combo Bar + Line Chart)', 4)
add_bullet('Loại biểu đồ: Kết hợp bar chart (số tiền nợ phí) + line chart (số tiền truy thu thành công).')
add_bullet('Tooltip hover: Tên chi nhánh + Số tiền nợ phí + Số tiền đã truy thu thành công.')
add_bullet('Tùy chỉnh mặc định: 7 chi nhánh (theo alphabet). Người dùng chọn tối đa 10 chi nhánh qua combobox (có tìm kiếm nhanh theo mã/tên). Placeholder sau khi chọn: "Đã chọn n chi nhánh".')
add_bullet('Dữ liệu: Nợ phí + truy thu từ đầu tháng đến ngày hệ thống. Tiêu đề: "Nợ phí và truy thu theo chi nhánh đến ngày DD/MM/YYYY".')

# Module 2: Dashboard theo Sản phẩm
add_heading('Module 2: Dashboard Theo Sản Phẩm', 3)

add_heading('2.1 Số giao dịch thu phí theo sản phẩm (Line Chart)', 4)
add_bullet('Loại biểu đồ: Line chart. Sắp xếp: sản phẩm có số giao dịch cao nhất lên top.')
add_bullet('Tooltip hover: Tên SPDV cấp 1 + Số giao dịch thu phí.')
add_bullet('Tùy chỉnh: Mặc định top 5 SPDV cấp 1, chọn n = 5/10/15.')
add_bullet('Dữ liệu: Tổng giao dịch thu phí theo SPDV cấp 1 tháng hiện tại (phí thực thu VND > 0, không tính reverse). Tiêu đề: "Số giao dịch thu phí theo sản phẩm tháng MM/YYYY".')

add_heading('2.2 Doanh thu phí theo sản phẩm dịch vụ (Stacked Horizontal Bar Chart)', 4)
add_bullet('Loại biểu đồ: Stacked horizontal bar chart. Trục Y = chi nhánh, trục X = doanh thu.')
add_bullet('Mỗi cột ngang = tổng doanh thu 1 chi nhánh, chia thành các màu theo SPDV.')
add_bullet('Giá trị tổng doanh thu hiển thị ở cuối mỗi cột.')
add_bullet('Tooltip hover: Tên chi nhánh + SPDV + Doanh thu phí giao dịch trong tháng của SPDV đó tại chi nhánh.')

p = doc.add_paragraph()
p.add_run('Tùy chỉnh phức tạp:').bold = True
add_bullet('Mặc định: Top 5 chi nhánh, mỗi cột chia 4 phần (3 SPDV cấp 1 đầu tiên + phần "Còn lại").')
add_bullet('Chọn top chi nhánh: n = 5/10/15.')
add_bullet('Dropdown "Chọn sản phẩm dịch vụ": Cho phép chọn cấp SPDV (tối đa theo tham số product_level), sau đó tìm kiếm nhanh theo mã/tên SPDV, chỉ tick SPDV thuộc cấp đã chọn, tối đa 10 SPDV cùng cấp.')
add_bullet('Sau khi chọn xong m SPDV: mỗi cột hiển thị (m+1) phần = m SPDV đã chọn + 1 phần "Còn lại".')
add_bullet('Dữ liệu: Tổng doanh thu phí (phí thực thu VND > 0, không tính reverse) theo SPDV, theo cấp SPDV người dùng chọn, thuộc từng chi nhánh. Tiêu đề: "Doanh thu phí theo sản phẩm dịch vụ MM/YYYY".')

# Module 3: Dashboard theo Khối
add_heading('Module 3: Dashboard Theo Khối', 3)

add_heading('3.1 Nợ phí và truy thu (Line Chart theo Khối)', 4)
add_bullet('Loại biểu đồ: Line chart. Trục X: Khối KHCN, Khối KHDNL, Khối KHDN.')
add_bullet('Line xanh lá: Số tiền nợ phí. Line cam: Số tiền đã truy thu thành công.')
add_bullet('Tooltip hover: Khối + Số tiền nợ phí + Số tiền đã truy thu thành công.')
add_bullet('Dữ liệu: Tổng Số tiền còn nợ + Số tiền đã truy thu theo từng Khối, từ đầu tháng đến ngày hệ thống. Tiêu đề: "Nợ phí và truy thu đến ngày DD/MM/YYYY".')

add_heading('3.2 Danh sách KH giao dịch nhiều nhất trong ngày (Bảng liệt kê)', 4)
add_bullet('Hiển thị top n KH có số lượng giao dịch thu phí nhiều nhất trong ngày hệ thống.')
add_bullet('Tùy chỉnh: Mặc định top 10. Chọn n = 10/20/30. Chọn 1 hoặc nhiều Khối (KHCN, KHDNL, KHDN) theo nguyên tắc phân quyền dữ liệu QTC-10.')

p = doc.add_paragraph()
p.add_run('Các cột bảng:').bold = True
add_bullet('CIF – Tên khách hàng: Hiển thị dạng "Mã CIF – Tên KH".')
add_bullet('Số lượng giao dịch: Tổng giao dịch thu phí trong ngày (phí thực thu VND > 0, không tính reverse).')
add_bullet('Doanh thu: Tổng doanh thu phí tương ứng với các giao dịch đã thống kê của KH.')

# ===== A.3 PRE-CONDITIONS =====
add_heading('A.3. Bảng Điều Kiện Tiên Quyết & Cấu Hình', 2)
add_bullet('Người dùng phải có quyền truy cập module Dashboard.')
add_bullet('Phân quyền dữ liệu theo Khối [QTC-10]: Dữ liệu hiển thị trên Dashboard phụ thuộc Khối của người dùng đăng nhập.')
add_bullet('Tham số hệ thống product_level: Xác định số cấp SPDV tối đa cho phép chọn trong dropdown "Chọn sản phẩm dịch vụ" (Module 2.2).')
add_bullet('Dữ liệu giao dịch phải đã được ghi nhận trong hệ thống để Dashboard có dữ liệu hiển thị.')
add_bullet('Không có Lưu đồ / Flowchart — tài liệu ghi N/A.')

# ===== A.4 QTC APPLIED =====
add_heading('A.4. Quy Tắc Chung Áp Dụng', 2)

qtc_items = [
    ('QTC-01.4 (Number):', 'Số tiền hiển thị phân cách hàng nghìn bằng dấu phẩy, thập phân bằng dấu chấm (2 chữ số). Đơn vị VND không có thập phân. Tuy nhiên Dashboard hiển thị Triệu VND — cần xác nhận quy tắc làm tròn.'),
    ('QTC-01.5 (Date):', 'Định dạng dd/mm/yyyy cho tiêu đề biểu đồ có hiển thị ngày.'),
    ('QTC-01.7 (Empty Data):', 'Nếu không có dữ liệu, hiển thị trống (blank) trên lưới/bảng.'),
    ('QTC-10 (Phân quyền Khối):', 'Dữ liệu Dashboard lọc theo Khối của người dùng. User thuộc Khối KHCN/KHDN → chỉ thấy dữ liệu thuộc Khối tương ứng.'),
]
for prefix, desc in qtc_items:
    add_bullet(desc, prefix + ' ')

# Save
output = '/Users/mac/antigravity-testing-kit/Feature_02_SA_Tham_So_He_Thong/US32_PartA_Summary.docx'
doc.save(output)
print(f'✅ Đã sinh file: {output}')
