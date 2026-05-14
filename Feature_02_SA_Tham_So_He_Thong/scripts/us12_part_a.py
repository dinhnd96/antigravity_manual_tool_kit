"""US12 Part A - Tóm Tắt Nghiệp Vụ Chuyên Sâu (Dành cho Tester)"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# Page setup: Landscape, Narrow margins
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 51, 102)
    return h

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run2 = p.add_run(text)
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
    return p

# ===== TITLE =====
title = doc.add_heading('US12 — PHẦN A: TÓM TẮT NGHIỆP VỤ CHUYÊN SÂU', level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)
    run.font.name = 'Times New Roman'

doc.add_paragraph('Feature: Khai báo CTƯĐ áp dụng cho danh sách khách hàng do người dùng định nghĩa')
doc.add_paragraph('Phiên bản: v1.0 | Ngày phân tích: 13/05/2026')

# ===== A.1 CORE BUSINESS VALUE =====
add_heading_styled('A.1. Thông Điệp Cốt Lõi (Core Business Value)', 1)
p = doc.add_paragraph()
p.add_run('Mục đích: ').bold = True
p.add_run('US12 cho phép người dùng (Maker) khai báo Chương trình ưu đãi (CTƯĐ) không đánh giá định kỳ, '
          'áp dụng cho một khách hàng cụ thể hoặc một danh sách khách hàng do người dùng tự upload/thêm mới thủ công. '
          'Điểm khác biệt cốt lõi so với US11 (CTƯĐ có đánh giá định kỳ): hệ thống KHÔNG tự động xác định danh sách KH '
          '— mà người dùng chủ động khai báo từng KH kèm mức ưu đãi riêng.')
p2 = doc.add_paragraph()
p2.add_run('Người dùng cuối: ').bold = True
p2.add_run('Cán bộ nghiệp vụ quản lý phí dịch vụ (Maker) và Cấp phê duyệt (Checker) theo luồng Maker-Checker [QTC-12].')
p3 = doc.add_paragraph()
p3.add_run('Phạm vi: ').bold = True
p3.add_run('Thêm mới và Chỉnh sửa CTƯĐ có khai báo danh sách khách hàng.')

# ===== A.2 FLOW STRUCTURE =====
add_heading_styled('A.2. Cấu Trúc Luồng Nghiệp Vụ & Phân Bổ Module', 1)

# --- Module 1: Thêm mới ---
add_heading_styled('Module 1: Thêm Mới CTƯĐ Có Danh Sách Khách Hàng', 2)
add_heading_styled('Luồng chính (Happy Path):', 3)
steps = [
    'Người dùng truy cập: Tham số >> Chương trình ưu đãi >> Quản lý CTƯĐ >> Thêm mới.',
    'Chọn "Ưu đãi không có đánh giá định kỳ" → FE hiển thị màn hình khai báo.',
    'Bật Toggle "Theo danh sách khách hàng" = On → FE hiển thị form đầy đủ với phân vùng Chi tiết ưu đãi (bảng KH).',
    'Khai báo Thông tin chung: Tên CTƯĐ, Ngày ban hành (≤ hôm nay), Ngày hiệu lực (≥ hôm nay), Ngày hết hiệu lực (≥ Ngày hiệu lực), '
    'Số văn bản, Tên văn bản, Số VB_Tên viết tắt (≤20 ký tự, format: [số]_[chữ viết hoa]), Đối tượng thu phí (KH/Merchant), Khối (KHCN/KHDN/KHDNL), '
    'Loại ưu đãi (Theo KH/TK/Thẻ). Trường không bắt buộc: Email CBNV đầu mối (@pvcombank.com.vn), Link iDoc.',
    'Khai báo Điều kiện áp dụng: Tham chiếu US11 — Thêm nhóm điều kiện, chọn Nguồn dữ liệu (ETL KH/TK/Thẻ/API giao dịch), '
    'Tên trường điều kiện (Combobox lọc theo Nguồn), Operator, Giá trị. Cặp (Trường + Operator) không trùng lặp trong cùng nhóm.',
    'Khai báo Chu kỳ áp dụng: Chọn "Theo chu kỳ" (Hàng tuần → chọn thứ; Hàng tháng → Ngày cụ thể 1-31 hoặc Ngày trong tháng + Tuần) '
    'HOẶC chọn "Liên tục" (mặc định áp dụng từ Ngày hiệu lực → Ngày hết hiệu lực).',
    'Khai báo Chi tiết ưu đãi: Toggle "Ưu đãi theo tỷ lệ" (On = %, Off = giá trị cố định). '
    'Upload file Excel (.xlsx) danh sách KH HOẶC "+ Thêm mới khách hàng" thủ công từng dòng. '
    'Mỗi dòng gồm: CIF, SPDV, Tỷ lệ/Giá trị ưu đãi, Loại tiền tối thiểu/tối đa, Số tiền phí min/max, '
    'Trường giá trị kiểm tra, Ngưỡng dừng ưu đãi, Ngày hiệu lực/hết hiệu lực (chỉ hiện khi Liên tục).',
    'Nhấn "Xác nhận" → BE validate toàn bộ → Lưu bản ghi trạng thái "Chờ duyệt" → Hiển thị toast thành công → Quay về lưới CTƯĐ.',
    'Checker phê duyệt (luồng Maker-Checker theo QTC-12/US25) → Mã CTƯĐ được sinh tự động (tham chiếu US04).'
]
for s in steps:
    add_bullet(s)

add_heading_styled('Các Luồng Rẽ Nhánh / Ngoại Lệ (Module 1):', 3)
exceptions = [
    'Nhấn "Đóng": Hệ thống không lưu dữ liệu, quay về màn hình CTƯĐ [QTC-15].',
    'Upload file KH — dữ liệu không hợp lệ (Bước 14.a): FE hiển thị thông báo lỗi, cho phép upload lại.',
    'Validate upload: CIF không tồn tại trong ETL KH → lỗi. SPDV trạng thái ≠ Hoạt động → lỗi. '
    'CIF+SPDV trùng trong cùng danh sách → lỗi. SPDV cha-con cùng CIF trong danh sách → lỗi.',
    'Thêm mới KH thủ công (Bước 13.a→13.e): FE gửi yêu cầu truy vấn CIF → BE trả danh sách → User chọn CIF → Khai báo chi tiết.',
    'Validate BE (Bước 16): Dữ liệu khai báo không hợp lệ → FE hiển thị thông báo lỗi, kết thúc.',
    'Toggle On/Off ưu đãi theo tỷ lệ: Hệ thống tự động clear toàn bộ giá trị đã nhập của lựa chọn trước đó.',
    'Tỷ lệ ưu đãi: Phải > 0 và ≤ 100. Giá trị ưu đãi: Phải > 0.',
    'Số tiền phí tối đa phải > Số tiền phí tối thiểu (nếu cả hai đều có giá trị).',
    'Ngưỡng dừng ưu đãi: Bắt buộc nhập nếu có chọn Trường giá trị kiểm tra.',
    'SPDV trong CTƯĐ: Không được chọn các SPDV có quan hệ cha-con.'
]
for e in exceptions:
    add_bullet(e)

# --- Module 2: Chỉnh sửa ---
add_heading_styled('Module 2: Chỉnh Sửa CTƯĐ Có Danh Sách Khách Hàng', 2)
add_heading_styled('Luồng chính (Happy Path):', 3)
edit_steps = [
    'Người dùng nhấn "Chỉnh sửa" trên lưới CTƯĐ.',
    'BE kiểm tra: Có tác vụ "Chờ duyệt" không? → Nếu CÓ → Hiển thị thông báo lỗi, không cho chỉnh sửa, kết thúc.',
    'BE kiểm tra trạng thái hoạt động CTƯĐ:',
    '   a) Hết hiệu lực → Thông báo lỗi "Trạng thái CTƯĐ không hợp lệ", kết thúc.',
    '   b) Chưa hiệu lực → Cho phép chỉnh sửa TOÀN BỘ trừ Mã CTƯĐ và Số VB_Tên viết tắt CTƯĐ. '
    'Validate tương tự thêm mới.',
    '   c) Đang hiệu lực → CHỈ cho phép chỉnh sửa: Ngày hết hiệu lực CTƯĐ + Chi tiết ưu đãi '
    '(Upload thay thế toàn bộ / Thêm mới / Xóa bản ghi). Validate tương tự thêm mới.',
    'Nhấn "Xác nhận" → BE validate → Tạo bản ghi Chờ duyệt → Toast thành công → Quay về lưới CTƯĐ.',
    'Checker phê duyệt → Dữ liệu chính thức được cập nhật trên lưới.'
]
for s in edit_steps:
    add_bullet(s)

add_heading_styled('Các Luồng Rẽ Nhánh / Ngoại Lệ (Module 2):', 3)
edit_exceptions = [
    'Nhấn "Đóng": Không lưu, quay về lưới [QTC-15].',
    'Validate BE (Bước 5): Không hợp lệ → Thông báo lỗi, kết thúc.',
    'Đang hiệu lực + Upload file thay thế: Toàn bộ Chi tiết ưu đãi cũ bị ghi đè bởi dữ liệu mới từ file.',
    'Chỉnh sửa không thay đổi gì → FE disable nút Xác nhận hoặc cảnh báo [QTC-14.1].'
]
for e in edit_exceptions:
    add_bullet(e)

# ===== A.3 PRE-CONDITIONS =====
add_heading_styled('A.3. Bảng Điều Kiện Tiên Quyết & Cấu Hình', 1)
table = doc.add_table(rows=1, cols=3, style='Light Grid Accent 1')
hdr = table.rows[0].cells
hdr[0].text = 'STT'
hdr[1].text = 'Điều Kiện'
hdr[2].text = 'Mô Tả'
preconditions = [
    ('1', 'Phân quyền Maker', 'Người dùng phải có quyền Thêm mới / Chỉnh sửa CTƯĐ.'),
    ('2', 'Phân quyền Checker', 'Cấp phê duyệt phải có quyền Phê duyệt / Từ chối [QTC-12].'),
    ('3', 'Dữ liệu mồi: SPDV', 'Danh mục SPDV phải tồn tại với trạng thái = Hoạt động.'),
    ('4', 'Dữ liệu mồi: ETL KH', 'Bảng ETL dữ liệu khách hàng phải có CIF hợp lệ.'),
    ('5', 'Dữ liệu mồi: Điều kiện tính phí', 'Các trường điều kiện tính phí (US khai báo điều kiện) phải tồn tại với trạng thái = Hoạt động.'),
    ('6', 'Dữ liệu mồi: Loại tiền tệ', 'Danh mục Loại tiền phải tồn tại trong hệ thống.'),
    ('7', 'Cấu hình Mã CTƯĐ', 'Nguyên tắc tự sinh Mã CTƯĐ theo US04 phải được cấu hình sẵn.'),
    ('8', 'File template upload', 'File Excel (.xlsx) upload KH phải đúng template được hệ thống quy định.'),
]
for row_data in preconditions:
    row = table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

# ===== A.4 QUY TẮC CHUNG =====
add_heading_styled('A.4. Quy Tắc Chung Áp Dụng (ProfiX Common Rules)', 1)
qtc_rules = [
    ('QTC-01.1', 'Combobox: Search-as-you-type cho CIF, SPDV, Tên trường điều kiện.'),
    ('QTC-01.2', 'Dropdown List: Đối tượng thu phí, Khối, Loại ưu đãi, Nguồn dữ liệu, Operator, Tuần, Loại tiền, Trường giá trị kiểm tra.'),
    ('QTC-01.3', 'Multiple Select Dropdown: Áp dụng vào thứ, Ngày trong tháng.'),
    ('QTC-01.4', 'Number: Giá trị ưu đãi, Tỷ lệ ưu đãi, Số tiền phí tối thiểu/tối đa, Ngưỡng dừng ưu đãi, Ngày cụ thể (1-31).'),
    ('QTC-01.5', 'Date: Ngày ban hành, Ngày hiệu lực, Ngày hết hiệu lực — format dd/mm/yyyy.'),
    ('QTC-01.6', 'Text mặc định: Tên CTƯĐ (50 ký tự), Mô tả (300 ký tự) — trừ khi US ghi khác (VD: Số VB_Tên viết tắt = 20 ký tự).'),
    ('QTC-05', 'Tải xuống: Danh sách KH theo format Excel .xlsx, tên file = "{Tên chức năng} - yyyymmddhhmmss".'),
    ('QTC-07', 'Upload file: Định dạng .xlsx; lỗi format/dung lượng → toast "Định dạng hoặc dung lượng không hợp lệ".'),
    ('QTC-11', 'Xử lý lỗi FE-First: FE chặn lỗi đầu tiên, BE là tầng cuối.'),
    ('QTC-12', 'Luồng Maker-Checker: Maker lưu → Chờ duyệt → Checker phê duyệt → Mã sinh chính thức.'),
    ('QTC-14.1', 'No-Change Guard: FE disable nút Xác nhận nếu không có thay đổi.'),
    ('QTC-14.2', 'Tên CTƯĐ không yêu cầu unique.'),
    ('QTC-14.4', 'Không hỗ trợ chỉnh sửa Tên/Mô tả ngoài luồng Chỉnh sửa → Chờ duyệt.'),
    ('QTC-15', 'Nút Đóng: Không popup xác nhận, không lưu, quay về màn hình trước.'),
]
table2 = doc.add_table(rows=1, cols=2, style='Light Grid Accent 1')
hdr2 = table2.rows[0].cells
hdr2[0].text = 'Mã QTC'
hdr2[1].text = 'Áp Dụng Trong US12'
for code, desc in qtc_rules:
    row = table2.add_row().cells
    row[0].text = code
    row[1].text = desc

# Save
out_dir = '/Users/mac/antigravity-testing-kit/Feature_02_SA_Tham_So_He_Thong'
out_path = os.path.join(out_dir, 'US12_PartA_Summary.docx')
doc.save(out_path)
print(f'Part A saved: {out_path}')
