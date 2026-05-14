"""US31 - Part A: Tóm tắt Nghiệp vụ (Requirements Breakdown)"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# Page setup: Landscape, narrow margins
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

# === TITLE ===
title = doc.add_heading("US31 – Báo Cáo Tổng Doanh Thu Phí", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading("PHẦN A: TÓM TẮT NGHIỆP VỤ (Requirements Breakdown)", level=2)

# === A.1 Core Business Value ===
doc.add_heading("A.1. Thông Điệp Cốt Lõi (Core Business Value)", level=3)
p = doc.add_paragraph()
p.add_run("Mục đích: ").bold = True
p.add_run("US31 cung cấp chức năng tra cứu và xuất báo cáo tổng hợp doanh thu phí dịch vụ theo khoảng thời gian, "
          "giúp người dùng (nhân viên ngân hàng, quản lý) nắm bắt được tổng doanh thu phí đã thu từ khách hàng "
          "trong một giai đoạn cụ thể.")
p2 = doc.add_paragraph()
p2.add_run("Người dùng cuối: ").bold = True
p2.add_run("Nhân viên ngân hàng, Quản lý nghiệp vụ phí — tùy theo phân quyền Khối (KHCN, KHDN, KHDNL, hoặc Khối khác).")
p3 = doc.add_paragraph()
p3.add_run("Navigation: ").bold = True
p3.add_run("Báo cáo >> Báo cáo tổng doanh thu phí")

# === A.2 Flow Structure & Module Mapping ===
doc.add_heading("A.2. Cấu Trúc Luồng Nghiệp Vụ & Phân Bổ Module", level=3)

doc.add_heading("Module 1: Tra cứu Báo cáo tổng doanh thu phí", level=4)
doc.add_paragraph("Luồng chính (Happy Path):", style="List Bullet")
steps = [
    "Bước 1: Người dùng chọn menu \"Báo cáo tổng doanh thu phí\".",
    "Bước 2: FE hiển thị màn hình Báo cáo tổng doanh thu phí gồm: vùng Điều kiện tìm kiếm (7 trường), "
    "lưới danh sách (rỗng ban đầu theo QTC-04), và các nút chức năng.",
    "Bước 3.1: Người dùng nhập/chọn các điều kiện tìm kiếm (Khối, Mã Chi nhánh, Biểu phí, Code phí, "
    "Loại tính phí, Từ ngày ★, Đến ngày ★).",
    "Bước 4.2: Nhấn nút \"Tra cứu\".",
    "Bước 6: FE gửi yêu cầu xử lý đến BE.",
    "Bước 7: BE nhận yêu cầu, xử lý logic tìm kiếm.",
    "Bước 8: FE hiển thị danh sách doanh thu phí trên lưới (12 cột dữ liệu + cột STT + dòng Tổng cộng).",
]
for s in steps:
    doc.add_paragraph(s, style="List Bullet 2")

doc.add_paragraph("Các Luồng Rẽ Nhánh / Ngoại lệ:", style="List Bullet")
exceptions = [
    "Xóa tra cứu (Bước 4.1 → 5): Người dùng nhấn \"Xóa tra cứu\" → FE xóa tất cả điều kiện đã nhập, "
    "lưới trở về trạng thái mặc định (theo QTC-04: không hiển thị kết quả).",
    "Bỏ trống trường bắt buộc: Từ ngày và Đến ngày là bắt buộc (★). Nếu bỏ trống → FE validate chặn, "
    "hiển thị \"Trường này bắt buộc\" (theo QTC-14.5, ngầm định FE tự validate).",
    "Từ ngày > Đến ngày: FE không cho phép chọn ngày tại field \"Từ ngày\" lớn hơn field \"Đến ngày\" "
    "(theo mô tả trường).",
    "Không có kết quả: Khi tra cứu không tìm thấy bản ghi → lưới hiển thị rỗng.",
    "Phân quyền Khối: User thuộc KHCN/KHDN/KHDNL → Khối auto-fill, không sửa được. "
    "User Khối khác → được chọn tự do hoặc để trống (tìm tất cả).",
]
for e in exceptions:
    doc.add_paragraph(e, style="List Bullet 2")

doc.add_heading("Module 2: Tải xuống báo cáo (Export)", level=4)
doc.add_paragraph("Luồng chính (Happy Path):", style="List Bullet")
export_steps = [
    "Bước 3.2: Người dùng nhấn nút \"Tải xuống\".",
    "Bước 9: FE hiển thị dropdownlist gồm 2 lựa chọn: Excel / PDF.",
    "Bước 10: Người dùng chọn Excel (10.1) hoặc PDF (10.2).",
    "Bước 11: Hệ thống tải xuống file theo quy tắc Tải xuống tại Quy tắc chung (QTC-05).",
]
for s in export_steps:
    doc.add_paragraph(s, style="List Bullet 2")

doc.add_paragraph("Các Luồng Rẽ Nhánh / Ngoại lệ:", style="List Bullet")
doc.add_paragraph("Tải xuống khi chưa có kết quả tra cứu: Theo QTC-05, luôn cho phép tải xuống dù có dữ liệu "
                   "hay không.", style="List Bullet 2")

doc.add_heading("Module 3: Tra cứu CIF", level=4)
doc.add_paragraph("Luồng chính (Happy Path):", style="List Bullet")
cif_steps = [
    "Bước 3.3: Người dùng nhấn nút \"Tra cứu CIF\".",
    "Bước 12: Hệ thống hiển thị popup Tra cứu CIF (tham chiếu QTC-09).",
    "Người dùng chọn Mã CIF → popup đóng, Mã CIF được fill tự động vào trường tương ứng.",
]
for s in cif_steps:
    doc.add_paragraph(s, style="List Bullet 2")

doc.add_paragraph("Lưu ý:", style="List Bullet")
doc.add_paragraph("Tài liệu US31 có nút \"Tra cứu CIF\" và mô tả tham chiếu QTC-09, nhưng trên lưới "
                   "kết quả và bảng mô tả trường KHÔNG có cột/trường \"Mã CIF\" nào. "
                   "Điều này tạo ra câu hỏi về mục đích sử dụng Mã CIF sau khi tra cứu "
                   "(xem Part B — Q&A).", style="List Bullet 2")

# === A.3 Pre-conditions ===
doc.add_heading("A.3. Bảng Điều Kiện Tiên Quyết & Cấu Hình", level=3)
precond_table = doc.add_table(rows=5, cols=2, style="Table Grid")
precond_table.rows[0].cells[0].text = "Điều kiện"
precond_table.rows[0].cells[1].text = "Chi tiết"
precond_data = [
    ("Phân quyền", "User phải có quyền truy cập menu Báo cáo >> Báo cáo tổng doanh thu phí"),
    ("Dữ liệu mồi", "Hệ thống cần có dữ liệu giao dịch thu phí đã được ghi nhận (Biểu phí, Code phí, Chi nhánh)"),
    ("Khối user", "Xác định Khối của user đăng nhập (KHCN/KHDN/KHDNL/Khác) để hiển thị đúng trường Khối"),
    ("Master data", "Danh sách Biểu phí, Code phí, Chi nhánh, Loại tính phí phải sẵn sàng trong hệ thống"),
]
for i, (k, v) in enumerate(precond_data):
    precond_table.rows[i+1].cells[0].text = k
    precond_table.rows[i+1].cells[1].text = v

# === A.4 Quy Tắc Chung Áp Dụng ===
doc.add_heading("A.4. Quy Tắc Chung Áp Dụng", level=3)
qtc_table = doc.add_table(rows=11, cols=2, style="Table Grid")
qtc_table.rows[0].cells[0].text = "QTC"
qtc_table.rows[0].cells[1].text = "Nội dung áp dụng cho US31"
qtc_data = [
    ("QTC-01.1", "Combobox: Khối, Mã Chi nhánh, Biểu phí, Code phí — cho phép gõ text tìm kiếm"),
    ("QTC-01.2", "Dropdown List: Loại tính phí — chỉ chọn 1 giá trị"),
    ("QTC-01.4", "Number: Doanh thu phí (Nguyên tệ), Doanh thu phí (VND), Tổng cộng — format phân cách hàng nghìn"),
    ("QTC-01.5", "Date: Từ ngày, Đến ngày — dd/mm/yyyy, Từ ngày=00:00:00.000, Đến ngày=23:59:59.999"),
    ("QTC-01.7", "Trường dữ liệu rỗng trên lưới → hiển thị blank"),
    ("QTC-04", "Tra cứu: Mặc định không hiển thị danh sách khi vào màn hình. Xóa tra cứu → về mặc định"),
    ("QTC-05", "Tải xuống: Excel/PDF, tên file = 'Báo cáo tổng doanh thu phí - yyyymmddhhmmss'. Luôn cho phép tải dù không có dữ liệu"),
    ("QTC-06", "Phân trang: 50 bản ghi/trang mặc định. Sort mặc định theo ngày update/tạo giảm dần"),
    ("QTC-09", "Tra cứu CIF: Popup tra cứu CIF theo Quy tắc chung"),
    ("QTC-10", "Phân quyền dữ liệu theo Khối: lọc dữ liệu trên lưới, không lọc dropdown"),
]
for i, (code, desc) in enumerate(qtc_data):
    qtc_table.rows[i+1].cells[0].text = code
    qtc_table.rows[i+1].cells[1].text = desc

# Bold header rows
for table in doc.tables:
    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True

out_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "US31_PartA_Summary.docx")
doc.save(out_path)
print(f"✅ Đã tạo: {out_path}")
