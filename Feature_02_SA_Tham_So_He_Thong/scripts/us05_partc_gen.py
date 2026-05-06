"""Generate US05 Part C - Test Case Coverage .docx"""
import sys, os
sys.path.insert(0, os.path.dirname(__file__))
from us05_partc_data import SCENARIOS

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUTPUT = os.path.join(os.path.dirname(__file__), "..", "US05_Part_C_TestCase_Coverage.docx")

doc = Document()

# Page setup: Landscape, Narrow margins
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

# Title
title = doc.add_heading("PHẦN C: BẢNG TỔNG HỢP TEST CASE ĐỀ XUẤT (Test Case Coverage)", level=1)
title.runs[0].font.size = Pt(16)

doc.add_paragraph(
    "US05 — Thiết lập Quy tắc tính phí cho Code phí\n"
    "Tổng số kịch bản: {} SC | Tổng TC dự kiến: {}".format(
        len(SCENARIOS), sum(s[5] for s in SCENARIOS)
    )
)

# Summary by group
from collections import Counter
group_counts = Counter(s[3] for s in SCENARIOS)
summary_p = doc.add_paragraph("Phân bổ theo nhóm: ")
for g, c in group_counts.items():
    summary_p.add_run(f"{g}: {c} SC  |  ")

doc.add_paragraph("")

# Table
HEADERS = ["Mã SC", "Feature", "Module", "Loại Test Case", "Tên Test Case / Kịch bản", "SL TC", "Trích dẫn tài liệu"]
COL_WIDTHS = [Cm(1.5), Cm(3.0), Cm(3.0), Cm(3.0), Cm(8.5), Cm(1.2), Cm(6.5)]

table = doc.add_table(rows=1, cols=len(HEADERS))
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
hdr = table.rows[0]
for i, h in enumerate(HEADERS):
    cell = hdr.cells[i]
    cell.text = h
    cell.width = COL_WIDTHS[i]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.bold = True

# Group colors
GROUP_COLORS = {
    "Happy Path": RGBColor(0x22, 0x8B, 0x22),
    "Negative Path": RGBColor(0xCC, 0x00, 0x00),
    "Boundary Value": RGBColor(0x00, 0x66, 0xCC),
    "UI/UX & Field Validation": RGBColor(0x33, 0x66, 0xFF),
    "Business Logic": RGBColor(0x99, 0x33, 0xCC),
    "Data Integrity": RGBColor(0xFF, 0x66, 0x00),
    "NFR": RGBColor(0xFF, 0x99, 0x00),
}

# Data rows
for sc in SCENARIOS:
    row = table.add_row()
    values = [sc[0], sc[1], sc[2], sc[3], sc[4], str(sc[5]), sc[6]]
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = val
        cell.width = COL_WIDTHS[i]
        p = cell.paragraphs[0]
        for run in p.runs:
            run.font.size = Pt(8)
        if i == 3:  # Color the test type column
            color = GROUP_COLORS.get(val)
            if color:
                for run in p.runs:
                    run.font.color.rgb = color
                    run.font.bold = True
        if i == 5:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Set column widths via XML
for row in table.rows:
    for i, cell in enumerate(row.cells):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = tcPr.find(qn('w:tcW'))
        if tcW is None:
            tcW = tc.makeelement(qn('w:tcW'), {})
            tcPr.append(tcW)
        tcW.set(qn('w:w'), str(int(COL_WIDTHS[i].emu / 914)))
        tcW.set(qn('w:type'), 'dxa')

doc.save(OUTPUT)
print(f"✅ File đã tạo: {os.path.abspath(OUTPUT)}")
print(f"📊 Tổng: {len(SCENARIOS)} scenarios")
