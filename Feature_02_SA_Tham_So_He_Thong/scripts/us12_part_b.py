"""US12 Part B - Q&A cho BA"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10)

title = doc.add_heading('US12 — PHẦN B: DANH SÁCH CẢNH BÁO & Q&A (Dành Cho BA)', level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)
    run.font.name = 'Times New Roman'
doc.add_paragraph('Feature: Khai báo CTƯĐ áp dụng cho danh sách khách hàng | Phiên bản: v1.0 | Ngày: 13/05/2026')

# Q&A Data
qa_data = [
    # Hạng mục 1: Nghiệp vụ / Luồng xử lý
    ("US12-QA-01.1", "Nghiệp vụ",
     'Mục "Yêu cầu nghiệp vụ", đoạn "Điều kiện áp dụng"',
     'Tài liệu ghi Nguồn dữ liệu cho phép chọn "ETL KH/TK/Thẻ/API giao dịch", nhưng Bảng Mô tả chi tiết các trường (phân vùng Điều kiện áp dụng) chỉ liệt kê 3 giá trị: ETL KH, ETL TK, ETL Thẻ — thiếu "API giao dịch". Vậy giá trị "API giao dịch" có được đưa vào dropdown Nguồn dữ liệu không?',
     'Đề xuất: Bổ sung "API giao dịch" vào dropdown Nguồn dữ liệu tại Bảng mô tả, hoặc BA làm rõ lý do loại trừ.'),

    ("US12-QA-01.2", "Nghiệp vụ",
     'Mục "Chỉnh sửa CTƯĐ có khai báo danh sách khách hàng", đoạn trạng thái "Đang hiệu lực"',
     'Khi CTƯĐ ở trạng thái "Đang hiệu lực", tài liệu cho phép chỉnh sửa "Ngày hết hiệu lực" và "Chi tiết ưu đãi". Flowchart (Bước 3.a) ghi "Hiển thị màn hình chỉ cho phép chỉnh Ngày hết hiệu lực và chi tiết ưu đãi". Tuy nhiên, tài liệu KHÔNG ghi rõ: khi chỉnh sửa Chi tiết ưu đãi ở trạng thái "Đang hiệu lực", người dùng có được thay đổi Toggle "Ưu đãi theo tỷ lệ" (On/Off) hay không?',
     'Đề xuất: Disable toggle Ưu đãi theo tỷ lệ khi trạng thái = Đang hiệu lực, chỉ cho sửa giá trị trong bản ghi hiện tại.'),

    ("US12-QA-01.3", "Nghiệp vụ",
     'Mục "Chỉnh sửa CTƯĐ", Flowchart Chỉnh sửa — Bước 3.1',
     'Bước 3.1 ghi "Hiển thị màn hình chỉnh sửa cho phép sửa toàn bộ CTƯĐ trừ mã CTƯĐ". Tuy nhiên, phần Yêu cầu nghiệp vụ ghi "trừ mã CTƯĐ VÀ trường Số văn bản_Tên viết tắt CTƯĐ". Flowchart thiếu ràng buộc readonly cho trường "Số VB_Tên viết tắt CTƯĐ" khi Chưa hiệu lực. Cần xác nhận trường "Số VB_Tên viết tắt CTƯĐ" có readonly ở cả 2 trạng thái (Chưa hiệu lực + Đang hiệu lực)?',
     'Đề xuất: Trường "Số VB_Tên viết tắt CTƯĐ" luôn readonly sau khi tạo thành công (đã duyệt), vì đây là mã định danh bổ trợ.'),

    ("US12-QA-01.4", "Nghiệp vụ",
     'Mục "Yêu cầu nghiệp vụ", đoạn "Chu kỳ áp dụng"',
     'Khi người dùng chọn "Theo chu kỳ" rồi khai báo xong, sau đó đổi sang "Liên tục" (hoặc ngược lại) — hệ thống có tự động clear các giá trị đã chọn trong phân vùng Chu kỳ áp dụng (Thứ, Ngày cụ thể, Ngày trong tháng, Tuần) hay không? Tài liệu chỉ ghi clear cho toggle Ưu đãi theo tỷ lệ, chưa đề cập clear cho Chu kỳ.',
     'Đề xuất: Hệ thống nên clear giá trị Chu kỳ cũ khi user chuyển đổi giữa "Theo chu kỳ" và "Liên tục" để tránh lưu dữ liệu rác.'),

    ("US12-QA-01.5", "Nghiệp vụ",
     'Mục "Chi tiết ưu đãi", đoạn "Upload file"',
     'Khi chỉnh sửa ở trạng thái "Đang hiệu lực", tài liệu ghi "Upload file để thay thế toàn bộ các thông tin Chi tiết ưu đãi cũ". Cần làm rõ: sau khi upload file mới, các bản ghi KH cũ đã có giao dịch áp dụng ưu đãi (đã tích lũy Ngưỡng dừng) có bị reset giá trị tích lũy hay không?',
     'Đề xuất: BA làm rõ xử lý dữ liệu tích lũy (Số lần ưu đãi / Số tiền giao dịch) khi upload thay thế danh sách KH.'),

    ("US12-QA-01.6", "Nghiệp vụ",
     'Mục "Yêu cầu nghiệp vụ", đoạn "Khi khai báo CTƯĐ..."',
     'Tài liệu yêu cầu "người dùng bắt buộc phải khai báo trước các thông tin: Đối tượng thu phí, Khối, Loại ưu đãi, Ngày hiệu lực, Ngày hết hiệu lực và áp dụng Chu kỳ hay Liên tục" trước khi khai báo Chi tiết ưu đãi. Hệ thống có enforce thứ tự này không (VD: disable phân vùng Chi tiết ưu đãi cho đến khi các trường trên đã điền)? Hay chỉ validate khi nhấn Xác nhận?',
     'Đề xuất: FE nên disable phân vùng Chi tiết ưu đãi cho đến khi các trường bắt buộc ở trên đã có giá trị, tránh user nhập thiếu thông tin.'),

    ("US12-QA-01.7", "Nghiệp vụ",
     'Mục "Lưu đồ Chỉnh sửa", Flowchart Bước 3.a',
     'Flowchart Chỉnh sửa — khi trạng thái "Đang hiệu lực", Bước 3.a ghi cho phép chỉnh "Ngày hết hiệu lực và chi tiết ưu đãi". Tuy nhiên Mockup UI chỉnh sửa KHÔNG được cung cấp trong tài liệu. Cần mockup để xác nhận: những trường nào ở trạng thái readonly/disabled trên giao diện khi CTƯĐ đang hiệu lực?',
     'Đề xuất: BA cung cấp mockup màn hình chỉnh sửa khi CTƯĐ ở trạng thái "Đang hiệu lực".'),

    ("US12-QA-01.8", "Nghiệp vụ",
     'Mục "Yêu cầu nghiệp vụ", đoạn "Ngày hiệu lực, Ngày hết hiệu lực" của Chi tiết ưu đãi',
     'Tài liệu ghi "Ngày hiệu lực, Ngày hết hiệu lực của mỗi KH... nếu có điền thì sẽ ưu tiên áp dụng hơn Ngày hiệu lực/Hết hiệu lực của CTƯĐ (Override)". Cần làm rõ: Ngày hiệu lực KH có được phép NHỎ HƠN Ngày hiệu lực CTƯĐ hay phải nằm trong khoảng [Ngày HL CTƯĐ, Ngày HHL CTƯĐ]?',
     'Đề xuất: Ràng buộc Ngày hiệu lực KH ≥ Ngày hiệu lực CTƯĐ và Ngày HHL KH ≤ Ngày HHL CTƯĐ để tránh mâu thuẫn logic.'),

    # Hạng mục 2: Giới hạn hệ thống
    ("US12-QA-02.1", "Giới hạn",
     'Mục "Kiểm tra Danh sách khách hàng upload"',
     'Tài liệu không quy định giới hạn số dòng tối đa trong file upload danh sách khách hàng. Nếu user upload file có hàng nghìn hoặc hàng chục nghìn bản ghi, hệ thống xử lý thế nào? Có giới hạn tối đa số KH trong 1 CTƯĐ không?',
     'Đề xuất: BA xác nhận giới hạn tối đa số bản ghi trong file upload (VD: 500/1000/5000 bản ghi) và hành vi khi vượt ngưỡng.'),

    ("US12-QA-02.2", "Giới hạn",
     'Mục "Điều kiện áp dụng", đoạn "Thêm nhóm điều kiện"',
     'Tài liệu không giới hạn số lượng Nhóm điều kiện và số dòng điều kiện chi tiết trong mỗi nhóm. Có giới hạn tối đa không?',
     'Đề xuất: BA xác nhận giới hạn tối đa số Nhóm điều kiện (VD: 10) và số dòng điều kiện trong mỗi nhóm (VD: 20).'),

    ("US12-QA-02.3", "Giới hạn",
     'Mục "Chi tiết ưu đãi", đoạn "Ngưỡng dừng ưu đãi"',
     'Trường "Ngưỡng dừng ưu đãi" chỉ yêu cầu "số dương > 0" nhưng không quy định giá trị tối đa. Nếu nhập giá trị cực lớn (VD: 999,999,999,999) thì hệ thống xử lý ra sao?',
     'Đề xuất: BA xác nhận giá trị tối đa cho phép của trường Ngưỡng dừng ưu đãi.'),

    # Hạng mục 3: Toàn vẹn dữ liệu
    ("US12-QA-03.1", "Toàn vẹn dữ liệu",
     'Mục "Chi tiết ưu đãi", đoạn SPDV',
     'Tài liệu yêu cầu SPDV phải có trạng thái = Hoạt động. Nếu sau khi CTƯĐ đã được phê duyệt thành công, một SPDV trong danh sách bị chuyển sang trạng thái "Không hoạt động" (do admin thay đổi ở danh mục SPDV), CTƯĐ có bị ảnh hưởng gì không?',
     'Đề xuất: BA làm rõ cascade behavior: CTƯĐ có tự động vô hiệu hóa dòng Chi tiết ưu đãi liên quan đến SPDV bị ngừng hoạt động hay tiếp tục áp dụng cho đến hết hiệu lực?'),

    ("US12-QA-03.2", "Toàn vẹn dữ liệu",
     'Mục "Kiểm tra Danh sách khách hàng upload", Ràng buộc CIF',
     'Tài liệu chỉ kiểm tra CIF "tồn tại trong bảng ETL dữ liệu khách hàng". Không kiểm tra thêm: CIF có trạng thái = Hoạt động / Tạm dừng hay không? Nếu CIF đã bị đóng/tạm dừng thì vẫn cho phép thêm vào CTƯĐ?',
     'Đề xuất: Bổ sung kiểm tra trạng thái CIF = "Hoạt động" khi upload hoặc thêm mới KH.'),

    ("US12-QA-03.3", "Toàn vẹn dữ liệu",
     'Mục "Kiểm tra Danh sách khách hàng upload", Ràng buộc SPDV cha-con',
     'Tài liệu ghi: "Không tồn tại bản ghi khác trong danh sách có cùng CIF và SPDV là SPDV cấp cha của bản ghi đang kiểm tra". Câu hỏi: Ràng buộc cha-con này chỉ kiểm tra trong phạm vi file upload hiện tại, hay bao gồm cả các bản ghi đã tồn tại trước đó trong CTƯĐ (khi chỉnh sửa thêm KH mới)?',
     'Đề xuất: Kiểm tra ràng buộc cha-con SPDV trên toàn bộ danh sách (bao gồm cả bản ghi cũ + mới) khi chỉnh sửa.'),

    ("US12-QA-03.4", "Toàn vẹn dữ liệu",
     'Mục "Kiểm tra Danh sách khách hàng upload", Ràng buộc CIF+SPDV',
     'Ràng buộc "SPDV và CIF không bị trùng với bản ghi khác trong Danh sách" — phạm vi kiểm tra chỉ trong 1 file upload hay bao gồm cả CIF+SPDV đã tồn tại ở các CTƯĐ khác đang hiệu lực? Một KH có thể tham gia nhiều CTƯĐ với cùng SPDV không?',
     'Đề xuất: BA xác nhận scope unique của cặp CIF+SPDV: chỉ unique trong 1 CTƯĐ hay unique trên toàn bộ CTƯĐ đang hiệu lực?'),

    # Hạng mục 4: UI/UX
    ("US12-QA-04.1", "UI-UX",
     'Mockup UI — Phân vùng Thông tin chung',
     'Mockup hiển thị trường "Loại ưu đãi" (Combobox) — nhưng Bảng mô tả chi tiết các trường dùng tên "Loại ưu đãi" (Dropdown list) với 3 giá trị: Theo KH, Theo TK, Theo Thẻ. Tuy nhiên tài liệu yêu cầu nghiệp vụ (mục Navigation) lại KHÔNG đề cập đến trường "Loại ưu đãi" mà dùng khái niệm "Loại ưu đãi" trong phần Điều kiện. Cần xác nhận: Mockup hiển thị label là "Loại ưu đãi" hay "Loại ưu đãi" đúng là trường trên form?',
     'Đề xuất: Thống nhất tên label trên Mockup và Bảng mô tả. BA xác nhận label chính xác.'),

    ("US12-QA-04.2", "UI-UX",
     'Mockup UI — Phân vùng Đối tượng thu phí',
     'Mockup hiển thị "Đối tượng thu phí" dạng Checkbox (☐ Khách hàng ☑ Merchant) — cho phép chọn nhiều. Nhưng Bảng mô tả ghi "Dropdown list, chọn 1 trong 2 giá trị". Đây là mâu thuẫn UI: Checkbox cho phép multi-select, Dropdown chỉ cho single-select.',
     'Đề xuất: BA xác nhận component là Dropdown (single-select) hay Checkbox (multi-select) cho trường Đối tượng thu phí.'),

    ("US12-QA-04.3", "UI-UX",
     'Mockup UI — Phân vùng Chi tiết ưu đãi',
     'Mockup hiển thị nút "Tải lên" và "Tải xuống" ở bên phải. Tuy nhiên Mockup "Theo chu kỳ" (image5) chỉ có nút "Tải xuống" mà KHÔNG có nút "Tải lên". Trong khi Mockup "Liên tục" (image6) lại có cả "Tải lên" + "Tải xuống". Đây có phải lỗi mockup không?',
     'Đề xuất: BA xác nhận: Nút "Tải lên" phải hiển thị ở CẢ HAI chế độ (Theo chu kỳ và Liên tục).'),

    ("US12-QA-04.4", "UI-UX",
     'Mockup UI — Phân vùng Điều kiện áp dụng',
     'Mockup hiển thị Dropdown "And/Or" giữa các Nhóm điều kiện (VD: giữa "Nhóm điều kiện số 1" và "Nhóm điều kiện số 2") — và hiển thị "And" / "Or" giữa các dòng điều kiện trong cùng một nhóm. Tuy nhiên, Bảng mô tả chi tiết các trường KHÔNG đề cập đến trường/logic "And/Or". Cần BA xác nhận hành vi kết hợp điều kiện.',
     'Đề xuất: BA bổ sung mô tả logic kết hợp giữa các nhóm điều kiện (AND/OR) và giữa các dòng điều kiện trong cùng nhóm.'),

    ("US12-QA-04.5", "UI-UX",
     'Bảng Mô tả chi tiết các trường — Trường "Số văn bản_Tên viết tắt CTƯĐ"',
     'Ràng buộc ghi "Phía trước dấu gạch dưới chỉ nhận các ký tự là số". Tuy nhiên tài liệu KHÔNG nói rõ: phía SAU dấu gạch dưới chỉ nhận ký tự chữ (alphabet) hay cho phép cả số+chữ? VD: "123_ABC" hợp lệ, nhưng "123_A1B" có hợp lệ không?',
     'Đề xuất: BA xác nhận quy tắc ký tự cho phần sau dấu gạch dưới trong trường Số VB_Tên viết tắt CTƯĐ.'),

    ("US12-QA-04.6", "UI-UX",
     'Bảng "Kiểm tra Danh sách khách hàng upload" — Cột Ngày hiệu lực/Ngày hết hiệu lực',
     'Bảng kiểm tra upload ghi định dạng Date là "dd-mm-yyyy" (dấu gạch ngang). Trong khi QTC-01.5 quy định format Date hệ thống là "dd/mm/yyyy" (dấu gạch chéo). Đây là mâu thuẫn format Date giữa file upload và hệ thống.',
     'Đề xuất: Thống nhất format Date trong file upload = dd/mm/yyyy theo QTC-01.5 để tránh nhầm lẫn cho user.'),
]

# Build table
table = doc.add_table(rows=1, cols=6, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['ID', 'Trích xuất (Reference)', 'Câu hỏi / Sự cố', 'Phân loại', 'Đề xuất từ QA', 'Trả lời của BA']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(9)

# Set column widths
from docx.shared import Inches
col_widths = [Inches(0.9), Inches(2.0), Inches(3.5), Inches(0.9), Inches(2.8), Inches(1.2)]
for row in table.rows:
    for i, w in enumerate(col_widths):
        row.cells[i].width = w

for qa in qa_data:
    row = table.add_row()
    row.cells[0].text = qa[0]
    row.cells[1].text = qa[2]
    row.cells[2].text = qa[3]
    row.cells[3].text = qa[1]
    row.cells[4].text = qa[4]
    row.cells[5].text = ''  # BA fills
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)

out_path = '/Users/mac/antigravity-testing-kit/Feature_02_SA_Tham_So_He_Thong/US12_PartB_QA.docx'
doc.save(out_path)
print(f'Part B saved: {out_path}')
