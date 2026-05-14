"""US32 Part B Merged - DOCX Generator"""
import sys, os
sys.path.insert(0, os.path.dirname(__file__))
from us32_partB_merged_data import HM1_DATA, HM2_DATA, HM3_DATA, HM4_DATA

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

doc.add_heading('US32 — PHẦN B TỔNG HỢP: DANH SÁCH CẢNH BÁO & Q&A', level=1)
doc.add_paragraph('Feature: Dashboard — Báo cáo tổng quan hoạt động thu phí')
doc.add_paragraph('Phiên bản: v1.0 (Merged AI + VA) | Ngày: 14/05/2026')
doc.add_paragraph('Ghi chú: Câu có tag [VA] là bổ sung từ Validation Analyst.')
doc.add_paragraph('─' * 80)

headers = ['ID', 'Trích xuất (Reference)', 'Câu hỏi / Sự cố', 'Phân loại', 'Đề xuất từ QA', 'Trả lời của BA']
col_widths = [Cm(2.0), Cm(4.5), Cm(7.0), Cm(2.5), Cm(6.0), Cm(5.0)]

sections_config = [
    ('🔶 Hạng mục 1: Vấn đề Nghiệp vụ / Luồng xử lý', HM1_DATA),
    ('🔴 Hạng mục 2: Giới hạn hệ thống & Exception Handling', HM2_DATA),
    ('🟠 Hạng mục 3: Toàn vẹn dữ liệu & Ràng buộc', HM3_DATA),
    ('🔵 Hạng mục 4: UI/UX & Giao diện', HM4_DATA),
]

def add_table(data):
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, w in enumerate(col_widths):
        table.columns[i].width = w
    # Header
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), '003366')
        cell._element.get_or_add_tcPr().append(shading)
    # Data rows
    for row_data in data:
        row = table.add_row()
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ''
            run = cell.paragraphs[0].add_run(text)
            run.font.size = Pt(8.5)
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
        # Empty BA column
        row.cells[5].text = ''
    table.autofit = True

for title, data in sections_config:
    doc.add_heading(title, level=2)
    add_table(data)
    doc.add_paragraph('')  # spacing

# Summary
doc.add_heading('Tổng kết', level=2)
total = len(HM1_DATA) + len(HM2_DATA) + len(HM3_DATA) + len(HM4_DATA)
doc.add_paragraph(f'Tổng số câu hỏi: {total} (Nghiệp vụ: {len(HM1_DATA)}, Giới hạn: {len(HM2_DATA)}, Toàn vẹn DL: {len(HM3_DATA)}, UI/UX: {len(HM4_DATA)})')

output = '/Users/mac/antigravity-testing-kit/Feature_02_SA_Tham_So_He_Thong/US32_PartB_QA_Merged.docx'
doc.save(output)
print(f'✅ Đã sinh file: {output}')
print(f'📊 Tổng: {total} câu hỏi')
