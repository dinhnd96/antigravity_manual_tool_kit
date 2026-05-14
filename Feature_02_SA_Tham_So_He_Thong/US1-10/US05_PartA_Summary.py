"""
US05 - Part A: Tóm Tắt Nghiệp Vụ (Requirements Breakdown)
Quy tắc tính phí cho Code phí
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Page setup: Landscape, Narrow margins
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(10)

# Title
title = doc.add_heading('US05 — PHẦN A: TÓM TẮT NGHIỆP VỤ (Dành cho Tester)', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')

# A.1. Core Business Value
h = doc.add_heading('A.1. Thông Điệp Cốt Lõi (Core Business Value)', level=2)
p = doc.add_paragraph()
p.add_run('Tính năng: ').bold = True
p.add_run('Định nghĩa Quy tắc tính phí cho từng Code phí.')
p = doc.add_paragraph()
p.add_run('Người dùng cuối: ').bold = True
p.add_run('Maker (người khai báo Code phí) — thuộc bộ phận quản lý sản phẩm/phí.')
p = doc.add_paragraph()
p.add_run('Mục đích: ').bold = True
p.add_run('Cho phép người dùng thiết lập quy tắc để tính mức phí áp dụng cho từng đối tượng khách hàng, tùy chỉnh theo từng thời kỳ. '
           'Đây là 1 phân vùng BẮT BUỘC khai báo khi thêm mới/chỉnh sửa Code phí.')
p = doc.add_paragraph()
p.add_run('Navigation: ').bold = True
p.add_run('Tham số >> Danh mục SPDV >> SPDV cấp cuối >> Thêm mới Code phí >> Quy tắc tính phí.')
p = doc.add_paragraph()
p.add_run('Lưu ý quan trọng: ').bold = True
p.add_run('US05 KHÔNG phải là màn hình độc lập — nó là một PHÂN VÙNG nằm trong màn hình Thêm mới/Chỉnh sửa Code phí (US02). '
           'Việc "Xác nhận" (Lưu) cuối cùng thuộc về US02, không phải US05.')

# A.2. Flow Structure & Module Mapping
doc.add_heading('A.2. Cấu Trúc Luồng Nghiệp Vụ & Phân Bổ Module', level=2)

# Module 1: Chọn Loại Quy Tắc
doc.add_heading('Module 1: Chọn Loại Quy Tắc Tính Phí', level=3)
doc.add_paragraph('Happy Path:', style='List Bullet')
items = [
    'Sau khi khai báo Thông tin chung của Code phí, người dùng chọn 1 trong 3 loại quy tắc: Số cố định / Thỏa thuận / Công thức.',
    'Giao diện phân vùng "Quy tắc tính phí" thay đổi tùy theo loại quy tắc được chọn.',
    'Nếu tham số "Khai báo theo nhóm khách hàng" = Có → hệ thống cho phép khai báo nhiều bản ghi, mỗi bản ghi tương ứng 1 nhóm KH.',
    'Nếu "Khai báo theo nhóm khách hàng" = Không → chỉ khai báo 1 bản ghi duy nhất.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet 2')

doc.add_paragraph('Luồng rẽ nhánh / Ngoại lệ:', style='List Bullet')
items_ex = [
    'Chọn loại quy tắc khác nhau → FE ẩn/hiện các trường tương ứng (dynamic form).',
    'Khi "Khai báo theo nhóm KH" = Có: xuất hiện nút "Thêm nhóm KH", "Xóa", "Sao chép" (chỉ ở Công thức).',
    'Nút "Xóa" chỉ hiển thị từ Nhóm KH thứ 2 trở đi (nhóm đầu tiên không được phép xóa).',
]
for item in items_ex:
    doc.add_paragraph(item, style='List Bullet 2')

# Module 2: Số cố định
doc.add_heading('Module 2: Quy Tắc — Số Cố Định', level=3)
doc.add_paragraph('Happy Path:', style='List Bullet')
items = [
    'FE hiển thị 1 trường duy nhất: "Số cố định" (Number, bắt buộc ★, giá trị >= 0).',
    'Người dùng nhập giá trị → FE ghi nhận → chuyển đến bước Xác nhận Code phí (US02).',
    'Nếu "Khai báo theo nhóm KH" = Có → mỗi nhóm KH có 1 trường "Số cố định" riêng.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet 2')

# Module 3: Thỏa thuận
doc.add_heading('Module 3: Quy Tắc — Thỏa Thuận', level=3)
doc.add_paragraph('Happy Path:', style='List Bullet')
items = [
    'FE hiển thị 2 trường: "Tối thiểu" (★, Number, >= 0) và "Tối đa" (★, Number, >= 0).',
    'Đây là mức phí tối thiểu / tối đa có thể thu — phí thực tế được thỏa thuận tại từng giao dịch.',
    'Nếu "Khai báo theo nhóm KH" = Có → mỗi nhóm KH có cặp Tối thiểu/Tối đa riêng.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet 2')

# Module 4: Công thức
doc.add_heading('Module 4: Quy Tắc — Công Thức (Phức tạp nhất)', level=3)
doc.add_paragraph('Happy Path:', style='List Bullet')
items = [
    'FE hiển thị các trường: Loại tiền tối thiểu/tối đa (★ Dropdown), Nguồn dữ liệu (★ Dropdown), Công thức (◎ readonly), Tối thiểu (－ Number), Tối đa (－ Number).',
    'Người dùng nhập Loại tiền tối thiểu/tối đa, Nguồn dữ liệu → nhấn "Thiết lập" để mở MÀN HÌNH THIẾT LẬP CÔNG THỨC.',
    'Tại màn hình Thiết lập công thức: mặc định 1 cấu phần, cho phép "Thêm cấu phần". Toán tử giữa cấu phần: + hoặc −.',
    'Nhấn icon "Thiết lập cấu phần" → mở MÀN HÌNH THIẾT LẬP CẤU PHẦN.',
    'Tại màn hình Thiết lập cấu phần: khai báo ít nhất 1 định dạng (Cố định / Tỷ lệ / Ngày). Toán tử giữa định dạng: * (nhân) mặc định.',
    'Nhấn "Xác nhận" ở cấu phần → quay về Thiết lập công thức → nhấn "Xác nhận" → quay về Code phí.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet 2')

doc.add_paragraph('Luồng rẽ nhánh / Ngoại lệ:', style='List Bullet')
items_ex = [
    'Nút "Đóng" ở Thiết lập cấu phần (Bước 7.1) → quay về Thiết lập công thức, KHÔNG lưu [QTC-15].',
    'Nút "Đóng" ở Thiết lập công thức (Bước 8.1) → quay về Khai báo Code phí, KHÔNG lưu [QTC-15].',
    'Nguồn dữ liệu = API Giao dịch (khi Loại tính phí = Theo giao dịch): trường readonly, không cho sửa.',
    'Nguồn dữ liệu = ETL Tài khoản → KHÔNG được khai báo điều kiện theo Thẻ.',
    'Nguồn dữ liệu = ETL Thẻ → KHÔNG được khai báo điều kiện theo Tài khoản.',
    'Nguồn dữ liệu = ETL Khách hàng → KHÔNG được khai báo điều kiện theo cả Tài khoản lẫn Thẻ.',
    'Cấu phần đầu tiên KHÔNG có nút Xóa (chỉ từ cấu phần thứ 2 trở đi).',
    'Định dạng đầu tiên KHÔNG có nút Xóa (chỉ từ định dạng thứ 2 trở đi).',
    'Phải khai báo ít nhất 1 định dạng mỗi cấu phần, không được bỏ trống.',
]
for item in items_ex:
    doc.add_paragraph(item, style='List Bullet 2')

# Sub-module: Định dạng trong cấu phần
doc.add_heading('Module 4.1: Chi Tiết Định Dạng Trong Cấu Phần', level=3)
# Bảng tóm tắt 3 loại định dạng
table = doc.add_table(rows=4, cols=4, style='Light List Accent 1')
headers = ['Định dạng', 'Trường hiển thị', 'Ràng buộc', 'Ghi chú']
for i, h_text in enumerate(headers):
    table.rows[0].cells[i].text = h_text
data = [
    ['Cố định', 'Giá trị (Number)', 'Số > 0', 'Khai báo 1 số cố định cho công thức'],
    ['Tỷ lệ', 'Tỷ lệ (Number), Tên trường (Combobox)', 'Tỷ lệ: số > 0 và <= 1.\nTên trường: chọn từ Điều kiện tính phí (Number, Hoạt động)', 'Khai báo cặp Tỷ lệ * Giá trị'],
    ['Ngày', 'Ngày bắt đầu, Ngày kết thúc (Combobox), Basis (Number)', 'Ngày BD ≠ Ngày KT.\nBasis: số nguyên dương > 0.\nĐiều kiện: Date, Hoạt động', 'Khoảng thời gian bị tính phí. VD: (Value_Date - Maturity_Date) / 365'],
]
for r_idx, row_data in enumerate(data, 1):
    for c_idx, val in enumerate(row_data):
        table.rows[r_idx].cells[c_idx].text = val

# Module 5: Nhóm KH
doc.add_heading('Module 5: Khai Báo Theo Nhóm Khách Hàng', level=3)
items = [
    'Điều kiện kích hoạt: Tham số "Khai báo theo nhóm KH" tại phân vùng Thông tin chung = Có.',
    'Mỗi nhóm KH = 1 bản ghi Quy tắc tính phí riêng biệt, có thể có mức phí / công thức khác nhau.',
    'Dropdown "Nhóm KH": hiển thị Danh mục Nhóm KH trạng thái Hoạt động (tham chiếu US27).',
    '"Sao chép": Sao chép bản ghi quy tắc tính phí từ nhóm KH này sang nhóm KH khác (chỉ có ở Công thức).',
    'Nút "Xóa" chỉ hiển thị ở nhóm KH thứ 2 trở đi.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# A.3. Pre-conditions
doc.add_heading('A.3. Điều Kiện Tiên Quyết & Cấu Hình', level=2)
items = [
    'Code phí đã được tạo tại US02 hoặc đang trong quá trình Thêm mới/Chỉnh sửa.',
    'Phân vùng "Thông tin chung" của Code phí đã được khai báo (đặc biệt: Loại tính phí, Loại tiền tệ, "Khai báo theo nhóm KH").',
    'Danh mục Nhóm khách hàng (US27) đã có dữ liệu mồi trạng thái Hoạt động.',
    'Danh mục Điều kiện tính phí đã có dữ liệu mồi với các Nguồn dữ liệu tương ứng (Trường điều kiện trạng thái Hoạt động).',
    'Luồng Maker-Checker áp dụng cho Code phí [QTC-12].',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# A.4. Quy Tắc Chung Áp Dụng
doc.add_heading('A.4. Quy Tắc Chung Áp Dụng', level=2)
rules = [
    '[QTC-01.2] Dropdown List: Chỉ chọn 1 giá trị (Nhóm KH, Loại tiền, Nguồn dữ liệu, Định dạng).',
    '[QTC-01.4] Number: Phân cách hàng nghìn bằng dấu phẩy, thập phân bằng dấu chấm, 2 chữ số thập phân. Ngoại lệ VND, JPY không có thập phân.',
    '[QTC-01.6] Text: Giới hạn ký tự mặc định — Mã 50 ký tự, Tên 50 ký tự, Ghi chú 300 ký tự.',
    '[QTC-11] Xử lý lỗi: FE-First — FE chặn lỗi trước, BE là tuyến phòng thủ thứ 2.',
    '[QTC-12] Luồng Maker-Checker: Áp dụng cho Code phí (US02). Mã sinh sau khi Last Checker duyệt.',
    '[QTC-14.1] No-Change Guard: Nếu không có thay đổi, FE disable nút Xác nhận hoặc cảnh báo.',
    '[QTC-15] Nút "Đóng": Không popup xác nhận, không lưu, quay về màn hình trước.',
]
for item in rules:
    doc.add_paragraph(item, style='List Bullet')

# Save
output_path = '/Users/mac/antigravity-testing-kit/Feature_02_SA_Tham_So_He_Thong/US05_PartA_Summary.docx'
doc.save(output_path)
print(f'✅ Part A saved: {output_path}')
