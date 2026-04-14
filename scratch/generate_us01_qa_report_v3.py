import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = docx.Document()

    # Title
    title = doc.add_heading('BÁO CÁO PHÂN TÍCH YÊU CẦU - US01 (KHAI BÁO CÂY DANH MỤC SPDV)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ─── PHẦN A ───────────────────────────────────────────────────────────────
    doc.add_heading('Phần A: Tóm Tắt Nghiệp Vụ Chuyên Sâu (Requirements Breakdown)', level=1)

    doc.add_heading('1. Thông điệp cốt lõi (Core Business Value)', level=2)
    doc.add_paragraph(
        'US01 giúp xây dựng "bộ khung" phân cấp cho toàn bộ sản phẩm dịch vụ của ngân hàng. '
        'Đây là nền tảng để gán các loại phí chi tiết ở các US sau (US02, US06).'
    )

    doc.add_heading('2. Luồng chính (Happy Path Workflow)', level=2)
    doc.add_paragraph(
        '1. Maker tạo mới Nghiệm vụ (Cấp 1) -> Mã 01-99 auto-gen.\n'
        '2. Maker tạo mới SPDV cấp chi tiết (dưới cấp cha) -> Mã = Mã cha + 01-99.\n'
        '3. Thiết lập Ngày hiệu lực/Hết hiệu lực.\n'
        '4. Checker phê duyệt yêu cầu (theo US25).\n'
        '5. Hệ thống tự động chuyển trạng thái Hoạt động/Không hoạt động hàng ngày dựa trên ngày hiệu lực.'
    )

    doc.add_heading('3. Các luồng rẽ nhánh / Ngoại lệ (Alternative & Exception Flows)', level=2)
    doc.add_paragraph(
        '• Chỉnh sửa: Chỉ cho phép sửa Ngày hiệu lực & Ngày hết hiệu lực của bản ghi đã duyệt.\n'
        '• Giới hạn mã: Mỗi cấp chỉ chứa tối đa 99 bản ghi con. Vượt quá sẽ báo lỗi.\n'
        '• Ràng buộc phân cấp: Ngày của con phải nằm "gói gọn" trong ngày của cha.'
    )

    doc.add_heading('4. Bảng Điều Kiện Tiên Quyết (Pre-conditions)', level=2)
    doc.add_paragraph(
        '• Phải có phân quyền Maker/Checker.\n'
        '• Danh sách các Nghiệp vụ cấp 1 không được trùng mã.'
    )

    # ─── PHẦN B ───────────────────────────────────────────────────────────────
    doc.add_heading('Phần B: Khai Quật Lỗ Hổng & Danh Sách Q&A (Dành Cho BA)', level=1)

    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['ID', 'Ref_ID', 'Nội dung Câu hỏi / Sự cố', 'Phân loại', 'Đề xuất xử lý', 'BA Answer']
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True

    qa_data = [
        (
            'QA-01', 'Line 16/32',
            'Hạn chế chỉnh sửa: Tại sao hệ thống CHỈ cho phép sửa Ngày hiệu lực/hết hiệu lực? '
            'Nếu Maker gõ sai "Tên Nghiệp vụ" và đã được duyệt, thì cách xử lý là gì? '
            'Xóa đi tạo lại hay phải sửa trực tiếp trong DB?',
            'Nghiệp vụ', 'Cho phép sửa Tên/Mô tả nếu bản ghi chưa được gán vào bất kỳ Code Phí nào.', ''
        ),
        (
            'QA-02', 'Line 18/33/34',
            'Ràng buộc ngày Cha-Con: Nếu Parent Date thay đổi dẫn tới xung đột với 100 Children Date bên dưới, '
            'hệ thống có cảnh báo danh sách các con bị lỗi hay chặn lưu cấp cha?',
            'UI/UX', 'Hiển thị popup danh sách các mã con đang bị conflict ngày để user điều chỉnh.', ''
        ),
        (
            'QA-03', 'Line 9/29',
            'Tái sử dụng Mã (Recycle Code): Nếu mã "05" bị Hủy hoặc Xóa, mã số 05 này có được '
            'cấp lại cho bản ghi tạo mới sau đó không (để tránh lãng phí tài nguyên mã từ 01-99)?',
            'Nghiệp vụ', 'Làm rõ cơ chế auto-gen: lấy Max+1 hay lấy số còn trống.', ''
        ),
        (
            'QA-04', 'Flowchart (Hình 1-2)',
            'Lỗi logic Flowchart: Hình vẽ mô tả "Kiểm tra thay đổi" (Bước 91) nhưng không thấy '
            'nhánh "Quay lại" nếu user không thay đổi gì mà vẫn nhấn Xác nhận.',
            'Document', 'Bổ sung nhánh xử lý nếu dữ liệu sửa trùng khớp hoàn toàn dữ liệu cũ.', ''
        ),
        (
            'QA-TC-01', 'Boundary / Date Edit',
            'Sửa bản ghi Đang hoạt động: Nếu Ngày hiệu lực là 01/01/2026 (quá khứ), hôm nay 14/04/2026 '
            'user muốn sửa Ngày hết hiệu lực. Theo Line 17 (Ngày hiệu lực sửa phải >= Ngày hệ thống), '
            'liệu hệ thống có chặn không cho lưu vì Ngày hiệu lực cũ < Ngày hôm nay?',
            'Test Case / Boundary', 'Cần BA làm rõ rule: Khi sửa, chỉ validate "Ngày thay đổi", không validate lại "Ngày cũ".', ''
        ),
        (
            'QA-TC-02', 'Acceptance Criteria / Limit',
            'Max limit 99: Khi tạo tới bản ghi thứ 100, message báo lỗi chính xác là gì? '
            'Màn hình có chặn nút "Xác nhận" ngay từ đầu không?',
            'Test Case / AC', 'Cung cấp Error Message code/text cụ thể cho trường hợp vượt limit 99.', ''
        ),
        (
            'QA-TC-03', 'Execution / Job timing',
            'Auto-update Status (Line 20): Job chạy tại thời điểm nào (00:00:01 hay theo manual run)? '
            'Nếu bản ghi hết hạn lúc 23:59:59 ngày T, thì giao dịch lúc 00:01:00 ngày T+1 sẽ áp trạng thái nào?',
            'Test Case / Job', 'Xác định thời điểm chuyển đổi trạng thái vật lý trong DB.', ''
        )
    ]

    for row_data in qa_data:
        row_cells = table.add_row().cells
        for i, text in enumerate(row_data):
            row_cells[i].text = text

    doc.save('US01_QnA_Report_V3.docx')
    print('✅ Đã tạo US01_QnA_Report_V3.docx thành công!')

create_report()
