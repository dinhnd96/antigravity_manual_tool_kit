import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = docx.Document()
    
    # Title
    title = doc.add_heading('BÁO CÁO PHÂN TÍCH YÊU CẦU - US15 (DANH MỤC BIỂU PHÍ)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Part A
    doc.add_heading('Phần A: Tóm Tắt Nghiệp Vụ Chuyên Sâu (Requirements Breakdown)', level=1)
    
    doc.add_heading('1. Thông điệp cốt lõi (Core Business Value)', level=2)
    doc.add_paragraph('US15 cung cấp giao diện dạng Data Grid để người dùng xem, tra cứu và lọc toàn bộ Danh sách các Biểu phí đang tồn tại trên hệ thống. Dữ liệu này đóng vai trò như một thư viện Biểu phí độc lập phục vụ cho việc gắn kết với Code Phí sau này.')
    
    doc.add_heading('2. Luồng chính (Happy Path Workflow)', level=2)
    doc.add_paragraph('Người dùng mở màn hình Danh mục biểu phí -> Hệ thống load Grid Data -> Người dùng có thể click Lọc Nâng Cao để tra cứu theo Tên, Trạng thái (Đang/Chưa/Hết hiệu lực), Ngày hiệu lực -> Hoặc bấm Tải xuống Excel để lấy báo cáo.')
    
    doc.add_heading('3. Các luồng rẽ nhánh / Ngoại lệ (Alternative & Exception Flows)', level=2)
    doc.add_paragraph('Flowchart chỉ định: Khi người dùng bấm Xóa Lọc, hệ thống có rẽ nhánh xử lý nhưng luồng này đang kết thúc bằng việc đóng cửa sổ thay vì refresh field.')
    
    doc.add_heading('4. Bảng Điều Kiện Tiên Quyết & Cấu Hình (Pre-conditions & Settings)', level=2)
    doc.add_paragraph('Mọi dữ liệu trên Grid thuộc về kết quả sau khi duyệt từ luồng Tạo mới/Sửa (Tham chiếu Ma trận phê duyệt). Các trường thông tin trống như Ngày sửa/Người duyệt sẽ hiển thị dấu "-" bảo đảm UX.')
    
    doc.add_heading('5. Ma trận Phân Quyền/Dữ Liệu', level=2)
    doc.add_paragraph('Đóng vai trò là màn View Only cho đa số Role, chỉ Maker mới có nút bấm Tạo Mới nhảy sang module khác.')
    
    # Part B
    doc.add_heading('Phần B: Khai Quật Lỗ Hổng & Danh Sách Q&A (Dành Cho BA)', level=1)
    
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['ID', 'Ref_ID (Mục)', 'Nội dung Câu hỏi / Sự cố', 'Phân loại', 'Đề xuất hướng xử lý từ QA', 'Câu trả lời của BA (Để trống)']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        
    qa_data = [
        (
            'QA-01', 
            'Khác biệt UI và Text về Cột hiển thị', 
            'Text mô tả Lưới dữ liệu (Từ Line 11-22) liệt kê các cột: Mã, Tên, Trạng thái, Ngày Hlực, Ngày Hết Hlực, Ngày/Người Tạo, Sửa, Duyệt. NHƯNG Hình UI Mockup 02 lại vẽ các cột: Tên văn bản, Ngày ban hành, Hành động (Và thiếu mất các cột Cấp duyệt/Ngày hết hiệu lực). Râu ông nọ cắm cằm bà kia.', 
            'UI/UX', 
            'BA cần chốt hạ danh sách cột Grid cuối cùng. Xóa hình Mockup cũ hoặc sửa Text URD cho khớp.',
            ''
        ),
        (
            'QA-02', 
            'Sai lệch bộ Trạng Thái ở Lưới', 
            'Theo URD Text (Line 14): Trạng thái chỉ gồm Chưa, Đang, Hết hiệu lực. NHƯNG UI Mockup Lưới dữ liệu (Ảnh 2) lại chình ình chữ "Chờ duyệt", "Ngừng hoạt động". Đây là lỗi sai trầm trọng về Status Flow.', 
            'Nghiệp vụ', 
            'Nếu Lưới này show cả các bản ghi Đang Pending thì URD phải update thêm trạng thái Chờ duyệt vào Text. Nếu không, phải sửa Hình UI nhặt các trạng thái sai ném đi.',
            ''
        ),
        (
            'QA-03', 
            'UX ngột ngạt của Nút Xóa Lọc', 
            'Đây là lỗi hệ thống lặp lại từ US15. Trên BPMN (Hình 1), luồng 8.1 Xóa Lọc -> 9. Đóng màn hình Lọc. Xóa chữ trong ô input mà lại sập luôn cả cửa sổ Popup khiến User phải click mở ra nhập lại.', 
            'UI/UX', 
            'Xóa lọc chỉ clear value ở Input, Dừng đóng Popup.',
            ''
        ),
        (
            'QA-04', 
            'Trùng mã Label ở Flowchart', 
            'Trên Lưu đồ BPMN (Hình 1): Có 2 ô hành động cùng bị đánh số thứ tự là (9). Ô trên là "9. Đóng màn hình lọc", ô dưới là "9. Hiển thị danh sách". Gây lỗi đánh index quy trình.', 
            'Document', 
            'Đánh số lại BPMN (VD: 9.1 và 9.2).',
            ''
        ),
        (
            'QA-05', 
            'Lọc ngày "1 chiều"', 
            'Trường "Từ ngày - Đến ngày" ở form Lọc được gán tĩnh chết cho "Ngày hiệu lực" (Line 8). Vậy muốn tìm các Biểu phí đã hoặc sắp bị HẾT HẠN trong tháng này thì User không thể nào tìm được vì thiếu filter cho "Ngày hết hiệu lực".', 
            'Nghiệp vụ', 
            'Đề xuất bổ sung thêm 1 bộ filter "Ngày hết hiệu lực (Từ-Đến)" vào form Lọc Nâng Cao để trọn vẹn scope tra cứu Master Data.',
            ''
        ),
        ('QA-TC-01', 'Validation / Input', 'Điều kiện filter Ngày: Nếu người dùng nhập (Từ ngày > Đến ngày) thì message lỗi hiển thị chính xác string là gì? Có chặn submit không?', 'Test Case / Error Msg', 'Bổ sung mã lỗi và message text cho UI validate.', '')
    ]
    
    for row_data in qa_data:
        row_cells = table.add_row().cells
        for i, text in enumerate(row_data):
            row_cells[i].text = text
            
    doc.save('US15_QnA_Report.docx')

create_report()
