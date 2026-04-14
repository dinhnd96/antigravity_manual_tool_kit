import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = docx.Document()

    # Title
    title = doc.add_heading('BÁO CÁO PHÂN TÍCH YÊU CẦU - US02 (KHAI BÁO CODE PHÍ)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ─── PHẦN A ───────────────────────────────────────────────────────────────
    doc.add_heading('Phần A: Tóm Tắt Nghiệp Vụ Chuyên Sâu (Requirements Breakdown)', level=1)

    doc.add_heading('1. Thông điệp cốt lõi (Core Business Value)', level=2)
    doc.add_paragraph(
        'US02 định nghĩa các thông số cụ thể cho một Code Phí (loại phí chi tiết nhất). '
        'Đây là nơi cấu hình quan trọng nhất để hệ thống biết Thu phí ai, Thu bao nhiêu, '
        'Thu khi nào và Cách tính phí theo từng phân khúc khách hàng.'
    )

    doc.add_heading('2. Luồng chính (Happy Path Workflow)', level=2)
    doc.add_paragraph(
        '1. Từ SPDV cấp cuối -> Nhấn Thêm mới Code Phí.\n'
        '2. Khai báo Thông tin chung: Tên phí, Loại tiền, VAT, Đối tượng (KH/Merchant), Phân khúc.\n'
        '3. Chọn Loại tính phí (Giao dịch/Định kỳ) và Tần suất (nếu là định kỳ).\n'
        '4. Thiết lập Điều kiện tính phí (lấy từ Danh mục điều kiện US26).\n'
        '5. Định nghĩa Công thức tính phí (Formula).\n'
        '6. Xác nhận -> Chờ duyệt -> Hoạt động.'
    )

    doc.add_heading('3. Các luồng rẽ nhánh / Ngoại lệ (Alternative & Exception Flows)', level=2)
    doc.add_paragraph(
        '• Chỉnh sửa: Chỉ cho phép sửa trạng thái của Code Phí sang trạng thái "Hủy".\n'
        '• Tần suất định kỳ: Quy định riêng cho tháng 2 (tối đa ngày 28) và xử lý ngày 31 cho các tháng thiếu.\n'
        '• VAT: Có logic tính gộp (đã gồm VAT) hoặc tính thêm (chưa gồm VAT).'
    )

    doc.add_heading('4. Bảng Điều Kiện Tiên Quyết & Cấu Hình (Pre-conditions)', level=2)
    doc.add_paragraph(
        '• Phải có SPDV cấp cuối đã được duyệt (US01).\n'
        '• Phải có Danh mục điều kiện tính phí đã định nghĩa (US26).'
    )

    # ─── PHẦN B ───────────────────────────────────────────────────────────────
    doc.add_heading('Phần B: Khai Quật Lỗ Hổng & Danh Sách Q&A (Dành Cho BA)', level=1)

    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['ID', 'Ref_ID', 'Nội dung Câu hỏi / Sự cố', 'Phân loại', 'Đề xuất xử lý', 'BA Answer']
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True

    qa_data = [
        (
            'QA-01', 'Line 28 vs 32',
            'Bất cập tên gọi Trạng thái: Line 28 ghi "Chưa gán", nhưng Line 32 lại ghi mặc định là "Chờ gán". '
            'Từ "Chưa" và "Chờ" có ý nghĩa khác nhau trong hệ thống quản lý.',
            'Document', 'Thống nhất dùng duy nhất 1 cụm từ (Gợi ý: "Chưa gán").', ''
        ),
        (
            'QA-02', 'Line 11 vs 14',
            'Mâu thuẫn Range giá trị VAT: Line 11 yêu cầu số < 100. Nhưng Line 14 lại cho phép số <= 100. '
            'Có trường hợp phí nào thuế VAT là 100% không?',
            'Nghiệp vụ', 'Chốt range chuẩn: (0, 100) hay [0, 100].', ''
        ),
        (
            'QA-03', 'Line 175',
            'Xử lý ngày 29/02 năm nhuận: Tài liệu quy định tháng 2 chỉ nhập tối đa ngày 28/02. '
            'Vậy các hợp đồng phí đăng ký thu vào ngày cuối cùng của tháng 2 năm nhuận sẽ ra sao?',
            'Nghiệp vụ', 'Cho phép nhập 29/02 và hệ thống tự động skip nếu năm đó không nhuận (về 28/02).', ''
        ),
        (
            'QA-04', 'Line 132',
            'Logic ẩn/hiện loại khách hàng: Khi Merchant = Đối tượng thu phí, ẩn CBNV. '
            'Nhưng bản liệt kê Text không nói rõ có ẩn luôn KHCN/KHTC/DNSN hay không? '
            'Vì Merchant thường là tổ chức.',
            'UI/UX', 'Làm rõ bộ list Loại khách hàng khi chọn Đối tượng = Merchant.', ''
        ),
        (
            'QA-05', 'Line 45',
            'Race condition Loại tiền: Hệ thống không cập nhật loại tiền tối đa/tối thiểu nếu chọn trường này TRƯỚC. '
            'Điều này dễ dẫn tới sai lệch giữa Loại tiền của Phí (USD) và Loại tiền limit (VND).',
            'UI/UX', 'Nên tự động update hoặc bôi đỏ cảnh báo nếu có sự lệch loại tiền giữa Phí và Limit.', ''
        ),
        (
            'QA-TC-01', 'Validation / Boundary',
            'Phí đã bao gồm VAT (Checkbox): Nếu chọn checkbox này nhưng bỏ trống ô VAT, hệ thống báo lỗi gì? '
            'Expected result: Báo lỗi "Trường VAT bắt buộc nhập khi chọn Phí đã bao gồm VAT".',
            'Test Case / Logic', 'Cung cấp message lỗi chính xác cho case này.', ''
        ),
        (
            'QA-TC-02', 'Yearly Cycle / Leap Year',
            'Tần suất Hàng năm: Nếu chọn 29/02 (năm nhuận) -> Năm sau 2027 không nhuận thì hệ thống sẽ thu phí vào ngày nào? '
            '28/02 hay 01/03? Cần rule cụ thể để viết Case.',
            'Test Case / Date', 'Xác định rule thu phí ngày cuối tháng ảo (EOM).', ''
        ),
        (
            'QA-TC-03', 'Mapping / Limit',
            'Loại tiền tối thiểu/tối đa vs VND/JPY: Nếu chọn VND nhưng nhập số thập phân (100.5), '
            'hệ thống chặn ở FE (masking) hay BE (validate error message)?',
            'Test Case / Input', 'Chốt phương án chặn: Masking input hay Show error msg.', ''
        )
    ]

    for row_data in qa_data:
        row_cells = table.add_row().cells
        for i, text in enumerate(row_data):
            row_cells[i].text = text

    doc.save('US02_QnA_Report.docx')
    print('✅ Đã tạo US02_QnA_Report.docx thành công!')

create_report()
