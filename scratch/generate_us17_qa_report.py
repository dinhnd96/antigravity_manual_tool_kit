import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = docx.Document()

    # Title
    title = doc.add_heading('BÁO CÁO PHÂN TÍCH YÊU CẦU - US17 (XEM CODE PHÍ THEO KHÁCH HÀNG)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ─── PHẦN A ───────────────────────────────────────────────────────────────
    doc.add_heading('Phần A: Tóm Tắt Nghiệp Vụ Chuyên Sâu (Requirements Breakdown)', level=1)

    doc.add_heading('1. Thông điệp cốt lõi (Core Business Value)', level=2)
    doc.add_paragraph(
        'US17 cung cấp một luồng tra cứu "2 bước" chuyên biệt:\n'
        '  • Bước 1: Tìm CIF khách hàng (qua số điện thoại, phân khúc, chi nhánh, tỉnh/xã, trạng thái).\n'
        '  • Bước 2: Sau khi chọn đúng CIF, hệ thống tự điền và tìm ra toàn bộ Code Phí đang '
        'được áp dụng cho CIF đó kèm theo Biểu Phí, SPDV, Công thức tính phí.\n'
        'Đây là màn "Tra cứu hợp đồng phí theo từng khách hàng" phục vụ Nhân viên tư vấn & Kiểm soát.'
    )

    doc.add_heading('2. Luồng chính (Happy Path Workflow)', level=2)
    doc.add_paragraph(
        '1. User chọn menu Tra cứu >> Xem code phí theo khách hàng.\n'
        '2. Nhấn [Tra cứu CIF] -> Popup "Tra cứu CIF" mở ra -> Nhập Số ĐT -> Nhấn Tra cứu.\n'
        '3. Kết quả Grid CIF hiển thị -> User click vào 1 dòng CIF.\n'
        '4. Hệ thống auto-fill Mã CIF vào ô tìm kiếm chính -> User nhấn [Tìm kiếm].\n'
        '5. Grid danh sách Code phí áp dụng cho CIF đó hiển thị đầy đủ.\n'
        '6. User có thể click vào mã Code Phí để xem chi tiết, hoặc [Tải xuống] Excel.'
    )

    doc.add_heading('3. Các luồng rẽ nhánh / Ngoại lệ (Alternative & Exception Flows)', level=2)
    doc.add_paragraph(
        '• User có thể bỏ qua bước Tra cứu CIF và tự gõ thẳng Mã CIF biết sẵn vào ô tìm kiếm.\n'
        '• Tìm kiếm SPDV/Code phí có hỗ trợ Multi-select dropdown (chọn nhiều SPDV/code phí).\n'
        '• Trường hợp Mã CIF không tồn tại: Tài liệu KHÔNG đề cập thông báo lỗi.'
    )

    doc.add_heading('4. Bảng Điều Kiện Tiên Quyết & Cấu Hình (Pre-conditions)', level=2)
    doc.add_paragraph(
        '• Mã CIF là trường BẮT BUỘC (Y) khi tìm kiếm Code phí.\n'
        '• Số điện thoại là trường BẮT BUỘC (Y) trong màn Tra cứu CIF.\n'
        '• Trường Tỉnh phải chọn TRƯỚC, rồi Phường/Xã mới được kích hoạt (Dependent Dropdown).'
    )

    doc.add_heading('5. Ma trận Phân Quyền / Dữ Liệu', level=2)
    doc.add_paragraph(
        'Màn hình này là Read-Only (Tra cứu). Không có chức năng Chỉnh sửa hay Xóa trực tiếp. '
        'User chỉ có thể drill-down vào Chi tiết Code phí.'
    )

    # ─── PHẦN B ───────────────────────────────────────────────────────────────
    doc.add_heading('Phần B: Khai Quật Lỗ Hổng & Danh Sách Q&A (Dành Cho BA)', level=1)

    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['ID', 'Ref_ID (Mục)', 'Nội dung Câu hỏi / Sự cố',
               'Phân loại', 'Đề xuất hướng xử lý từ QA', 'Câu trả lời của BA (Để trống)']
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True

    qa_data = [
        (
            'QA-01',
            'Lệch cột lưới Grid và Text (Line 8 vs Hình 2)',
            'Text (Line 8) mô tả cột lưới Code phí gồm: STT, CIF áp dụng, Code, Tên, Biểu phí, SPDV, '
            'Công thức tính phí, Người tạo, Ngày tạo, Người sửa, Ngày sửa, Người duyệt, Ngày duyệt.\n'
            'NHƯNG trên Hình UI Mockup 2, lưới Grid được vẽ CHỈ CÓ: #, CIF áp dụng, Code, Tên, '
            'Biểu phí, Sản phẩm DV, Công thức tính phí, Người tạo, Ngày tạo, Người sửa. '
            'Thiếu hoàn toàn: Ngày sửa, Người duyệt, Ngày duyệt.',
            'UI/UX',
            'Chốt hạ danh sách cột cuối cùng. Cập nhật lại ảnh Mockup hoặc sửa Text URD cho đồng nhất.',
            ''
        ),
        (
            'QA-02',
            'Bắt buộc Số ĐT trong Tra cứu CIF nhưng UI cho phép bỏ trống (Line 279 vs Hình 3)',
            'Bảng Field Definition (Line 279) đánh dấu "Y" (Bắt buộc) cho trường Số điện thoại. '
            'Tuy nhiên Hình Mockup Tra cứu CIF (Hình 3) KHÔNG hiển thị dấu (*) hay màu viền đỏ '
            'nào chỉ dẫn bắt buộc. Lại không có mô tả thông báo lỗi khi để trống.',
            'Nghiệp vụ',
            'Bổ sung dấu (*) bắt buộc trên UI và mô tả message lỗi validate khi Số ĐT bỏ trống.',
            ''
        ),
        (
            'QA-03',
            'Không có filter Trạng Thái trong Tra cứu CIF chính (Hình 2 vs Line 3)',
            'Tại màn hình chính "Xem code phí theo khách hàng" (Hình 2), điều kiện tìm kiếm chỉ '
            'gồm Mã CIF + Sản phẩm DV. Người dùng không thể biết CIF đó đang Hoạt động hay Tạm dừng '
            'trước khi tìm Code Phí. Cột "Trạng thái" CŨNG KHÔNG xuất hiện trong Grid kết quả.',
            'Nghiệp vụ',
            'Đề xuất thêm cột Trạng thái CIF vào Grid kết quả hoặc hiển thị badge cảnh báo '
            'nếu CIF đang ở trạng thái Tạm dừng.',
            ''
        ),
        (
            'QA-04',
            'Bỏ qua STT 7 trong bảng trường Tra cứu CIF (Line 314)',
            'Bảng mô tả trường của màn Tra cứu CIF đánh số: 1 -> 2 -> 3 -> 4 -> 5 -> 6 '
            'rồi nhảy thẳng lên 8 (STT 7 biến mất). Gây khó chịu khi kiểm đếm số trường '
            'và ánh xạ sang bộ Test Case.',
            'Document',
            'Đánh số lại tuần tự 1 -> 2 -> ... -> 7 -> 8 trong bảng mô tả trường.',
            ''
        ),
        (
            'QA-05',
            'Sai kiểu dữ liệu (Data Type) của "Người duyệt" và "Ngày duyệt" (Line 237, Line 243)',
            'Trường "Người duyệt" (Line 237) được ghi kiểu dữ liệu là "Date" – rõ ràng SAI, '
            'phải là "Text" (username). Ngược lại, trường "Ngày duyệt" (Line 243) lại ghi kiểu '
            '"Text" – cũng SAI, phải là "Date time". Hai trường bị hoán đổi Data Type cho nhau.',
            'Document',
            'Sửa: Người duyệt -> Text, Ngày duyệt -> Date time (hh:mm:ss – dd/mm/yyyy).',
            ''
        ),
        (
            'QA-06',
            'Lỗi điều hướng trong Flowchart: thiếu connector sau bước 10 (Hình 1)',
            'Trong BPMN (Hình 1), sau bước 9 "Click vào 1 CIF" -> Bước 10 "Hệ thống fill CIF vào box" '
            '-> mũi tên ghi "=> bước 3.2". NHƯNG trên BPMN không vẽ đường nối rõ ràng từ 10 về '
            'diamond "3.2 Nhập điều kiện tìm kiếm". Cả luồng Tra cứu CIF trả về màn chính bị '
            'thiếu connector vật lý trên sơ đồ.',
            'Document',
            'Thêm mũi tên nối từ bước [10] quay về diamond [3.2] trên Hình BPMN.',
            ''
        ),
        (
            'QA-07',
            'Chưa mô tả hành vi khi Mã CIF không tồn tại (Gap nghiệp vụ)',
            'Tài liệu mô tả Happy Path: User nhập Mã CIF -> Tìm kiếm -> Kết quả hiển thị. '
            'Nhưng HOÀN TOÀN KHÔNG đề cập: Nếu Mã CIF không tồn tại trong hệ thống thì hệ thống '
            'phản hồi thế nào? Empty State? Toast error? Hay Grid trắng trơn im lặng?',
            'Nghiệp vụ',
            'Bổ sung mô tả Exception: "Nếu Mã CIF không tồn tại, hiển thị thông báo [Không tìm thấy dữ liệu]".',
            ''
        ),
        ('QA-TC-01', 'Validation / Phone', 'Trường Số ĐT (Tra cứu CIF): Hệ thống có validate format số ĐT (ví dụ start=0, length=10-11) hay chỉ quan tâm không bỏ trống? Message lỗi khi nhập chữ cái là gì?', 'Test Case / Input Validate', 'Nêu rõ regex của phone number.', '')
    ]

    for row_data in qa_data:
        row_cells = table.add_row().cells
        for i, text in enumerate(row_data):
            row_cells[i].text = text

    doc.save('US17_QnA_Report.docx')
    print('✅ Đã tạo US17_QnA_Report.docx thành công!')

create_report()
