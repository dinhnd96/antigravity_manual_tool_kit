import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = docx.Document()
    
    # Title
    title = doc.add_heading('BÁO CÁO PHÂN TÍCH YÊU CẦU - US05 (ĐỊNH NGHĨA CÔNG THỨC)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Part A
    doc.add_heading('Phần A: Tóm Tắt Nghiệp Vụ Chuyên Sâu (Requirements Breakdown)', level=1)
    
    doc.add_heading('1. Thông điệp cốt lõi (Core Business Value)', level=2)
    doc.add_paragraph('US05 mang tới công cụ "Formula Engine", giúp người dùng tự xây dựng công thức tính toán phí theo biến số động. Tính năng này cho phép cá nhân hóa mức thu phí theo từng Phân khúc khách hàng, sử dụng các hàm toán học kinh điển (Math, Date, Rounding).')
    
    doc.add_heading('2. Luồng chính (Happy Path Workflow)', level=2)
    doc.add_paragraph('Maker bật form Cấu hình Công thức -> Quyết định có phân tách Nhóm khách hàng hay áp dụng chung -> Bấm icon cây bút -> Click chọn các trường biến số (được bọc trong #Field#) và các toán tử/hàm -> Điền giới hạn Min/Max -> Bấm Xác nhận -> Hệ thống chạy Validate cú pháp (Syntax Check) -> Về màn chờ duyệt.')
    
    doc.add_heading('3. Các luồng rẽ nhánh / Ngoại lệ (Alternative & Exception Flows)', level=2)
    doc.add_paragraph('Hệ thống kiểm tra cực kỳ gay gắt cú pháp. Nếu: kết hợp MIN/MAX chung mâm với DATE_DIFF; hoặc quên đóng ngoặc; hoặc điền 2 hàm đụng nhau -> Đều bắn ra thông báo lỗi và chặn lưu.')
    
    doc.add_heading('4. Bảng Điều Kiện Tiên Quyết & Cấu Hình (Pre-conditions & Settings)', level=2)
    doc.add_paragraph('• Biến động (Dynamic Variables): Phụ thuộc vào Loại tính phí ở US02. Thu Theo Giao dịch thì lấy nguồn API. Thu Định kỳ thì lấy nguồn Bảng DB nội bộ (ETL).\n• Tiền Min/Max: Có thể set loại tiền tệ tách biệt hoàn toàn so với loại tiền gốc của Công thức.')
    
    doc.add_heading('5. Ma trận Phân Quyền/Dữ Liệu', level=2)
    doc.add_paragraph('Sau khi Công thức được Approve lần đầu, Maker BỊ CẤM sửa đổi cấu trúc các hàm toán học. Maker chỉ được phép đổi các "Giá trị số học" thông qua cơ chế Upload file Excel (US07).')
    
    # Part B
    doc.add_heading('Phần B: Khai Quật Lỗ Hổng & Danh Sách Q&A (Dành Cho BA)', level=1)
    
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['ID', 'Ref_ID (Mục)', 'Nội dung Câu hỏi / Sự cố', 'Phân loại', 'Đề xuất hướng xử lý từ QA', 'Câu trả lời của BA (Để trống)']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        
    qa_data = [
        (
            'QA-01', 
            'Luồng Flowchart "Cắt cầu" (Bản vẽ 1)', 
            'Trong BPMN, khi Xác thực điều kiện trả "No" -> "Hiển thị thông báo lỗi" -> "End" luôn. Nghĩa là Maker vừa gõ sai công thức 1 phát bị văng khỏi form? Trải nghiệm UX thực tế phải là báo lỗi và Maker được stay lại form để sửa cú pháp.', 
            'UI/UX', 
            'Sửa lại net vẽ BPMN nối mũi tên báo lỗi quay vòng lại bước Thiết lập công thức (B.9).',
            ''
        ),
        (
            'QA-02', 
            'Thiếu phím bấm MIN/MAX ở UI (Ảnh Mockup 4)', 
            'Text (Line 216) yêu cầu cung cấp hàm MIN và MAX, nhưng trên ảnh thiết kế Bàn phím bấm ảo ở màn hình nhập liệu KHÔNG HỀ có 2 nút MIN, MAX này.', 
            'UI/UX', 
            'Chèn thêm nút Bấm MIN, MAX vào Mockup / Design bàn phím máy tính.',
            ''
        ),
        (
            'QA-03', 
            'Missing ký tự Phần trăm (%) (Line 217 vs Ảnh 4)', 
            'Trong ảnh UI, công thức có chứa "0.15%". NHƯNG danh sách các Text được phép gõ vào ở Line 217 KHÔNG nhắc đến ký tự % (Chỉ có ) ( , . #). Nếu Dev code chặn theo Line 217 thì user không gõ mốc % được.', 
            'Nghiệp vụ', 
            'Bổ sung "%" vào danh sách các ký tự được phép sử dụng.',
            ''
        ),
        (
            'QA-04', 
            'Khoản trống Văn Vở (Line 207)', 
            'Tiếp tục xuất hiện lỗi drop text của BA: "Loại tính phí = Định kỳ và Ngày thu cố định...". Dòng này chưa viết hết ý, để lửng lơ rồi nhảy sang Line 208.', 
            'Nghiệp vụ', 
            'Yêu cầu BA bù đoạn mô tả còn thiếu.',
            ''
        ),
        (
            'QA-05', 
            'Logic chặn Cấu phần rủi ro rớt hệ thống (Line 38)', 
            'Luồng Validate: "Khi công thức nhập 2 hoặc nhiều cấu phần, trong đó có 2 cấu là các giá trị số học cụ thể, hệ thống báo lỗi". Định nghĩa 1 cấu phần là gì? VD công thức: "1000 * #Amount# + 5000", đây là 2 số học cụ thể. Vậy có bị tính là lỗi không? Rule này đang tả rất mơ hồ và dễ khóa nhầm.', 
            'Nghiệp vụ', 
            'Làm rõ định nghĩa "Cấu phần" hoặc mô tả Exception cụ thể bằng ví dụ để Developer viết validate chặn Regex được chuẩn.',
            ''
        ),
        (
            'QA-06', 
            'Lộn xộn Headers (Image 2 vs 3)', 
            'Trên ảnh Mockup 2: Cột là Tối thiểu -> Tối đa. Xuống Mockup 3: Cột bị đảo lại thành Tối đa -> Tối thiểu. Gây thiết kế mâu thuẫn cho Team UI.', 
            'UI/UX', 
            'Khớp lại vị trí cột xuyên suốt các mockup.',
            ''
        ),
        ('QA-TC-01', 'Validation / Logic', 'Lỗi cú pháp Formula: Nếu user nhập công thức thiếu đóng ngoặc kép, sai biến thì hệ thống Validate báo lỗi message bằng chữ gì? Mã Code HTTP trả về là 400?', 'Test Case / Error handling', 'Yêu cầu Mapping bảng Error Messages cho công thức tính phí.', '')
    ]
    
    for row_data in qa_data:
        row_cells = table.add_row().cells
        for i, text in enumerate(row_data):
            row_cells[i].text = text
            
    doc.save('US05_QnA_Report.docx')

create_report()
