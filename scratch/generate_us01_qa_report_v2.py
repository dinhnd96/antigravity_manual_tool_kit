from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('BÁO CÁO PHÂN TÍCH YÊU CẦU - US01', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Part A
    doc.add_heading('Phần A: Tóm Tắt Nghiệp Vụ Chuyên Sâu (Requirements Breakdown)', level=1)
    
    doc.add_heading('1. Thông điệp cốt lõi (Core Business Value)', level=2)
    doc.add_paragraph('Tính năng cho phép khởi tạo và quản lý hệ thống phân tầng các Sản phẩm, Dịch vụ (SPDV) theo mô hình cây thư mục (từ cấp 1 - Nghiệp vụ đến các cấp SPDV chi tiết). Đây là dữ liệu nền (Master Data) để hệ thống ProfiX sau này dựa vào đó để gắn các Biểu phí, Code phí tương ứng.')
    
    doc.add_heading('2. Luồng chính (Happy Path Workflow)', level=2)
    doc.add_paragraph('Thêm mới SPDV: Maker chọn Thêm mới -> Chọn cấp (Nghiệp vụ hoặc SPDV con) -> Điền Tên, Mô tả, Ngày hiệu lực, Ngày hết hiệu lực -> Hệ thống tự sinh Mã theo quy tắc -> Lưu bản ghi (Trạng thái "Chờ duyệt") -> Checker tiến hành duyệt -> Bản ghi kích hoạt và hiện trên lưới. Trạng thái hoạt động sẽ do Job đầu ngày tự động bật/tắt (Hoạt động/Không hoạt động) dựa trên Ngày hệ thống so với Ngày hiệu lực/hết hiệu lực.')
    
    doc.add_heading('3. Các luồng rẽ nhánh / Ngoại lệ (Alternative & Exception Flows)', level=2)
    doc.add_paragraph('• Chỉnh sửa: Khi SPDV đã được duyệt, Maker chỉ được phép chỉnh sửa "Ngày hiệu lực" và "Ngày hết hiệu lực", nhưng bản ghi phải qua rào cản Chờ duyệt mới cập nhật.\n• Vượt hạn mức mã: Hệ thống giới hạn mỗi nhánh chỉ nở được 99 cấp con (Mã từ 01->99). Sinh nhánh thứ 100 hệ thống sẽ chặn không cho lưu.')
    
    doc.add_heading('4. Bảng Điều Kiện Tiên Quyết & Cấu Hình (Pre-conditions & Settings)', level=2)
    doc.add_paragraph('• Ràng buộc thời gian: Ngày hệ thống <= Ngày hiệu lực <= Ngày hết hiệu lực.\n• Ràng buộc cây phân cấp: Ngày hiệu lực Cấp Cha <= Ngày hiệu lực Cấp Con <= Ngày hết hiệu lực Cấp Con <= Ngày hết hiệu lực Cấp Cha.')
    
    doc.add_heading('5. Ma trận Phân Quyền/Dữ Liệu', level=2)
    doc.add_paragraph('• Maker: Quyền thêm mới, Sửa ngày hiệu lực của SPDV.\n• Checker: Quyền duyệt các tác vụ do Maker trình lên chặn bởi Ma trận phê duyệt (US25).')
    
    # Part B
    doc.add_heading('Phần B: Khai Quật Lỗ Hổng & Danh Sách Q&A (Dành Cho BA)', level=1)
    
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['ID', 'Ref_ID (Mục)', 'Nội dung Câu hỏi / Sự cố', 'Phân loại', 'Đề xuất hướng xử lý từ QA', 'Câu trả lời của BA (Để trống)']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        
    qa_data = [
        (
            'QA-01', 
            'Chỉnh sửa SPDV (Line 16, 32)', 
            'Tài liệu ghi: "hệ thống chỉ cho phép chỉnh sửa Ngày hiệu lực, Ngày hết hiệu lực". Vậy nếu Maker gõ sai chính tả ở "Tên SPDV" hoặc "Mô tả" có được sửa không? Nếu bắt buộc phải Hủy để tạo mới sẽ gây rác dữ liệu mã SPDV.', 
            'Nghiệp vụ', 
            'Đề xuất cho phép sửa Tên/Mô tả (hoặc khi chưa phát sinh Code phí/CTƯĐ đính kèm).',
            ''
        ),
        (
            'QA-02', 
            'Chỉnh sửa Ngày HL (Line 17)', 
            'Ràng buộc khi sửa: "Ngày hiệu lực, Ngày hết HL chỉ được sửa >= Ngày hệ thống". Nếu SPDV này đã active 1 năm (Ngày HL ở quá khứ), và user chỉ muốn gia hạn "Ngày hết HL", thì khi lưu lại bị lỗi do "Ngày HL < Ngày hệ thống" hay sao?', 
            'Nghiệp vụ', 
            'Chỉ áp dụng check "Ngày HL >= Ngày hệ thống" với các bản ghi Trạng thái = Chưa hiệu lực. Bản ghi đang Hoạt động thì bỏ qua check Ngày HL.',
            ''
        ),
        (
            'QA-03', 
            'Cấp sâu giới hạn (Line 29)', 
            'Mã tự sinh nối +02 ký tự số. Vậy cây SPDV có giới hạn độ sâu tối đa không (bao nhiêu node con)? Tránh trường hợp UI bị tràn hoặc đệ quy vượt mức.', 
            'UI/UX', 
            'Nên hard-limit độ sâu của cây tối đa (VD: 5 Tầng) trong DB và UI.',
            ''
        ),
        (
            'QA-04', 
            'Auto switch status (Line 22)', 
            'Job đầu ngày quét Update Trạng thái = Không hoạt động (nếu đã quá hạn). Câu hỏi: Nếu Cấp Cha bị Không hoạt động, hệ thống có tự động cascade trạng thái Không hoạt động xuống toàn bộ Cấp Con không?', 
            'Nghiệp vụ', 
            'Bắt buộc phải cascade down xuống tất cả cấp con, hoặc chặn SPDV Cha tự hết độ dài hiệu lực trước SPDV Con để không bị tình trạng "Cha chết Con sống".',
            ''
        ),
        ('QA-TC-01', 'Validation / Boundary', 'Code quy định độ dài (Max-length) của Mã và Tên SPDV cấp 1/2/3 là bao nhiêu ký tự? Ký tự đặc biệt nào bị cấm? Message lỗi khi nhập trùng Mã SPDV là gì?', 'Test Case / Boundary', 'Cần BA cung cấp message lỗi chính xác và rule Regex cho Input.', '')
    ]
    
    for row_data in qa_data:
        row_cells = table.add_row().cells
        for i, text in enumerate(row_data):
            row_cells[i].text = text
            
    doc.save('US01_QnA_Report_V2.docx')

create_report()
