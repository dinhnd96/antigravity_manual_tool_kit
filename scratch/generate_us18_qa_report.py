import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = docx.Document()

    # Title
    title = doc.add_heading('BÁO CÁO PHÂN TÍCH YÊU CẦU - US18 (XEM LỊCH SỬ THU PHÍ THEO KHÁCH HÀNG)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ─── PHẦN A ───────────────────────────────────────────────────────────────
    doc.add_heading('Phần A: Tóm Tắt Nghiệp Vụ Chuyên Sâu (Requirements Breakdown)', level=1)

    doc.add_heading('1. Thông điệp cốt lõi (Core Business Value)', level=2)
    doc.add_paragraph(
        'US18 là màn hình Tra cứu lịch sử từng giao dịch thu phí đã thực hiện cho 1 CIF cụ thể. '
        'Đây là công cụ thiết yếu dành cho Ban Kiểm soát/Kế toán để đối chiếu số tiền phí đã thu, '
        'kiểm tra VAT, và xác minh tài khoản thu phí theo từng nghiệp vụ phát sinh trong kỳ.'
    )

    doc.add_heading('2. Luồng chính (Happy Path Workflow)', level=2)
    doc.add_paragraph(
        '1. User vào màn "Tra cứu lịch sử thu phí theo khách hàng".\n'
        '2. Nhập Mã CIF (bắt buộc) và tuỳ chọn thêm SPDV, Biểu phí, Code phí, Từ ngày - Đến ngày.\n'
        '3. Có thể dùng [Tra cứu CIF] (popup tái sử dụng từ US18) để tra và auto-fill Mã CIF.\n'
        '4. Nhấn [Tìm kiếm] -> Grid lịch sử thu phí hiển thị.\n'
        '5. Tải xuống Excel kết quả tìm kiếm.'
    )

    doc.add_heading('3. Các luồng rẽ nhánh / Ngoại lệ (Alternative & Exception Flows)', level=2)
    doc.add_paragraph(
        '• User có thể lọc kết hợp nhiều điều kiện (multi-select: SPDV, Biểu phí, Code phí).\n'
        '• Bộ lọc ngày tháng: Hai trường "Từ ngày - Đến ngày" lọc theo ngày thu phí thực tế.\n'
        '• Tài liệu KHÔNG mô tả: Trường hợp không có lịch sử thu phí cho CIF, hệ thống phản hồi gì.'
    )

    doc.add_heading('4. Bảng Điều Kiện Tiên Quyết & Cấu Hình (Pre-conditions)', level=2)
    doc.add_paragraph(
        '• Lịch sử hiển thị từ các giao dịch thu phí đã được hệ thống US33/US34 thực hiện thành công.\n'
        '• Số tiền phí Nguyên tệ và Quy đổi VNĐ là 2 cột riêng biệt – phục vụ đối chiếu '
        'cho trường hợp phí thu bằng ngoại tệ.'
    )

    doc.add_heading('5. Ma trận Phân Quyền / Dữ Liệu', level=2)
    doc.add_paragraph(
        'Màn hình này hoàn toàn Read-Only. Không có tính năng chỉnh sửa hay hủy giao dịch. '
        'Chỉ Xem + Tải xuống báo cáo.'
    )

    # ─── PHẦN B ───────────────────────────────────────────────────────────────
    doc.add_heading('Phần B: Khai Quật Lỗ Hổng & Danh Sách Q&A (Dành Cho BA)', level=1)

    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['ID', 'Ref_ID (Mục)', 'Nội dung Câu hỏi / Sự cố',
               'Phân loại', 'Đề xuất hướng xử lý từ QA', 'Câu trả lời của BA (Để trống)']
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True

    qa_data = [
        (
            'QA-01',
            'Tham chiếu sai User Story (Line 9 và 125)',
            'Tại Line 9 và Line 125, BA viết: "tham khảo tại mục Tra cứu CIF của US 19". '
            'Nhưng US18 CHÍNH LÀ tài liệu đang đọc! Rõ ràng BA muốn tham chiếu sang '
            'US18 (chứa định nghĩa màn Tra cứu CIF), nhưng đã copy số US sai.',
            'Document',
            'Sửa "US 19" thành "US 18" tại Line 9 và Line 125.',
            ''
        ),
        (
            'QA-02',
            'Lệch Cột Grid: Text vs Hình UI Mockup (Line 6 vs Hình 2)',
            'Text (Line 6) định nghĩa cột Grid gồm: STT, CIF áp dụng, Tài khoản thu phí, '
            'Code phí, Loại tiền code phí, Biểu phí, SPDV, VAT, Số tiền phí nguyên tệ, '
            'Số tiền phí quy đổi VNĐ, Ngày thu.\n'
            'NHƯNG Hình UI Mockup (Hình 2) lại bổ sung thêm cột "Tên" (Tên code phí) và '
            '"Công thức tính phí" mà Text URD KHÔNG đề cập. '
            'Ngược lại "Loại tiền code phí" và "Tài khoản thu phí" có trong Text nhưng '
            'Grid Mockup không vẽ.',
            'UI/UX',
            'Chốt danh sách cột cuối cùng và cập nhật đồng bộ cả Text URD lẫn Hình Mockup.',
            ''
        ),
        (
            'QA-03',
            'Bỏ qua STT số 7 trong bảng trường (Line 171)',
            'Bảng mô tả các trường lưới kết quả đánh số: 1 -> 2 -> 3 -> 4 -> 5 -> 6 '
            'rồi nhảy thẳng lên 8 (STT 7 hoàn toàn biến mất, giống lỗi di truyền từ US18).',
            'Document',
            'Đánh số lại tuần tự đúng: 7 -> 8 -> ... -> 17.',
            ''
        ),
        (
            'QA-04',
            'Lỗi chính tả nghiêm trọng trên tên trường (Line 172, 178)',
            '"CIF áp dụg" (thiếu chữ n) tại Line 172, và "Tải khoản thu phí" (sai thành "Tải" '
            'thay vì "Tài") tại Line 178. Nếu Dev dùng text này làm Label UI thì sẽ hiển thị '
            'sai chính tả trực tiếp lên sản phẩm.',
            'Document',
            'Sửa: "CIF áp dụng" và "Tài khoản thu phí".',
            ''
        ),
        (
            'QA-05',
            'Lỗi chính tả tên field điều kiện tìm kiếm (Line 158)',
            'Trường tra cứu được ghi là "Từ ngày - đến ngay" (thiếu dấu "à" ở "ngày"). '
            'Tuy nhỏ nhưng nếu mapping sang DB column hay label thì sẽ sai.',
            'Document',
            'Sửa: "Từ ngày - Đến ngày".',
            ''
        ),
        (
            'QA-06',
            'Thiếu mô tả VAT hiển thị dưới dạng nào (Line 212)',
            'Cột VAT định nghĩa kiểu "Number" và mô tả "Hiển thị số tiền thuế VAT đã thu". '
            'Nhưng không rõ: VAT hiển thị dưới dạng Số tiền tuyệt đối (VD: 1,000 VND) '
            'hay Tỷ lệ phần trăm (VD: 10%)? Cũng không rõ nguyên tệ hay quy đổi VNĐ?',
            'Nghiệp vụ',
            'BA cần làm rõ: VAT = số tiền tuyệt đối (đồng) hay % và theo loại tiền nào.',
            ''
        ),
        (
            'QA-07',
            'Không có Mã CIF bắt buộc nhưng Flowchart cho phép tìm không cần CIF (Line 128 vs Hình 1)',
            'Bảng trường (Line 128-129) đánh "Y" (Bắt buộc) cho Mã CIF. '
            'Nhưng trên Hình Flowchart (Hình 1), bước 3.2 ghi "Nhập/CHỌN điều kiện" (không khẳng định CIF '
            'là bắt buộc). Hình UI Mockup (Hình 2) cũng chỉ hiển thị placeholder "Vui lòng nhập" bình thường, '
            'không có dấu * bắt buộc. Hai nguồn mâu thuẫn nhau.',
            'Nghiệp vụ',
            'Chốt rõ: Mã CIF có THỰC SỰ bắt buộc không? Nếu có, bổ sung dấu (*) trên UI Mockup '
            'và validate message khi để trống.',
            ''
        ),
        (
            'QA-08',
            'Thiếu connector Flowchart từ bước [10] quay về [3.2] (Hình 1)',
            'Giống lỗi US18: Sau bước [10] "Hệ thống fill CIF vào box", text mô tả ghi '
            '"=> bước 3.2" nhưng trên BPMN (Hình 1) không có đường nối vật lý từ bước [10] '
            'về diamond [3.2]. Luồng bị đứt ngang trên sơ đồ.',
            'Document',
            'Thêm mũi tên nối từ [10] về [3.2] trên Hình BPMN.',
            ''
        ),
        ('QA-TC-01', 'Data Format / UI', 'Cột Số tiền phí/quy đổi VNĐ: Có áp dụng format dấu phẩy phân cách hàng nghìn (1,000,000) không? Nếu ra tiền lẻ VNĐ thì hệ thống quy tròn (Round) lên hay xuống?', 'Test Case / Display', 'Chốt quy tắc làm tròn (Round/Ceil/Floor) để so sánh expected result.', '')
    ]

    for row_data in qa_data:
        row_cells = table.add_row().cells
        for i, text in enumerate(row_data):
            row_cells[i].text = text

    doc.save('US18_QnA_Report.docx')
    print('✅ Đã tạo US18_QnA_Report.docx thành công!')

create_report()
