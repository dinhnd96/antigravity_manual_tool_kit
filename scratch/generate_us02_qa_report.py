import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = docx.Document()
    
    # Title
    title = doc.add_heading('BÁO CÁO PHÂN TÍCH YÊU CẦU - US02 (KHAI BÁO CODE PHÍ)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Part A
    doc.add_heading('Phần A: Tóm Tắt Nghiệp Vụ Chuyên Sâu (Requirements Breakdown)', level=1)
    
    doc.add_heading('1. Thông điệp cốt lõi (Core Business Value)', level=2)
    doc.add_paragraph('US02 định nghĩa công cụ để User (Maker) cấu hình cốt lõi của bảng phí (Code Phí). Nó đóng vai trò "Bộ nội logic" chứa các điều kiện (If-Else), các rào cản tính phí (Min/Max, VAT), đối tượng áp dụng khách hàng (KHCN, KHTC...) và luồng hoạt động (Định kỳ vs Theo giao dịch). Kết quả của Code phí sẽ được map nội bộ để tính ra lượng tiền thu cuối cùng theo công thức (US06).')
    
    doc.add_heading('2. Luồng chính (Happy Path Workflow)', level=2)
    doc.add_paragraph('Cấu hình Thông tin chung (Tên, Loại tính phí, Đối tượng, v.v) -> Điền Config VAT -> Cấu hình Điều kiện kích hoạt (Ví dụ: Loại KH = VIP, Giao dịch > 10M...) -> Submit. Hệ thống gen Mã tự động, đẩy sang Pending. Sau duyệt, Code Phí mặc định Trạng thái = "Chờ gán". Khi được ghim vào Biểu Phí, nó chuyển thành "Hoạt động".')
    
    doc.add_heading('3. Các luồng rẽ nhánh / Ngoại lệ (Alternative & Exception Flows)', level=2)
    doc.add_paragraph('• Form động đậy theo Loại tính phí: Nếu chọn "Theo giao dịch" -> Show điều kiện Giao dịch/KH/Account. Nếu chọn "Định kỳ" -> Bắt buộc hỏi Tần suất & Ngày thu, hiện điều kiện KH/Account/Thẻ.\n• Form động theo Phân khúc: Nếu click Checkbox "Khai báo theo phân khúc" -> Quy tắc tính phí sẽ bị tách ra riêng lẻ.\n• Update trạng thái tự động ngầm: Đầu ngày, CronJob quét Biểu Phí cha. Nếu Cha hết hiệu lực, toàn bộ Code Phí rụng thành "Ngừng hoạt động".')
    
    doc.add_heading('4. Bảng Điều Kiện Tiên Quyết & Cấu Hình (Pre-conditions & Settings)', level=2)
    doc.add_paragraph('• Nguồn dữ liệu (Điều kiện): Combo Box Trường điều kiện kế thừa từ chức năng US26 (Mạch Data API hoặc ETL Khách hàng/Tài khoản).\n• Ràng buộc Thứ tự nhập: Phải nhập Loại KH, Loại Tiền, Loại Tính phí đầu tiên trước khi thiết lập các tab bên dưới.')
    
    doc.add_heading('5. Ma trận Phân Quyền/Dữ Liệu', level=2)
    doc.add_paragraph('• Quyền sửa: Maker KHÔNG ĐƯỢC PHÉP SỬA thông tin nghiệp vụ/tiền đài sau khi tạo. Maker chỉ có duy nhất quyền chuyển sang trạng thái "HỦY".\n• Quyền sinh: Mã Code tự sinh, User cấm nhập.')
    
    # Part B
    doc.add_heading('Phần B: Khai Quật Lỗ Hổng & Danh Sách Q&A (Dành Cho BA)', level=1)
    
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['ID', 'Ref_ID (Mục)', 'Nội dung Câu hỏi / Sự cố', 'Phân loại', 'Đề xuất hướng xử lý từ QA', 'Câu trả lời của BA (Để trống)']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        
    qa_data = [
        (
            'QA-01', 
            'Công thức Tính VAT (Line 13)', 
            'Tài liệu viết: "Số tiền thu = Số HT tính + VAT*Số tiền HT tính". Trong khi Field VAT lại quy định nhập "Số dương <100" (Vd: nhập 10 nếu là 10%). Nếu app công thức thô trong URD sẽ bị lỗi nhân 10 lần phí thay vì nhân 0.1 (thiếu phép chia 100).', 
            'Nghiệp vụ', 
            'Cần chốt rõ logic Code và Note lại công thức chuẩn: Tiền phí * (1 + VAT/100)',
            ''
        ),
        (
            'QA-02', 
            'Luồng Chỉnh Sửa Khắc Nghiệt (Line 49, 50)', 
            'Tài liệu ghi: "Hệ thống chỉ cho phép chỉnh sửa Trạng thái Code phí... Chỉ được sửa thành trạng thái Hủy". Tức là gõ sai tên, nhầm Loại KH, sai check VAT... thì User đều phải xóa bỏ làm lại từ zero? Nếu Form quá dài sẽ gây ức chế lớn.', 
            'UI/UX', 
            'Cho phép sửa toàn bộ nếu Trạng thái = "Chờ gán". Chỉ lock không cho sửa khi Trạng thái = "Hoạt động".',
            ''
        ),
        (
            'QA-03', 
            'Validation Ngày Thu (Line 175-177)', 
            '"Thu hàng năm: Nhập DD/MM... Thu hàng tháng: Nhập số 1-31. Các tháng ko có ngày 31 hiểu là ngày cuối cùng". Vậy nhập số 31 vào tháng 2 thì hệ thống chạy Job vào ngày 28 hay bị gãy ngầm ở dưới DB do Invalid Date?', 
            'Nghiệp vụ', 
            'Confirm kỹ cơ chế map "Ngày cuối tháng" cho hệ thống Backend bằng Cron expression (như L thay cho 31) để tránh bug DateTime exception.',
            ''
        ),
        (
            'QA-04', 
            'Đứt gánh URD (Line 200)', 
            'URD đang bị VĂNG CHỮ, đoạn: "Khi Loại tính phí=Định kỳ... hệ thống kiểm tra trường Bảng DL tại quy tắc chung và..." (CÂU CHƯA VIẾT HẾT VÀ BỊ NGẮT DỞ TỪ BA).', 
            'Nghiệp vụ', 
            'Yêu cầu BA viết nốt câu còn thiếu ở line 200: Hệ thống kiểm tra ... và gì tiếp theo??',
            ''
        ),
        (
            'QA-05', 
            'Mâu thuẫn Flowchart vs Text', 
            'Ảnh Flowchart BPMN vẽ luồng "Tham chiếu tới US05 về Định nghĩa công thức tính phí". Nhưng Text lại ghi tham chiếu tới "US06 về Khai báo Biểu phí và Quy tắc". BA đang gắn lộn reference?', 
            'Nghiệp vụ', 
            'BA cần đính chính lại Flowchart để Designer/Dev đi theo đúng US.',
            ''
        ),
        (
            'QA-06', 
            'Thứ tự UX trên Flowchart', 
            'Ảnh Flowchart xếp khối "Định nghĩa Công thức tính phí" nằm TRƯỚC "Khai báo Thông tin chung" và "Điều kiện tính phí". Logic UX này rất ngược, vì form thường đi từ Thông tin chung -> Điều kiện -> Công thức.', 
            'UI/UX', 
            'Đổi vị trí box trên lưu đồ để logic thiết kế màn hình tự nhiên hơn.',
            ''
        ),
        ('QA-TC-01', 'Validation / Value', 'Tính năng VAT: Nếu chọn VAT là Có, tỉ lệ % giới hạn Min-Max là bao nhiêu? Cho phép lẻ tới mấy chữ số thập phân? (VD: 8.5%). Error msg khi nhập sai là gì?', 'Test Case / Boundary', 'Bổ sung Data Limit cho trường VAT %.', '')
    ]
    
    for row_data in qa_data:
        row_cells = table.add_row().cells
        for i, text in enumerate(row_data):
            row_cells[i].text = text
            
    doc.save('US02_QnA_Report.docx')

create_report()
