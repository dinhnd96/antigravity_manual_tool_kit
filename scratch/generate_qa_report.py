from docx import Document
from docx.shared import Pt
import os

doc = Document()

# Tiêu đề chính
heading = doc.add_heading('BÁO CÁO PHÂN TÍCH TÀI LIỆU & Q&A CHO MANUAL TESTER', 0)

doc.add_paragraph('Tài liệu phân tích: US01 - Khai báo Sản phẩm dịch vụ theo mô hình cây phân cấp')
doc.add_paragraph('Kỹ năng áp dụng: Manual Requirement Analyzer & ProfiX Domain Expert')
doc.add_paragraph('')

# Phần A
doc.add_heading('Phần A: Tóm Tắt Nghiệp Vụ Chuyên Sâu (Dành cho Tester)', level=1)

doc.add_heading('1. Thông điệp cốt lõi (Core Business Value)', level=2)
doc.add_paragraph('Chức năng cho phép quản lý danh mục Sản phẩm dịch vụ (SPDV) của ngân hàng theo dạng cây phân cấp (Nghiệp vụ là cấp 1, SPDV chi tiết là các cấp con 2, 3...) phục vụ cho việc đối soát và cấu hình các loại phí. Đây là nền tảng cốt lõi (master data) để cấu hình Biểu phí & Code Phí.')

doc.add_heading('2. Luồng chính (Happy Path Workflow)', level=2)
doc.add_paragraph('- Tạo Nghiệp vụ (Cấp 1): Maker thêm mới -> Nhập thông tin & Ngày hiệu lực/hết hiệu lực -> Lưu chờ duyệt -> Checker duyệt -> Hiển thị trên lưới Danh mục SPDV.')
doc.add_paragraph('- Tạo SPDV chi tiết (Cấp 2, 3...): Maker click dấu "+" tại bản ghi cha liền trước trên lưới -> Nhập thông tin -> Lưu chờ duyệt -> Checker duyệt.')
doc.add_paragraph('- Chỉnh sửa: Maker nhấn Chỉnh sửa (hệ thống chặn sửa các trường khác, chỉ cho phép mở trường Ngày hiệu lực & Ngày hết hiệu lực) -> Lưu chờ duyệt -> Checker duyệt.')
doc.add_paragraph('- Auto-Job (Batch xử lý tự động đầu ngày): Đầu mỗi ngày, hệ thống chạy Job tự động cập nhật trạng thái hoạt động dựa theo Ngày hiện tại và Ngày hiệu lực/hết hiệu lực.')

doc.add_heading('3. Các luồng rẽ nhánh / Ngoại lệ (Alternative & Exception Flows)', level=2)
doc.add_paragraph('- Ràng buộc tự động sinh Mã: Số tự tăng 01-99 cho mỗi cấp. Hệ thống tự động catch lỗi và chặn lưu nếu số thứ tự vượt ngưỡng 99.')
doc.add_paragraph('- Ràng buộc Logic ngày phân cấp: SPDV cấp con phải nằm hoàn toàn trong khoảng giới hạn thời gian (ngày hiệu lực -> hết hiệu lực) của SPDV cấp cha.')
doc.add_paragraph('- Nếu Maker nhập sai định dạng hoặc nhập ngày quá khứ -> Hệ thống hiển thị cảnh báo lỗi (Validation error).')

doc.add_heading('4. Bảng Điều Kiện Tiên Quyết & Cấu Hình (Pre-conditions & Settings)', level=2)
doc.add_paragraph('- Cần thiết lập giá trị tham số hệ thống PRODUCT_LEVEL (quy định số lượng cấp SPDV tối đa trong cây phân cấp).')
doc.add_paragraph('- Cấu hình Nhóm quyền & Ma trận phê duyệt cho tính năng Danh mục SPDV (US25) trước khi có thể Test duyệt yêu cầu.')

doc.add_heading('5. Ma trận Phân Quyền & Dữ Liệu', level=2)
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Role'
hdr_cells[1].text = 'Thao tác'
hdr_cells[2].text = 'Trạng thái dữ liệu'
row_1 = table.add_row().cells
row_1[0].text = 'Maker'
row_1[1].text = 'Thêm / Sửa / Khởi tạo dữ liệu SPDV'
row_1[2].text = 'Chờ duyệt (Pending)'
row_2 = table.add_row().cells
row_2[0].text = 'Checker'
row_2[1].text = 'Phê duyệt / Từ chối'
row_2[2].text = 'Hoạt động / Không hoạt động (nếu duyệt thành công)'

doc.add_paragraph('')

# Phần B
doc.add_heading('Phần B: Danh Sách Cảnh Báo & Q&A (Dành cho BA)', level=1)

table_qa = doc.add_table(rows=1, cols=5)
table_qa.style = 'Table Grid'
hdr_qa = table_qa.rows[0].cells
hdr_qa[0].text = 'ID '
hdr_qa[1].text = 'Mục tham chiếu (Ref_ID)'
hdr_qa[2].text = 'Nội dung Câu hỏi/Sự cố'
hdr_qa[3].text = 'Phân loại'
hdr_qa[4].text = 'Đề xuất hướng xử lý từ QA'

data_qa = [
    ("Q01", "Khai báo Sản phẩm, dịch vụ chi tiết - Nhãn cấp cuối cùng", "Field label 'Đây là cấp cuối cùng trong cây sơ đồ SPDV' sẽ tự động hiển thị dựa vào nhận diện tham số hệ thống PRODUCT_LEVEL hay là một flag (Checkbox) do người dùng chọn khi nhập?", "Mơ hồ về UI/UX", "BA làm rõ cơ chế hiển thị label. Đề xuất: Hệ thống tự động tính toán dựa trên độ sâu cấp bậc và tham số thiết lập PRODUCT_LEVEL để auto-show label."),
    ("Q02", "Giao diện sửa SPDV (Ràng buộc ngày quá khứ)", "Ràng buộc trên form tạo mới yêu cầu 'Ngày hiệu lực -> Không chọn ngày quá khứ'. Khi Maker 'SỬA' bản ghi SPDV TRONG KHI bản ghi đang có hiệu lực (từ tháng trước), thì hệ thống có bắt lỗi valid ngày hiệu lực đang nằm ở quá khứ hay không?", "Mâu thuẫn Nghiệp vụ", "BA bổ sung ngoại lệ (exception rule): Trường hợp SPDV đang ở trạng thái Hoạt động, khi Sửa cho phép vô hiệu hoá validate 'Không được là ngày quá khứ' với field Ngày hiệu lực, hoặc làm mờ field này chỉ cho sửa Ngày hết hiệu lực."),
    ("Q03", "Từ chối duyệt theo Workflow", "Khi Checker 'Từ chối duyệt' bản ghi dạng Sửa, thì Maker nhận lại trạng thái gì ở màn Pending? Maker có được quyền edit tiếp để trình duyệt lại không hay phải hủy tác vụ đi tạo lại? Nếu hủy tác vụ thì bản ghi gốc đang Live (Active) có bị ảnh hưởng?", "Thiếu Luồng Lỗi / Ngoại Lệ", "BA bổ sung luồng: Khi từ chối duyệt Sửa, tác vụ quay về 'Từ chối', Maker có thể edit gửi lại hoặc delete luồng tác vụ đó, KHÔNG ảnh hưởng bản ghi gốc đang hoạt động trên hệ thống."),
    ("Q04", "Luồng Xóa / Hủy bản ghi lỗi", "Tài liệu hoàn toàn không đề cập đến luồng Xóa (Delete) / Hủy. Nếu Maker nhập sai test data rồi lưu sinh rác, thay vì bắt buộc phải duyệt và chờ 'Hết hiệu lực', có được quyền Xóa cứng/Hủy SPDV không?", "Thiếu Luồng Ngoại Lệ", "Cho phép Xóa SPDV thao tác nhầm nếu SPDV chưa từng được duyệt, hoặc thêm trạng thái 'Hủy' nếu SPDV đó chưa gắn với bất kỳ SPDV con / Code phí nào."),
    ("Q05", "Ràng buộc thao tác phân cấp khi Chờ duyệt", "Giả sử SPDV A (cấp 1) đang 'Chờ duyệt', Maker có được nhấn dấu [+] tại SPDV A để tạo tiếp SPDV B (cấp 2) hay không?", "Mâu thuẫn Nghiệp vụ", "Đề xuất chặt chẽ: Nút [+] sinh SPDV con CHỈ xuất hiện khi SPDV cha đã ở trạng thái đã duyệt (Hoạt động / Không hoạt động). Chặn tạo con trên cha đang Pending.")
]

for item in data_qa:
    r_cells = table_qa.add_row().cells
    r_cells[0].text = item[0]
    r_cells[1].text = item[1]
    r_cells[2].text = item[2]
    r_cells[3].text = item[3]
    r_cells[4].text = item[4]

output_path = '/Users/mac/antigravity-testing-kit/Feature_02_SA_Tham_So_He_Thong/tài liệu/US01_Analysis_Report.docx'
doc.save(output_path)
print(f"File Docx created successfully at: {output_path}")
