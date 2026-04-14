import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = docx.Document()

    # Title
    title = doc.add_heading('BÁO CÁO PHÂN TÍCH YÊU CẦU - US19 (TRA CỨU LỊCH THU PHÍ DỰ KIẾN THEO KHÁCH HÀNG)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ─── PHẦN A ───────────────────────────────────────────────────────────────
    doc.add_heading('Phần A: Tóm Tắt Nghiệp Vụ Chuyên Sâu (Requirements Breakdown)', level=1)

    doc.add_heading('1. Thông điệp cốt lõi (Core Business Value)', level=2)
    doc.add_paragraph(
        'US19 là màn hình Tra cứu LỊCH THU PHÍ DỰ KIẾN (Forecast/Scheduled) — khác biệt hoàn toàn '
        'với US19 là lịch sử thu phí ĐÃ THỰC HIỆN. Màn hình này giúp người dùng xem trước số tiền '
        'phí sẽ bị thu trong tương lai (trong vòng 1 năm tính từ ngày hôm nay) để chủ động chuẩn bị '
        'thanh khoản hoặc kiểm soát rủi ro phát sinh.'
    )

    doc.add_heading('2. Luồng chính (Happy Path Workflow)', level=2)
    doc.add_paragraph(
        '1. User vào màn "Tra cứu lịch thu phí dự kiến theo khách hàng".\n'
        '2. Nhập Mã CIF (bắt buộc) và tuỳ chọn thêm SPDV, Biểu phí, Code phí, '
        'Từ ngày - Đến ngày (khung thời gian tối đa 1 năm).\n'
        '3. Có thể dùng [Tra cứu CIF] popup (tái sử dụng từ US18) để tra và điền Mã CIF.\n'
        '4. Nhấn [Tìm kiếm] -> Grid lịch thu phí dự kiến hiển thị.\n'
        '5. Tải xuống Excel.'
    )

    doc.add_heading('3. Các luồng rẽ nhánh / Ngoại lệ (Alternative & Exception Flows)', level=2)
    doc.add_paragraph(
        '• Khoảng thời gian lọc bị giới hạn CỨNG ở 1 năm (Từ ngày hôm nay).\n'
        '• Tài liệu không mô tả: Nếu Mã CIF không có lịch thu phí dự kiến nào thì phản hồi ra sao.\n'
        '• Tài liệu không mô tả: Điều gì xảy ra nếu User chọn "Đến ngày" trước ngày hôm nay.'
    )

    doc.add_heading('4. Bảng Điều Kiện Tiên Quyết & Cấu Hình (Pre-conditions)', level=2)
    doc.add_paragraph(
        '• Dữ liệu gốc là kết quả từ Job tính phí định kỳ (US35) đã tính trước.\n'
        '• Cột "Số thẻ" là điểm khác biệt so với US19 — US19 thêm Số thẻ để tra cứu '
        'phí liên quan đến thẻ tín dụng/ghi nợ.'
    )

    doc.add_heading('5. Ma trận Phân Quyền / Dữ Liệu', level=2)
    doc.add_paragraph(
        'Màn hình hoàn toàn Read-Only. Không có thao tác chỉnh sửa hay hủy kế hoạch thu phí.'
    )

    # ─── PHẦN B ───────────────────────────────────────────────────────────────
    doc.add_heading('Phần B: Khai Quật Lỗ Hổng & Danh Sách Q&A (Dành Cho BA)', level=1)

    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['ID', 'Ref_ID (Mục)', 'Nội dung Câu hỏi / Sự cố',
               'Phân loại', 'Đề xuất hướng xử lý từ QA', 'Câu trả lời của BA (Để trống)']
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True

    qa_data = [
        (
            'QA-01',
            'Tham chiếu "US 19" SAI — lặp lại lỗi từ US19 (Line 126)',
            'Tại Line 126: "màn hình Tra cứu CIF tham khảo tại mục Tra cứu CIF của US 19". '
            'Màn Tra cứu CIF được định nghĩa tại US18, không phải US19. '
            'Lỗi copy-paste này lặp lại y hệt US19 chưa được sửa. '
            'Cả chuỗi US18 -> US19 -> US19 đều chứa tham chiếu sai.',
            'Document',
            'Sửa "US 19" thành "US 18" tại Line 126. Rà soát toàn bộ các US khác có cùng lỗi.',
            ''
        ),
        (
            'QA-02',
            'Lệch Cột Lưới Grid: Text vs Hình UI Mockup (Line 7 vs Hình 2)',
            'Text (Line 7) định nghĩa 12 cột: STT, CIF áp dụng, Tài khoản thu phí, Số thẻ, '
            'Code phí, Loại tiền Code phí, Biểu phí, SPDV, VAT, Số tiền phí nguyên tệ, '
            'Số tiền phí quy đổi VNĐ, Ngày thu dự kiến.\n'
            'NHƯNG Hình UI Mockup (Hình 2) lại hiển thị: Số tài khoản, Số thẻ, Code, Tên, '
            'Biểu phí, SPDV, Công thức tính phí, Số tiền tính phí, Thuế, Số tiền phí, Ngày thu.\n'
            'Sai lệch nghiêm trọng: (1) Thiếu cột "Loại tiền Code phí"; '
            '(2) Thêm cột "Tên" và "Công thức tính phí" không có trong Text; '
            '(3) Cột "Thuế" trên UI ≠ cột "VAT" trong Text — tên khác nhau.',
            'UI/UX',
            'Chốt danh sách cột cuối cùng. Đồng bộ Text URD và Hình Mockup. '
            'Làm rõ "Thuế" hay "VAT" là tên chính thức trên UI.',
            ''
        ),
        (
            'QA-03',
            'Thiếu cột "CIF áp dụng" trên Mockup UI (Line 7 vs Hình 2)',
            'Text yêu cầu cột "CIF áp dụng" là cột thứ 2 trên gird. '
            'Tuy nhiên lướt toàn bộ Header Mockup (Hình 2) không tìm thấy cột "CIF áp dụng". '
            'Nếu Dev code theo Mockup UI thì người dùng sẽ không biết giao dịch dự kiến '
            'thuộc về CIF nào khi xem tổng hợp.',
            'UI/UX',
            'Bổ sung cột "CIF áp dụng" vào Grid Mockup, hoặc làm rõ lý do ẩn cột này.',
            ''
        ),
        (
            'QA-04',
            'Nhảy số STT trên bảng trường (Line 170)',
            'Bảng mô tả trường Grid nhảy từ STT 6 lên 8, bỏ qua STT 7 — lỗi lặp di truyền '
            'từ US18 và US19. Gây khó đọc và sai khi viết Test Case theo index.',
            'Document',
            'Đánh số lại tuần tự đúng: 7 -> 8 -> ... -> 18.',
            ''
        ),
        (
            'QA-05',
            'Sai chính tả trên tên trường (Line 157, 171, 177)',
            '"Từ ngày - đến ngay" (Line 157, thiếu "à"), '
            '"CIF áp dụg" (Line 171, thiếu chữ n), '
            '"Tải khoản thu phí" (Line 177, sai thành "Tải" — copy-paste chưa sửa từ US19).',
            'Document',
            'Sửa: "Từ ngày - Đến ngày", "CIF áp dụng", "Tài khoản thu phí".',
            ''
        ),
        (
            'QA-06',
            'Ràng buộc 1 năm cho Từ ngày - Đến ngày thiếu mô tả validate (Line 6 và Line 162)',
            'Text (Line 6 và 162) nêu rõ khoảng thời gian tối đa là 1 năm bao gồm ngày hôm nay. '
            'Nhưng KHÔNG mô tả:\n'
            '(1) Nếu User chọn "Đến ngày" VÀO QUÁ KHỨ (trước hôm nay) thì hệ thống báo lỗi gì?\n'
            '(2) Nếu User chọn khoảng > 1 năm thì validate message là gì?\n'
            '(3) Default value của "Từ ngày" có phải là ngày hôm nay không?',
            'Nghiệp vụ',
            'BA bổ sung mô tả validate cho cả 3 case trên, kèm nội dung thông báo lỗi cụ thể.',
            ''
        ),
        (
            'QA-07',
            'Thiếu connector Flowchart từ bước [10] về [3.2] — lỗi di truyền (Hình 1)',
            'Lỗi giống hệt US18 và US19: Sau bước [10] "Hệ thống fill CIF vào box", '
            'không có mũi tên nối vật lý về diamond [3.2] trên BPMN (Hình 1). '
            'Luồng quay trở lại màn hình chính bị đứt gãy trên sơ đồ.',
            'Document',
            'Thêm mũi tên nối từ [10] về [3.2] trên BPMN. Kiểm tra tất cả US theo mẫu cùng template này.',
            ''
        ),
        (
            'QA-08',
            'Tên Button không nhất quán: "Tra cứu" vs "Tìm kiếm" (Line 103 vs Line 31)',
            'Tại Diễn giải Flowchart (Line 31): "Nhập/chọn điều kiện và nhấn Tìm kiếm". '
            'Nhưng tại Bảng Nút Chức Năng (Line 103) lại ghi tên Button là "Tìm kiếm" — '
            'mâu thuẫn nội bộ vì nhiều màn hình khác dùng "Tra cứu" (US18, US19). '
            'Trên UI Mockup (Hình 2) cũng vẽ nút là "Tìm kiếm". '
            'Cần thống nhất toàn bộ hệ thống dùng "Tra cứu" hay "Tìm kiếm".',
            'UI/UX',
            'Thống nhất tên button xuyên suốt (gợi ý: "Tìm kiếm" cho tra cứu dữ liệu, '
            '"Áp dụng" cho lọc nâng cao). Cập nhật Text URD và Mockup đồng bộ.',
            ''
        ),
        ('QA-TC-01', 'Boundary / Date', 'Ràng buộc 1 năm cho khoảng ngày: Nếu chọn khoảng > 1 năm, validation message lỗi hiển thị chính xác từng chữ là gì? Nếu Từ ngày ở tương lai thì sao?', 'Test Case / Boundary', 'Cung cấp Error Message cụ thể để Tester viết Expected Result.', '')
    ]

    for row_data in qa_data:
        row_cells = table.add_row().cells
        for i, text in enumerate(row_data):
            row_cells[i].text = text

    doc.save('US19_QnA_Report.docx')
    print('✅ Đã tạo US19_QnA_Report.docx thành công!')

create_report()
