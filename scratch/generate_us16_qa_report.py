import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = docx.Document()
    
    # Title
    title = doc.add_heading('BÁO CÁO PHÂN TÍCH YÊU CẦU - US16 (QUẢN LÝ CHƯƠNG TRÌNH ƯU ĐÃI)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Part A
    doc.add_heading('Phần A: Tóm Tắt Nghiệp Vụ Chuyên Sâu (Requirements Breakdown)', level=1)
    
    doc.add_heading('1. Thông điệp cốt lõi (Core Business Value)', level=2)
    doc.add_paragraph('US16 là giao diện Master Data giúp người dùng tra cứu toàn bộ các Chương Trình Ưu Đãi (CTƯĐ - Promotion Programs) trong hệ thống. Đây là nơi quản lý tổng hợp các chiến dịch liên quan đến việc giảm phí/cắt phí bằng cách gắn với 1 Code phí hoặc 1 Biểu phí cụ thể.')
    
    doc.add_heading('2. Luồng chính (Happy Path Workflow)', level=2)
    doc.add_paragraph('Người dùng mở màn Danh mục CTƯĐ -> Chọn xem nhóm "Ưu đãi đánh giá định kỳ" hoặc "Không đánh giá định kỳ" -> Hệ thống load Grid Data -> Người dùng có thể click Lọc Nâng Cao để chiết xuất dữ liệu -> Click vào mã CTƯĐ để xem thông tin lịch sử chi tiết.')
    
    doc.add_heading('3. Các luồng rẽ nhánh / Ngoại lệ (Alternative & Exception Flows)', level=2)
    doc.add_paragraph('Trường hợp user muốn tìm theo khoảng thời gian: Hệ thống hiện tại KHÔNG hỗ trợ Từ ngày - Đến ngày, mà bắt buộc chọn duy nhất 1 "Ngày hiệu lực" cụ thể để tìm chính xác (Exact match).')
    
    doc.add_heading('4. Bảng Điều Kiện Tiên Quyết & Cấu Hình (Pre-conditions & Settings)', level=2)
    doc.add_paragraph('Mọi CTƯĐ trên Grid đều nằm trong trạng thái Đã được Approved (Hoạt động/Tạm dừng). Không hiển thị các bản ghi ngáp ngáp (Chờ Nháp/Từ chối). Không có logic xóa vật lý CTƯĐ.')
    
    doc.add_heading('5. Ma trận Phân Quyền/Dữ Liệu', level=2)
    doc.add_paragraph('Chế độ Read-Only đối với lưới dữ liệu. Có nút Xem Lịch Sử Tác Động (Audit Trail) để truy vết User nào đã thao túng CTƯĐ.')
    
    # Part B
    doc.add_heading('Phần B: Khai Quật Lỗ Hổng & Danh Sách Q&A (Dành Cho BA)', level=1)
    
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['ID', 'Ref_ID (Mục)', 'Nội dung Câu hỏi / Sự cố', 'Phân loại', 'Đề xuất hướng xử lý từ QA', 'Câu trả lời của BA (Để trống)']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        
    qa_data = [
        (
            'QA-01', 
            'Đứt gãy Text và UI (Thiếu Dropdown Phân Loại CTƯĐ)', 
            'Tại (Line 5-6) của chữ: BA quy định trên UI phải có trường Đổi qua lại giữa "Ưu đãi có đánh giá định kỳ" và "Không đánh giá định kỳ". NHƯNG trên Ảnh vẽ Mockup UI (Hình số 2) TRẮNG TRƠN, hoàn toàn bốc hơi Component Cấu hình này! Cả trên Flowchart cũng MẤT TÍCH bước chọn Phân loại.', 
            'UI/UX', 
            'Bắt buộc Design lại Hình Mockup UI số 2, bổ sung thêm Radio Button hoặc Dropdown "Có ĐGĐK / Không ĐGĐK" và Móc nối vào bản vẽ Lưu đồ BPMN.',
            ''
        ),
        (
            'QA-02', 
            'Lệch Cột Lưới Dữ Liệu', 
            'Text mô tả Lưới dữ liệu (Từ Line 18-28) quy định phải có: Ngày duyệt, Người duyệt. Nhưng Ảnh Thiết kế Lưới (Hình số 2) lại bốc hơi hoàn toàn 2 cột này và được nhét thêm cột "Hành Động".', 
            'UI/UX', 
            'Thiết kế lại ảnh Mockup số 2 để đồng bộ 100% với Text URD.',
            ''
        ),
        (
            'QA-03', 
            'Vết Xe Đổ của "Xóa Lọc"', 
            'Bug này lây lan từ US15, US16 sang tận US16. Nhấn Xóa Lọc [8.1 trên Flowchart] -> Hệ thống lại chạy đi "Đóng màn hình Lọc [9]" và Kết thúc. Muốn nhập Filter cực kỳ ức chế.', 
            'Nghiệp vụ', 
            'Cần sửa Flowchart, "Xóa lọc" chỉ reset TextBox, không đóng Modal.',
            ''
        ),
        (
            'QA-04', 
            'Trùng bộ đếm Index', 
            'Lại một lỗi copy-paste từ URD trước: Diễn giải text (Line 93) và (Line 97) đều đánh số thứ tự là "10". Dẫn tới Step "Tham chiếu tải xuống" và Step "Hiển thị danh sách CTƯĐ đều chập làm 1 cục số 10". Rất khó để ánh xạ sang Test Case Flow.', 
            'Document', 
            'Đánh số thứ tự tăng dần chuẩn xác lại.',
            ''
        ),
        (
            'QA-05', 
            'Tìm kiếm Ngày theo phong cách Gây Khó Dễ', 
            'Lọc Ngày Hiệu Lực/Ngày Hết Hiệu Lực (Hình 3) chỉ cung cấp 1 ô điền. Tức là Exact Match (Bắt buộc phải gõ đúng ngày 15/05/2026). Để tìm 1 CTYĐ khởi tạo trong Tháng đó, User phải mò mẫm gõ thủ công từng ngày? Khá thiếu thực tế.', 
            'Nghiệp vụ', 
            'Suggest: Sửa Component chọn ngày gốc thành dạng "Từ ngày - Đến ngày" (Date Range Picker) cho các trường Date filter.',
            ''
        ),
        ('QA-TC-01', 'Validation / Data trim', 'Ô tìm kiếm Mã CTƯĐ: Nếu search bằng cách paste dư khoảng trắng ở đầu/cuối, hệ thống có tự động Trim() khi truy vấn không?', 'Test Case / Data', 'Đề xuất Trim() toàn bộ khoảng trắng thừa để UX tốt hơn.', '')
    ]
    
    for row_data in qa_data:
        row_cells = table.add_row().cells
        for i, text in enumerate(row_data):
            row_cells[i].text = text
            
    doc.save('US16_QnA_Report.docx')

create_report()
