import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = docx.Document()
    
    # Title
    title = doc.add_heading('BÁO CÁO PHÂN TÍCH YÊU CẦU - US14 (XEM DANH MỤC SPDV)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Part A
    doc.add_heading('Phần A: Tóm Tắt Nghiệp Vụ Chuyên Sâu (Requirements Breakdown)', level=1)
    
    doc.add_heading('1. Thông điệp cốt lõi (Core Business Value)', level=2)
    doc.add_paragraph('US14 mang tới hai góc nhìn toàn cảnh cho hệ thống Sản phẩm Dịch vụ (SPDV). Người dùng có thể nhìn dưới dạng Lưới Dữ liệu phẳng (Dành cho việc quản lý Tham số, xem Code phí chưa sử dụng) và dạng Sơ đồ Cây Thư Mục (Tra cứu để xem mối liên hệ cha-con giữa SPDV, và truy vết ngược về Biểu Phí hoặc Code Phí đang gán).')
    
    doc.add_heading('2. Luồng chính (Happy Path Workflow)', level=2)
    doc.add_paragraph('Sẽ có 2 menu chính:\n- Menu Tham số: User xem Danh mục SPDV (Data Grid) và List Code phí chưa sử dụng. Tìm kiếm, Lọc nâng cao, xem Chi tiết.\n- Menu Tra cứu: User xem "Danh mục toàn hàng" dạng Cây (Tree View). Bấm vào Node cha -> Thấy link "Xem Biểu phí đi kèm". Bấm vào Node lá (Cấp cuối) -> Thấy link "Xem Code phí đi kèm".')
    
    doc.add_heading('3. Các luồng rẽ nhánh / Ngoại lệ (Alternative & Exception Flows)', level=2)
    doc.add_paragraph('Các màn hình đều có bộ lọc ngày tháng. Tại Tab SPDV, cột Ngày filter theo "Ngày hiệu lực". Tại Tab Code phí, Cột ngày filter theo "Ngày tạo". Tại Tab Biểu phí, lọc theo "Ngày hiệu lực".')
    
    doc.add_heading('4. Bảng Điều Kiện Tiên Quyết & Cấu Hình (Pre-conditions & Settings)', level=2)
    doc.add_paragraph('• SPDV cấp cuối: Mới có liên kết với Code phí. Trạng thái Hoạt động/Không hoạt động phụ thuộc vào Ngày Hệ Thống vs Ngày Hiệu Lực.\n• Code phí: Trạng thái của Code phí phụ thuộc vào trạng thái của Biểu phí chứa nó.')
    
    doc.add_heading('5. Ma trận Phân Quyền/Dữ Liệu', level=2)
    doc.add_paragraph('Màn hình chỉ thuần quyền "Tra cứu/Xem". Các thao tác chỉnh sửa dữ liệu nếu có sẽ bật Popup đẩy về luồng của US khác.')
    
    # Part B
    doc.add_heading('Phần B: Khai Quật Lỗ Hổng & Danh Sách Q&A (Dành Cho BA)', level=1)
    
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['ID', 'Ref_ID (Mục)', 'Nội dung Câu hỏi / Sự cố', 'Phân loại', 'Đề xuất hướng xử lý từ QA', 'Câu trả lời của BA (Để trống)']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        
    qa_data = [
        (
            'QA-01', 
            'Khái niệm "Chưa sử dụng" vô lý (Line 30-42)', 
            'Tab tên là "Danh sách code phí CHƯA SỬ DỤNG". Nhưng trong mô tả Trạng thái, lại liệt kê có cả trạng thái "Hoạt động" (đang gán với biểu phí hiệu lực) và "Ngừng hoạt động" (gán với biểu phí hết hiệu lực). Nếu ĐÃ GÁN với Biểu phí thì sao gọi là "Chưa sử dụng" được? Chưa sử dụng chỉ có thể là "Chờ gán".', 
            'Nghiệp vụ', 
            'Đổi tên Tab thành "Danh sách toàn bộ Code phí" hoặc giới hạn data chỉ load các code phí có trạng thái = Chờ Gán.',
            ''
        ),
        (
            'QA-02', 
            'Sự khó chịu của nút Xóa Lọc (Trong Flowchart 1 & 2)', 
            'Trên ảnh BPMN Lưu đồ 1 và 2: "10.1 Xoá Lọc" -> "11. Đóng màn hình Lọc". Khi User nhấn Xóa Lọc, bản chất họ muốn Reset Form để nhập Text mới. Nếu Đóng Form, họ lại phải click nút "Lọc" 1 lần nữa ở bên ngoài để mở lên lại. Cực kỳ kém UX.', 
            'UI/UX', 
            'Xóa lọc chỉ làm thao tác Clear Data Input Field, không Đóng Popup.',
            ''
        ),
        (
            'QA-03', 
            'Missing Step Flowchart 2', 
            'Trên ảnh Flowchart Số 2 (Tìm kiếm code phí chưa sử dụng): Đoạn kết nối đánh số [4] nối thẳng tới [6]. Bước Số 5 đã bốc hơi mất khỏi ảnh.', 
            'Document', 
            'BA cần đánh số lại Flowchart tránh gây rối cho QA khi viết Test Case theo step.',
            ''
        ),
        (
            'QA-04', 
            'Phân trang Cây Thư Mục vô lý (Line 69)', 
            'Tài liệu viết: "Hệ thống hiển thị 20 SPDV cấp 1 trên 1 trang... nhấn nút tiếp theo để xem". Component Cây (Tree View) hiển thị dạng Hierarchy, nhúng Phân trang (Pagination) vào Tree là một thiết kế UX phi tiêu chuẩn và gây gãy cấu trúc cây khi expand.', 
            'UI/UX', 
            'Tree View nên dùng Scroll bar (Lazy load) thay vì đập hẳn bộ cục Phân Trang (Pagination 1, 2, 3...) vào Component Tree.',
            ''
        ),
        (
            'QA-05', 
            'Nhãn Text Field Lọc dễ nhầm lẫn (Line 15 vs 47)', 
            'Cùng tên Field trên UI là "Từ ngày - Đến ngày", nhưng Tab 1 thì móc vào SQL gọi "Ngày hiệu lực", Tab 2 thì móc SQL lấy "Ngày tạo". Dev và User sẽ rất dễ nhầm.', 
            'UI/UX', 
            'Nên đổi tên nhãn tĩnh thành "Ngày hiệu lực (Từ-Đến)" ở Tab 1 và "Ngày tạo (Từ-Đến)" ở tab 2 để rạch ròi.',
            ''
        ),
        ('QA-TC-01', 'Validation / Boundary', 'Pagination: Lưới danh sách tối đa show bao nhiêu bản ghi mỗi trang? Dropdown SPDV có giới hạn load bao nhiêu item để tránh lag không?', 'Test Case / Boundary', 'Cần BA chốt size cụ thể để bổ sung vào Expected Result.', '')
    ]
    
    for row_data in qa_data:
        row_cells = table.add_row().cells
        for i, text in enumerate(row_data):
            row_cells[i].text = text
            
    doc.save('US14_QnA_Report.docx')

create_report()
