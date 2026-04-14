import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = docx.Document()

    # Title
    title = doc.add_heading('BÁO CÁO PHÂN TÍCH YÊU CẦU - US26 (QUẢN TRỊ DANH MỤC ĐIỀU KIỆN TÍNH PHÍ)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ─── PHẦN A ───────────────────────────────────────────────────────────────
    doc.add_heading('Phần A: Tóm Tắt Nghiệp Vụ Chuyên Sâu (Requirements Breakdown)', level=1)

    doc.add_heading('1. Thông điệp cốt lõi (Core Business Value)', level=2)
    doc.add_paragraph(
        'US26 cung cấp công cụ quản trị giúp định nghĩa các "biến số" (Điều kiện tính phí) lấy từ API '
        'hoặc ETL để đưa vào công thức tính phí. Đây là "nguyên liệu" đầu vào cho Formula Engine.'
    )

    doc.add_heading('2. Luồng chính (Happy Path Workflow)', level=2)
    doc.add_paragraph(
        '1. Admin vào Danh mục điều kiện tính phí.\n'
        '2. Nhấn [Thêm mới] -> Chọn Nguồn (API/ETL).\n'
        '   • Nếu API: Nhập Mapping note msg.\n'
        '   • Nếu ETL: Chọn Bảng dữ liệu (Customer/Account/Card) -> Chọn Trường dữ liệu tương ứng.\n'
        '3. Chọn Kiểu dữ liệu (Number/String/Date/Time).\n'
        '4. Nhấn Xác nhận -> Hệ thống lưu và hiển thị trên Lưới.'
    )

    doc.add_heading('3. Các luồng rẽ nhánh / Ngoại lệ (Alternative & Exception Flows)', level=2)
    doc.add_paragraph(
        '• Chỉnh sửa trạng thái: Hệ thống chặn chuyển trạng thái (Active/Inactive) nếu Điều kiện '
        'đang được gán vào 1 CTƯĐ hoặc Code Phí đang có hiệu lực.\n'
        '• Tải xuống: Xuất danh sách điều kiện hiện tại ra file Excel.'
    )

    doc.add_heading('4. Bảng Điều Kiện Tiên Quyết (Pre-conditions)', level=2)
    doc.add_paragraph(
        '• Mã điều kiện phải DUY NHẤT và luôn được viết hoa.\n'
        '• Dữ liệu ETL phải được tham chiếu từ các schema chuẩn (Khách hàng, Tài khoản, Thẻ).'
    )

    # ─── PHẦN B ───────────────────────────────────────────────────────────────
    doc.add_heading('Phần B: Khai Quật Lỗ Hổng & Danh Sách Q&A (Dành Cho BA)', level=1)

    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['ID', 'Ref_ID', 'Nội dung Câu hỏi / Sự cố', 'Phân loại', 'Đề xuất xử lý', 'BA Answer']
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True

    qa_data = [
        (
            'QA-01', 'Hình 1 (Flowchart)',
            'Lỗi logic Xoá lọc: BPMN vẽ bước 8.1 (Xoá lọc) dẫn tới bước 9 (Đóng màn hình lọc). '
            'User nhấn clear text mà lại bị sập luôn popup tìm kiếm.',
            'UI/UX', 'Xoá lọc chỉ clear input, KHÔNG đóng popup.', ''
        ),
        (
            'QA-02', 'Line 18-22 vs Hình 3',
            'Mâu thuẫn Filter: Text mô tả Lọc nâng cao chỉ có 4 trường (Mã, Nguồn, Trạng thái, Mapping msg). '
            'Nhưng UI Mockup (Hình 3) lại vẽ thêm "Mô tả" và "Ghi chú".',
            'UI/UX', 'BA cần chốt lại bộ lọc chính xác. Ưu tiên theo UI Mockup vì data này có sẵn.', ''
        ),
        (
            'QA-03', 'Line 40',
            'Trạng thái Code phí "Chờ gán" (Line 40)? Theo Domain Expert, Code phí thường có trạng thái '
            'Chờ duyệt/Hoạt động. "Chờ gán" là trạng thái mới phát sinh hay BA viết nhầm?',
            'Nghiệp vụ', 'Làm rõ bộ trạng thái chuẩn của Code Phí để chặn sửa trạng thái Điều kiện.', ''
        ),
        (
            'QA-04', 'Lưới danh sách (Hình 2)',
            'Thiếu trường dữ liệu ETL trên Grid: Khi Nguồn = ETL, grid chỉ show "Bảng dữ liệu" (Account/Card...). '
            'Nên show thêm cột "Trường dữ liệu" (Field name) để Admin dễ phân biệt các điều kiện cùng bảng.',
            'UI/UX', 'Bổ sung cột "Trường dữ liệu" vào Grid kết quả.', ''
        ),
        (
            'QA-05', 'Mapping note msg (Hình 2)',
            'Trên UI Mockup Hình 2 ghi "Maping note msg" (Thiếu 1 chữ p).',
            'UI/UX', 'Sửa label thành "Mapping note msg" cho đúng chính tả.', ''
        ),
        (
            'QA-TC-01', 'Validation / Case-Sensitive',
            'Mã điều kiện (Line 27) auto viết hoa. Vậy nếu đã có mã "LIMIT" (viết hoa), user gõ "limit" '
            '(viết thường) thì hệ thống báo trùng ngay khi gõ hay nhấn Xác nhận mới báo?',
            'Test Case / Logic', 'Hệ thống nên chặn trùng ngay khi blur khỏi trường Mã điều kiện (không phân biệt hoa/thường).', ''
        ),
        (
            'QA-TC-02', 'Acceptance Criteria / ETL',
            'Dependent Dropdown: Khi chọn Bảng = Khách hàng, dropdown "Trường dữ liệu" có tự động tải '
            'danh sách field tương ứng không? Nếu user đổi Bảng sau khi đã chọn Trường, hệ thống có clear Trường cũ không?',
            'Test Case / AC', 'Phải clear Trường dữ liệu cũ khi đổi Bảng để tránh râu ông nọ cắm cằm bà kia.', ''
        ),
        (
            'QA-TC-03', 'Error Handling / Special Chars',
            'Trường Mapping note Msg (API field): Có giới hạn chỉ chứa chữ cái/số/dấu gạch dưới hay không? '
            'Nếu user nhập tiếng Việt có dấu thì message báo lỗi định dạng là gì?',
            'Test Case / Error Msg', 'Quy định Regex chuẩn cho Mapping field (thường là SnakeCase hoặc CamelCase).', ''
        )
    ]

    for row_data in qa_data:
        row_cells = table.add_row().cells
        for i, text in enumerate(row_data):
            row_cells[i].text = text

    doc.save('US26_QnA_Report.docx')
    print('✅ Đã tạo US26_QnA_Report.docx thành công!')

create_report()
