import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = docx.Document()

    # Title
    title = doc.add_heading('BÁO CÁO PHÂN TÍCH YÊU CẦU - US01 (MASTER CONSOLIDATED)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ─── PHẦN A ───────────────────────────────────────────────────────────────
    doc.add_heading('Phần A: Tóm Tắt Nghiệp Vụ Chuyên Sâu (Requirements Breakdown)', level=1)

    doc.add_heading('1. Thông điệp cốt lõi (Core Business Value)', level=2)
    doc.add_paragraph(
        'US01 giúp xây dựng "bộ khung" phân cấp cho toàn bộ sản phẩm dịch vụ của ngân hàng. '
        'Đây là nền tảng để gán các loại phí chi tiết ở các US sau (US02, US06).'
    )

    doc.add_heading('2. Luồng chính (Happy Path Workflow)', level=2)
    doc.add_paragraph(
        '1. Maker tạo mới Nghiệm vụ (Cấp 1) -> Mã 01-99 auto-gen.\n'
        '2. Maker tạo mới SPDV cấp chi tiết (dưới cấp cha) -> Mã = Mã cha + 01-99.\n'
        '3. Thiết lập Ngày hiệu lực/Hết hiệu lực.\n'
        '4. Checker phê duyệt yêu cầu (theo US25).\n'
        '5. Hệ thống tự động chuyển trạng thái Hoạt động/Không hoạt động hàng ngày dựa trên ngày hiệu lực.'
    )

    doc.add_heading('3. Các luồng rẽ nhánh / Ngoại lệ (Alternative & Exception Flows)', level=2)
    doc.add_paragraph(
        '• Chỉnh sửa: Chỉ cho phép sửa Ngày hiệu lực & Ngày hết hiệu lực của bản ghi đã duyệt.\n'
        '• Giới hạn mã: Mỗi cấp chỉ chứa tối đa 99 bản ghi con. Vượt quá sẽ báo lỗi.\n'
        '• Ràng buộc phân cấp: Ngày của con phải nằm "gói gọn" trong ngày của cha.'
    )

    doc.add_heading('4. Bảng Điều Kiện Tiên Quyết (Pre-conditions)', level=2)
    doc.add_paragraph(
        '• Phải có phân quyền Maker/Checker.\n'
        '• Danh sách các Nghiệp vụ cấp 1 không được trùng mã.'
    )

    # ─── PHẦN B ───────────────────────────────────────────────────────────────
    doc.add_heading('Phần B: Khai Quật Lỗ Hổng & Danh Sách Q&A (Dành Cho BA)', level=1)

    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['ID', 'Ref_ID', 'Nội dung Câu hỏi / Sự cố', 'Phân loại', 'Đề xuất xử lý', 'BA Answer']
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True

    qa_data = [
        (
            'QA-01', 'Line 16/32',
            'Hạn chế chỉnh sửa: Tại sao hệ thống CHỈ cho phép sửa Ngày hiệu lực/hết hiệu lực? '
            'Nếu Maker gõ sai "Tên Nghiệp vụ" và đã được duyệt, thì cách xử lý là gì? Xóa đi tạo lại hay sửa DB?',
            'Nghiệp vụ', 'Cho phép sửa Tên/Mô tả nếu bản ghi chưa được gán vào Code Phí nào.', ''
        ),
        (
            'QA-02', 'Line 18/33/34',
            'Ràng buộc ngày Cha-Con: Nếu Parent Date thay đổi dẫn tới xung đột ngày với 100 Children Date '
            'bên dưới, hệ thống hiển thị cảnh báo lỗi liệt kê ds con hay chặn lưu cha?',
            'UI/UX', 'Hiển thị popup danh sách các mã con đang conflict ngày để user điều chỉnh.', ''
        ),
        (
            'QA-03', 'Line 9/29',
            'Tái sử dụng Mã (Recycle Code): Nếu mã "05" bị Hủy/Xóa, mã số 05 này có được '
            'cấp lại cho bản ghi tạo mới sau đó không (để tránh lãng phí 99 mã)?',
            'Nghiệp vụ', 'Làm rõ cơ chế auto-gen: lấy Max+1 hay quét số nhỏ nhất còn trống.', ''
        ),
        (
            'QA-04', 'Flowchart',
            'Lỗi logic Flowchart: Hình vẽ mô tả "Kiểm tra thay đổi" nhưng không có nhánh "Quay lại" '
            'nếu user không thay đổi gì mà vẫn nhấn Xác nhận.',
            'Document', 'Bổ sung nhánh Validation bắt lỗi "Không có thay đổi dữ liệu".', ''
        ),
        (
            'QA-05', 'Line 29',
            'Sâu giới hạn của Cây phân cấp (Tree depth): Mã tự sinh nối mã Cha + 02 ký tự. '
            'Việc đệ quy này giới hạn sâu tối đa bao nhiêu cấp (Node con) để tránh lỗi tràn UI và treo DB?',
            'UI/UX', 'Cần Hard-limit độ sâu của cây danh mục tối đa (Ví dụ: 3 đến 5 cấp) trong Logic.', ''
        ),
        (
            'QA-06', 'Line 20-22',
            'Job Auto Switch - Hiệu ứng Domino: Nếu Job quét Cấp Cha đã Hết hiệu lực và sửa sang "Không hoạt động", '
            'hệ thống có tự động (Cascade) cho tất cả Cấp Con bên dưới thành "Không hoạt động" theo không?',
            'Nghiệp vụ', 'Action bắt buộc của Job: Phải Cascade down xuống các child node để không xảy ra "Cha chết Con sống".', ''
        ),
        (
            'QA-TC-01', 'Test Boundary',
            'Sửa bản ghi Đang hoạt động: Ngày HL ở quá khứ (01/01/2026), nay sửa Ngày kết thúc. Khi Valid, '
            'hệ thống có báo lỗi oan "Ngày HL phải >= Hôm nay" không?',
            'Test Case / Boundary', 'Rule cho Tester: Chỉ bắt error với data bị modify, skip check cho "Ngày HL cũ".', ''
        ),
        (
            'QA-TC-02', 'Test Boundary',
            'Max limit 99: Khi tạo bản ghi thứ 100, message lỗi hiển thị chính xác text là gì? '
            'Chặn ngay màn hình hay nút lưu?',
            'Test Case / AC', 'Cung cấp Error Message text cụ thể cho Test Case.', ''
        ),
        (
            'QA-TC-03', 'Test Job',
            'Auto-update Status Job: Job tự động chuyển trạng thái có fixed timestamp không (VD: 00:00:01)? '
            'Quyết định mốc thời gian giao dịch nằm vùng tranh chấp biên.',
            'Test Case / Job', 'Cung cấp timing của Job Batch để cover Integration test.', ''
        )
    ]

    for row_data in qa_data:
        row_cells = table.add_row().cells
        for i, text in enumerate(row_data):
            row_cells[i].text = text

    doc.save('US01_QnA_Report_Master.docx')
    print('✅ Đã tạo US01_QnA_Report_Master.docx thành công!')

create_report()
