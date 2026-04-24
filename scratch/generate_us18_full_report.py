
import docx
from docx.shared import Inches, Cm
from docx.enum.section import WD_ORIENT

def create_review_report():
    doc = docx.Document()
    
    # Set Landscape orientation and Narrow margins
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)

    doc.add_heading('BÁO CÁO TỔNG HỢP QA REVIEW & Q&A - US18', 0)
    doc.add_paragraph('PHIÊN BẢN: MASTER FINAL (19 ITEMS)')

    doc.add_heading('PHẦN A – TÓM TẮT NGHIỆP VỤ (DÀNH CHO TESTER)', level=1)
    
    doc.add_heading('1. Thông điệp cốt lõi (Core Business Value)', level=2)
    doc.add_paragraph('Tính năng này cho phép người dùng xem lại toàn bộ lịch sử thu phí của một khách hàng cụ thể (theo mã CIF). Giúp đối soát dữ liệu thu phí, kiểm tra các khoản phí đã thu, số tiền VAT, và các ưu đãi hoặc override phí đã thực hiện.')

    doc.add_heading('2. Cấu trúc Luồng Nghiệp Vụ (Flow Structure)', level=2)
    doc.add_paragraph('2.1. Tra cứu lịch sử thu phí - Luồng chính (Happy Path)', style='List Bullet')
    doc.add_paragraph('Bước 1: Truy cập menu Tra cứu >> Xem lịch sử thu phí theo khách hàng.', style='List Bullet 2')
    doc.add_paragraph('Bước 2: Nhập Mã CIF (Bắt buộc) và các điều kiện lọc khác.', style='List Bullet 2')
    doc.add_paragraph('Bước 3: Nhấn nút "Tra cứu". Hệ thống hiển thị danh sách kết quả.', style='List Bullet 2')
    doc.add_paragraph('Bước 4: Nhấn nút "Tải xuống" để xuất dữ liệu ra file Excel.', style='List Bullet 2')

    doc.add_paragraph('2.2. Các Luồng Rẽ Nhánh / Ngoại lệ', style='List Bullet')
    doc.add_paragraph('Trường hợp không nhập Mã CIF: Báo lỗi trường bắt buộc [QTC-03].', style='List Bullet 2')
    doc.add_paragraph('Chức năng "Tra cứu CIF": Mở popup US19 để chọn khách hàng.', style='List Bullet 2')

    doc.add_heading('3. Quy Tắc Chung Áp Dụng (ProfiX Common Rules)', level=2)
    qtc_table = doc.add_table(rows=1, cols=3)
    qtc_table.style = 'Table Grid'
    hdr_cells = qtc_table.rows[0].cells
    hdr_cells[0].text = 'Mã QTC'
    hdr_cells[1].text = 'Tên Quy Tắc'
    hdr_cells[2].text = 'Nội dung áp dụng cho US18'
    
    qtc_data = [
        ['QTC-01.3', 'Multiple Select Dropdown', 'Áp dụng cho Sản phẩm dịch vụ, Biểu phí, Code phí.'],
        ['QTC-01.5', 'Date Format', 'Định dạng dd/mm/yyyy. Từ ngày 00:00:00, Đến ngày 23:59:59.'],
        ['QTC-03', 'Advanced Filter', 'Kết hợp điều kiện AND. Mã CIF là bắt buộc.'],
        ['QTC-05', 'Tải xuống', 'Định dạng .xlsx, tên file theo chuẩn hệ thống.'],
        ['QTC-09', 'Phân quyền Khối', 'Áp dụng mặc định theo Khối của User (Cần làm rõ ngoại lệ tra cứu CIF).']
    ]
    for m, t, n in qtc_data:
        row = qtc_table.add_row().cells
        row[0].text = m
        row[1].text = t
        row[2].text = n

    doc.add_page_break()

    doc.add_heading('PHẦN B – DANH SÁCH CẢNH BÁO & Q&A (DÀNH CHO BA)', level=1)
    
    qa_table = doc.add_table(rows=1, cols=6)
    qa_table.style = 'Table Grid'
    hdr_cells = qa_table.rows[0].cells
    hdr_cells[0].text = 'ID'
    hdr_cells[1].text = 'Vị trí trích xuất (Trong US18.docx)'
    hdr_cells[2].text = 'Nội dung Câu hỏi / Sự cố chi tiết'
    hdr_cells[3].text = 'Phân loại'
    hdr_cells[4].text = 'Đề xuất xử lý cụ thể từ QA'
    hdr_cells[5].text = 'BA Reply'

    qa_items = [
        ['QA-01', 'Mục "3. Danh sách chức năng" & Hình 2. Giao diện', 'Tên menu/chức năng không đồng nhất: Text ghi "Xem lịch sử thu phí theo khách hàng", Mockup menu ghi "Tra cứu lịch sử thu".', 'UX', 'Thống nhất tên chức năng trên menu và tiêu đề trang theo đúng Text đặc tả.'],
        ['QA-02', 'Bảng "Mô tả chi tiết các trường" (STT 15, 17-20, 22) & Hình 2', 'Lưới lịch sử trên Mockup thiếu 6 cột: VAT, CTƯĐ, Override, Thực thu nguyên tệ, Quy đổi VNĐ và Ngày đến hạn.', 'Business', 'Cập nhật Mockup hiển thị đủ 15 cột dữ liệu như đặc tả text.'],
        ['QA-03', 'Bảng "Mô tả chi tiết các trường" (STT 15, 16, 19, 20) & Hình 2', 'Các trường số tiền hiện chỉ hiển thị con số, không rõ Loại tiền (VND, USD...).', 'UX', 'Bổ sung ký hiệu ISO code (VND/USD...) hoặc ghi rõ đơn vị tính tại tiêu đề cột.'],
        ['QA-04', 'Bảng "Mô tả chi tiết các trường" (STT 9 - CIF áp dụng)', 'Cột "CIF áp dụng" hiển thị dạng Hyperlink nhưng thiếu đặc tả hành động khi click.', 'Business', 'Thiết lập click vào Mã CIF sẽ mở màn hình US19 (Chi tiết khách hàng) dưới dạng Popup.'],
        ['QA-05', 'Bảng "Mô tả chi tiết các trường" (STT 21, 22) - Cột Ngày/Giờ', 'Dữ liệu Giờ (hh:mm:ss) lấy từ nguồn nào (Backend hay giờ chốt Batch)?', 'Business', 'Lấy Server Time tại thời điểm ghi nhận giao dịch thành công để đảm bảo Audit chính xác.'],
        ['QA-06', 'Bảng "Mô tả điều kiện lọc" (STT 1 - Mã CIF)', 'Trường Mã CIF chỉ là ô nhập Text đơn thuần, dễ sai sót.', 'UX', 'Tích hợp tính năng gợi ý (Autocomplete) khi user gõ từ 3 ký tự đầu.'],
        ['QA-07', 'Bảng "Mô tả chi tiết các trường" (STT 20)', 'Chưa quy định quy tắc lấy tỉ giá tại trường "Số tiền phí thực thu quy đổi VNĐ".', 'Business', 'Hệ thống chốt tỉ giá tại thời điểm giao dịch thành công và lưu cố định vào bản ghi.'],
        ['QA-08', 'Bảng "Mô tả điều kiện lọc" (STT 4 - Code phí)', 'Dropdown "Code phí" có tự động lọc theo "Sản phẩm dịch vụ" đã chọn ở trên không?', 'UI/UX', 'Thiết lập lọc phân cấp (Cascading Filter): Chọn SPDV thì chỉ hiện Code phí thuộc SPDV đó.'],
        ['QA-09', 'Bảng "Mô tả điều kiện lọc" (STT 2 - SPDV)', 'Behavior cây thư mục SPDV: Khi chọn Node cha, có tự động chọn toàn bộ các Node con không?', 'UI/UX', 'Cần hỗ trợ tính năng "Auto-select child nodes" khi chọn Parent node.'],
        ['QA-10', 'Mục "4. Lưới dữ liệu" (Toàn mục)', 'Chưa định nghĩa tiêu chí và thứ tự sắp xếp mặc định của lưới dữ liệu.', 'Business', 'Mặc định sắp xếp theo cột "Ngày thu" với thứ tự giảm dần.'],
        ['QA-11', 'Mục "5. Tải xuống"', 'Chưa quy định giới hạn số dòng tối đa khi xuất file Excel.', 'Performance', 'Thiết lập giới hạn tối đa 50,000 bản ghi để tránh API timeout.'],
        ['QA-12', 'Mục "3. Danh sách các nút bấm" & Hình 2', 'Văn bản dùng từ "Tra cứu" nhưng Mockup dùng nhãn "Tìm kiếm".', 'UX', 'Thống nhất dùng nhãn "Tra cứu" để đồng bộ với bộ thuật ngữ hệ thống.'],
        ['QA-13', 'Bảng "Mô tả chi tiết các trường" (STT 9)', 'Lỗi chính tả (Typo) trong văn bản: "CIF áp dụg".', 'Typo', 'Sửa lại thành "CIF áp dụng".'],
        ['QA-14', 'Mục "2. Lưu đồ" (Bước 3.1) & Hình 2', 'Nút "Tra cứu CIF" có trong Lưu đồ nhưng không xuất hiện trên Mockup UI.', 'UI/UX', 'Bổ sung nút "Tra cứu CIF" cạnh trường nhập liệu Mã CIF.'],
        ['QA-15', 'Mục "Mô tả điều kiện lọc" & QTC-09', 'Màn hình không có trường "Khối". Có chặn dữ liệu theo Khối của User khi tra cứu CIF không?', 'Business', 'Xác nhận cho phép tra cứu xuyên suốt các Khối vì đây là chức năng tra cứu lịch sử theo khách hàng.'],
        ['QA-16', 'Bảng "Mô tả chi tiết các trường" (STT 22)', 'Logic hiển thị cột "Ngày đến hạn thu phí định kỳ" đối với các loại phí thu ngay?', 'Business', 'Nếu không phải phí định kỳ, hiển thị giá trị mặc định là "N/A" hoặc để trống.'],
        ['QA-17', 'Bảng "Mô tả chi tiết các trường" (STT 18)', 'Cột "Override" chỉ hiện trạng thái Sửa tăng/giảm. Có cần lý do hoặc số tiền chênh lệch không?', 'UX/Business', 'Bổ sung "Lý do override" để phục vụ công tác đối soát tốt hơn.'],
        ['QA-18', 'Bảng "Mô tả điều kiện lọc" (STT 5, 6)', 'Chưa quy định giới hạn tối đa của khoảng thời gian tra cứu (Từ ngày - Đến ngày).', 'Performance', 'Quy định giới hạn tra cứu trong vòng tối đa 01 năm.'],
        ['QA-19', 'Mục "4. Lưới dữ liệu" & Hình 2', 'Lưới dữ liệu thiếu dòng "Tổng cộng" (Summary row) ở cuối lưới.', 'UX', 'Bổ sung dòng tổng cộng để tổng hợp số tiền phí, VAT, thực thu của danh sách hiển thị.'],
    ]

    for item in qa_items:
        row = qa_table.add_row().cells
        row[0].text = item[0]
        row[1].text = item[1]
        row[2].text = item[2]
        row[3].text = item[3]
        row[4].text = item[4]
        row[5].text = '' # BA Reply empty

    # Set column widths
    widths = [Cm(1), Cm(4), Cm(8), Cm(2), Cm(6), Cm(3)]
    for row in qa_table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = width

    save_path = "/Users/mac/antigravity-testing-kit/Feature_02_SA_Tham_So_He_Thong/tài liệu/QA_Review_US18_Master_Final_19Items.docx"
    doc.save(save_path)
    print(f"File saved to: {save_path}")

if __name__ == "__main__":
    create_review_report()
