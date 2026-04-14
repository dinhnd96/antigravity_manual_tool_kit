import docx
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import docx.enum.section

def create_report():
    doc = docx.Document()
    
    # Set narrow margins & Landscape
    sections = doc.sections
    for section in sections:
        section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
        # Swap width and height for landscape
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height
        
        # Set narrow margins (0.5 inches)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Title
    title = doc.add_heading('BÁO CÁO PHÂN TÍCH YÊU CẦU - US02 (MASTER CONSOLIDATED)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ─── PHẦN A ───────────────────────────────────────────────────────────────
    doc.add_heading('Phần A: Tóm Tắt Nghiệp Vụ Chuyên Sâu (Requirements Breakdown)', level=1)

    doc.add_heading('1. Thông điệp cốt lõi (Core Business Value)', level=2)
    doc.add_paragraph(
        'US02 định nghĩa cấu trúc chi tiết cho 1 Code Phí (Bao nhiêu tiền, thuộc SPDV nào, tính '
        'theo Giao dịch hay Định kỳ, điều kiện tính là gì). Đây là "trái tim" cấu hình phí của hệ thống.'
    )

    doc.add_heading('2. Luồng chính (Happy Path Workflow)', level=2)
    doc.add_paragraph(
        '1. Từ Cây SPDV, chọn Thêm mới Code phí.\n'
        '2. Khai báo Thông tin chung (Mã tự sinh, Tên, VAT, Đối tượng, Tần suất).\n'
        '3. Chọn Loại tính phí (Giao dịch hoặc Định kỳ) để kích hoạt danh sách điều kiện lọc tương ứng.\n'
        '4. Cấu hình Điều kiện tính phí (Thêm/Sửa/Xóa các Rule logic =, <, >, IN).\n'
        '5. Xác nhận -> Lưu bản ghi trạng thái "Chờ duyệt" -> Đợi Checker duyệt.'
    )

    doc.add_heading('3. Các luồng rẽ nhánh / Ngoại lệ (Alternative & Exception Flows)', level=2)
    doc.add_paragraph(
        '• Chỉnh sửa: Hệ thống khóa toàn bộ các trường nghiệp vụ khi đã được duyệt, '
        'chỉ cho phép chuyển trạng thái thành "Hủy".\n'
        '• Logic Loại tiền: Chọn Loại tiền tệ Code phí -> Loại tiền tối đa/tối thiểu auto thay đổi theo.'
    )

    doc.add_heading('4. Bảng Điều Kiện Tiên Quyết (Pre-conditions)', level=2)
    doc.add_paragraph(
        '• Đã khai báo Cây danh mục SPDV (US01) và Danh mục điều kiện (US26).\n'
        '• Đối tượng thu phí (Merchant/Khách hàng) sẽ quết định danh sách Loại khách hàng (hiển thị hay ẩn CBNV).'
    )

    # ─── PHẦN B ───────────────────────────────────────────────────────────────
    doc.add_heading('Phần B: Khai Quật Lỗ Hổng & Danh Sách Q&A (Dành Cho BA)', level=1)

    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['ID', 'Chỉ dẫn Trích Xuất', 'Nội dung Câu hỏi / Sự cố', 'Phân loại', 'Đề xuất xử lý', 'BA Answer']
    
    # Set table width properties
    table.autofit = False
    table.allow_autofit = False
    widths = [Cm(1.5), Cm(3.0), Cm(9.0), Cm(3.0), Cm(5.0), Cm(4.0)]

    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].width = widths[i]

    qa_data = [
        (
            'QA-01', 'Mục "Thông tin chung", trường hợp Loại tính phí định kỳ',
            'Tài liệu viết bị Đứt Đoạn (Incomplete sentences): "về thống cần thực hiện kiểm tra trường '
            'Bảng dữ liệu tại phân vùng Quy tắc chung và …" -> Chữ "và" cụt lủn chưa viết xong.',
            'Document', 'BA viết bổ sung hoàn thiện spec đoạn này để Dev/QA có hướng đi tiếp.', ''
        ),
        (
            'QA-02', 'Đoạn diễn giải quy tắc Chỉnh sửa',
            'Cấm sửa toàn tập: Khi Code phí hoạt động, hệ thống "chỉ cho phép chỉnh sửa trạng thái thành Hủy". '
            'Nếu User gõ sai Tên phí, hoặc sai VAT (VD 10% thành 100%), thao tác bắt buộc là Hủy Code cũ -> '
            'Tạo lại Code mới từ đầu? Nếu đúng, đây sẽ là trải nghiệm UX cực kỳ mệt mỏi.',
            'Nghiệp vụ', 'Làm rõ cơ chế Versioning hoặc cho phép tạo Request Sửa thông tin thay vì khóa cứng.', ''
        ),
        (
            'QA-03', 'Bảng Chi tiết, mục VAT',
            'Mâu thuẫn Range giá trị VAT: Quy tắc trước yêu cầu số < 100. Đoạn sau lại cho phép số <= 100. '
            'Ngân hàng có trường hợp thuế VAT đạt 100% không?',
            'Nghiệp vụ', 'Chốt range chuẩn: (0, 100) hay (0, 100].', ''
        ),
        (
            'QA-04', 'Flowchart (Hình 1), Bước 6.1',
            'Mã lỗi "Râu ông nọ cắm cằm bà kia": Ở bước 6.1 báo lỗi, flowchart ghi mã lỗi là (PR.02.01), '
            'bước 6.b báo thành công ghi (PR.02.02). PR là viết tắt của Promotion. Đây là tham số (SA).',
            'Document', 'Cập nhật lại chuẩn mã message Hệ Thống (SA) thay vì copy từ hệ thống Promotion.', ''
        ),
        (
            'QA-TC-01', 'Bảng Chi tiết, mục Tần suất',
            'Năm nhuận (Leap Year) Hole: Tần suất Hàng năm chỉ cho nhập tối đa ngày 28/02. '
            'Vậy nếu User muốn thu vào ngày cuối tháng (EoM) của các tháng 2 năm nhuận (29/02/2024) thì bó tay? ',
            'Test Case / Date Valid', 'Nên cho cấu hình "Thu vào ngày cuối cùng của tháng" thay vì fix cứng 28/02.', ''
        ),
        (
            'QA-TC-02', 'Bảng Chi tiết, mục Operator',
            'Toán tử IN / NOT IN: Khi User chọn Toán tử (Operator) = IN hoặc NOT IN, input box "Giá trị" '
            'sẽ hiển thị dưới dạng Multiselect Dropdown hay Textbox ngăn cách bởi dấu phẩy (,)?',
            'Test Case / Input', 'Cần BA mô tả rõ format nhập liệu của danh sách (List values) để test Regex.', ''
        ),
        (
            'QA-TC-03', 'Mục "Trạng thái"',
            'Sai khác text Label: Ở đoạn text đầu ghi "Chưa gán", sau đó lại ghi "Chờ gán", tại bảng chi tiết lại là "Chờ gắn". '
            'Tester base vào chữ nào để làm expected result trên UI?',
            'UI/UX', 'Chốt đúng 1 cụm Text (Ví dụ chuẩn: "Chưa gán").', ''
        )
    ]

    for row_data in qa_data:
        row_cells = table.add_row().cells
        for i, text in enumerate(row_data):
            row_cells[i].text = text
            row_cells[i].width = widths[i]

    doc.save('US02_QnA_Report.docx')
    print('✅ Đã cập nhật (Overwrite) US02_QnA_Report.docx vơi Traceability Rule (Không dùng Line Number)')

create_report()
