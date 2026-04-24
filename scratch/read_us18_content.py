
import docx
import sys

def read_docx(file_path):
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Read tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                full_text.append(" | ".join(row_text))
                
        return "\n".join(full_text)
    except Exception as e:
        return f"Error reading {file_path}: {str(e)}"

if __name__ == "__main__":
    us18_path = "/Users/mac/antigravity-testing-kit/Feature_02_SA_Tham_So_He_Thong/tài liệu/US18.docx"
    review_path = "/Users/mac/antigravity-testing-kit/Feature_02_SA_Tham_So_He_Thong/tài liệu/QA_Review_US18_Master_Final_Full_14Items.docx"
    
    print("--- US18 CONTENT ---")
    print(read_docx(us18_path))
    print("\n\n--- REVIEW CONTENT ---")
    print(read_docx(review_path))
