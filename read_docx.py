import docx
import sys

def extract_docx(file_path, output_path):
    doc = docx.Document(file_path)
    with open(output_path, 'w', encoding='utf-8') as f:
        for para in doc.paragraphs:
            if para.text.strip():
                f.write(para.text + '\n')
        
        for table in doc.tables:
            f.write("\n--- TABLE ---\n")
            for row in table.rows:
                row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                f.write(" | ".join(row_data) + "\n")
            f.write("-------------\n")

if __name__ == '__main__':
    extract_docx(
        "/Users/mac/antigravity-testing-kit/Tài liệu toàn hệ thống Profix/tài liệu/SRS - US07 sinh lời chủ động.docx",
        "/Users/mac/antigravity-testing-kit/US07_extracted.txt"
    )
    print("Done extracting US07 docx.")
