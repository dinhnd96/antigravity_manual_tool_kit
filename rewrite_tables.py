from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

doc = Document('/Users/mac/antigravity-testing-kit/Feature_02_SA_Tham_So_He_Thong/tài liệu/QA_Review_US01_Analysis.docx')

# Print headings and following tables
for i, p in enumerate(doc.paragraphs):
    if p.style.name.startswith('Heading'):
        print(f"[{i}] {p.text}")

print("Tables in doc:")
for i, t in enumerate(doc.tables):
    print(f"Table {i}: {len(t.rows)} rows, first cell: {t.cell(0,0).text[:20]}")

