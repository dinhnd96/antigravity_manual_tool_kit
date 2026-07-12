#!/usr/bin/env python3
"""Extract text from FR02 docx file."""
from docx import Document
import sys

doc = Document("/Users/mac/antigravity-testing-kit/FR02- Mua trái phiếu PVCB Bond - Duyệt.docx")

# Extract paragraphs
print("=" * 80)
print("PARAGRAPHS")
print("=" * 80)
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        style = para.style.name if para.style else "None"
        print(f"[P{i}|{style}] {para.text}")

# Extract tables
print("\n" + "=" * 80)
print("TABLES")
print("=" * 80)
for ti, table in enumerate(doc.tables):
    print(f"\n--- Table {ti+1} ---")
    for ri, row in enumerate(table.rows):
        cells = [cell.text.strip().replace('\n', ' | ') for cell in row.cells]
        print(f"  Row {ri}: {' || '.join(cells)}")
