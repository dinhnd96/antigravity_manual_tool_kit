# US39 Part A - Summary Generator
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "output")
os.makedirs(OUTPUT_DIR, exist_ok=True)

doc = Document()

# Page setup: Landscape, narrow margins
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
    return h

def add_para(text, bold=False, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    return p

# ===== TITLE =====
title = doc.add_heading('US39 — Phần A: Tóm Tắt Nghiệp Vụ Chuyên Sâu', level=1)
add_para('Feature: SA — Tham Số Hệ Thống', bold=True)
add_para('User Story: US39 — Là người dùng, khi khách hàng phát sinh giao dịch trả nợ trước hạn trên kênh offline (T24), tôi muốn hệ thống ProfiX tự động tính toán số tiền phí cần thu')
add_para('Dành cho: Manual Tester — Đọc hiểu luồng nghiệp vụ trước khi viết Test Case.', bold=True)

# ===== A.1 CORE BUSINESS VALUE =====
add_heading('A.1. Thông Điệp Cốt Lõi (Core Business Value)', level=2)
add_para('Mục đích: Tự động tính phí trả nợ trước hạn khi khách hàng thực hiện giao dịch trả nợ trước hạn trên kênh offline (T24). Hệ thống T24 call API đến ProfiX để tính phí, sau đó cho phép user sửa số tiền phí trước khi commit giao dịch.')
add_para('Người dùng cuối: User Maker (khởi tạo giao dịch trên T24) và User Checker (phê duyệt giao dịch).')
add_para('Hệ thống liên quan: T24 (hệ thống giao dịch nội bộ ngân hàng), ProfiX (hệ thống quản lý phí), DWH (Data Warehouse — nguồn dữ liệu ETL).')

# ===== A.2 FLOW STRUCTURE =====
add_heading('A.2. Cấu Trúc Luồng Nghiệp Vụ & Phân Bổ Module', level=2)

add_heading('Module 0: Đồng bộ dữ liệu T-1 (ETL)', level=3)
add_para('Luồng chính (Happy Path):', bold=True)
add_para('• Ngân hàng hoàn tất khóa sổ cuối ngày → DWH thực hiện ETL dữ liệu T-1 (Tiền vay/LD) vào bảng dữ liệu Tài khoản trên ProfiX.', indent=0.5)
add_para('• Dữ liệu ETL bao gồm: Loại tài khoản (mặc định LD), CIF, Số LD, Loại tiền, Thời hạn khoản vay, Đơn vị thời hạn khoản vay, % thời hạn hợp đồng, Số tháng tồn tại khoản vay (THKV), Mã hạch toán, Ngày hiệu lực, Ngày đáo hạn, Trạng thái (chỉ lấy tài khoản còn hoạt động).', indent=0.5)
add_para('Luồng rẽ nhánh / Ngoại lệ:', bold=True)
add_para('• ETL fail giữa chừng → dữ liệu LD chưa sẵn sàng trên ProfiX → ảnh hưởng bước tính phí.', indent=0.5)
add_para('• Dữ liệu LD không đồng bộ kịp (T24 chạy giao dịch trước khi DWH ETL xong).', indent=0.5)

add_heading('Module 1: Tính phí trả nợ trước hạn (Luồng chính)', level=3)
add_para('Luồng chính (Happy Path):', bold=True)
add_para('Bước 1: Maker khởi tạo giao dịch trả nợ trước hạn trên T24.', indent=0.5)
add_para('Bước 2: Maker nhập số tài khoản LD + số tiền trả nợ trước hạn.', indent=0.5)
add_para('Bước 3: T24 call API đến ProfiX để tính phí (tham chiếu US33 — xử lý tính phí theo giao dịch).', indent=0.5)
add_para('Bước 4: ProfiX thực hiện tính toán phí dựa trên điều kiện Code phí (THKV, Mã hạch toán, % thời hạn hợp đồng).', indent=0.5)
add_para('Bước 5: T24 hiển thị thông tin phí cho user.', indent=0.5)
add_para('Bước 6 (Decision): User lựa chọn:', indent=0.5)
add_para('  → Sửa số tiền phí thực thu (Bước 6.1) → T24 lưu số tiền phí mới → Nhấn commit.', indent=1)
add_para('  → Nhấn commit trực tiếp (không sửa).', indent=1)
add_para('Bước 7: Checker phê duyệt giao dịch trên T24.', indent=0.5)
add_para('Bước 8: Sau khi duyệt thành công, T24 call API đến ProfiX để ghi nhận thông tin thu phí thành công.', indent=0.5)
add_para('Bước 9: ProfiX lưu giao dịch thu phí thành công.', indent=0.5)

add_para('Luồng rẽ nhánh / Ngoại lệ:', bold=True)
add_para('• User sửa số tiền phí nằm NGOÀI khoảng [Phí tối thiểu, Phí tối đa] của Code phí → T24 báo lỗi.', indent=0.5)
add_para('• Checker từ chối giao dịch → giao dịch không được duyệt → ProfiX không ghi nhận.', indent=0.5)
add_para('• API call từ T24 đến ProfiX thất bại (timeout, lỗi mạng).', indent=0.5)
add_para('• ProfiX không tìm thấy Code phí / Biểu phí phù hợp cho LD.', indent=0.5)

# ===== A.3 PRECONDITIONS =====
add_heading('A.3. Điều Kiện Tiên Quyết & Cấu Hình', level=2)

data_precond = [
    ['#', 'Điều kiện', 'Mô tả'],
    ['1', 'ETL dữ liệu LD', 'DWH phải hoàn tất ETL dữ liệu T-1 vào bảng Tài khoản trên ProfiX trước khi T24 gọi API tính phí.'],
    ['2', 'Code phí đã khai báo', 'Các Code phí trả nợ trước hạn phải được khai báo trên ProfiX với điều kiện: THKV, Mã hạch toán, % thời hạn hợp đồng.'],
    ['3', 'Biểu phí đã gắn Code phí', 'Biểu phí dịch vụ (chung) hoặc Biểu phí chương trình ưu đãi cho vay phải được gắn các Code phí tương ứng.'],
    ['4', 'Tài khoản LD còn hoạt động', 'Chỉ LD có trạng thái hoạt động mới được đồng bộ ETL.'],
    ['5', 'Tham chiếu US33', 'Logic xử lý tính phí theo giao dịch của ProfiX tuân theo US33 (kênh online). US39 tái sử dụng logic này cho kênh offline T24.'],
    ['6', 'Phân quyền', 'Maker: khởi tạo và commit giao dịch. Checker: phê duyệt giao dịch.'],
]
table = doc.add_table(rows=len(data_precond), cols=3, style='Table Grid')
for i, row_data in enumerate(data_precond):
    for j, cell_text in enumerate(row_data):
        cell = table.cell(i, j)
        cell.text = cell_text
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                if i == 0:
                    run.bold = True

# ===== A.4 QTC =====
add_heading('A.4. Quy Tắc Chung Áp Dụng (Common Rules)', level=2)
qtc_list = [
    'QTC-01.4 (Number): Định dạng số tiền phí — phân cách hàng nghìn bằng dấu phẩy, 2 chữ số thập phân (ngoại trừ VND/JPY).',
    'QTC-01.5 (Date): Định dạng ngày hiệu lực / đáo hạn — dd/mm/yyyy.',
    'QTC-11 (FE-First Error Handling): T24 (FE) chặn lỗi trước: validate số tiền phí nằm trong khoảng [tối thiểu, tối đa].',
    'QTC-12 (Maker-Checker): Giao dịch phải qua Maker → Checker. Sau khi Checker duyệt, T24 mới call API ghi nhận thu phí.',
]
for item in qtc_list:
    add_para(f'• {item}', indent=0.5)

# ===== A.5 BUSINESS RULES =====
add_heading('A.5. Nghiệp Vụ Phí Trả Nợ Trước Hạn — Bảng Bậc Thang Minh Họa', level=2)
add_para('Ví dụ minh họa từ tài liệu (KH tham gia Chương trình cho vay ưu đãi có TSBĐ là BĐS):')

fee_data = [
    ['Điều kiện THKV', 'Tỷ lệ phí'],
    ['THKV ≤ 12 tháng', '3% × Số tiền trả trước hạn'],
    ['12M < THKV ≤ 18M', '2.5% × Số tiền trả trước hạn'],
    ['18M < THKV ≤ 24M', '2% × Số tiền trả trước hạn'],
    ['24M < THKV ≤ 36M', '1% × Số tiền trả trước hạn'],
    ['THKV > 36M hoặc THKV > 70% thời hạn HĐ', 'Miễn phí'],
]
table2 = doc.add_table(rows=len(fee_data), cols=2, style='Table Grid')
for i, row_data in enumerate(fee_data):
    for j, cell_text in enumerate(row_data):
        cell = table2.cell(i, j)
        cell.text = cell_text
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                if i == 0:
                    run.bold = True

add_para('')
add_para('Lưu ý quan trọng:', bold=True)
add_para('• KH KHÔNG tham gia Chương trình ưu đãi → phí theo Biểu phí dịch vụ chung trong từng thời kỳ.', indent=0.5)
add_para('• KH CÓ tham gia Chương trình ưu đãi → phí theo Code phí gắn với Biểu phí riêng cho CTƯĐ cho vay, phân biệt bằng Mã hạch toán.', indent=0.5)
add_para('• Code phí gom nhóm các bậc thang có cùng THKV + tỷ lệ phí → mỗi nhóm = 1 Code phí, với điều kiện Mã hạch toán tương ứng.', indent=0.5)

# ===== SAVE =====
output_path = os.path.join(OUTPUT_DIR, "US39_PartA_Summary.docx")
doc.save(output_path)
print(f"✅ Part A saved: {output_path}")
