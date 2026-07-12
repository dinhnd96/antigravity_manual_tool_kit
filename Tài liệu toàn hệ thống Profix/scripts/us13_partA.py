"""US13 Part A - Business Summary Generator"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# --- Page Setup: Landscape + Narrow margins ---
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10)

# === TITLE ===
title = doc.add_heading('US13 — Phần A: Tóm Tắt Nghiệp Vụ Chuyên Sâu (Dành Cho Tester)', level=0)
doc.add_paragraph('Feature: Khai báo CTƯĐ không xác định sẵn danh sách khách hàng áp dụng')
doc.add_paragraph('Phiên bản: v1.0 | Ngày phân tích: 13/05/2026')
doc.add_paragraph('')

# === A.1 Core Business Value ===
doc.add_heading('A.1. Thông Điệp Cốt Lõi (Core Business Value)', level=1)
p = doc.add_paragraph()
p.add_run('Mục đích: ').bold = True
p.add_run(
    'US13 cho phép người dùng (Maker) khai báo và quản lý các Chương trình ưu đãi (CTƯĐ) '
    'không xác định sẵn danh sách khách hàng áp dụng. Khác với US12 (có danh sách KH cụ thể), '
    'US13 áp dụng cho các chương trình ưu đãi mang tính phổ quát — áp dụng cho BẤT KỲ khách hàng nào '
    'thỏa mãn Điều kiện áp dụng tại thời điểm giao dịch, mà không cần chỉ định danh sách CIF trước.'
)
p2 = doc.add_paragraph()
p2.add_run('Người dùng cuối: ').bold = True
p2.add_run('Maker (nhân viên khai báo CTƯĐ) và Checker (cấp phê duyệt) theo luồng Maker-Checker [QTC-12].')
p3 = doc.add_paragraph()
p3.add_run('Mối quan hệ với US12: ').bold = True
p3.add_run(
    'US13 tái sử dụng gần như toàn bộ logic của US12 (Thông tin chung, Chu kỳ áp dụng, Điều kiện áp dụng, '
    'Chi tiết ưu đãi). Điểm khác biệt DUY NHẤT: Toggle "Theo danh sách khách hàng" = OFF → '
    'KHÔNG có trường CIF/Tên KH trong bảng Chi tiết ưu đãi, KHÔNG có nút "Tải lên" (Upload), '
    'KHÔNG có nút "Thêm mới khách hàng".'
)

# === A.2 Flow Structure & Module Mapping ===
doc.add_heading('A.2. Cấu Trúc Luồng Nghiệp Vụ & Phân Bổ Module', level=1)

# --- Module 1 ---
doc.add_heading('Module 1: Thêm Mới CTƯĐ Không Xác Định Sẵn Danh Sách Khách Hàng', level=2)
doc.add_heading('Luồng chính (Happy Path):', level=3)
steps = [
    'Người dùng truy cập: Tham số >> Chương trình ưu đãi >> Quản lý CTƯĐ >> Thêm mới.',
    'Chọn "Ưu đãi không có đánh giá định kỳ" → FE hiển thị màn hình khai báo CTƯĐ không đánh giá định kỳ.',
    'Đảm bảo Toggle "Theo danh sách khách hàng" = OFF → FE hiển thị form khai báo KHÔNG CÓ trường CIF/Tên tại phân vùng Chi tiết ưu đãi.',
    'Khai báo Thông tin chung: Tên CTƯĐ, Ngày ban hành (≤ hôm nay), Ngày hiệu lực (≥ hôm nay), Ngày hết hiệu lực (≥ Ngày hiệu lực), Số văn bản, Tên văn bản, Số VB_Tên viết tắt (≤20 ký tự), Email CBNV đầu mối, Link iDoc, Đối tượng thu phí, Khối, Loại ưu đãi. [Tham chiếu US12]',
    'Khai báo Điều kiện áp dụng: Thêm nhóm điều kiện, chọn Nguồn dữ liệu (ETL KH/TK/Thẻ/API giao dịch), Tên trường điều kiện, Operator, Giá trị. [Tham chiếu US12/US11]',
    'Khai báo Chu kỳ áp dụng: Chọn "Theo chu kỳ" (Hàng tuần/Hàng tháng) HOẶC "Liên tục". [Tham chiếu US12]',
    'Khai báo Chi tiết ưu đãi: Toggle "Ưu đãi theo tỷ lệ" (On = %, Off = giá trị cố định). Mỗi dòng gồm: SPDV, Giá trị ưu đãi/Tỷ lệ ưu đãi, Loại tiền tối thiểu/tối đa, Số tiền phí tối thiểu, Số tiền phí tối đa, Trường giá trị kiểm tra, Ngưỡng dừng ưu đãi. Có thể thêm nhiều dòng SPDV bằng nút "+ Thêm SPDV áp dụng". [Tham chiếu US12, LOẠI TRỪ CIF/Tên]',
    'Nhấn "Xác nhận" → BE validate toàn bộ → Lưu bản ghi trạng thái "Chờ duyệt" → Toast thành công → Quay về lưới CTƯĐ.',
    'Checker phê duyệt (luồng Maker-Checker [QTC-12/US25]) → Mã CTƯĐ được sinh tự động (tham chiếu US04).',
]
for s in steps:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('Các Luồng Rẽ Nhánh / Ngoại Lệ (Module 1):', level=3)
exceptions = [
    'Nhấn "Đóng": Hệ thống không lưu dữ liệu, quay về màn hình CTƯĐ [QTC-15].',
    'Validate BE (Bước 10): Dữ liệu khai báo không hợp lệ → FE hiển thị thông báo lỗi, kết thúc.',
    'Toggle On/Off ưu đãi theo tỷ lệ: Hệ thống tự động clear toàn bộ giá trị đã nhập của lựa chọn trước đó.',
    'Tỷ lệ ưu đãi: Phải > 0 và ≤ 100. Giá trị ưu đãi: Phải > 0.',
    'Số tiền phí tối đa phải > Số tiền phí tối thiểu (nếu cả hai đều có giá trị).',
    'Ngưỡng dừng ưu đãi: Bắt buộc nhập nếu có chọn Trường giá trị kiểm tra.',
    'SPDV trong CTƯĐ: Không được chọn các SPDV có quan hệ cha-con.',
    'Chỉnh sửa không thay đổi gì → FE disable nút Xác nhận hoặc cảnh báo [QTC-14.1].',
]
for e in exceptions:
    doc.add_paragraph(e, style='List Bullet')

# --- Module 2 ---
doc.add_heading('Module 2: Chỉnh Sửa CTƯĐ Không Xác Định Sẵn Danh Sách Khách Hàng', level=2)
doc.add_heading('Luồng chính (Happy Path):', level=3)
edit_steps = [
    'Người dùng nhấn "Chỉnh sửa" trên lưới CTƯĐ.',
    'BE kiểm tra: Có tác vụ "Chờ duyệt" không? → Nếu CÓ → Hiển thị thông báo lỗi, không cho chỉnh sửa, kết thúc.',
    'BE kiểm tra trạng thái hoạt động CTƯĐ:',
    '   a) Hết hiệu lực → Thông báo lỗi "Trạng thái CTƯĐ không hợp lệ", kết thúc.',
    '   b) Chưa hiệu lực → Cho phép chỉnh sửa TOÀN BỘ trừ Mã CTƯĐ và Số VB_Tên viết tắt CTƯĐ. Validate tương tự thêm mới.',
    '   c) Đang hiệu lực → CHỈ cho phép chỉnh sửa:',
    '      • Ngày hết hiệu lực CTƯĐ',
    '      • Chi tiết ưu đãi: Thêm mới bản ghi + Xóa bản ghi (KHÔNG cho phép chỉnh sửa bản ghi hiện hữu).',
    'Nhấn "Xác nhận" → BE validate → Tạo bản ghi Chờ duyệt → Toast thành công → Quay về lưới CTƯĐ.',
    'Checker phê duyệt → Dữ liệu chính thức được cập nhật trên lưới.',
]
for s in edit_steps:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('Các Luồng Rẽ Nhánh / Ngoại Lệ (Module 2):', level=3)
edit_exceptions = [
    'Nhấn "Đóng": Không lưu, quay về lưới [QTC-15].',
    'Validate BE (Bước 5 US12): Không hợp lệ → Thông báo lỗi, kết thúc.',
    'Chỉnh sửa không thay đổi gì → FE disable nút Xác nhận hoặc cảnh báo [QTC-14.1].',
    'Hệ thống kiểm tra các điều kiện và ràng buộc tương tự như khi thêm mới CTƯĐ.',
]
for e in edit_exceptions:
    doc.add_paragraph(e, style='List Bullet')

# === KHÁC BIỆT US13 vs US12 ===
doc.add_heading('Bảng Tổng Hợp Khác Biệt US13 vs US12', level=2)
diff_table = doc.add_table(rows=7, cols=3, style='Light Grid Accent 1')
diff_header = ['Tiêu chí', 'US12 (Có danh sách KH)', 'US13 (Không có danh sách KH)']
for i, h in enumerate(diff_header):
    diff_table.rows[0].cells[i].text = h
    for p in diff_table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True
diffs = [
    ['Toggle "Theo danh sách KH"', 'ON', 'OFF'],
    ['Trường CIF trong Chi tiết ưu đãi', 'CÓ (★ bắt buộc)', 'KHÔNG CÓ'],
    ['Trường Tên KH trong Chi tiết ưu đãi', 'CÓ (◎ auto-fill theo CIF)', 'KHÔNG CÓ'],
    ['Nút "Tải lên" (Upload file KH)', 'CÓ', 'KHÔNG CÓ (Mockup xác nhận)'],
    ['Nút "Thêm mới khách hàng"', 'CÓ', 'KHÔNG CÓ'],
    ['Nút "Tải xuống" (Download)', 'CÓ', 'CÓ (Mockup chu kỳ: có. Mockup liên tục: KHÔNG)'],
]
for i, row_data in enumerate(diffs):
    for j, val in enumerate(row_data):
        diff_table.rows[i+1].cells[j].text = val

# === A.3 Pre-conditions ===
doc.add_heading('A.3. Bảng Điều Kiện Tiên Quyết & Cấu Hình', level=1)
pre_table = doc.add_table(rows=8, cols=3, style='Light Grid Accent 1')
pre_header = ['STT', 'Điều Kiện', 'Mô Tả']
for i, h in enumerate(pre_header):
    pre_table.rows[0].cells[i].text = h
    for p in pre_table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True
preconditions = [
    ['1', 'Phân quyền Maker', 'Người dùng phải có quyền Thêm mới / Chỉnh sửa CTƯĐ.'],
    ['2', 'Phân quyền Checker', 'Cấp phê duyệt phải có quyền Phê duyệt / Từ chối [QTC-12].'],
    ['3', 'Dữ liệu mồi: SPDV', 'Danh mục SPDV phải tồn tại với trạng thái = Hoạt động.'],
    ['4', 'Dữ liệu mồi: Điều kiện tính phí', 'Các trường điều kiện tính phí phải tồn tại với trạng thái = Hoạt động.'],
    ['5', 'Dữ liệu mồi: Loại tiền tệ', 'Danh mục Loại tiền phải tồn tại trong hệ thống.'],
    ['6', 'Cấu hình Mã CTƯĐ', 'Nguyên tắc tự sinh Mã CTƯĐ theo US04 phải được cấu hình sẵn.'],
    ['7', 'Cấu hình Toggle', 'Toggle "Theo danh sách khách hàng" mặc định = OFF trên màn hình thêm mới.'],
]
for i, row_data in enumerate(preconditions):
    for j, val in enumerate(row_data):
        pre_table.rows[i+1].cells[j].text = val

# === A.4 QTC Applied ===
doc.add_heading('A.4. Quy Tắc Chung Áp Dụng (ProfiX Common Rules)', level=1)
qtc_table = doc.add_table(rows=13, cols=2, style='Light Grid Accent 1')
qtc_header = ['Mã QTC', 'Áp Dụng Trong US13']
for i, h in enumerate(qtc_header):
    qtc_table.rows[0].cells[i].text = h
    for p in qtc_table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True
qtcs = [
    ['QTC-01.1', 'Combobox: Search-as-you-type cho SPDV, Tên trường điều kiện.'],
    ['QTC-01.2', 'Dropdown List: Đối tượng thu phí, Khối, Loại ưu đãi, Nguồn dữ liệu, Operator, Tuần, Loại tiền, Trường giá trị kiểm tra.'],
    ['QTC-01.3', 'Multiple Select Dropdown: Áp dụng vào thứ, Ngày trong tháng.'],
    ['QTC-01.4', 'Number: Giá trị ưu đãi, Tỷ lệ ưu đãi, Số tiền phí tối thiểu/tối đa, Ngưỡng dừng ưu đãi, Ngày cụ thể (1-31).'],
    ['QTC-01.5', 'Date: Ngày ban hành, Ngày hiệu lực, Ngày hết hiệu lực — format dd/mm/yyyy.'],
    ['QTC-01.6', 'Text mặc định: Tên CTƯĐ (50 ký tự), Mô tả (300 ký tự) — trừ khi US ghi khác.'],
    ['QTC-05', 'Tải xuống: Format Excel .xlsx, tên file = "{Tên chức năng} - yyyymmddhhmmss".'],
    ['QTC-11', 'Xử lý lỗi FE-First: FE chặn lỗi đầu tiên, BE là tầng cuối.'],
    ['QTC-12', 'Luồng Maker-Checker: Maker lưu → Chờ duyệt → Checker phê duyệt → Mã sinh chính thức.'],
    ['QTC-14.1', 'No-Change Guard: FE disable nút Xác nhận nếu không có thay đổi.'],
    ['QTC-14.2', 'Tên CTƯĐ không yêu cầu unique.'],
    ['QTC-15', 'Nút Đóng: Không popup xác nhận, không lưu, quay về màn hình trước.'],
]
for i, row_data in enumerate(qtcs):
    for j, val in enumerate(row_data):
        qtc_table.rows[i+1].cells[j].text = val

# === Save ===
out_dir = '/Users/mac/antigravity-testing-kit/Feature_02_SA_Tham_So_He_Thong/US11-20'
os.makedirs(out_dir, exist_ok=True)
out_path = os.path.join(out_dir, 'US13_PartA_Summary.docx')
doc.save(out_path)
print(f"✅ Saved: {out_path}")
