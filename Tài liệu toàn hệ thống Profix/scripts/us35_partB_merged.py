"""US35 Part B Merged - Tổng hợp AI + VA Q&A"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(9)

title = doc.add_heading('US35 – Part B Merged: AI + VA Q&A Tổng Hợp', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Ghi chú: Nguồn = AI (sinh bởi AI) | VA (từ Validation Analyst) | BOTH (cả 2 nguồn phát hiện)')
doc.add_paragraph('')

# Merged Q&A data: (ID, Trích xuất, Câu hỏi, Phân loại, Đề xuất, Nguồn)
qa_data = [
    # === HM1: Nghiệp vụ ===
    ("US35-QA-01.1", 
     'Mục "Yêu cầu nghiệp vụ", đoạn TK thu phí mặc định',
     'Khi đối tượng tính phí = Customer, nếu TK thu phí mặc định không đủ ĐK VÀ không tìm được TK thay thế nào đủ ĐK → khoản phí xử lý thế nào? Tài liệu chỉ mô tả "tìm TK thay thế" nhưng không nêu nhánh thất bại.',
     'Nghiệp vụ',
     'Đề xuất: Khoản phí không ghi vào Topic, trạng thái = "Chưa thanh toán" hoặc ghi log riêng. Cần BA xác nhận.',
     'BOTH'),

    ("US35-QA-01.2",
     'Mục "Yêu cầu nghiệp vụ", đoạn timeout Kafka',
     'Trạng thái "Đang xử lý" duy trì vô thời hạn nếu Kafka message bị mất hoặc T24 không phản hồi. Có cơ chế timeout/retry không?',
     'Nghiệp vụ',
     'Đề xuất: Cần timeout (VD: 24h → tự chuyển "Chưa thanh toán" + ghi log) hoặc retry. Cần BA xác nhận.',
     'BOTH'),

    ("US35-QA-01.3",
     'Bảng "Xử lý tính phí định kỳ", Bước 1',
     'Với đối tượng = Account/Card, bước 2 (kiểm tra TK) bị bỏ qua. TK/Thẻ đó có cần kiểm tra trạng thái + CA_PRODUCT không? Tài liệu chỉ mô tả cho Customer.',
     'Nghiệp vụ',
     'Đề xuất: Account → kiểm tra trạng thái + CA_PRODUCT tương tự Customer. Card → kiểm tra trạng thái thẻ. Cần BA xác nhận.',
     'AI'),

    ("US35-QA-01.4",
     'Mục "Yêu cầu nghiệp vụ", đoạn "thứ tự ưu tiên nghiệp vụ phí"',
     'Thứ tự ưu tiên được cài đặt theo đơn vị nào: từng Job, từng Code phí, hay Nhóm nghiệp vụ? Nếu 2 Code phí cùng ưu tiên thì thứ tự ghi Kafka theo tiêu chí nào?',
     'Nghiệp vụ',
     'Đề xuất: Cần BA xác nhận đơn vị cài đặt ưu tiên và tiebreak khi cùng mức.',
     'BOTH'),

    ("US35-QA-01.5",
     'Bảng "Xử lý tính phí định kỳ", Bước 1, Account/Card',
     'Nếu TK thỏa ĐK (1) nhưng CIF không thỏa ĐK (2) → bỏ qua hoàn toàn hay chỉ bỏ ĐK KH? Logic AND hay OR?',
     'Nghiệp vụ',
     'Đề xuất: Phải thỏa CẢ HAI (AND). Nếu CIF không thỏa → không sinh phí. Cần BA xác nhận.',
     'BOTH'),

    ("US35-QA-01.6",
     'Mục "Yêu cầu nghiệp vụ", đoạn chu kỳ thu',
     'Phí thu hàng tháng vào "ngày cố định". Nếu ngày cố định = 31 nhưng tháng chỉ có 28/29/30 ngày → Job chạy ngày nào?',
     'Nghiệp vụ',
     'Đề xuất: Ngày cố định > số ngày tháng → chạy ngày cuối tháng. Cần BA xác nhận.',
     'AI'),

    ("US35-QA-01.7",
     'Bảng "Xử lý tính phí định kỳ", Bước 5.2',
     'Khi nhiều CTƯĐ cùng Số tiền ƯĐ và cùng hiệu lực → thiếu tiebreak thứ 3. Và "hiệu lực xa nhất" = End Date xa nhất hay Start Date sớm nhất?',
     'Nghiệp vụ',
     'Đề xuất: Cần BA xác nhận "hiệu lực xa nhất" nghĩa là gì + tiebreak cuối cùng khi vẫn bằng nhau.',
     'BOTH'),

    ("US35-QA-01.8",
     'Mục "Yêu cầu nghiệp vụ", trạng thái "Xóa nợ"',
     'Trạng thái "Xóa nợ" được tạo bởi luồng nào? Thủ công hay tự động? US nào mô tả luồng chuyển trạng thái này?',
     'Nghiệp vụ',
     'Đề xuất: Cần BA xác nhận US mô tả luồng xóa nợ để Tester trace.',
     'BOTH'),

    ("US35-QA-01.9",
     'Mục "Yêu cầu nghiệp vụ", bước 4 xử lý kết quả',
     'Khi kết quả thu = "Chưa thanh toán": bước 4 chỉ nêu xử lý cho "thanh toán toàn bộ/một phần". Khi "Chưa thanh toán" → có ghi lịch sử? Trạng thái kỳ nợ cập nhật thế nào?',
     'Nghiệp vụ',
     'Đề xuất: Cần BA xác nhận ProfiX có ghi lịch sử + cập nhật trạng thái khi "Chưa thanh toán".',
     'VA'),

    ("US35-QA-01.10",
     'Bảng "Xử lý tính phí định kỳ", Bước 5.1/5.2',
     'Nếu Số tiền ƯĐ >= Số tiền phí cần thu (ưu đãi 100%+), Phí sau ƯĐ = 0 hoặc âm. ProfiX có gửi giao dịch 0 đồng sang T24 không? Hay bỏ qua và trạng thái nào?',
     'Nghiệp vụ',
     'Đề xuất: Nếu phí sau ƯĐ <= 0 → không gửi T24, trạng thái = "Đã miễn phí" hoặc tương đương. Cần BA xác nhận.',
     'VA'),

    ("US35-QA-01.11",
     'Bảng "Xử lý tính phí định kỳ", Bước 3.4 vs 5.4',
     'Xung đột min/max: Sau áp min CTƯĐ (5.4), phí có thể vượt max Code phí (3.4). Ngược lại cũng vậy. Bộ min/max nào ưu tiên áp dụng cuối cùng?',
     'Nghiệp vụ',
     'Đề xuất: Cần BA xác nhận thứ tự ưu tiên min/max giữa Code phí và CTƯĐ.',
     'VA'),

    ("US35-QA-01.12",
     'Mục "Yêu cầu nghiệp vụ", đoạn truy thu/tận thu',
     'Sau khi thu "Thanh toán một phần", kỳ nợ tiếp theo sinh phí = toàn bộ phí gốc hay chỉ phần còn thiếu? Tài liệu nói "mô tả tại US truy thu/tận thu" nhưng chưa có tham chiếu US cụ thể.',
     'Nghiệp vụ',
     'Đề xuất: Ghi rõ tham chiếu US truy thu/tận thu để QA trace và thiết kế TC liên kết.',
     'VA'),

    ("US35-QA-01.13",
     'Bảng "Xử lý tính phí định kỳ", Bước 6 - VAT',
     '"Số tiền phí sau ưu đãi" dùng tính VAT là giá trị trước hay sau khi áp min/max CTƯĐ (bước 5.4)? Cần xác định đầu vào bước 6.',
     'Nghiệp vụ',
     'Đề xuất: Đầu vào bước 6 = "Số tiền phí thực thu" (sau toàn bộ xử lý ưu đãi + min/max). Cần BA xác nhận.',
     'VA'),

    ("US35-QA-01.14",
     'Mục "Yêu cầu nghiệp vụ", đoạn Job thu phí',
     'Job thu phí chạy vào thời điểm nào trong ngày T (giờ cụ thể)? Nếu ngày T rơi vào nghỉ lễ/cuối tuần → có chạy không hay dời sang ngày làm việc tiếp theo?',
     'Nghiệp vụ',
     'Đề xuất: Cần BA xác nhận thời điểm chạy Job và xử lý ngày nghỉ.',
     'VA'),

    ("US35-QA-01.15",
     'Bảng "Diễn giải lưu đồ", Bước 2 - T24 tận thu',
     'T24 "tận thu (thu một phần nếu TK không đủ số dư)": Nếu số dư = 0 → kết quả = "Thanh toán một phần" (0 đồng) hay "Chưa thanh toán"?',
     'Nghiệp vụ',
     'Đề xuất: Số dư = 0 → "Chưa thanh toán". 0 < số dư < phí → "Thanh toán một phần". Cần BA xác nhận.',
     'AI'),

    # === HM2: Giới hạn ===
    ("US35-QA-02.1",
     'Mục "Yêu cầu nghiệp vụ", ghi danh sách vào Topic Kafka',
     'Nếu Kafka bị down khi ProfiX ghi dữ liệu → xử lý thế nào? Có retry policy (số lần, backoff)? Volume lớn (hàng triệu KH) có batch size limit?',
     'Giới hạn',
     'Đề xuất: Cần định nghĩa retry policy + batch size + trạng thái cuối khi Kafka fail.',
     'BOTH'),

    ("US35-QA-02.2",
     'Bảng "Xử lý tính phí định kỳ", Bước 3.3 - tỷ giá',
     'Tỷ giá quy đổi min/max là của ngày T hay T-1? Tần suất đồng bộ? Nếu job chạy 00:01 nhưng tỷ giá ngày T chưa có → dùng tỷ giá nào? Nếu không có tỷ giá cho loại tiền hiếm → xử lý thế nào?',
     'Giới hạn',
     'Đề xuất: Xác định rõ nguồn/thời điểm tỷ giá + fallback khi không có. Cần BA xác nhận.',
     'BOTH'),

    ("US35-QA-02.3",
     'Bảng "Xử lý tính phí định kỳ", Bước 6 - VAT config',
     'Nếu Code phí KHÔNG cấu hình trường VAT (null/undefined, khác chuỗi rỗng "") → xử lý giống "không có VAT" hay báo lỗi cấu hình?',
     'Giới hạn',
     'Đề xuất: VAT = null → mặc định = không có VAT. Cần BA xác nhận.',
     'AI'),

    ("US35-QA-02.4",
     'Mục "Yêu cầu nghiệp vụ", luồng batch job tổng thể',
     'Job crash giữa chừng (OOM, network) → các khoản đã ghi Kafka bị duplicate khi retry? Có idempotency/deduplication key không?',
     'Giới hạn',
     'Đề xuất: Cần idempotency key (mã khoản phí + ngày T) để tránh thu trùng khi retry. Cần BA/SA xác nhận.',
     'BOTH'),

    ("US35-QA-02.5",
     'Bảng "Xử lý tính phí định kỳ", Bước 3.3-3.4',
     'Nếu Code phí chỉ khai báo 1 trong 2 (chỉ Min hoặc chỉ Max) → logic clamping xử lý thế nào?',
     'Giới hạn',
     'Đề xuất: Chỉ có Min → chỉ so Min. Chỉ có Max → chỉ so Max. Tương tự bước 5.3-5.4. Cần BA xác nhận.',
     'AI'),

    ("US35-QA-02.6",
     'Bảng "Xử lý tính phí định kỳ", Bước 3.3/3.4 - cấu hình sai',
     'Nếu cấu hình Code phí có Số tiền tối thiểu > Số tiền tối đa (dữ liệu sai) → hệ thống phát hiện/xử lý thế nào tại runtime?',
     'Giới hạn',
     'Đề xuất: Validate Min <= Max tại thời điểm cài đặt Code phí (không phải runtime). Cần BA xác nhận.',
     'VA'),

    # === HM3: Toàn vẹn dữ liệu ===
    ("US35-QA-03.1",
     'Mục "Yêu cầu nghiệp vụ", trạng thái TK thu phí',
     'Dấu "…" trong danh sách trạng thái TK được phép thu phí ám chỉ còn trạng thái nào khác? Danh sách đầy đủ là gì? Tham số hệ thống nào quy định?',
     'Toàn vẹn dữ liệu',
     'Đề xuất: Liệt kê đầy đủ trạng thái TK được phép thu phí hoặc tham chiếu tham số hệ thống.',
     'VA'),

    ("US35-QA-03.2",
     'Mục "Yêu cầu nghiệp vụ", đoạn đồng bộ T-1',
     'Nếu trong khoảng T-1→T, TK bị đóng hoặc KH bị vô hiệu trên Core nhưng dữ liệu ProfiX chưa cập nhật → ProfiX vẫn sinh phí cho TK/KH không còn hợp lệ. T24 xử lý thế nào?',
     'Toàn vẹn dữ liệu',
     'Đề xuất: T24 trả "Chưa thanh toán" + lý do TK không hợp lệ. ProfiX cập nhật trạng thái. Cần BA xác nhận.',
     'AI'),

    ("US35-QA-03.3",
     'Bảng "Xử lý tính phí định kỳ", Bước 2 - TK thay thế',
     'TK thay thế chỉ dùng cho phí đúng hạn, không dùng cho truy thu/tận thu. Trong lịch sử thu phí, trường "TK thu phí" lưu TK mặc định hay TK thay thế thực tế? Nợ phí theo dõi trên TK nào?',
     'Toàn vẹn dữ liệu',
     'Đề xuất: Lưu cả 2 (TK mặc định + TK thực tế) cho kiểm toán. Nợ phí theo dõi trên TK mặc định. Cần BA xác nhận.',
     'BOTH'),

    ("US35-QA-03.4",
     'Mục "Yêu cầu nghiệp vụ", trạng thái "Thêm mới"/"Đang xử lý"',
     'Job chạy lại (retry/reschedule) → có risk sinh trùng khoản phí? Khoản ở "Đang xử lý" → Job chạy lại có bỏ qua không?',
     'Toàn vẹn dữ liệu',
     'Đề xuất: Khi Job retry → check cùng Code phí + KH/TK/Thẻ + kỳ thu → bỏ qua, không sinh trùng. Cần BA xác nhận.',
     'AI'),

    ("US35-QA-03.5",
     'Bảng "Xử lý tính phí định kỳ", Bước 4 - ngưỡng ƯĐ',
     'Điều kiện "KH chưa chạm ngưỡng ưu đãi" kiểm tra theo tiêu chí nào (số lần, tổng tiền, theo kỳ)? Nếu đã chạm → CTƯĐ không vào danh sách D hay vào nhưng ƯĐ = 0? Race condition trong cùng batch?',
     'Toàn vẹn dữ liệu',
     'Đề xuất: Làm rõ logic ngưỡng + kiểm tra snapshot đầu Job (không cập nhật giữa batch). Cần BA xác nhận.',
     'BOTH'),

    # === HM4: UI/UX ===
    ("US35-QA-04.1",
     'Phần "Giao diện" và "Mô tả chi tiết các trường" ghi N/A',
     'US35 backend thuần tuý, không có UI. Nhưng cần theo dõi kết quả: Lịch sử chạy Job, danh sách khoản phí, kết quả thu. Hiển thị ở US nào?',
     'UI-UX',
     'Đề xuất: Tham chiếu rõ US quản lý kết quả (Lịch sử thu phí, Quản lý nợ phí) để QA biết phạm vi E2E.',
     'BOTH'),

    ("US35-QA-04.2",
     'Mục "Lưu đồ" - Flowchart (image1.png)',
     'Flowchart chỉ 4 bước high-level, không có nhánh xử lý lỗi (Kafka fail, T24 timeout, không tìm TK).',
     'UI-UX',
     'Đề xuất: Ghi nhận. Tester đọc bảng mô tả chi tiết, không dựa flowchart. Nếu có thể, bổ sung nhánh lỗi.',
     'BOTH'),
]

# Build table
doc.add_heading('Bảng Q&A Tổng Hợp', level=2)
p = doc.add_paragraph()
p.add_run('Tổng: ').bold = True
p.add_run(f'{len(qa_data)} câu hỏi | ')
ai_only = sum(1 for q in qa_data if q[5] == 'AI')
va_only = sum(1 for q in qa_data if q[5] == 'VA')
both = sum(1 for q in qa_data if q[5] == 'BOTH')
p.add_run(f'AI only: {ai_only} | VA only: {va_only} | BOTH: {both}')

table = doc.add_table(rows=1, cols=7)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

col_widths = [Cm(2.2), Cm(3.5), Cm(7.0), Cm(2.0), Cm(5.5), Cm(4.0), Cm(1.5)]
for i, w in enumerate(col_widths):
    table.columns[i].width = w

headers = ['ID', 'Trích xuất', 'Câu hỏi / Sự cố', 'Phân loại', 'Đề xuất từ QA', 'Trả lời của BA', 'Nguồn']
hdr = table.rows[0].cells
for i, h in enumerate(headers):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(8)

category_map = {
    '01': '🔶 Hạng mục 1: Vấn đề Nghiệp vụ / Luồng xử lý',
    '02': '🔴 Hạng mục 2: Giới hạn hệ thống & Exception Handling',
    '03': '🟠 Hạng mục 3: Toàn vẹn dữ liệu & Ràng buộc',
    '04': '🔵 Hạng mục 4: UI/UX & Giao diện',
}

current_cat = None
for qa in qa_data:
    cat = qa[0].split('-QA-')[1][:2]
    if cat != current_cat:
        current_cat = cat
        crow = table.add_row().cells
        crow[0].merge(crow[6])
        crow[0].text = category_map.get(cat, '')
        for p in crow[0].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0, 0, 128)

    row = table.add_row().cells
    for i, val in enumerate(qa):
        row[i].text = val
        for p in row[i].paragraphs:
            for r in p.runs:
                r.font.size = Pt(8)

for r in table.rows:
    for i, c in enumerate(r.cells):
        c.width = col_widths[i]

# Summary
doc.add_paragraph('')
doc.add_heading('Thống kê', level=2)
doc.add_paragraph(f'• Hạng mục 1 (Nghiệp vụ): 15 câu')
doc.add_paragraph(f'• Hạng mục 2 (Giới hạn): 6 câu')
doc.add_paragraph(f'• Hạng mục 3 (Toàn vẹn dữ liệu): 5 câu')
doc.add_paragraph(f'• Hạng mục 4 (UI/UX): 2 câu')
doc.add_paragraph(f'• TỔNG: {len(qa_data)} câu')

out = '/Users/mac/antigravity-testing-kit/Feature_02_SA_Tham_So_He_Thong/output/US35_PartB_QA_Merged.docx'
os.makedirs(os.path.dirname(out), exist_ok=True)
doc.save(out)
print(f'✅ Saved: {out}')
