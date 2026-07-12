"""US31 - Part C: Generate Test Case Coverage DOCX"""
import sys, os
sys.path.insert(0, os.path.dirname(__file__))
from us31_partc_data import TC_DATA

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Page setup: Landscape, narrow margins
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

# Title
title = doc.add_heading("US31 – Báo Cáo Tổng Doanh Thu Phí", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_heading("PHẦN C: BẢNG TỔNG HỢP TEST CASE ĐỀ XUẤT (Test Case Coverage)", level=2)

# Summary
p = doc.add_paragraph()
p.add_run(f"Tổng số Test Case: {len(TC_DATA)}").bold = True

# Category summary
from collections import Counter
cat_counts = Counter(tc[3] for tc in TC_DATA)
cat_order = ["Happy Path", "Negative Path", "Boundary Value",
             "UI/UX & Field Validation", "Business Logic", "Data Integrity", "NFR"]
cat_icons = {"Happy Path": "🟢", "Negative Path": "🔴", "Boundary Value": "📐",
             "UI/UX & Field Validation": "🎨", "Business Logic": "🧠",
             "Data Integrity": "🔗", "NFR": "⚡"}

for cat in cat_order:
    count = cat_counts.get(cat, 0)
    if count > 0:
        icon = cat_icons.get(cat, "")
        doc.add_paragraph(f"{icon} {cat}: {count} TC", style="List Bullet")

doc.add_paragraph()  # spacer

# Table
HEADERS = ["Mã SC", "Feature", "Module", "Loại Test Case", "Tên Test Case / Kịch bản", "Trích dẫn tài liệu (Traceability)"]
COL_WIDTHS = [Cm(1.5), Cm(3.5), Cm(2.5), Cm(2.8), Cm(9.0), Cm(8.0)]

table = doc.add_table(rows=1 + len(TC_DATA), cols=6, style="Table Grid")
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False

# Set widths
for row in table.rows:
    for i, w in enumerate(COL_WIDTHS):
        row.cells[i].width = w

# Header
for i, h in enumerate(HEADERS):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(8.5)

# Data rows
for idx, (sc_id, feature, module, category, tc_title, trace) in enumerate(TC_DATA):
    row = table.rows[idx + 1]
    data = [sc_id, feature, module, category, tc_title, trace]
    for i, val in enumerate(data):
        cell = row.cells[i]
        cell.text = val
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(8)

out_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "US31_PartC_TestCoverage.docx")
doc.save(out_path)
print(f"✅ Đã tạo: {out_path}")
print(f"📊 Tổng số TC: {len(TC_DATA)}")
for cat in cat_order:
    c = cat_counts.get(cat, 0)
    if c > 0:
        print(f"   {cat_icons.get(cat,'')} {cat}: {c}")
