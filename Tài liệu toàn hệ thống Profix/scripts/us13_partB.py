"""US13 Part B - Q&A / Loopholes Discovery Generator"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# --- Page Setup: Landscape + Narrow margins ---
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10)

# === TITLE ===
doc.add_heading('US13 — Phần B: Danh Sách Cảnh Báo & Q&A (Dành Cho BA)', level=0)
doc.add_paragraph('Feature: Khai báo CTƯĐ không xác định sẵn danh sách khách hàng áp dụng')
doc.add_paragraph('Phiên bản: v1.1 | Ngày phân tích: 13/05/2026')
doc.add_paragraph('')

# === Q&A TABLE ===
headers = ['ID', 'Hạng mục', 'Trích xuất (Reference)', 'Câu hỏi / Sự cố', 'Phân loại', 'Đề xuất từ QA', 'Trả lời của BA']

# Column widths in cm
col_widths = [Cm(2.2), Cm(2), Cm(4), Cm(5.5), Cm(2), Cm(5.5), Cm(5.5)]

qa_data = [
    # === Hạng mục 1: Nghiệp vụ / Luồng xử lý ===
    [
        'US13-QA-01.1',
        '🔶 Nghiệp vụ',
        'Mục "Thêm mới CTƯĐ không xác định sẵn danh sách KH", đoạn "Các Thông tin chung, Chu kỳ áp dụng, Điều kiện áp dụng tương tự với yêu cầu tại US12"',
        'US13 tham chiếu toàn bộ logic US12 cho 3 phân vùng (Thông tin chung, Chu kỳ, Điều kiện). Tuy nhiên US12 có trường "Loại ưu đãi" (Theo Khách hàng / Theo Tài khoản / Theo Thẻ) — trường này có ý nghĩa gì khi CTƯĐ KHÔNG có danh sách KH cụ thể? Giá trị "Loại ưu đãi" ảnh hưởng đến logic xử lý nào của hệ thống trong trường hợp US13?',
        'Nghiệp vụ',
        'Đề xuất BA làm rõ: Trường "Loại ưu đãi" trong US13 có cùng danh sách giá trị và hành vi như US12 không? Nếu có, mô tả tác động nghiệp vụ cụ thể (VD: hệ thống dùng giá trị này để quyết định scope áp dụng ưu đãi tại thời điểm giao dịch).',
        ''
    ],
    [
        'US13-QA-01.2',
        '🔶 Nghiệp vụ',
        'Mục "Chỉnh sửa CTƯĐ", đoạn "Nếu trạng thái hoạt động CTƯĐ = Đang hiệu lực, hệ thống chỉ cho phép chỉnh sửa"',
        'Khi trạng thái = Đang hiệu lực, US13 nói chỉ cho phép chỉnh sửa "Ngày hết hiệu lực" và "Chi tiết ưu đãi" (Thêm mới/Xóa bản ghi). Câu hỏi: Bản ghi Chi tiết ưu đãi hiện hữu (đã tồn tại) có được phép CHỈNH SỬA nội dung (VD: thay đổi Giá trị ưu đãi, Tỷ lệ ưu đãi, Loại tiền...) hay chỉ cho phép Thêm mới và Xóa?',
        'Nghiệp vụ',
        'Theo text US13, chỉ đề cập "Thêm mới" và "Xóa" bản ghi → QA hiểu là KHÔNG cho phép chỉnh sửa bản ghi hiện hữu. Đề xuất BA xác nhận.',
        ''
    ],
    [
        'US13-QA-01.3',
        '🔶 Nghiệp vụ',
        'Mục "Chỉnh sửa CTƯĐ", đoạn "Nếu trạng thái hoạt động CTƯĐ = Hết hiệu lực/Chưa hiệu lực, hệ thống xử lý tương tự US12"',
        'Tài liệu ghi: "Nếu trạng thái hoạt động CTƯĐ = Hết hiệu lực/Chưa hiệu lực, hệ thống xử lý tương tự US12". Tuy nhiên, trong US12, trạng thái "Hết hiệu lực" sẽ hiển thị thông báo lỗi "Trạng thái CTƯĐ không hợp lệ" và không cho chỉnh sửa. Vậy US13 khi gom "Hết hiệu lực/Chưa hiệu lực" vào cùng một nhánh, có phải ý nghĩa là: Hết hiệu lực → báo lỗi, Chưa hiệu lực → cho sửa toàn bộ (tương tự US12)?',
        'Nghiệp vụ',
        'Đề xuất BA xác nhận: logic xử lý cho từng trạng thái giống US12 (Hết hiệu lực = báo lỗi kết thúc, Chưa hiệu lực = cho sửa toàn bộ trừ Mã CTƯĐ và Số VB_Tên viết tắt).',
        ''
    ],
    [
        'US13-QA-01.4',
        '🔶 Nghiệp vụ',
        'Lưu đồ Thêm mới, Bước 7 "Tham chiếu tới luồng khai báo các Thông tin chung, Điều kiện đánh giá và Lịch ưu đãi tại US12"',
        'Lưu đồ Bước 7 ghi "Điều kiện đánh giá" và "Lịch ưu đãi" nhưng trong phần Yêu cầu nghiệp vụ (text) gọi là "Điều kiện áp dụng" và "Chu kỳ áp dụng". Đây là sự không nhất quán về tên gọi giữa Flowchart và Text. Tên chính thức nào sẽ được sử dụng?',
        'Nghiệp vụ',
        'Đề xuất: Thống nhất sử dụng tên "Điều kiện áp dụng" và "Chu kỳ áp dụng" theo text yêu cầu nghiệp vụ. Flowchart cần cập nhật lại cho nhất quán.',
        ''
    ],
    [
        'US13-QA-01.5',
        '🔶 Nghiệp vụ',
        'Mục "Yêu cầu nghiệp vụ", đoạn "Hệ thống chỉ cho phép chỉnh sửa khi CTƯĐ có trạng thái hoạt động = Chưa hiệu lực/Đang hiệu lực"',
        'Tài liệu mô tả điều kiện cho phép chỉnh sửa nhưng KHÔNG đề cập đến việc kiểm tra có tác vụ "Chờ duyệt" hay không trước khi cho phép chỉnh sửa. Trong US12, BE kiểm tra nếu đã có tác vụ Chờ duyệt → không cho chỉnh sửa. US13 có áp dụng cùng logic này không?',
        'Nghiệp vụ',
        'Đề xuất: Áp dụng tương tự US12 — nếu đã tồn tại tác vụ "Chờ duyệt" cho CTƯĐ này → hiển thị lỗi, không cho chỉnh sửa.',
        ''
    ],
    [
        'US13-QA-01.6',
        '🔶 Nghiệp vụ',
        'Mục "Navigation" và Lưu đồ Thêm mới Bước 2',
        'Navigation chỉ đề cập truy cập chức năng "Ưu đãi không đánh giá định kỳ" để khai báo. Tuy nhiên Flowchart (Bước 2) hiển thị 2 lựa chọn: Ưu đãi CÓ định kỳ và KHÔNG định kỳ. Câu hỏi: Nếu người dùng muốn tạo CTƯĐ "CÓ đánh giá định kỳ" nhưng "KHÔNG theo danh sách khách hàng" thì luồng này thuộc phạm vi US13 hay US khác? Nếu thuộc US13 thì Mockup/Text chưa mô tả.',
        'Nghiệp vụ',
        'Đề xuất BA xác nhận: Phạm vi US13 chỉ bao gồm "Ưu đãi KHÔNG đánh giá định kỳ". Nếu user chọn "Có đánh giá định kỳ" → rẽ nhánh sang US khác (VD: US11).',
        ''
    ],

    # === Hạng mục 2: Giới hạn hệ thống & Exception ===
    [
        'US13-QA-02.1',
        '🔴 Giới hạn',
        'Mục "Chi tiết ưu đãi", Mockup liên tục (image2.png)',
        'Mockup liên tục hiển thị 6 dòng SPDV với nút "+ Thêm SPDV áp dụng". Có giới hạn tối đa số lượng SPDV có thể thêm trong 1 CTƯĐ không? Nếu có, khi đạt giới hạn, hệ thống xử lý như thế nào (ẩn nút thêm, báo lỗi)?',
        'Giới hạn',
        'Đề xuất BA xác nhận giới hạn tối đa (nếu có). Nếu không giới hạn, cần cân nhắc hiệu năng khi danh sách quá dài.',
        ''
    ],
    [
        'US13-QA-02.2',
        '🔴 Giới hạn',
        'Mục "Chi tiết ưu đãi", tham chiếu US12 về "Ngày hiệu lực" và "Ngày hết hiệu lực" tại dòng Chi tiết ưu đãi',
        'Theo US12, mỗi dòng Chi tiết ưu đãi có Ngày hiệu lực và Ngày hết hiệu lực riêng (chỉ hiển thị khi chọn ưu đãi Liên tục). US13 không đề cập trường này nhưng Mockup liên tục (image2.png) cũng KHÔNG hiển thị cột Ngày hiệu lực / Ngày hết hiệu lực trong bảng Chi tiết ưu đãi. Mockup chu kỳ (image3.png) cũng không. Vậy US13 có bao gồm Ngày hiệu lực / Ngày hết hiệu lực ở cấp dòng Chi tiết ưu đãi không?',
        'Giới hạn',
        'Đề xuất BA xác nhận: (a) US13 có kế thừa cột Ngày hiệu lực / Ngày hết hiệu lực từ US12 tại bảng Chi tiết ưu đãi không? (b) Nếu không, ưu đãi cấp dòng SPDV sẽ tuân theo ngày hiệu lực cấp CTƯĐ?',
        ''
    ],

    # === Hạng mục 3: Toàn vẹn dữ liệu ===
    [
        'US13-QA-03.1',
        '🟠 Toàn vẹn DL',
        'Mục "Chi tiết ưu đãi", tham chiếu US12 về ràng buộc SPDV cha-con',
        'US12 quy định: Không được chọn các SPDV có quan hệ cha-con trong cùng CTƯĐ. Trong US13 (không có CIF), ràng buộc cha-con SPDV vẫn áp dụng tại cấp toàn bộ CTƯĐ? Cụ thể: nếu dòng 1 chọn SPDV cha và dòng 2 chọn SPDV con (cùng nhánh) → hệ thống có chặn không?',
        'Toàn vẹn dữ liệu',
        'Đề xuất: Ràng buộc cha-con SPDV áp dụng tương tự US12 — không cho phép chọn SPDV có quan hệ cha-con trong cùng danh sách Chi tiết ưu đãi.',
        ''
    ],
    [
        'US13-QA-03.2',
        '🟠 Toàn vẹn DL',
        'Mục "Chi tiết ưu đãi", tham chiếu US12 về trùng lặp SPDV',
        'US12 quy định không cho phép CIF+SPDV trùng nhau. Trong US13 (không có CIF), ràng buộc unique SPDV là gì? Cụ thể: Có được phép khai báo cùng 1 SPDV trên 2 dòng khác nhau trong bảng Chi tiết ưu đãi không? (VD: cùng SPDV nhưng Giá trị ưu đãi khác nhau)',
        'Toàn vẹn dữ liệu',
        'Đề xuất: SPDV phải unique trong danh sách Chi tiết ưu đãi của 1 CTƯĐ. Nếu user chọn SPDV đã tồn tại → FE chặn hoặc hiển thị lỗi.',
        ''
    ],
    [
        'US13-QA-03.3',
        '🟠 Toàn vẹn DL',
        'Mục "Chỉnh sửa CTƯĐ", đoạn "Xóa một hoặc nhiều bản ghi Chi tiết ưu đãi đã được khai báo"',
        'Khi chỉnh sửa CTƯĐ Đang hiệu lực, người dùng được phép xóa bản ghi Chi tiết ưu đãi. Nếu người dùng xóa TẤT CẢ bản ghi Chi tiết ưu đãi hiện có (danh sách rỗng) rồi nhấn Xác nhận — hệ thống xử lý thế nào? CTƯĐ không có bất kỳ SPDV ưu đãi nào có hợp lệ không?',
        'Toàn vẹn dữ liệu',
        'Đề xuất: Bảng Chi tiết ưu đãi phải có ít nhất 1 bản ghi. Nếu user xóa hết → FE chặn (disable Xác nhận) hoặc hiển thị cảnh báo.',
        ''
    ],

    # === Hạng mục 4: UI/UX ===
    [
        'US13-QA-04.1',
        '🔵 UI/UX',
        'Mockup Thêm mới liên tục (image2.png), phân vùng "Chi tiết ưu đãi"',
        'Mockup liên tục hiển thị nút "Tải lên" (Upload) và "Tải xuống" (Download) tại phân vùng Chi tiết ưu đãi. Tuy nhiên, US13 tham chiếu text nói "tương tự US12 nhưng không bao gồm trường CIF và Tên". Câu hỏi: Nút "Tải lên" có hiển thị trên US13 không? Vì không có CIF, việc Upload danh sách SPDV (không gắn KH) có logic khác với US12.',
        'UI-UX',
        'Đề xuất: Nút "Tải lên" KHÔNG hiển thị (vì không có CIF/danh sách KH). Nút "Tải xuống" hiển thị để export danh sách SPDV đã khai báo. BA xác nhận.',
        ''
    ],
    [
        'US13-QA-04.2',
        '🔵 UI/UX',
        'Mockup Thêm mới liên tục (image2.png) vs Mockup theo chu kỳ (image3.png), phân vùng "Chi tiết ưu đãi"',
        'So sánh 2 Mockup: Mockup liên tục có cả nút "Tải lên" + "Tải xuống", nhưng Mockup theo chu kỳ CHỈ có nút "Tải xuống" (không có "Tải lên"). Sự khác biệt này là cố ý (logic khác nhau giữa 2 chế độ chu kỳ) hay lỗi Mockup?',
        'UI-UX',
        'Đề xuất BA xác nhận: Nút "Tải lên" có/không hiển thị phải nhất quán giữa 2 chế độ (Liên tục vs Theo chu kỳ). QA nghiêng về: KHÔNG có "Tải lên" ở cả 2.',
        ''
    ],
    [
        'US13-QA-04.3',
        '🔵 UI/UX',
        'Mockup Thêm mới liên tục (image2.png), toggle "Ưu đãi theo tỷ lệ"',
        'Mockup liên tục hiển thị toggle "Ưu đãi theo tỷ lệ" = ON (xanh). Khi ON, bảng Chi tiết ưu đãi hiển thị cột "Giá trị ưu đãi" (không thấy cột "Tỷ lệ ưu đãi" rõ ràng). Theo US12, khi toggle ON → hiển thị "Tỷ lệ ưu đãi" (%), ẩn "Giá trị ưu đãi"; khi OFF → hiển thị "Giá trị ưu đãi", ẩn "Tỷ lệ ưu đãi". Cột trên Mockup ghi "Giá trị ưu đãi" nhưng toggle đang ON — đây có phải lỗi label trên Mockup?',
        'UI-UX',
        'Đề xuất: Đây khả năng cao là lỗi label trên Mockup. Khi toggle ON → cột phải là "Tỷ lệ ưu đãi (%)". BA xác nhận.',
        ''
    ],
    [
        'US13-QA-04.4',
        '🔵 UI/UX',
        'Mô tả chi tiết các trường, đoạn "tương tự tại US12 nhưng không bao gồm trường CIF và Tên tại Chi tiết ưu đãi"',
        'US13 không cung cấp bảng mô tả chi tiết các trường riêng mà tham chiếu hoàn toàn US12. Với việc loại bỏ CIF/Tên, trường "Thêm mới khách hàng" (Button) trong US12 cũng không còn. Vậy nút "+ Thêm SPDV áp dụng" (hiển thị trên Mockup) tương đương với chức năng gì? Nó thêm 1 dòng trống vào bảng Chi tiết ưu đãi để user chọn SPDV?',
        'UI-UX',
        'Đề xuất: Nút "+ Thêm SPDV áp dụng" thêm 1 dòng trống vào bảng Chi tiết ưu đãi, cho phép user chọn SPDV từ Combobox và nhập các trường liên quan.',
        ''
    ],
    [
        'US13-QA-04.5',
        '🔵 UI/UX',
        'Mô tả chi tiết các trường, tham chiếu US12 trường Giá trị ưu đãi',
        'Khi Toggle "Ưu đãi theo tỷ lệ" = OFF, cột hiển thị là "Giá trị ưu đãi" (số tiền cố định, VD: 50.000). Tuy nhiên, trên UI không có trường "Loại tiền" để gắn với Giá trị ưu đãi này (VD: VND hay USD). Cột "Loại tiền tối thiểu/tối đa" chỉ phục vụ cho 2 cột số tiền phí tối thiểu/tối đa. Vậy Giá trị ưu đãi dùng loại tiền nào?',
        'UI-UX',
        'Đề xuất: Bổ sung dropdown "Loại tiền" đi kèm với cột "Giá trị ưu đãi", hoặc ngầm định sử dụng Loại tiền tệ chung của giao dịch. BA làm rõ logic.',
        ''
    ],

    # === Bổ sung từ VA Review ===
    [
        'US13-QA-01.7',
        '🔶 Nghiệp vụ',
        'Mục "Chi tiết ưu đãi", đoạn: "không bao gồm trường CIF và Tên"',
        'Khi không có danh sách KH cứng, hệ thống xác định KH được hưởng ưu đãi bằng cách nào? Cơ chế là: (1) Quét định kỳ toàn bộ KH khớp Điều kiện áp dụng để gán ưu đãi, hay (2) Đánh giá real-time tại thời điểm phát sinh giao dịch của KH?',
        'Nghiệp vụ',
        'Đề xuất BA làm rõ cơ chế hệ thống áp dụng ưu đãi cho KH để Tester thiết kế Test data tương ứng.',
        ''
    ],
    [
        'US13-QA-02.3',
        '🔴 Giới hạn',
        'Mục "Chỉnh sửa CTƯĐ", thao tác "Xóa bản ghi Chi tiết ưu đãi"',
        'Trong trường hợp CTƯĐ Đang hiệu lực, Maker xóa 1 SPDV khỏi Chi tiết ưu đãi và được Checker duyệt. Nếu tại thời điểm duyệt, hệ thống đang xử lý dở dang các giao dịch của SPDV đó (batch/job), thì ưu đãi của các giao dịch đó xử lý thế nào?',
        'Giới hạn',
        'Đề xuất: Các giao dịch phát sinh trước thời điểm duyệt vẫn hưởng ưu đãi; giao dịch sau thời điểm duyệt không được hưởng.',
        ''
    ],
    [
        'US13-QA-03.4',
        '🟠 Toàn vẹn DL',
        'Mục "Yêu cầu nghiệp vụ", Mã CTƯĐ',
        'Mã CTƯĐ của US13 có được tự động sinh theo cùng quy tắc với US12 không? Có sử dụng prefix (tiền tố) riêng biệt nào để nhận diện đây là loại CTƯĐ "không xác định KH" so với loại "có danh sách KH" không?',
        'Toàn vẹn dữ liệu',
        'Đề xuất: Dùng chung 1 sequence tự sinh nhưng có thể khác Prefix để dễ nhận diện. BA xác nhận.',
        ''
    ],
    [
        'US13-QA-04.6',
        '🔵 UI/UX',
        'Mockup Thiết lập chương trình (image2)',
        'Trường "Theo chu kỳ" và "Liên tục" đang được thể hiện dưới dạng Checkbox (ô vuông) trên Mockup, nhưng về logic 2 lựa chọn này loại trừ lẫn nhau (chỉ được chọn 1).',
        'UI-UX',
        'Đề xuất: Chuyển 2 control này thành Radio Button (ô tròn) để đúng chuẩn UI/UX và logic. BA xác nhận.',
        ''
    ],
    [
        'US13-QA-04.7',
        '🔵 UI/UX',
        'Mockup Thiết lập chương trình (image2)',
        'Trường "Ngày ban hành" trên Mockup đang không có dấu (*) bắt buộc. Theo nghiệp vụ, ngày ban hành văn bản/chương trình thường là thông tin trọng yếu. Trường này có bắt buộc không?',
        'UI-UX',
        'Đề xuất: Trường Ngày ban hành BẮT BUỘC nhập (có dấu *). BA xác nhận.',
        ''
    ],
    [
        'US13-QA-04.8',
        '🔵 UI/UX',
        'Mockup Thiết lập chương trình (image2)',
        'Trường "Mã CTƯĐ" trên Mockup hiển thị dấu (*) bắt buộc nhập. Tuy nhiên, theo quy tắc hệ thống, Mã CTƯĐ được tự động sinh SAU khi Checker duyệt (QTC-12).',
        'UI-UX',
        'Đề xuất: Trường Mã CTƯĐ phải ở trạng thái Disabled (◎) ngay từ màn hình Thêm mới, không cho user nhập thủ công và không có dấu (*).',
        ''
    ]
]

# Create table
table = doc.add_table(rows=len(qa_data)+1, cols=len(headers), style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Set headers
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9)

# Fill data
for row_idx, row_data in enumerate(qa_data):
    for col_idx, val in enumerate(row_data):
        cell = table.rows[row_idx+1].cells[col_idx]
        cell.text = val
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

# Set column widths
for row in table.rows:
    for i, cell in enumerate(row.cells):
        cell.width = col_widths[i]

# === Save ===
out_dir = '/Users/mac/antigravity-testing-kit/Feature_02_SA_Tham_So_He_Thong/US11-20'
os.makedirs(out_dir, exist_ok=True)
out_path = os.path.join(out_dir, 'US13_PartB_QA.docx')
doc.save(out_path)
print(f"✅ Saved: {out_path}")
