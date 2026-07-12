"""US35 Part A - Tóm tắt Nghiệp vụ Chuyên Sâu (Dành cho Tester)"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# Page setup: Landscape, narrow margins
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# ============ TITLE ============
title = doc.add_heading('US35 – Phần A: Tóm Tắt Nghiệp Vụ Chuyên Sâu', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Tính năng: Tự động thu các loại phí định kỳ cần thu theo lịch')
doc.add_paragraph('Dành cho: Manual Tester | Phiên bản: v1.0')
doc.add_paragraph('')

# ============ A.1. CORE BUSINESS VALUE ============
doc.add_heading('A.1. Thông Điệp Cốt Lõi (Core Business Value)', level=2)
p = doc.add_paragraph()
p.add_run('Mục đích:').bold = True
p.add_run(' Hệ thống ProfiX tự động sinh dữ liệu các khoản phí định kỳ đến hạn cần thu, gửi yêu cầu thu phí tới Core Banking T24 qua Kafka, và cập nhật kết quả thu phí trở lại ProfiX.')
doc.add_paragraph('Người dùng cuối: Hệ thống (batch job tự động) — không có giao diện người dùng (UI = N/A).')
doc.add_paragraph('Đặc điểm quan trọng: Đây là tính năng backend thuần tuý (batch processing), không có màn hình khai báo hay tương tác trực tiếp từ User. Toàn bộ logic chạy tự động đầu ngày theo lịch đã cài đặt.')

# ============ A.2. FLOW STRUCTURE & MODULE MAPPING ============
doc.add_heading('A.2. Cấu Trúc Luồng Nghiệp Vụ & Phân Bổ Module', level=2)

doc.add_heading('Module 1: Sinh dữ liệu phí định kỳ đến hạn (ProfiX)', level=3)
doc.add_paragraph('Luồng chính (Happy Path):', style='List Bullet')
items = [
    'Đầu ngày T, hệ thống xác định các Job có lịch chạy trong ngày T.',
    'Với mỗi Job A, xác định Danh sách Code phí B gắn với Job đó.',
    'Bước 1 — Xác định đối tượng tính phí & Code phí áp dụng (Danh sách C):',
    '   • Nếu đối tượng = Customer → đối chiếu bảng Customer với Danh sách Code phí B.',
    '   • Nếu đối tượng = Account → đối chiếu bảng Account + CIF của Account.',
    '   • Nếu đối tượng = Card → đối chiếu bảng Card + CIF của Card.',
    'Bước 2 — Kiểm tra TK thu phí mặc định (chỉ khi đối tượng = Customer):',
    '   • Điều kiện: TK cùng loại tiền Code phí, sản phẩm TK thuộc CA_PRODUCT, trạng thái TK hợp lệ (Hoạt động/Tạm ngừng/Tạm khóa ghi có…).',
    '   • Nếu TK mặc định không đủ ĐK → tìm TK thay thế: ưu tiên số dư lớn nhất, cùng số dư thì lấy ngẫu nhiên.',
    '   • LƯU Ý: TK thay thế chỉ dùng cho phí định kỳ đúng hạn, KHÔNG dùng cho truy thu/tận thu.',
    'Bước 3 — Xác định số tiền phí cần thu cho mỗi khoản trong Danh sách C:',
    '   3.1. Xác định công thức tính phí (nếu Code phí có "Khai báo theo Nhóm KH" = Có → dựa trên Nhóm KH của CIF).',
    '   3.2. Tính Số tiền phí theo công thức đã thiết lập.',
    '   3.3. Quy đổi tỷ giá cho Số tiền tối thiểu/tối đa (nếu Code phí có khai báo):',
    '       — Nếu Loại tiền phí tối thiểu/tối đa = Loại tiền Code phí → không quy đổi.',
    '       — Nếu khác Loại tiền và Code phí = VND → dùng Tỷ giá bán giao ngay (T).',
    '       — Nếu khác Loại tiền và Code phí ≠ VND → dùng Tỷ giá chéo T3 = T1/T2.',
    '   3.4. So sánh clamping: Min ≤ Số tiền phí ≤ Max → xác định Số tiền phí cần thu.',
    'Bước 4 — Xác định CTƯĐ áp dụng (Danh sách D):',
    '   • TH1: CTƯĐ không đánh giá định kỳ/liên tục → kiểm tra SPDV cấp cuối + ĐK KH + ĐK TK/Thẻ + Danh sách KH (chưa chạm ngưỡng).',
    '   • TH2: CTƯĐ có đánh giá định kỳ → kiểm tra SPDV cấp cuối + Danh sách KH (đã xác định theo chu kỳ đánh giá).',
    'Bước 5 — Tính toán mức ưu đãi:',
    '   5.1. Tính: Ưu đãi theo tỷ lệ = Yes → Số tiền ƯĐ = Tỷ lệ × Phí cần thu. Nếu No → Số tiền ƯĐ = Số tiền giảm.',
    '   5.2. Nếu nhiều CTƯĐ → lấy CTƯĐ có Số tiền ƯĐ lớn nhất, bằng nhau thì lấy hiệu lực xa nhất (CTƯĐ P).',
    '   5.3. Clamping theo CTƯĐ: tương tự bước 3.3-3.4 nhưng dùng Min/Max từ bản ghi ưu đãi theo SPDV cấp cuối.',
    '   5.4. So sánh: Min ≤ Phí sau ƯĐ ≤ Max → xác định Số tiền phí thực thu.',
    'Bước 6 — Tính VAT:',
    '   • Code phí không có VAT → để trống VAT.',
    '   • Code phí có VAT + "Phí đã bao gồm VAT" = Có → VAT = Phí sau ƯĐ / 110 × 10.',
    '   • Code phí có VAT + "Phí đã bao gồm VAT" = Không → VAT = Phí sau ƯĐ / 100 × 10.',
    '   • Làm tròn: VND/JPY → số nguyên; các loại tiền khác → 2 chữ số thập phân.',
    'Ghi danh sách khoản phí vào Topic Kafka theo thứ tự ưu tiên nghiệp vụ đã cài đặt.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('')
doc.add_heading('Module 2: Hạch toán thu phí (Core T24)', level=3)
items2 = [
    'Core T24 đọc Topic thu phí định kỳ từ Kafka.',
    'Thực hiện hạch toán thu phí — lưu ý tận thu (thu một phần nếu TK không đủ số dư).',
    'Ghi kết quả thu phí vào Topic kết quả thu phí.',
]
for item in items2:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('')
doc.add_heading('Module 3: Đọc message & Update kết quả (ProfiX)', level=3)
items3 = [
    'ProfiX đọc message từ Topic kết quả thu phí.',
    'Ghi nhận lịch sử thu phí theo kết quả.',
    'Cập nhật trạng thái kỳ nợ phí tương ứng.',
]
for item in items3:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('')
doc.add_paragraph('Luồng rẽ nhánh / Ngoại lệ chung:', style='List Bullet')
exceptions = [
    'TK thu phí mặc định không đủ điều kiện → fallback TK thay thế.',
    'Không tìm được TK thay thế → khoản phí xử lý như thế nào? (chưa rõ trong US).',
    'Code phí không match với bất kỳ KH/TK/Thẻ nào → không sinh khoản thu.',
    'Kafka message timeout → trạng thái duy trì "Đang xử lý".',
    'Khoản phí trạng thái "Thanh toán một phần" / "Chưa thanh toán" → theo dõi nợ phí (chi tiết ở US khác).',
]
for e in exceptions:
    doc.add_paragraph(e, style='List Bullet 2')

# ============ A.2.1 TRẠNG THÁI KHOẢN PHÍ ============
doc.add_heading('Vòng đời trạng thái khoản phí định kỳ:', level=3)
states = [
    ('Thêm mới', 'Khoản phí mới được sinh (chưa gửi yêu cầu đi).'),
    ('Đang xử lý', 'Đã ghi vào Topic Kafka nhưng chưa nhận phản hồi kết quả.'),
    ('Thanh toán toàn bộ', 'Đã thu được toàn bộ số tiền phí cần thu.'),
    ('Thanh toán một phần', 'Đã thu được một phần (tận thu do TK không đủ số dư).'),
    ('Chưa thanh toán', 'Thu không thành công toàn bộ.'),
    ('Xóa nợ', 'Đã miễn giảm — không tiếp tục theo dõi nợ phí.'),
]
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Trạng thái'
hdr[1].text = 'Mô tả'
for s, d in states:
    row = table.add_row().cells
    row[0].text = s
    row[1].text = d

# ============ A.2.2 ĐỐI TƯỢNG TÍNH PHÍ ============
doc.add_paragraph('')
doc.add_heading('Đối tượng tính phí định kỳ:', level=3)
objects = [
    ('Khách hàng (Customer)', 'Phí dịch vụ SMS, Phí quản lý Merchant, Phí PVconnect…', 'Hàng tháng (ngày cố định)', 'TK thanh toán mặc định (đồng bộ T-1)'),
    ('Tài khoản (Account)', 'Phí quản lý TK thanh toán', 'Hàng tháng (ngày cố định)', 'Chính TK thanh toán đó'),
    ('Thẻ (Card)', 'Phí thường niên thẻ', 'Hàng năm (tròn năm từ ngày kích hoạt, đồng bộ T-1)', 'Chính số thẻ đó'),
]
table2 = doc.add_table(rows=1, cols=4)
table2.style = 'Light Grid Accent 1'
hdr2 = table2.rows[0].cells
hdr2[0].text = 'Đối tượng'
hdr2[1].text = 'Ví dụ loại phí'
hdr2[2].text = 'Chu kỳ thu'
hdr2[3].text = 'TK thu phí'
for o in objects:
    row = table2.add_row().cells
    for i, v in enumerate(o):
        row[i].text = v

# ============ A.3. PRE-CONDITIONS ============
doc.add_paragraph('')
doc.add_heading('A.3. Bảng Điều Kiện Tiên Quyết & Cấu Hình', level=2)
preconds = [
    'Code phí định kỳ đã được thiết lập và ở trạng thái Hiệu lực trên ProfiX.',
    'Biểu phí/Công thức tính phí đã được khai báo cho các Code phí.',
    'Job thu phí định kỳ đã được cài đặt lịch chạy và gắn Code phí tương ứng.',
    'Dữ liệu KH/TK/Thẻ đã được đồng bộ T-1 từ Core.',
    'Tham số hệ thống CA_PRODUCT đã được cấu hình (danh sách sản phẩm TK được phép trích thu phí).',
    'Tỷ giá đã được đồng bộ từ Core (nếu Code phí đa tệ).',
    'CTƯĐ (nếu có) đã được khai báo và ở trạng thái Hiệu lực.',
    'Kafka Topics thu phí & kết quả thu phí đã được thiết lập và sẵn sàng.',
    'Thứ tự ưu tiên nghiệp vụ phí đã được cài đặt.',
]
for pc in preconds:
    doc.add_paragraph(pc, style='List Bullet')

# ============ A.4. QUY TẮC CHUNG ÁP DỤNG ============
doc.add_heading('A.4. Quy Tắc Chung Áp Dụng (QTC)', level=2)
qtcs = [
    '[QTC-01.4] Định dạng Number: phân cách nghìn bằng dấu phẩy, thập phân bằng dấu chấm, 2 chữ số. VND/JPY không có thập phân.',
    '[QTC-01.5] Định dạng Date: dd/mm/yyyy.',
    '[QTC-01.7] Dữ liệu rỗng hiển thị blank (không hiển thị "-", "N/A").',
    '[QTC-06] Phân trang mặc định 50 bản ghi/trang (nếu áp dụng cho màn hình tra cứu nợ phí liên quan).',
    '[QTC-10] Phân quyền dữ liệu theo Khối — áp dụng cho Code phí (lọc dữ liệu trả về, không lọc dropdown).',
    'Ghi chú: US này không có UI (N/A), nên các QTC liên quan đến giao diện (QTC-02, 03, 05, 07, 08, 09, 14, 15) không áp dụng trực tiếp.',
]
for q in qtcs:
    doc.add_paragraph(q, style='List Bullet')

# Save
output_dir = '/Users/mac/antigravity-testing-kit/Feature_02_SA_Tham_So_He_Thong/output'
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, 'US35_PartA_Summary.docx')
doc.save(output_path)
print(f'✅ Saved: {output_path}')
