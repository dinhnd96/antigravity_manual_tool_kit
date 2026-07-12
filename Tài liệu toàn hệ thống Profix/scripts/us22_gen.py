import os, sys
sys.path.insert(0, os.path.dirname(__file__))
from us22_data import *
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def setup_doc():
    doc = Document()
    for s in doc.sections:
        s.orientation = WD_ORIENT.LANDSCAPE
        s.page_width, s.page_height = s.page_height, s.page_width
        s.top_margin = Cm(1.27); s.bottom_margin = Cm(1.27)
        s.left_margin = Cm(1.27); s.right_margin = Cm(1.27)
    style = doc.styles['Normal']
    style.font.name = 'Arial'; style.font.size = Pt(10)
    return doc

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs: r.font.color.rgb = RGBColor(0,0,0)
    return h

def set_cell(cell, text, bold=False, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.font.name = 'Arial'

def shade_row(row, color="D9E2F3"):
    for c in row.cells:
        tc = c._tc; shading = tc.get_or_add_tcPr()
        s = shading.makeelement(qn('w:shd'), {qn('w:fill'): color, qn('w:val'): 'clear'})
        shading.append(s)

def gen_part_a():
    doc = setup_doc()
    add_heading(doc, f"PHẦN A: TÓM TẮT NGHIỆP VỤ – {FEATURE}")
    # A.1 Core Value
    add_heading(doc, "A.1. Thông Điệp Cốt Lõi (Core Business Value)", 2)
    doc.add_paragraph(CORE_VALUE)
    # A.2 Modules
    add_heading(doc, "A.2. Cấu Trúc Luồng Nghiệp Vụ & Module Mapping", 2)
    for m in MODULES:
        add_heading(doc, m["name"], 3)
        doc.add_paragraph("Luồng chính (Happy Path):", style='List Bullet')
        for s in m["happy"]:
            p = doc.add_paragraph(s); p.paragraph_format.left_indent = Cm(1.5)
        doc.add_paragraph("Luồng rẽ nhánh / Ngoại lệ:", style='List Bullet')
        for s in m["alt"]:
            p = doc.add_paragraph(s); p.paragraph_format.left_indent = Cm(1.5)
    # A.3 Preconditions
    add_heading(doc, "A.3. Điều Kiện Tiên Quyết & Cấu Hình", 2)
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell(t.rows[0].cells[0], "Hạng mục", True); set_cell(t.rows[0].cells[1], "Mô tả", True)
    shade_row(t.rows[0])
    for item, desc in PRECONDITIONS:
        row = t.add_row(); set_cell(row.cells[0], item); set_cell(row.cells[1], desc)
    # A.4 QTC
    add_heading(doc, "A.4. Quy Tắc Chung Áp Dụng", 2)
    doc.add_paragraph("US22 là tính năng đơn giản (đăng xuất), không liên quan đến tìm kiếm, lọc, phân trang, upload, download, Maker-Checker. Các QTC áp dụng tối thiểu:")
    t2 = doc.add_table(rows=1, cols=2)
    t2.style = 'Table Grid'; t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell(t2.rows[0].cells[0], "QTC", True); set_cell(t2.rows[0].cells[1], "Nội dung áp dụng", True)
    shade_row(t2.rows[0])
    for qtc, desc in QTC_APPLIED:
        row = t2.add_row(); set_cell(row.cells[0], qtc); set_cell(row.cells[1], desc)
    # Flowchart note
    add_heading(doc, "A.5. Ghi Chú Flowchart & Mockup", 2)
    doc.add_paragraph("Flowchart: Chỉ vẽ luồng Đăng xuất thủ công (7 bước, 3 swimlane: User/Frontend/Backend). Không có Flowchart cho luồng tự động đăng xuất.")
    doc.add_paragraph("Mockup 1: Khu vực thông tin user ở sidebar trái dưới cùng – hiển thị tên, đơn vị, nút \"Đăng xuất\".")
    doc.add_paragraph("Mockup 2: Popup xác nhận – Title \"Đăng xuất\", nội dung \"Bạn có chắc chắn muốn đăng xuất không?\", 2 nút: \"Hủy bỏ\" (outline) và \"Đồng ý\" (primary, màu vàng).")
    doc.add_paragraph("Lưu ý: Không có mockup cho popup tự động đăng xuất (\"Phiên đăng nhập hết hạn\").")
    path = os.path.join(OUTPUT_DIR, PART_A_FILE)
    doc.save(path); print(f"Saved: {path}")

def gen_part_b():
    doc = setup_doc()
    add_heading(doc, f"PHẦN B: DANH SÁCH CẢNH BÁO & Q&A – {FEATURE}")
    doc.add_paragraph("Tài liệu phân tích: US22.docx | Ngày phân tích: 2026-05-18 | Trạng thái: Chờ BA trả lời")
    # QA Table
    t = doc.add_table(rows=1, cols=6)
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["ID", "Trích xuất (Vị trí + Nội dung)", "Câu hỏi / Sự cố", "Phân loại", "Đề xuất từ QC", "Trả lời của BA"]
    widths = [Cm(2.2), Cm(6.5), Cm(6), Cm(2), Cm(5.5), Cm(3.5)]
    for i, h in enumerate(headers):
        set_cell(t.rows[0].cells[i], h, True, 9)
        t.rows[0].cells[i].width = widths[i]
    shade_row(t.rows[0])
    cat_colors = {"Nghiệp vụ": "FFF2CC", "Giới hạn": "FCE4EC", "Toàn vẹn dữ liệu": "E8F5E9", "UI-UX": "E3F2FD"}
    prev_cat = ""
    for qid, ref, q, cat, sug in QA_DATA:
        if cat != prev_cat:
            cat_row = t.add_row()
            cat_labels = {"Nghiệp vụ": "🔶 Hạng mục 1: Vấn đề Nghiệp vụ / Luồng xử lý",
                          "Giới hạn": "🔴 Hạng mục 2: Giới hạn hệ thống & Exception Handling",
                          "Toàn vẹn dữ liệu": "🟠 Hạng mục 3: Toàn vẹn dữ liệu & Ràng buộc",
                          "UI-UX": "🔵 Hạng mục 4: UI/UX & Giao diện"}
            merged = cat_row.cells[0].merge(cat_row.cells[5])
            set_cell(merged, cat_labels.get(cat, cat), True, 10)
            shade_row(cat_row, cat_colors.get(cat, "FFFFFF"))
            prev_cat = cat
        row = t.add_row()
        vals = [qid, ref, q, cat, sug, ""]
        for i, v in enumerate(vals):
            set_cell(row.cells[i], v, False, 8)
            row.cells[i].width = widths[i]
    path = os.path.join(OUTPUT_DIR, PART_B_FILE)
    doc.save(path); print(f"Saved: {path}")

if __name__ == "__main__":
    gen_part_a()
    gen_part_b()
    print("Done!")
