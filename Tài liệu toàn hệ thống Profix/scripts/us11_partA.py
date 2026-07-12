"""Script sinh file US11_PartA_Summary.docx — Tóm tắt Nghiệp vụ Chuyên Sâu."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Page setup: Landscape, narrow margins
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(10)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.size = Pt(14 if level == 1 else 12 if level == 2 else 11)
    return h

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

# ======= PART A =======
add_heading("PHẦN A: TÓM TẮT NGHIỆP VỤ CHUYÊN SÂU — US11", level=1)
doc.add_paragraph("Chương trình ưu đãi có đánh giá định kỳ khách hàng được áp dụng ưu đãi")

# A.1
add_heading("A.1. Thông Điệp Cốt Lõi (Core Business Value)", level=2)
p = doc.add_paragraph()
p.add_run("Mục đích: ").bold = True
p.add_run("Cho phép người dùng (Maker) thiết lập và quản lý các Chương trình ưu đãi (CTƯĐ) có cơ chế đánh giá định kỳ khách hàng. Hệ thống sẽ tự động xác định tập khách hàng đủ điều kiện và áp dụng ưu đãi (giảm phí) theo các tiêu chí đã cấu hình.")
p2 = doc.add_paragraph()
p2.add_run("Người dùng cuối: ").bold = True
p2.add_run("Cán bộ quản lý tham số phí (Maker) và Người phê duyệt (Checker) tại PVCB.")
p3 = doc.add_paragraph()
p3.add_run("Navigation: ").bold = True
p3.add_run("Tham số >> Chương trình ưu đãi >> Thêm mới >> Ưu đãi có đánh giá định kỳ | Chỉnh sửa bản ghi CTƯĐ.")

# A.2
add_heading("A.2. Cấu Trúc Luồng Nghiệp Vụ & Phân Bổ Module", level=2)

# Module 1
add_heading("Module 1: Thêm mới CTƯĐ có đánh giá định kỳ", level=3)
doc.add_paragraph().add_run("Luồng chính (Happy Path):").bold = True
add_bullet("Bước 1-4: Maker chọn Thêm mới → chọn 'Ưu đãi có đánh giá định kỳ' → FE hiển thị form khai báo.")
add_bullet("Bước 5: Khai báo Thông tin chung (Tên CTƯĐ, Ngày ban hành, Số văn bản, Tên văn bản, Số VB_Tên viết tắt, Ngày hiệu lực, Ngày hết hiệu lực, Link iDoc, Đối tượng thu phí, Khối, Loại ưu đãi, Email CBNV).")
add_bullet("Bước 6-7: Khai báo Điều kiện đánh giá — chọn Nguồn dữ liệu, Trường điều kiện, Operator, Giá trị. Hỗ trợ nhiều nhóm điều kiện (AND/OR).")
add_bullet("Bước 8-9: Khai báo Lịch đánh giá — thêm/sửa/xóa lần đánh giá. Lần 1 mặc định = Ngày hiệu lực.")
add_bullet("Bước 10: Khai báo Chi tiết ưu đãi — chọn Kênh, Kỳ, SPDV, toggle Ưu đãi theo tỷ lệ/giá trị.")
add_bullet("Bước 10.a-11: Nhấn Xác nhận → BE validate → kiểm tra trùng mã CTƯĐ → Lưu thành công → trạng thái Chờ duyệt → hiển thị tại Tác vụ chờ duyệt [QTC-12].")

doc.add_paragraph().add_run("Các Luồng Rẽ Nhánh / Ngoại lệ:").bold = True
add_bullet("Nhấn Đóng → hủy thao tác, quay về màn hình Quản lý CTƯĐ [QTC-15].")
add_bullet("Validate lỗi (BE) → FE hiển thị thông báo lỗi, không lưu.")
add_bullet("Mã CTƯĐ trùng → báo lỗi, không cho lưu.")
add_bullet("Thêm lịch khi Ngày hiệu lực/Ngày hết hiệu lực trống → báo lỗi.")
add_bullet("SPDV có quan hệ cha-con trong cùng CTƯĐ → không cho phép chọn.")

# Module 2
add_heading("Module 2: Chỉnh sửa CTƯĐ có đánh giá định kỳ", level=3)
doc.add_paragraph().add_run("Luồng chính (Happy Path):").bold = True
add_bullet("Bước 1-2: Maker nhấn Chỉnh sửa → BE kiểm tra tác vụ Chờ duyệt đang tồn tại.")
add_bullet("Bước 3: Kiểm tra trạng thái hoạt động CTƯĐ:")
add_bullet("  • Chưa hiệu lực → cho sửa toàn bộ trừ Mã CTƯĐ và Số VB_Tên viết tắt.")
add_bullet("  • Đang hiệu lực → chỉ cho sửa: Ngày hết hiệu lực (> Ngày HT & > Ngày HL), Lịch đánh giá (thêm mới/sửa/xóa lần có Ngày đánh giá > Ngày HT), Chi tiết ưu đãi (thêm mới).")
add_bullet("Bước 4.a-5: Nhấn Xác nhận → BE validate → Lưu → trạng thái Chờ duyệt hành động Sửa → Tác vụ chờ duyệt [QTC-12].")

doc.add_paragraph().add_run("Các Luồng Rẽ Nhánh / Ngoại lệ:").bold = True
add_bullet("Đã tồn tại tác vụ Chỉnh sửa Chờ duyệt → chặn, không cho chỉnh sửa.")
add_bullet("Trạng thái = Hết hiệu lực → báo lỗi 'Trạng thái hoạt động của CTƯĐ không hợp lệ'.")
add_bullet("Nhấn Đóng → hủy, quay về [QTC-15].")

# Module 3
add_heading("Module 3: Xử lý tự động xác định danh sách KH tham gia CTƯĐ", level=3)
doc.add_paragraph().add_run("Luồng chính (Batch Job — 3 thời điểm):").bold = True

add_bullet("Thời điểm 1 — Đầu ngày hiệu lực CTƯĐ:", bold_prefix="")
add_bullet("  Bước 1: Lấy các bộ tham số (Danh sách A) KHÔNG bao gồm ĐK 'Trạng thái KH = NEW'.")
add_bullet("  Bước 2: Kiểm tra dữ liệu ETL T-1 theo Loại ưu đãi (KH/TK/Thẻ) → Khởi tạo Danh sách B gồm: Mã CTƯĐ, ID Bộ tham số, CIF, Kênh, Ưu đãi từ ngày/đến ngày, SPDV cấp cuối, Phương thức ưu đãi.")
add_bullet("  Nếu Loại ưu đãi = Theo TK → bổ sung 'Số tài khoản'; Theo Thẻ → bổ sung 'Số thẻ/Token thẻ'.")

add_bullet("Thời điểm 2 — Trong ngày, khi có KH mở mới:", bold_prefix="")
add_bullet("  Chỉ áp dụng Loại ưu đãi = Theo KH. KH mở mới qua Topic đồng bộ → kiểm tra ĐK có 'Trạng thái KH = NEW' → nếu thỏa → bổ sung vào Danh sách B.")

add_bullet("Thời điểm 3 — Đầu ngày tái đánh giá:", bold_prefix="")
add_bullet("  Bước 1: Tái đánh giá đối tượng trong Danh sách B → nếu không thỏa → hủy ưu đãi các kỳ còn hiệu lực.")
add_bullet("  Bước 2: Tái đánh giá đối tượng KHÔNG thuộc Danh sách B → tương tự xử lý đầu ngày hiệu lực.")

doc.add_paragraph().add_run("Cơ chế tính Ưu đãi từ ngày / đến ngày:").bold = True
add_bullet("Ưu đãi từ ngày = Ngày đánh giá.")
add_bullet("Ưu đãi đến ngày = Ngày đánh giá + Khoảng thời gian áp dụng ưu đãi.")
add_bullet("Khoảng thời gian đầu tiên = Kỳ đầu tiên; Khoảng thời gian tiếp theo = Kỳ hiện tại – Kỳ liền trước.")

# A.3
add_heading("A.3. Bảng Điều Kiện Tiên Quyết & Cấu Hình", level=2)
table = doc.add_table(rows=6, cols=2, style='Light Grid Accent 1')
table.cell(0, 0).text = "Điều kiện tiên quyết"
table.cell(0, 1).text = "Mô tả"
data = [
    ("Phân quyền Maker/Checker", "User được phân quyền tạo/chỉnh sửa CTƯĐ (Maker) và phê duyệt (Checker) theo QTC-12."),
    ("Dữ liệu SPDV", "Danh mục SPDV phải có sẵn với trạng thái Hoạt động để chọn trong Chi tiết ưu đãi."),
    ("Cấu hình Điều kiện tính phí", "Các Trường điều kiện phải được cấu hình sẵn với Nguồn dữ liệu tương ứng (ETL KH/TK/Thẻ)."),
    ("Dữ liệu ETL T-1", "Dữ liệu khách hàng/tài khoản/thẻ đồng bộ từ Core banking (T-1) để phục vụ đánh giá."),
    ("Ma trận phê duyệt (US25)", "Cấu hình cấp duyệt cho hành động Thêm mới/Chỉnh sửa CTƯĐ."),
]
for i, (k, v) in enumerate(data, 1):
    table.cell(i, 0).text = k
    table.cell(i, 1).text = v

# A.4
add_heading("A.4. Quy Tắc Chung Áp Dụng", level=2)
qtc_data = [
    ("QTC-01.2", "Dropdown List — chọn 1 giá trị (Đối tượng thu phí, Khối, Loại ưu đãi, Kênh, Operator, Nguồn dữ liệu)."),
    ("QTC-01.4", "Number — phân cách hàng nghìn bằng dấu phẩy, 2 số sau dấu chấm (Giá trị ưu đãi, Tỷ lệ ưu đãi)."),
    ("QTC-01.5", "Date — định dạng dd/mm/yyyy (Ngày ban hành, Ngày hiệu lực, Ngày hết hiệu lực, Ngày đánh giá)."),
    ("QTC-01.6", "Text mặc định: Mã/Tên 50 ký tự, Ghi chú/Mô tả 300 ký tự (trừ Số VB_Tên viết tắt = 20 ký tự theo US)."),
    ("QTC-12", "Luồng Maker-Checker: Thêm mới/Chỉnh sửa → Chờ duyệt → Phê duyệt/Từ chối. Mã sinh sau khi Last Checker duyệt."),
    ("QTC-14.1", "No-Change Guard — FE disable nút Xác nhận nếu không có thay đổi."),
    ("QTC-14.2", "Tên CTƯĐ không yêu cầu unique."),
    ("QTC-15", "Nút Đóng: không popup xác nhận, không lưu, quay về màn hình trước."),
]
table2 = doc.add_table(rows=len(qtc_data) + 1, cols=2, style='Light Grid Accent 1')
table2.cell(0, 0).text = "Mã QTC"
table2.cell(0, 1).text = "Nội dung áp dụng"
for i, (code, desc) in enumerate(qtc_data, 1):
    table2.cell(i, 0).text = code
    table2.cell(i, 1).text = desc

# Save
output_path = "/Users/mac/antigravity-testing-kit/Feature_02_SA_Tham_So_He_Thong/output/US11_PartA_Summary.docx"
import os
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f"✅ Saved: {output_path}")
