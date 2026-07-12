"""US35 Part B Merged - DOCX Generator (imports data from 2 data files)"""
import sys, os
sys.path.insert(0, os.path.dirname(__file__))
from us35_merged_data import QA_DATA
from us35_merged_data2 import QA_DATA_2
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

ALL_QA = QA_DATA + QA_DATA_2
doc = Document()
for s in doc.sections:
    s.orientation = WD_ORIENT.LANDSCAPE
    s.page_width, s.page_height = Cm(29.7), Cm(21.0)
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Cm(1.27)

doc.styles['Normal'].font.name = 'Arial'
doc.styles['Normal'].font.size = Pt(9)

t = doc.add_heading('US35 – Part B Merged: AI + VA Q&A Tổng Hợp (v2 - Fixed References)', level=1)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Nguồn = AI | VA | BOTH. Cột Trích xuất đã sửa theo chuẩn verbatim.')
doc.add_paragraph('')

ai_n = sum(1 for q in ALL_QA if q[5]=='AI')
va_n = sum(1 for q in ALL_QA if q[5]=='VA')
bo_n = sum(1 for q in ALL_QA if q[5]=='BOTH')
p = doc.add_paragraph()
p.add_run(f'Tổng: {len(ALL_QA)} câu | AI only: {ai_n} | VA only: {va_n} | BOTH: {bo_n}').bold = True

table = doc.add_table(rows=1, cols=7)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
cw = [Cm(2.2), Cm(5.5), Cm(5.5), Cm(1.8), Cm(5.0), Cm(4.0), Cm(1.5)]
for i, w in enumerate(cw):
    table.columns[i].width = w

for i, h in enumerate(['ID','Trích xuất','Câu hỏi / Sự cố','Phân loại','Đề xuất từ QA','Trả lời của BA','Nguồn']):
    c = table.rows[0].cells[i]
    c.text = h
    for p in c.paragraphs:
        for r in p.runs:
            r.bold = True; r.font.size = Pt(8)

cat_map = {
    '01': '🔶 HM1: Vấn đề Nghiệp vụ / Luồng xử lý',
    '02': '🔴 HM2: Giới hạn hệ thống & Exception',
    '03': '🟠 HM3: Toàn vẹn dữ liệu & Ràng buộc',
    '04': '🔵 HM4: UI/UX & Giao diện',
}
cur = None
for qa in ALL_QA:
    cc = qa[0].split('-QA-')[1][:2]
    if cc != cur:
        cur = cc
        cr = table.add_row().cells
        cr[0].merge(cr[6])
        cr[0].text = cat_map.get(cc,'')
        for p in cr[0].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.bold=True; r.font.size=Pt(9); r.font.color.rgb=RGBColor(0,0,128)
    row = table.add_row().cells
    for i, v in enumerate(qa):
        row[i].text = v
        for p in row[i].paragraphs:
            for r in p.runs:
                r.font.size = Pt(8)

for r in table.rows:
    for i, c in enumerate(r.cells):
        c.width = cw[i]

out = '/Users/mac/antigravity-testing-kit/Feature_02_SA_Tham_So_He_Thong/output/US35_PartB_QA_Merged.docx'
doc.save(out)
print(f'✅ Saved: {out} ({len(ALL_QA)} questions)')
