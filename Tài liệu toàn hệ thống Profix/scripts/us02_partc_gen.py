#!/usr/bin/env python3
"""Generate US02 Part C docx - Test Case Coverage Table."""
import subprocess, sys
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

from us02_partc_data import DATA, HEADERS

OUT = "/Users/mac/antigravity-testing-kit/Feature_02_SA_Tham_So_He_Thong/tài liệu/US02_PartC_TestCase_Coverage.docx"

doc = Document()
# Landscape + narrow margins
for sec in doc.sections:
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Cm(29.7)
    sec.page_height = Cm(21.0)
    sec.top_margin = Cm(1.27)
    sec.bottom_margin = Cm(1.27)
    sec.left_margin = Cm(1.27)
    sec.right_margin = Cm(1.27)

# Title
title = doc.add_heading("PHẦN C: BẢNG TỔNG HỢP TEST CASE ĐỀ XUẤT (Test Case Coverage)", level=1)
for run in title.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph("US02 - Khai báo Code phí với thông tin chi tiết về mức phí cho từng Phân khúc Khách hàng")
doc.add_paragraph(f"Tổng số kịch bản: {len(DATA)} | Tổng TC dự kiến: {sum(r[5] for r in DATA)}")
doc.add_paragraph("")

# Group by type for summary
from collections import Counter
type_cnt = Counter(r[3] for r in DATA)
type_map = {
    "Happy Path": "🟢 Happy Path",
    "Negative": "🔴 Negative & Exception",
    "Boundary": "📐 Boundary Value",
    "UI/UX": "🎨 UI/UX & Field Validation",
    "Business Logic": "🧠 Business Logic & State",
    "Data Integrity": "🔗 Data Integrity",
    "NFR": "⚡ NFR (Spam click / Concurrency)",
}
summary_para = doc.add_paragraph("Phân bổ theo loại: ")
for k, label in type_map.items():
    if k in type_cnt:
        summary_para.add_run(f"{label}: {type_cnt[k]}  |  ")
doc.add_paragraph("")

# Table
COL_WIDTHS = [Cm(1.8), Cm(3.0), Cm(3.5), Cm(3.0), Cm(8.0), Cm(1.5), Cm(6.5)]
table = doc.add_table(rows=1, cols=len(HEADERS))
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
hdr = table.rows[0]
for i, h in enumerate(HEADERS):
    cell = hdr.cells[i]
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(255, 255, 255)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Dark blue background
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): "003366",
        qn("w:val"): "clear",
    })
    shading.append(shd)

# Color map for test type
TYPE_COLORS = {
    "Happy Path": "E6F4EA",
    "Negative": "FCE4EC",
    "Boundary": "FFF3E0",
    "UI/UX": "E3F2FD",
    "Business Logic": "F3E5F5",
    "Data Integrity": "FFF8E1",
    "NFR": "EFEBE9",
}

# Data rows
for row_data in DATA:
    row = table.add_row()
    test_type = row_data[3]
    bg = TYPE_COLORS.get(test_type, "FFFFFF")
    for i, val in enumerate(row_data):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(val))
        run.font.size = Pt(7.5)
        if i == 0:
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif i == 5:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Apply bg color
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn("w:shd"), {
            qn("w:fill"): bg,
            qn("w:val"): "clear",
        })
        shading.append(shd)

# Set column widths
for row in table.rows:
    for i, w in enumerate(COL_WIDTHS):
        row.cells[i].width = w

# Legend
doc.add_paragraph("")
legend = doc.add_paragraph("Chú thích phân loại:")
legend.runs[0].bold = True
for k, label in type_map.items():
    doc.add_paragraph(f"  {label}", style="List Bullet")

doc.add_paragraph("")
doc.add_paragraph("Ghi chú: Các câu hỏi Q01→Q07 đã được BA trả lời. Q01 (Drop - UI đang work), Q02 (Drop - chỉ tiện ích), Q03 (Drop - cứng 28), Q04 (BA update: check trùng Trường+Operator), Q05 (BA update US), Q06 (giữ nguyên giá trị VAT), Q07 (Drop - chưa chốt giờ Batch).")

doc.save(OUT)
print(f"OK: {OUT}")
