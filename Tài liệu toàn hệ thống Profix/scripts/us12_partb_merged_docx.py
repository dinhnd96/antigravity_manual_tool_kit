"""US12 Part B Merged - LOGIC ONLY"""
import sys, os
sys.path.insert(0, os.path.dirname(__file__))
from us12_partb_merged_data import QA_DATA
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10)

title = doc.add_heading('US12 — PHẦN B TỔNG HỢP: Q&A (AI + VA Merged)', level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)
doc.add_paragraph('Feature: Khai báo CTƯĐ áp dụng cho danh sách KH | Phiên bản: v1.1 (Merged) | Ngày: 13/05/2026')
doc.add_paragraph(f'Tổng số câu hỏi: {len(QA_DATA)} | Nguồn: AI (12 câu riêng) + VA (10 câu riêng) + 6 câu trùng')

# Group by category
categories = [
    ('🔶 Hạng mục 1: Nghiệp vụ / Luồng xử lý', 'Nghiệp vụ'),
    ('🔴 Hạng mục 2: Giới hạn hệ thống & Exception Handling', 'Giới hạn'),
    ('🟠 Hạng mục 3: Toàn vẹn dữ liệu & Ràng buộc', 'Toàn vẹn dữ liệu'),
    ('🔵 Hạng mục 4: UI/UX & Giao diện', 'UI-UX'),
]

col_widths = [Inches(0.8), Inches(1.8), Inches(3.2), Inches(0.7), Inches(2.5), Inches(0.6), Inches(1.2)]
headers = ['ID', 'Trích xuất', 'Câu hỏi / Sự cố', 'Phân loại', 'Đề xuất từ QA', 'Nguồn', 'Trả lời BA']

for cat_title, cat_key in categories:
    doc.add_heading(cat_title, level=2)
    items = [q for q in QA_DATA if q[1] == cat_key]
    if not items:
        doc.add_paragraph('(Không có)')
        continue
    table = doc.add_table(rows=1, cols=7, style='Light Grid Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(8)
    for row in table.rows:
        for i, w in enumerate(col_widths):
            row.cells[i].width = w
    for qa in items:
        row = table.add_row()
        vals = [qa[0], qa[2], qa[3], qa[1], qa[4], qa[5], '']
        for i, v in enumerate(vals):
            row.cells[i].text = v
            for p in row.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
    for row in table.rows:
        for i, w in enumerate(col_widths):
            row.cells[i].width = w

out = '/Users/mac/antigravity-testing-kit/Feature_02_SA_Tham_So_He_Thong/US12_PartB_QA_Merged.docx'
doc.save(out)
print(f'Merged Part B saved: {out}')
print(f'Total Q&A: {len(QA_DATA)}')
