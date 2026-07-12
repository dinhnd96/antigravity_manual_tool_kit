"""Script sinh file US11_PartB_QA.docx — Danh sách Cảnh Báo & Q&A."""
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(9)

doc.add_heading("PHẦN B: DANH SÁCH CẢNH BÁO & Q&A — US11", level=1)
doc.add_paragraph("Chương trình ưu đãi có đánh giá định kỳ khách hàng được áp dụng ưu đãi")

# === Q&A DATA ===
qa_items = [
    # Hạng mục 1: Nghiệp vụ / Luồng xử lý
    ("US11-QA-01.1", 
     "Mục 'Thêm mới CTƯĐ', đoạn Bộ tham số",
     "Tài liệu cho phép nhiều bộ tham số. Khi 2 bộ tham số có điều kiện đánh giá overlap (cùng 1 KH thỏa cả 2 bộ), KH đó được áp dụng ưu đãi của bộ nào? Có cho phép cộng dồn ưu đãi hay chỉ áp dụng bộ có giá trị cao nhất?",
     "Nghiệp vụ",
     "Đề xuất: Hệ thống áp dụng ưu đãi của bộ tham số có mức ưu đãi cao nhất (most favorable), không cộng dồn. Cần BA xác nhận logic xử lý chồng lấn."),

    ("US11-QA-01.2",
     "Mục 'Chỉnh sửa CTƯĐ', đoạn trạng thái Đang hiệu lực",
     "Khi CTƯĐ Đang hiệu lực, tài liệu chỉ cho sửa Ngày hết hiệu lực, Lịch đánh giá, Chi tiết ưu đãi. Nhưng mục 'Điều kiện đánh giá' không được nhắc đến. Xác nhận: Khi Đang hiệu lực, Điều kiện đánh giá bị khóa hoàn toàn không cho sửa?",
     "Nghiệp vụ",
     "Đề xuất: Khi Đang hiệu lực, Điều kiện đánh giá ở trạng thái readonly. Chỉ cho sửa khi Chưa hiệu lực."),

    ("US11-QA-01.3",
     "Mục 'Chỉnh sửa CTƯĐ', đoạn trạng thái Đang hiệu lực, phần 'Chi tiết ưu đãi'",
     "Khi Đang hiệu lực, tài liệu ghi 'cho phép thêm mới chi tiết ưu đãi'. Vậy có cho phép SỬA hoặc XÓA các chi tiết ưu đãi đã tồn tại không?",
     "Nghiệp vụ",
     "Đề xuất: Khi Đang hiệu lực, các chi tiết ưu đãi đã tồn tại ở trạng thái readonly. Chỉ cho phép THÊM MỚI dòng ưu đãi mới."),

    ("US11-QA-01.4",
     "Mục 'Xử lý tự động', phần 'Đầu ngày tái đánh giá', Bước 1",
     "Khi tái đánh giá phát hiện KH không còn thỏa → hủy ưu đãi các kỳ còn hiệu lực. Vậy kỳ ưu đãi đang được áp dụng tại thời điểm tái đánh giá có bị hủy ngay không, hay chỉ hủy các kỳ TƯƠNG LAI?",
     "Nghiệp vụ",
     "Đề xuất: Kỳ đang áp dụng (Ưu đãi từ ngày <= Ngày đánh giá <= Ưu đãi đến ngày) bị hủy ngay lập tức. Cần BA xác nhận."),

    ("US11-QA-01.5",
     "Mục 'Xử lý tự động', phần 'Trong ngày khi phát sinh KH mới'",
     "Cơ chế KH mở mới chỉ áp dụng với Loại ưu đãi = Theo KH. Vậy khi Loại ưu đãi = Theo TK hoặc Theo Thẻ, KH mở mới TK/Thẻ trong ngày có được tự động đánh giá và bổ sung vào Danh sách B không? Hay phải chờ đến lần tái đánh giá tiếp theo?",
     "Nghiệp vụ",
     "Đề xuất: Theo tài liệu, cơ chế real-time chỉ áp dụng cho Loại ưu đãi = Theo KH. Các loại TK/Thẻ chờ đến lần tái đánh giá kế tiếp. Cần BA xác nhận."),

    ("US11-QA-01.6",
     "Mục 'Chỉnh sửa CTƯĐ', đoạn kiểm tra tác vụ Chờ duyệt",
     "Tài liệu chỉ kiểm tra tác vụ 'Chỉnh sửa' Chờ duyệt. Nếu tồn tại tác vụ 'Thêm mới' Chờ duyệt (CTƯĐ vừa tạo chưa duyệt), có cho phép Chỉnh sửa không?",
     "Nghiệp vụ",
     "Đề xuất: CTƯĐ chưa duyệt (Thêm mới Chờ duyệt) thì không hiển thị trên lưới chính thức → không có nút Chỉnh sửa. Tuy nhiên cần xác nhận hành vi tại Tác vụ Pending nếu Maker muốn sửa bản ghi Thêm mới bị Từ chối."),

    ("US11-QA-01.7",
     "Mục 'Thêm mới CTƯĐ', Lưu đồ Thêm mới, Bước 11",
     "Diễn giải lưu đồ tại Bước 11 ghi: 'Các thông tin khai báo hợp lệ chuyển xuống Bước 11.a' VÀ 'Các thông tin khai báo không hợp lệ chuyển xuống Bước 11.a'. Cả 2 nhánh đều trỏ về 11.a — rõ ràng là lỗi copy-paste. Nhánh không hợp lệ phải trỏ về Bước nào?",
     "Nghiệp vụ",
     "Đề xuất: Nhánh không hợp lệ phải trỏ về Bước 11.1 (Hiển thị thông báo lỗi). Đồng thời, trong diễn giải các bước sau đó, Bước '10.1' và '10.a' bị đánh số trùng với Bước 10 phía trên — cần chỉnh thành '11.1' và '11.a' tương ứng. Đây là lỗi copy-paste trong tài liệu."),

    # Hạng mục 2: Giới hạn hệ thống & Exception
    ("US11-QA-02.1",
     "Mục 'Thêm mới CTƯĐ', đoạn Bộ tham số",
     "Tài liệu không giới hạn số lượng bộ tham số tối đa trong 1 CTƯĐ. Có giới hạn tối đa không? Nếu có thì bao nhiêu?",
     "Giới hạn",
     "Đề xuất: Thiết lập giới hạn tối đa (ví dụ: 10 bộ tham số/CTƯĐ) để tránh ảnh hưởng hiệu năng Batch Job đánh giá."),

    ("US11-QA-02.2",
     "Mục 'Thêm mới CTƯĐ', phần Lịch đánh giá",
     "Tài liệu không giới hạn số lần đánh giá tối đa trong 1 bộ tham số. Có giới hạn tối đa không?",
     "Giới hạn",
     "Đề xuất: Thiết lập giới hạn (ví dụ: 50 lần đánh giá) hoặc tính toán tự động dựa trên khoảng cách Ngày hiệu lực – Ngày hết hiệu lực."),

    ("US11-QA-02.3",
     "Mục 'Thêm mới CTƯĐ', phần Chi tiết ưu đãi",
     "Tài liệu không giới hạn số dòng SPDV áp dụng tối đa trong 1 bộ tham số. Có giới hạn tối đa không?",
     "Giới hạn",
     "Đề xuất: Giới hạn số dòng SPDV hoặc để không giới hạn. Cần BA xác nhận."),

    ("US11-QA-02.4",
     "Mục 'Thêm mới CTƯĐ', phần Điều kiện đánh giá",
     "Tài liệu không giới hạn số nhóm điều kiện tối đa và số dòng điều kiện chi tiết tối đa trong 1 nhóm. Có giới hạn không?",
     "Giới hạn",
     "Đề xuất: Thiết lập giới hạn hợp lý cho cả 2 cấp (nhóm điều kiện, dòng trong nhóm) để tránh query quá nặng khi Batch Job đánh giá."),

    ("US11-QA-02.5",
     "Mục 'Xử lý tự động', phần 'Đầu ngày hiệu lực'",
     "Batch Job đánh giá chạy 'đầu ngày hiệu lực'. Nếu Job fail giữa chừng (lỗi hệ thống, timeout), hệ thống xử lý thế nào? Có cơ chế retry/rollback không?",
     "Giới hạn",
     "Đề xuất: Cần cơ chế retry tự động (ví dụ: 3 lần) và ghi log lỗi. Nếu vẫn fail → gửi alert cho Admin/CBNV đầu mối."),

    # Hạng mục 3: Toàn vẹn dữ liệu & Ràng buộc
    ("US11-QA-03.1",
     "Mục 'Thêm mới CTƯĐ', phần 'Số văn bản_Tên viết tắt CTƯĐ'",
     "Trường này bắt buộc có dấu gạch dưới, phía trước là số, phía sau là chữ viết hoa. Tài liệu ghi tối đa 20 ký tự. Ràng buộc '20 ký tự' này tính cả dấu gạch dưới không? Và có giới hạn TỐI THIỂU không (ví dụ: ít nhất 3 ký tự)?",
     "Toàn vẹn dữ liệu",
     "Đề xuất: 20 ký tự bao gồm cả dấu gạch dưới. Tối thiểu: 1 số + '_' + 1 chữ = 3 ký tự."),

    ("US11-QA-03.2",
     "Mục 'Thêm mới CTƯĐ', phần Chi tiết ưu đãi, trường 'SPDV'",
     "Tài liệu ghi: 'Người dùng không được phép chọn các SPDV có quan hệ cha con trong cùng 1 CTƯĐ'. Ràng buộc này áp dụng trong phạm vi 1 bộ tham số hay toàn bộ CTƯĐ (cross-bộ tham số)?",
     "Toàn vẹn dữ liệu",
     "Đề xuất: Theo đoạn văn, phạm vi là 'trong cùng 1 CTƯĐ' → cross-bộ tham số. Cần BA xác nhận."),

    ("US11-QA-03.3",
     "Mục 'Thêm mới CTƯĐ', phần Điều kiện đánh giá, mô tả ràng buộc cặp (Trường điều kiện + Operator)",
     "Bảng mô tả trường ghi: 'một cặp (Trường điều kiện + Operator) không được trùng lặp'. Ràng buộc unique này áp dụng trong phạm vi 1 nhóm điều kiện, hay toàn bộ bộ tham số?",
     "Toàn vẹn dữ liệu",
     "Đề xuất: Phạm vi unique = trong 1 nhóm điều kiện (vì các nhóm có thể cần kiểm tra cùng trường với cùng operator nhưng giá trị khác). Cần BA xác nhận."),

    ("US11-QA-03.4",
     "Mục 'Thêm mới CTƯĐ', phần Điều kiện đánh giá, trường 'Nguồn dữ liệu'",
     "Bảng mô tả trường có mâu thuẫn nội bộ: Dòng STT Nguồn dữ liệu ghi 'Khi Loại ưu đãi = Theo Tài khoản, nguồn dữ liệu = ETL Khách hàng/ETL Tài khoản' VÀ tiếp theo lại ghi 'Khi Loại ưu đãi = Theo Tài khoản, nguồn dữ liệu = ETL Khách hàng/ETL Thẻ'. Dòng thứ 2 rõ ràng phải là 'Theo Thẻ' chứ không phải 'Theo Tài khoản'.",
     "Toàn vẹn dữ liệu",
     "Đề xuất: Đây là lỗi copy-paste. Dòng cuối phải đọc là: 'Khi Loại ưu đãi = Theo Thẻ, nguồn dữ liệu = ETL Khách hàng/ETL Thẻ'. Cần BA xác nhận và sửa tài liệu."),

    ("US11-QA-03.5",
     "Mục 'Thêm mới CTƯĐ', phần Thông tin chung, trường 'Ngày ban hành'",
     "Bảng mô tả trường ghi ràng buộc 'Không chọn ngày tương lai'. Đồng thời, yêu cầu nghiệp vụ ghi 'Ngày hiệu lực và Ngày hết hiệu lực phải >= Ngày ban hành (nếu có)'. Tuy nhiên, tài liệu không nói rõ: Nếu người dùng KHÔNG nhập Ngày ban hành, ràng buộc >= Ngày ban hành có bị bỏ qua hoàn toàn không? (Trường Ngày ban hành có ★ = bắt buộc, nhưng mệnh đề 'nếu có' lại ngầm cho phép bỏ trống).",
     "Toàn vẹn dữ liệu",
     "Đề xuất: Ngày ban hành là ★ (bắt buộc) theo bảng mô tả trường, do đó mệnh đề '(nếu có)' mâu thuẫn với ★. Cần BA làm rõ: Ngày ban hành bắt buộc hay không bắt buộc?"),

    ("US11-QA-03.6",
     "Mục 'Thêm mới CTƯĐ', phần Chi tiết ưu đãi, trường 'Kỳ'",
     "Tài liệu ghi: 'Trường hợp nhập = Bản ghi liền trước, hệ thống kiểm tra các ưu đãi trong cùng 1 kỳ phải có Kênh hoặc SPDV khác nhau'. Toán tử 'hoặc' ở đây nghĩa là: (a) Khác Kênh HOẶC khác SPDV (chỉ cần 1), hay (b) Khác cả Kênh VÀ SPDV?",
     "Toàn vẹn dữ liệu",
     "Đề xuất: Hiểu theo nghĩa (a): chỉ cần khác ít nhất 1 trong 2 (Kênh hoặc SPDV) là hợp lệ. Cần BA xác nhận."),

    # Hạng mục 4: UI/UX & Giao diện
    ("US11-QA-04.1",
     "Mockup UI, khu vực 'Bộ tham số (1)' và 'Bộ tham số (2)'",
     "Mockup hiển thị cột 'Tần suất' ở bảng Chi tiết ưu đãi trong Bộ tham số (2) nhưng không xuất hiện ở Bộ tham số (1). Đồng thời, cột 'Tần suất' KHÔNG được liệt kê trong bảng 'Mô tả chi tiết các trường'. Trường này là gì và có cần khai báo không?",
     "UI-UX",
     "Đề xuất: Đây có thể là lỗi mockup hoặc trường bị thiếu trong mô tả text. Cần BA xác nhận trường 'Tần suất' có tồn tại không, nếu có thì bổ sung vào bảng mô tả trường."),

    ("US11-QA-04.2",
     "Mockup UI, khu vực 'Thiết lập chương trình'",
     "Mockup hiển thị trường 'Loại ưu đãi' nằm ở vùng giữa form (giữa Ngày ban hành và Ngày hiệu lực), nhưng bảng mô tả trường liệt kê 'Loại ưu đãi' ở cuối nhóm THÔNG TIN CHUNG (sau Khối). Layout thực tế sẽ theo mockup hay theo bảng mô tả?",
     "UI-UX",
     "Đề xuất: Ưu tiên bố cục theo mockup (vì mockup là giao diện đã thiết kế). BA xác nhận vị trí chính thức."),

    ("US11-QA-04.3",
     "Mockup UI, khu vực Điều kiện đánh giá",
     "Mockup hiển thị cột '#' (STT dòng) ở bảng Điều kiện, nhưng bảng mô tả trường không liệt kê cột '#'. Cột STT này có tồn tại trên giao diện thực tế không?",
     "UI-UX",
     "Đề xuất: Cột '#' là STT tự động, không cần user nhập, chỉ hiển thị. Cần BA xác nhận."),

    ("US11-QA-04.4",
     "Bảng 'Mô tả chi tiết các trường', trường 'Email CBNV đầu mối'",
     "Bảng mô tả ghi ràng buộc: 'Có ký tự @' và 'Có domain pvcombank.com.vn'. Nhưng không rõ: (a) FE validate format email real-time hay chỉ khi nhấn Xác nhận? (b) Cho phép nhập nhiều email (ngăn cách bằng dấu phẩy/chấm phẩy) hay chỉ 1 email?",
     "UI-UX",
     "Đề xuất: (a) FE validate format email khi blur ra khỏi trường. (b) Cho phép 1 email duy nhất (vì tên trường dùng số ít 'Email CBNV đầu mối'). Cần BA xác nhận."),

    ("US11-QA-04.5",
     "Mục 'Chỉnh sửa CTƯĐ', Flowchart Chỉnh sửa, Bước 3.a",
     "Flowchart ghi Bước 3.a: 'Hiển thị màn hình chỉ cho phép chỉnh Ngày hết hiệu lực và chi tiết ưu đãi'. Tuy nhiên yêu cầu nghiệp vụ (text) ghi cho phép sửa cả 'Lịch đánh giá'. Flowchart bị thiếu 'Lịch đánh giá' so với text — đây là lệch pha giữa Flowchart và Text.",
     "UI-UX",
     "Đề xuất: Text yêu cầu nghiệp vụ là chuẩn. Flowchart Bước 3.a cần bổ sung 'Lịch đánh giá'. Cần BA xác nhận và sửa flowchart."),

    ("US11-QA-04.6",
     "Mục 'Chỉnh sửa CTƯĐ', yêu cầu nghiệp vụ, đoạn 'Chưa hiệu lực'",
     "Tài liệu ghi: 'cho phép chỉnh sửa toàn bộ thông tin trừ mã CTƯĐ và trường Số văn bản_Tên viết tắt CTƯĐ'. Nhưng Flowchart Bước 3.1 chỉ ghi: 'cho phép Chỉnh sửa toàn bộ CTƯĐ trừ Mã CTƯĐ' — thiếu đề cập 'Số VB_Tên viết tắt'. Lệch pha giữa Text và Flowchart.",
     "UI-UX",
     "Đề xuất: Text yêu cầu nghiệp vụ là chuẩn (trừ cả Mã CTƯĐ và Số VB_Tên viết tắt). Flowchart cần bổ sung. BA xác nhận."),
]

# Build table
HEADERS = ["ID", "Trích xuất (Reference)", "Câu hỏi / Sự cố", "Phân loại", "Đề xuất từ QA", "Trả lời của BA"]
COL_WIDTHS = [Cm(2.2), Cm(4.0), Cm(8.0), Cm(2.0), Cm(7.0), Cm(4.0)]

table = doc.add_table(rows=1 + len(qa_items), cols=6, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False

# Header
for i, h in enumerate(HEADERS):
    cell = table.cell(0, i)
    cell.text = h
    cell.width = COL_WIDTHS[i]
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(9)

# Data
for row_idx, (qid, ref, question, cat, suggestion) in enumerate(qa_items, 1):
    vals = [qid, ref, question, cat, suggestion, ""]
    for col_idx, val in enumerate(vals):
        cell = table.cell(row_idx, col_idx)
        cell.text = val
        cell.width = COL_WIDTHS[col_idx]
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(8)

import os
output_path = "/Users/mac/antigravity-testing-kit/Feature_02_SA_Tham_So_He_Thong/output/US11_PartB_QA.docx"
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f"✅ Saved: {output_path}")
print(f"📊 Total Q&A items: {len(qa_items)}")
