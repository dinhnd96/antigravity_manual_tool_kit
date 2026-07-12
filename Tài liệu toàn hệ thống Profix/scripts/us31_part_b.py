"""US31 - Part B: Danh sách Cảnh Báo & Q&A (Loopholes Discovery & BA Queries)"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# Page setup: Landscape, narrow margins
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

# === TITLE ===
title = doc.add_heading("US31 – Báo Cáo Tổng Doanh Thu Phí", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_heading("PHẦN B: DANH SÁCH CẢNH BÁO & Q&A", level=2)

# Column widths (total ~27cm for landscape narrow)
COL_WIDTHS = [Cm(2.2), Cm(4.5), Cm(6.0), Cm(1.8), Cm(5.5), Cm(5.5)]
HEADERS = ["ID", "Trích xuất (Reference)", "Câu hỏi / Sự cố", "Phân loại", "Đề xuất từ QA", "Trả lời của BA"]

# Q&A Data
qa_items = [
    # --- Hạng mục 1: Nghiệp vụ / Luồng xử lý ---
    ("US31-QA-01.1",
     "Mục \"Yêu cầu nghiệp vụ\", đoạn mô tả Button Tra cứu CIF và Bảng mô tả trường STT 4 (Tra cứu CIF)",
     "Tài liệu mô tả nút \"Tra cứu CIF\" tham chiếu QTC-09 (popup tra cứu CIF → chọn Mã CIF → auto-fill). "
     "Tuy nhiên, cả vùng Điều kiện tìm kiếm lẫn lưới danh sách đều KHÔNG có trường \"Mã CIF\". "
     "Vậy sau khi chọn CIF từ popup, giá trị CIF được fill vào đâu? "
     "Nút Tra cứu CIF phục vụ mục đích gì trên màn hình này?",
     "Nghiệp vụ",
     "Đề xuất BA làm rõ: (1) Bổ sung trường Mã CIF vào vùng Điều kiện tìm kiếm nếu cần lọc theo CIF, "
     "hoặc (2) Xác nhận loại bỏ nút Tra cứu CIF khỏi US31 nếu không cần thiết.",
     ""),
    ("US31-QA-01.2",
     "Mục \"Yêu cầu nghiệp vụ\", đoạn mô tả Button Tra cứu — \"hệ thống trả kết quả trên lưới danh sách dự thu phí dịch vụ kỳ tiếp theo\"",
     "Mô tả nút Tra cứu ghi: \"hệ thống trả kết quả trên lưới danh sách dự thu phí dịch vụ kỳ tiếp theo\". "
     "Tuy nhiên, US31 là \"Báo cáo tổng DOANH THU phí\" (doanh thu thực tế đã thu), "
     "không phải \"dự thu phí kỳ tiếp theo\" (US30). "
     "Đây có phải lỗi copy-paste từ tài liệu US30 không?",
     "Nghiệp vụ",
     "Đề xuất BA xác nhận: Nội dung đúng phải là \"lưới danh sách tổng doanh thu phí\", "
     "không phải \"dự thu phí dịch vụ kỳ tiếp theo\".",
     ""),
    ("US31-QA-01.3",
     "Bảng Diễn giải (Table 0), R3 và R4: Bước 3.1 → \"Nhập/Chọn điều kiện tìm kiếm\" nhưng Step ID ghi là 3.2",
     "Tại bảng Diễn giải, nhánh \"Nhập/Chọn thông tin tại các trường điều kiện tìm kiếm\" "
     "được ghi là Bước 3.1 trong mô tả nhánh, nhưng dòng tiếp theo (R4) lại ghi Step ID = 3.2. "
     "Đây có phải lỗi đánh số bước? Nội dung R4 mô tả \"Nhập/chọn thông tin\" giống nhánh 3.1, "
     "nhưng Step ID 3.2 lại trùng với nhánh \"Tải xuống\" ở R12.",
     "Nghiệp vụ",
     "Đề xuất: Sửa Step ID tại R4 từ \"3.2\" thành \"3.1\" để khớp với flowchart. "
     "Nhánh \"Tải xuống\" mới là 3.2.",
     ""),
    ("US31-QA-01.4",
     "Mục \"Yêu cầu nghiệp vụ\", đoạn \"Doanh thu phí (nguyên tệ)\" và \"Doanh thu phí (VND)\"",
     "Doanh thu phí (Nguyên tệ) mô tả là \"doanh thu phí nguyên tệ đã thu, chưa bao gồm VAT\". "
     "Doanh thu phí (VND) mô tả là \"doanh thu phí chuyển sang VND đã thu, chưa bao gồm VAT\". "
     "Câu hỏi: Khi Loại tiền = VND, cả 2 cột Nguyên tệ và VND sẽ hiển thị cùng một giá trị "
     "(Nguyên tệ VND = VND)? Hay cột Nguyên tệ không hiển thị khi loại tiền là VND?",
     "Nghiệp vụ",
     "Đề xuất: Khi Loại tiền = VND, cả 2 cột hiển thị cùng giá trị. "
     "Xin BA xác nhận hoặc đề xuất cách hiển thị phân biệt.",
     ""),
    ("US31-QA-01.5",
     "Mục \"Yêu cầu nghiệp vụ\", đoạn \"Tổng cộng\" (R28 Bảng mô tả trường)",
     "\"Tổng cộng\" mô tả = tổng tất cả Doanh thu phí VND. Câu hỏi: "
     "(1) Dòng Tổng cộng hiển thị ở đâu trên lưới — cuối trang hiện tại hay cuối toàn bộ dữ liệu? "
     "(2) Nếu dữ liệu phân trang (2+ trang), Tổng cộng là tổng của trang đang xem hay tổng toàn bộ? "
     "(3) Tổng cộng có hiển thị khi không có bản ghi nào không (= 0 hay ẩn)?",
     "Nghiệp vụ",
     "Đề xuất: Tổng cộng nên hiển thị ở footer cố định (sticky) và tính trên toàn bộ dữ liệu "
     "(không chỉ trang hiện tại). Khi không có bản ghi → hiển thị 0.",
     ""),
    ("US31-QA-01.6",
     "Mục \"Yêu cầu nghiệp vụ\", trường \"Từ ngày\" — \"Hệ thống tìm kiếm theo ngày giao dịch thu phí thực tế\"",
     "\"Từ ngày\" mô tả: \"Hệ thống tìm kiếm theo ngày giao dịch thu phí thực tế\". "
     "Nhưng \"Đến ngày\" cũng mô tả: \"Tìm kiếm theo field 'Ngày thu'\". "
     "\"Ngày giao dịch thu phí thực tế\" và \"Ngày thu\" có phải cùng một field dữ liệu không?",
     "Nghiệp vụ",
     "Đề xuất BA xác nhận: \"Ngày giao dịch thu phí thực tế\" = \"Ngày thu\" (cùng 1 field). "
     "Nếu đúng, nên thống nhất tên gọi trong tài liệu.",
     ""),
    ("US31-QA-01.7",
     "Bảng mô tả trường, STT 3 (Biểu phí) và STT 4 (Code phí) — liên quan đến STT 1 (Khối) và STT 2 (Mã Chi nhánh)",
     "Khi người dùng chọn Khối hoặc Chi nhánh, các dropdown Biểu phí và Code phí "
     "có tự động được lọc liên đới (cascade filter) theo Khối/Chi nhánh tương ứng không? "
     "Nếu không lọc, người dùng có thể chọn Code phí không thuộc Khối/Chi nhánh đã chọn, "
     "dẫn đến kết quả tra cứu rỗng hoặc sai lệch.",
     "Nghiệp vụ",
     "Đề xuất: Có lọc liên đới (cascade) để đảm bảo dữ liệu hợp lệ, "
     "tránh người dùng chọn sai Code phí/Biểu phí không thuộc Khối. "
     "Xin BA xác nhận.",
     ""),
    ("US31-QA-01.8",
     "Bảng mô tả trường, STT 18 (Từ ngày) và STT 19 (Đến ngày) trên lưới & Mockup UI",
     "Bảng mô tả trường ghi cột Từ ngày/Đến ngày trên lưới là 'Hiển thị ngày người dùng đã chọn' "
     "(nghĩa là mọi dòng sẽ hiển thị giống hệt ngày filter đầu vào). "
     "Tuy nhiên, trên Mockup UI, mỗi dòng lại hiển thị một khoảng thời gian KHÁC NHAU "
     "(ví dụ: dòng 37 = 08/03/2025 - 15/09/2025, dòng 38 = 28/03/2025 - 01/09/2025). "
     "Vậy thực chất 2 cột này lấy dữ liệu từ đâu? "
     "(1) Lặp lại ngày filter người dùng nhập (theo Text)? "
     "Hay (2) Ngày phát sinh giao dịch đầu/cuối của code phí đó (theo Mockup)?",
     "Nghiệp vụ",
     "Đề xuất: Làm rõ logic. Nếu chỉ lặp lại ngày filter thì 2 cột này dư thừa và có thể bỏ. "
     "Nếu đây là ngày phát sinh giao dịch đầu/cuối → cần sửa lại Text mô tả cho chính xác.",
     ""),

    # --- Hạng mục 2: Giới hạn hệ thống & Exception ---
    ("US31-QA-02.1",
     "Mục \"Yêu cầu nghiệp vụ\", toàn bộ trường Điều kiện tìm kiếm và Lưới kết quả",
     "Tài liệu không đề cập giới hạn khoảng thời gian tìm kiếm (Từ ngày – Đến ngày). "
     "Nếu user chọn khoảng thời gian rất lớn (ví dụ: 5 năm), "
     "có thể gây quá tải hệ thống hoặc timeout. "
     "Có giới hạn tối đa cho khoảng thời gian tra cứu không?",
     "Giới hạn",
     "Đề xuất: Giới hạn khoảng thời gian tối đa (ví dụ: 1 năm hoặc 365 ngày). "
     "Nếu vượt quá → FE hiển thị cảnh báo.",
     ""),
    ("US31-QA-02.2",
     "Mục \"Yêu cầu nghiệp vụ\", trường \"Doanh thu phí (VND)\" và \"Tổng cộng\"",
     "Với tỷ giá chuyển đổi, cột \"Doanh thu phí (VND)\" có thể đạt giá trị rất lớn. "
     "Hệ thống có giới hạn số tối đa cho cột Number này không "
     "(ví dụ: overflow khi vượt 999,999,999,999.99)? "
     "Tương tự, dòng \"Tổng cộng\" tổng hợp toàn bộ bản ghi có thể rất lớn.",
     "Giới hạn",
     "Đề xuất BA xác nhận giới hạn Number tối đa mà hệ thống hỗ trợ hiển thị, "
     "hoặc xác nhận không giới hạn (backend xử lý BigDecimal).",
     ""),

    # --- Hạng mục 3: Toàn vẹn dữ liệu ---
    ("US31-QA-03.1",
     "Mục \"Yêu cầu nghiệp vụ\", trường \"Doanh thu phí (VND)\" — \"chuyển sang VND\"",
     "Cột \"Doanh thu phí (VND)\" mô tả: \"doanh thu phí được chuyển sang VND\". "
     "Câu hỏi: Tỷ giá quy đổi được áp dụng tại thời điểm nào? "
     "(1) Tỷ giá tại ngày giao dịch thu phí? "
     "(2) Tỷ giá tại thời điểm tra cứu? "
     "(3) Tỷ giá cố định theo tham số hệ thống?",
     "Toàn vẹn dữ liệu",
     "Đề xuất: Sử dụng tỷ giá tại ngày giao dịch (đã lưu trong hệ thống khi thu phí) "
     "để đảm bảo tính nhất quán. Xin BA xác nhận.",
     ""),
    ("US31-QA-03.2",
     "Mục \"Yêu cầu nghiệp vụ\", lưới kết quả — Mức độ tổng hợp (Aggregation Level)",
     "Lưới kết quả hiển thị các cột: Mã Chi nhánh, Tên chi nhánh, Loại tiền, Loại tính phí, "
     "Biểu phí, Code phí, Tên phí. Câu hỏi: Dữ liệu được tổng hợp (group by) "
     "theo tổ hợp nào? Có phải mỗi dòng = 1 tổ hợp duy nhất "
     "(Chi nhánh + Loại tiền + Loại tính phí + Biểu phí + Code phí)? "
     "Hay mỗi dòng = 1 giao dịch thu phí riêng lẻ?",
     "Toàn vẹn dữ liệu",
     "Đề xuất BA xác nhận cấp độ tổng hợp. Nếu là group by → "
     "mỗi dòng đại diện cho tổng doanh thu của 1 tổ hợp. "
     "Nếu là chi tiết → mỗi dòng là 1 giao dịch.",
     ""),

    # --- Hạng mục 4: UI/UX ---
    ("US31-QA-04.1",
     "Mockup UI (image2.png), vùng Điều kiện tìm kiếm — trường \"Khối\" hiển thị dạng Combobox",
     "Tại Bảng mô tả trường, trường \"Khối\" ghi Định dạng = Combobox (chọn + search). "
     "Tuy nhiên, tài liệu mô tả Khối của user thuộc KHCN/KHDN/KHDNL sẽ \"không cho phép sửa\". "
     "Khi Combobox bị khóa (disabled), nó nên hiển thị như thế nào trên UI? "
     "Giữ Combobox nhưng disabled, hay đổi thành Text readonly?",
     "UI-UX",
     "Đề xuất: Giữ Combobox nhưng ở trạng thái disabled (greyed out, không cho thao tác). "
     "Đây là pattern phổ biến. Xin BA xác nhận.",
     ""),
    ("US31-QA-04.2",
     "Mockup UI (image2.png), lưới danh sách — cột \"Mã chi nhánh\" và header lưới",
     "Trên Mockup, lưới hiển thị các cột header: #, Mã chi nhánh, Tên chi nhánh, Loại tiền, "
     "Loại tính phí, Biểu phí, Code phí, Tên phí, Doanh thu phí (nguyên tệ), "
     "Doanh thu phí (VND), Từ ngày, Đến ngày. "
     "Mockup có cột \"#\" (số thứ tự tự động trên lưới). "
     "Bảng mô tả trường ghi STT 8 = \"STT\" (Text). "
     "Xác nhận: Cột # trên Mockup chính là cột STT trong bảng mô tả?",
     "UI-UX",
     "Đề xuất: Xác nhận cột # = STT. Cần thống nhất tên header (#, STT, hay Số thứ tự) "
     "trên Mockup và tài liệu.",
     ""),
    ("US31-QA-04.3",
     "Mockup UI (image2.png), dòng cuối cùng — \"Tổng cộng: 1.000.000 VND\"",
     "Trên Mockup, dòng \"Tổng cộng\" hiển thị ở footer dưới cùng, ngoài lưới (sticky footer). "
     "Bảng mô tả trường STT 20 ghi \"Tổng cộng\" là Number trên lưới. "
     "Câu hỏi: Tổng cộng là 1 dòng cuối trong lưới (row cuối) hay là 1 footer riêng biệt cố định?",
     "UI-UX",
     "Đề xuất: Dựa trên Mockup, \"Tổng cộng\" nên là footer cố định (sticky), "
     "không phải row trong lưới (để không bị mất khi cuộn trang). Xin BA xác nhận.",
     ""),
    ("US31-QA-04.4",
     "Bảng mô tả trường, trường \"Khối\" (STT 1) — Mô tả ghi \"KHDN hoặc KHDNL\"",
     "Tài liệu US31 liệt kê 3 khối bị khóa: KHCN, KHDN, KHDNL. "
     "Tuy nhiên, QTC-10 chỉ đề cập KHCN và KHDN (không có KHDNL). "
     "\"KHDNL\" (Khách hàng doanh nghiệp lớn?) có phải là tên mới/bổ sung so với QTC-10 không? "
     "Hay đây là lỗi đánh máy (KHDNL thực chất = KHDN)?",
     "UI-UX",
     "Đề xuất BA xác nhận: KHDNL là Khối riêng biệt hay là tên gọi khác của KHDN. "
     "Nếu là Khối mới, cần cập nhật QTC-10.",
     ""),
    ("US31-QA-04.5",
     "Mockup UI (image2.png) & Bảng mô tả nút STT 4 (Tra cứu CIF)",
     "Trong Mockup UI KHÔNG CÓ nút \"Tra cứu CIF\" (không xuất hiện cạnh nút Tra cứu/Xóa tra cứu "
     "hay nút Tải xuống). Tuy nhiên, Bảng mô tả nút chức năng (STT 4) lại yêu cầu có nút này. "
     "Đây là sự không đồng nhất giữa Mockup và Text.",
     "UI-UX",
     "Đề xuất: BA xác nhận có cần nút Tra cứu CIF không. Nếu có → cần bổ sung vào Mockup UI. "
     "Nếu không → xóa mô tả nút này khỏi bảng mô tả trường.",
     ""),
    ("US31-QA-04.6",
     "Mockup UI (image2.png), lưới danh sách — thiếu cột \"Khối\"",
     "Trường \"Khối\" là điều kiện lọc quan trọng. User Hội sở (không thuộc KHCN/KHDN/KHDNL) "
     "có thể tra cứu tất cả Khối. Tuy nhiên, trên lưới kết quả KHÔNG CÓ cột \"Khối\", "
     "khiến user khó nhận biết dòng doanh thu thuộc Khối nào nếu chỉ dựa vào Mã chi nhánh.",
     "UI-UX",
     "Đề xuất: Bổ sung cột \"Khối\" vào lưới hiển thị kết quả báo cáo, "
     "đặt trước cột \"Mã Chi nhánh\" để user dễ phân biệt.",
     ""),
]

# Create table
table = doc.add_table(rows=1 + len(qa_items), cols=6, style="Table Grid")
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False

# Set column widths
for row in table.rows:
    for i, width in enumerate(COL_WIDTHS):
        row.cells[i].width = width

# Header row
for i, h in enumerate(HEADERS):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(9)

# Add category headers as merged rows
def add_category_row(table, row_idx, category_text):
    """Style the first cell of a category row."""
    pass  # Categories are embedded in data, we'll add them as group separators

# Data rows
current_category = ""
for idx, (qid, ref, question, cat, suggestion, ba_answer) in enumerate(qa_items):
    row = table.rows[idx + 1]
    data = [qid, ref, question, cat, suggestion, ba_answer]
    for i, val in enumerate(data):
        cell = row.cells[i]
        cell.text = val
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(8.5)

out_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "US31_PartB_QA.docx")
doc.save(out_path)
print(f"✅ Đã tạo: {out_path}")
