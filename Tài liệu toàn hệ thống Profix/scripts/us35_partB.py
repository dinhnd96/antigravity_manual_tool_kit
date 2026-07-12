"""US35 Part B - Danh Sách Cảnh Báo & Q&A (Dành cho BA)"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# Page setup: Landscape, narrow margins
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(10)

# ============ TITLE ============
title = doc.add_heading('US35 – Phần B: Danh Sách Cảnh Báo & Q&A', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Tính năng: Tự động thu các loại phí định kỳ cần thu theo lịch')
doc.add_paragraph('Dành cho: Business Analyst | Phiên bản: v1.0')
doc.add_paragraph('')

# ============ Q&A DATA ============
# Columns: ID, Trích xuất, Câu hỏi/Sự cố, Phân loại, Đề xuất từ QA, Trả lời BA
qa_data = [
    # === Hạng mục 1: Nghiệp vụ / Luồng xử lý ===
    (
        "US35-QA-01.1",
        'Mục "Yêu cầu nghiệp vụ", đoạn mô tả TK thu phí mặc định',
        'Khi đối tượng tính phí = Customer, nếu TK thu phí mặc định không đủ ĐK VÀ không tìm được TK thay thế nào đủ ĐK trong danh sách TK của KH → khoản phí này xử lý như thế nào? Tài liệu chỉ mô tả "tìm TK thay thế" nhưng không nêu nhánh thất bại khi không có TK nào hợp lệ.',
        'Nghiệp vụ',
        'Đề xuất: Khoản phí không được ghi vào Topic Kafka, trạng thái = "Chưa thanh toán" hoặc ghi log lỗi riêng, chờ truy thu/tận thu sau. Cần BA xác nhận hành vi cụ thể.',
        ''
    ),
    (
        "US35-QA-01.2",
        'Mục "Yêu cầu nghiệp vụ", đoạn "Trong khoảng thời gian ProfiX chưa nhận được kết quả thu phí"',
        'Trạng thái "Đang xử lý" được duy trì cho đến khi nhận kết quả từ Kafka. Nếu Kafka message bị mất (lost message) hoặc T24 không phản hồi vô thời hạn → có cơ chế timeout/retry không? Khoản phí sẽ mãi ở trạng thái "Đang xử lý" hay có threshold time tự động chuyển trạng thái?',
        'Nghiệp vụ',
        'Đề xuất: Cần có cơ chế timeout (VD: sau 24h không nhận kết quả → tự động chuyển "Chưa thanh toán" + ghi log). Hoặc có cơ chế retry gửi lại message. Cần BA xác nhận.',
        ''
    ),
    (
        "US35-QA-01.3",
        'Bảng "Xử lý tính phí định kỳ", Bước 1',
        'Với đối tượng tính phí = Account/Card, bước 2 (kiểm tra TK thu phí mặc định) được bỏ qua → phí được thu trực tiếp từ TK/Thẻ đó. Nhưng TK đó có cần kiểm tra trạng thái (Hoạt động/Tạm ngừng…) và sản phẩm thuộc CA_PRODUCT không? Tài liệu chỉ mô tả điều kiện TK cho trường hợp Customer, không rõ Account/Card có áp dụng tương tự hay không.',
        'Nghiệp vụ',
        'Đề xuất: Đối với Account → kiểm tra trạng thái TK + CA_PRODUCT tương tự Customer. Đối với Card → kiểm tra trạng thái thẻ (Hoạt động?) + TK gắn với thẻ phải hợp lệ. Cần BA xác nhận.',
        ''
    ),
    (
        "US35-QA-01.4",
        'Mục "Yêu cầu nghiệp vụ", đoạn "thứ tự ưu tiên nghiệp vụ phí"',
        'Tài liệu nêu "hệ thống cho phép cài đặt thứ tự ưu tiên của nghiệp vụ phí cần thu" nhưng không mô tả: (a) Ai cài đặt thứ tự ưu tiên này? (b) Ở màn hình/module nào? (c) Nếu 2 Code phí cùng ưu tiên thì thứ tự ghi Kafka theo tiêu chí nào?',
        'Nghiệp vụ',
        'Đề xuất: Thứ tự ưu tiên được cài đặt ở cấu hình Job hoặc tham số hệ thống. Nếu cùng ưu tiên → sắp xếp theo Mã Code phí ASC. Cần BA xác nhận.',
        ''
    ),
    (
        "US35-QA-01.5",
        'Bảng "Xử lý tính phí định kỳ", Bước 1, mô tả Account/Card',
        'Với đối tượng tính phí = Account: tài liệu yêu cầu kiểm tra "(1) TK thỏa ĐK theo TK và (2) CIF của Account thỏa ĐK theo KH". Tương tự với Card. Nếu TK thỏa ĐK (1) nhưng CIF không thỏa ĐK (2) → khoản phí này có bị bỏ qua hoàn toàn hay chỉ bỏ ĐK khách hàng? Logic AND hay OR giữa 2 ĐK chưa rõ.',
        'Nghiệp vụ',
        'Đề xuất: Phải thỏa mãn CẢ HAI điều kiện (AND). Nếu CIF không thỏa → khoản phí không được sinh. Cần BA xác nhận.',
        ''
    ),
    (
        "US35-QA-01.6",
        'Mục "Yêu cầu nghiệp vụ", đoạn "thu định kỳ hàng tháng vào một ngày cố định"',
        'Đối với Customer/Account: phí thu hàng tháng vào "ngày cố định". Nếu ngày cố định là ngày 31 nhưng tháng chỉ có 28/29/30 ngày → Job chạy vào ngày nào? Tài liệu không mô tả edge case ngày không tồn tại.',
        'Nghiệp vụ',
        'Đề xuất: Nếu ngày cố định > số ngày trong tháng → chạy vào ngày cuối tháng (VD: ngày 31 → chạy ngày 28/29/30 tùy tháng). Cần BA xác nhận logic xử lý.',
        ''
    ),
    (
        "US35-QA-01.7",
        'Bảng "Xử lý tính phí định kỳ", Bước 5.2, logic chọn CTƯĐ ưu tiên',
        'Khi có nhiều CTƯĐ, tài liệu nêu "Lấy CTƯĐ có Số tiền ƯĐ lớn nhất. Nếu bằng nhau thì lấy CTƯĐ có hiệu lực xa nhất." Nếu cả Số tiền ƯĐ lẫn ngày hiệu lực đều bằng nhau → lấy CTƯĐ nào? Thiếu tiebreak thứ 3.',
        'Nghiệp vụ',
        'Đề xuất: Nếu vẫn bằng nhau → lấy CTƯĐ có Mã nhỏ nhất (ASC) hoặc lấy ngẫu nhiên. Cần BA xác nhận.',
        ''
    ),
    (
        "US35-QA-01.8",
        'Mục "Yêu cầu nghiệp vụ", đoạn trạng thái khoản phí "Xóa nợ"',
        'Trạng thái "Xóa nợ" — Đã miễn giảm, không theo dõi nợ. Ai có quyền thực hiện thao tác xóa nợ? Qua màn hình nào? Có cần luồng Maker-Checker không? US35 chỉ liệt kê trạng thái mà không mô tả luồng nghiệp vụ dẫn đến trạng thái này.',
        'Nghiệp vụ',
        'Đề xuất: Luồng xóa nợ được mô tả tại US khác (truy thu/tận thu nợ phí). Nếu đúng → cần BA xác nhận tên US liên quan để Tester trace.',
        ''
    ),
    (
        "US35-QA-01.9",
        'Bảng "Diễn giải lưu đồ", Bước 2, Core T24 xử lý',
        'Core T24 "hạch toán thu phí, lưu ý tận thu (thu một phần phí nếu TK không đủ số dư)". Nhưng tài liệu không mô tả logic tận thu cụ thể: (a) Tận thu toàn bộ số dư còn lại hay giữ lại một khoản tối thiểu? (b) Nếu số dư = 0 thì kết quả trả về là "Thanh toán một phần" (0 đồng) hay "Chưa thanh toán"?',
        'Nghiệp vụ',
        'Đề xuất: Nếu số dư = 0 → "Chưa thanh toán". Nếu 0 < số dư < phí cần thu → "Thanh toán một phần" (thu hết phần có). Cần BA xác nhận logic tận thu của T24.',
        ''
    ),

    # === Hạng mục 2: Giới hạn hệ thống & Exception ===
    (
        "US35-QA-02.1",
        'Mục "Yêu cầu nghiệp vụ", đoạn "ghi danh sách vào Topic Kafka"',
        'Nếu Job A sinh ra khối lượng rất lớn khoản phí (VD: hàng triệu KH) và ghi vào Topic Kafka cùng lúc → có giới hạn batch size hoặc rate limit không? Kafka consumer (T24) có xử lý kịp không? Tài liệu không đề cập volume/throughput constraint.',
        'Giới hạn',
        'Đề xuất: Cần BA/SA xác nhận batch size tối đa ghi vào Kafka mỗi lần, và cơ chế phân batch nếu volume quá lớn.',
        ''
    ),
    (
        "US35-QA-02.2",
        'Bảng "Xử lý tính phí định kỳ", Bước 3.3, quy đổi tỷ giá',
        'Khi cần quy đổi tỷ giá (bước 3.3 và 5.3), nếu ProfiX chưa nhận được dữ liệu tỷ giá từ Core (dữ liệu tỷ giá chưa đồng bộ hoặc lỗi đồng bộ) → hệ thống xử lý thế nào? Dùng tỷ giá cũ nhất có sẵn? Bỏ qua khoản phí? Dừng toàn bộ Job?',
        'Giới hạn',
        'Đề xuất: Nếu không có tỷ giá → bỏ qua khoản phí đó + ghi log lỗi, không dừng toàn bộ Job. Cần BA xác nhận.',
        ''
    ),
    (
        "US35-QA-02.3",
        'Bảng "Xử lý tính phí định kỳ", Bước 6, tính VAT',
        'Bước 6 tính VAT: "Code phí có VAT = """ (phí không có VAT) → "để trống VAT trong response". Nếu Code phí KHÔNG cấu hình trường VAT (null/undefined, khác với chuỗi rỗng "") → hệ thống xử lý giống "không có VAT" hay báo lỗi cấu hình?',
        'Giới hạn',
        'Đề xuất: Nếu VAT = null/undefined → mặc định xử lý giống không có VAT (để trống). Cần BA xác nhận.',
        ''
    ),
    (
        "US35-QA-02.4",
        'Mục "Yêu cầu nghiệp vụ", luồng batch job tổng thể',
        'Nếu Job A đang chạy giữa chừng và gặp lỗi hệ thống (crash, OOM, network failure) → các khoản phí đã ghi vào Kafka sẽ bị xử lý trùng lặp (duplicate) khi Job retry không? Có cơ chế idempotency/deduplication không?',
        'Giới hạn',
        'Đề xuất: Cần cơ chế idempotency key (VD: mã khoản phí + ngày T) để tránh thu phí trùng lặp khi retry. Cần BA/SA xác nhận.',
        ''
    ),
    (
        "US35-QA-02.5",
        'Bảng "Xử lý tính phí định kỳ", Bước 3.3-3.4',
        'Nếu Code phí KHÔNG khai báo Số tiền tối thiểu/tối đa (cả hai đều trống) → bước 3.3-3.4 được bỏ qua, Số tiền phí cần thu = Số tiền phí đã tính. Nhưng nếu chỉ khai báo 1 trong 2 (VD: chỉ có Min, không có Max, hoặc ngược lại) → logic clamping xử lý thế nào?',
        'Giới hạn',
        'Đề xuất: Nếu chỉ có Min → chỉ so sánh Min. Nếu chỉ có Max → chỉ so sánh Max. Tương tự cho bước 5.3-5.4 (CTƯĐ). Cần BA xác nhận.',
        ''
    ),

    # === Hạng mục 3: Toàn vẹn dữ liệu & Ràng buộc ===
    (
        "US35-QA-03.1",
        'Mục "Yêu cầu nghiệp vụ", đoạn "dữ liệu đồng bộ T-1"',
        'Dữ liệu KH/TK/Thẻ được đồng bộ T-1. Nếu trong khoảng T-1 đến đầu ngày T, một TK bị đóng hoặc KH bị vô hiệu hóa trên Core nhưng dữ liệu ProfiX chưa cập nhật → ProfiX vẫn sẽ sinh phí cho TK/KH đã không còn hợp lệ. T24 xử lý thế nào khi nhận yêu cầu thu phí cho TK đã đóng?',
        'Toàn vẹn dữ liệu',
        'Đề xuất: T24 sẽ trả kết quả = "Chưa thanh toán" với lý do TK không hợp lệ. ProfiX cập nhật trạng thái. Cần BA xác nhận luồng xử lý.',
        ''
    ),
    (
        "US35-QA-03.2",
        'Bảng "Xử lý tính phí định kỳ", Bước 2, TK thay thế',
        'TK thay thế "chỉ được sử dụng để ghi thông tin phí định kỳ đúng hạn cần thu vào Topic, không áp dụng cho truy thu/tận thu nợ phí." Nếu khoản phí từ TK thay thế → kết quả = "Thanh toán một phần" hoặc "Chưa thanh toán" → nợ phí theo dõi trên TK nào? TK mặc định (gốc) hay TK thay thế?',
        'Toàn vẹn dữ liệu',
        'Đề xuất: Nợ phí được theo dõi trên TK mặc định (gốc) của KH, không phải TK thay thế. Khi truy thu/tận thu, hệ thống sẽ kiểm tra lại TK mặc định. Cần BA xác nhận.',
        ''
    ),
    (
        "US35-QA-03.3",
        'Mục "Yêu cầu nghiệp vụ", đoạn trạng thái "Thêm mới" và "Đang xử lý"',
        'Khoản phí ở trạng thái "Thêm mới" (chưa gửi) → nếu Job chạy lại (retry/reschedule) → có risk sinh trùng khoản phí? Khoản phí đã ở "Đang xử lý" → nếu Job chạy lại → có bỏ qua các khoản đang xử lý không? Tài liệu không mô tả cơ chế kiểm tra trạng thái khi Job chạy lại.',
        'Toàn vẹn dữ liệu',
        'Đề xuất: Khi Job chạy lại → kiểm tra: nếu đã tồn tại khoản phí cùng Code phí + cùng KH/TK/Thẻ + cùng kỳ thu → bỏ qua, không sinh trùng. Cần BA xác nhận.',
        ''
    ),
    (
        "US35-QA-03.4",
        'Bảng "Xử lý tính phí định kỳ", Bước 4, Danh sách CTƯĐ',
        'CTƯĐ có "ngưỡng ưu đãi" — KH chưa chạm ngưỡng thì được ưu đãi. Thời điểm kiểm tra ngưỡng là khi nào? Nếu Job A chạy cho KH_1 và KH_1 vừa chạm ngưỡng CTƯĐ do khoản phí trước đó trong cùng Job A → khoản phí sau đó có còn được ưu đãi không? (Race condition trong cùng 1 batch)',
        'Toàn vẹn dữ liệu',
        'Đề xuất: Ngưỡng ưu đãi kiểm tra tại thời điểm bắt đầu Job (snapshot). Trong cùng 1 batch, ngưỡng không được cập nhật giữa chừng. Cần BA xác nhận.',
        ''
    ),

    # === Hạng mục 4: UI/UX & Giao diện ===
    (
        "US35-QA-04.1",
        'Mục "Giao diện" và "Mô tả chi tiết các trường"',
        'Cả mục Giao diện và Mô tả chi tiết các trường đều ghi "N/A". Tuy nhiên, mục Yêu cầu nghiệp vụ mô tả 6 trạng thái khoản phí + lịch sử thu phí → liệu có màn hình Tra cứu/Xem danh sách khoản phí định kỳ riêng không? Hay dữ liệu chỉ hiển thị ở các US khác (Lịch sử thu phí, Nợ phí)?',
        'UI-UX',
        'Đề xuất: Cần BA xác nhận: (a) US35 thuần backend, không có UI; (b) Dữ liệu khoản phí định kỳ được xem tại US nào (VD: US Lịch sử thu phí, US Tra cứu nợ phí).',
        ''
    ),
    (
        "US35-QA-04.2",
        'Lưu đồ (Flowchart) — Hình ảnh trong tài liệu',
        'Flowchart chỉ mô tả 4 bước ở mức high-level (Sinh dữ liệu → Kafka → T24 hạch toán → Kafka kết quả → ProfiX update). Không có flowchart chi tiết cho logic tính phí 6 bước trong bảng "Xử lý tính phí định kỳ". Các nhánh rẽ (TK không hợp lệ, không có CTƯĐ, nhiều CTƯĐ…) không được thể hiện trên flowchart.',
        'UI-UX',
        'Đề xuất: Đây là ghi nhận về độ chi tiết flowchart. Không bắt buộc BA bổ sung nhưng lưu ý Tester cần đọc bảng mô tả chi tiết để nắm đầy đủ logic, không dựa vào flowchart.',
        ''
    ),
]

# ============ BUILD TABLE ============
doc.add_heading('Bảng Q&A — Danh sách vấn đề cần BA xác nhận', level=2)
doc.add_paragraph('Ghi chú: Mỗi dòng là 1 vấn đề riêng biệt. Cột "Trả lời của BA" để trống để BA điền.')
doc.add_paragraph('')

# Add category headers within table
table = doc.add_table(rows=1, cols=6)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Set column widths
col_widths = [Cm(2.5), Cm(4.0), Cm(8.0), Cm(2.0), Cm(6.0), Cm(4.5)]
for i, width in enumerate(col_widths):
    table.columns[i].width = width

# Header row
headers = ['ID', 'Trích xuất', 'Câu hỏi / Sự cố', 'Phân loại', 'Đề xuất từ QA', 'Trả lời của BA']
hdr_cells = table.rows[0].cells
for i, h in enumerate(headers):
    hdr_cells[i].text = h
    for paragraph in hdr_cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)

# Category labels
current_category = None
category_map = {
    '01': '🔶 Hạng mục 1: Vấn đề Nghiệp vụ / Luồng xử lý',
    '02': '🔴 Hạng mục 2: Giới hạn hệ thống & Exception Handling',
    '03': '🟠 Hạng mục 3: Toàn vẹn dữ liệu & Ràng buộc',
    '04': '🔵 Hạng mục 4: UI/UX & Giao diện',
}

for qa in qa_data:
    qa_id = qa[0]
    cat_code = qa_id.split('-QA-')[1][:2]
    
    if cat_code != current_category:
        current_category = cat_code
        cat_row = table.add_row().cells
        cat_label = category_map.get(cat_code, '')
        cat_row[0].merge(cat_row[5])
        cat_row[0].text = cat_label
        for paragraph in cat_row[0].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 128)
    
    row_cells = table.add_row().cells
    for i, val in enumerate(qa):
        row_cells[i].text = val
        for paragraph in row_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

# Set cell widths
for row in table.rows:
    for i, cell in enumerate(row.cells):
        cell.width = col_widths[i]

# Summary
doc.add_paragraph('')
doc.add_paragraph(f'Tổng số câu hỏi: {len(qa_data)}')
doc.add_paragraph('• Hạng mục 1 (Nghiệp vụ): 9 câu')
doc.add_paragraph('• Hạng mục 2 (Giới hạn): 5 câu')
doc.add_paragraph('• Hạng mục 3 (Toàn vẹn dữ liệu): 4 câu')
doc.add_paragraph('• Hạng mục 4 (UI/UX): 2 câu')

# Save
output_dir = '/Users/mac/antigravity-testing-kit/Feature_02_SA_Tham_So_He_Thong/output'
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, 'US35_PartB_QA.docx')
doc.save(output_path)
print(f'✅ Saved: {output_path}')
