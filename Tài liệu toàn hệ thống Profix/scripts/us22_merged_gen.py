import os, sys
sys.path.insert(0, os.path.dirname(__file__))
from us22_merged_data import *
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def setup_doc():
    doc = Document()
    for s in doc.sections:
        s.orientation = WD_ORIENT.LANDSCAPE
        s.page_width, s.page_height = s.page_height, s.page_width
        s.top_margin = Cm(1.27); s.bottom_margin = Cm(1.27)
        s.left_margin = Cm(1.27); s.right_margin = Cm(1.27)
    style = doc.styles['Normal']
    style.font.name = 'Arial'; style.font.size = Pt(10)
    return doc

def set_cell(cell, text, bold=False, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.font.name = 'Arial'

def shade_row(row, color="D9E2F3"):
    for c in row.cells:
        tc = c._tc; pr = tc.get_or_add_tcPr()
        s = pr.makeelement(qn('w:shd'), {qn('w:fill'): color, qn('w:val'): 'clear'})
        pr.append(s)

def main():
    doc = setup_doc()
    h = doc.add_heading(f"PHẦN B TỔNG HỢP (MERGED): Q&A – {FEATURE}", level=1)
    for r in h.runs: r.font.color.rgb = RGBColor(0,0,0)
    doc.add_paragraph("Tổng hợp từ: AI Part B + VA Part B | Đã loại trùng + kiểm tra QTC | Ngày: 2026-05-18")
    doc.add_paragraph("Cột \"Nguồn\" cho biết câu hỏi phát hiện bởi AI, VA, hoặc cả hai (AI+VA).")

    t = doc.add_table(rows=1, cols=7)
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["ID", "Trích xuất", "Câu hỏi / Sự cố", "Phân loại", "Đề xuất từ QC", "Trả lời của BA", "Nguồn"]
    widths = [Cm(2), Cm(5.5), Cm(5.5), Cm(1.8), Cm(5), Cm(3), Cm(2)]
    for i, h in enumerate(headers):
        set_cell(t.rows[0].cells[i], h, True, 9)
        t.rows[0].cells[i].width = widths[i]
    shade_row(t.rows[0])

    cat_colors = {"Nghiệp vụ": "FFF2CC", "Giới hạn": "FCE4EC", "Toàn vẹn dữ liệu": "E8F5E9", "UI-UX": "E3F2FD"}
    cat_labels = {"Nghiệp vụ": "🔶 Hạng mục 1: Vấn đề Nghiệp vụ / Luồng xử lý",
                  "Giới hạn": "🔴 Hạng mục 2: Giới hạn hệ thống & Exception Handling",
                  "Toàn vẹn dữ liệu": "🟠 Hạng mục 3: Toàn vẹn dữ liệu & Ràng buộc",
                  "UI-UX": "🔵 Hạng mục 4: UI/UX & Giao diện"}
    prev_cat = ""
    for qid, ref, q, cat, sug, src in QA_MERGED:
        if cat != prev_cat:
            cr = t.add_row()
            merged = cr.cells[0].merge(cr.cells[6])
            set_cell(merged, cat_labels.get(cat, cat), True, 10)
            shade_row(cr, cat_colors.get(cat, "FFFFFF"))
            prev_cat = cat
        row = t.add_row()
        for i, v in enumerate([qid, ref, q, cat, sug, "", src]):
            set_cell(row.cells[i], v, False, 8)
            row.cells[i].width = widths[i]

    path = os.path.join(OUTPUT_DIR, MERGED_FILE)
    doc.save(path)
    print(f"Saved: {path}")

if __name__ == "__main__":
    main()
