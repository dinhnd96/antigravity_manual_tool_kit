#!/usr/bin/env python3
"""US06 Part C: Generate Test Case Coverage docx from data."""
import sys, os
sys.path.insert(0, os.path.dirname(__file__))
from us06_partc_data import DATA

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()
for s in doc.sections:
    s.orientation = WD_ORIENT.LANDSCAPE
    s.page_width, s.page_height = Cm(29.7), Cm(21.0)
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Cm(1.27)

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(9)

doc.add_heading("US06 – Phần C: Bảng Tổng Hợp Test Case Đề Xuất (Test Case Coverage)", level=1)
doc.add_paragraph(
    "Tài liệu gốc: US06-v2 – Khai báo Biểu phí\n"
    "Phiên bản: v2 (đã tích hợp BA feedback)\n"
    f"Tổng số kịch bản: {len(DATA)} | Tổng TC ước tính: {sum(d[5] for d in DATA)}"
)

# Group by test type
type_order = [
    ("Happy Path", "🟢 Happy Path (Positive Cases – Luồng cơ bản)"),
    ("Negative Path", "🔴 Negative Path & Exception Handling"),
    ("Boundary Value", "📐 Boundary Value Analysis (Giá trị biên)"),
    ("Field Validation", "🎨 UI/UX & Field Validation"),
    ("Business Logic", "🧠 Business Logic & State Transition"),
    ("Data Integrity", "🔗 Data Integrity & Integration"),
    ("NFR", "⚡ NFR (Non-Functional Requirements)"),
]

headers = ["Mã SC", "Feature", "Module", "Loại TC", "Tên Kịch Bản", "Số TC", "Trích dẫn tài liệu"]
col_widths = [Cm(1.5), Cm(2.5), Cm(3.5), Cm(2), Cm(8), Cm(1.2), Cm(8.5)]

for type_key, type_label in type_order:
    items = [d for d in DATA if d[3] == type_key]
    if not items:
        continue
    
    doc.add_heading(type_label, level=2)
    
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # Header
    hdr = table.rows[0]
    for i, (h, w) in enumerate(zip(headers, col_widths)):
        cell = hdr.cells[i]
        cell.width = w
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(8)
        shading = cell._element.get_or_add_tcPr()
        elm = shading.makeelement(qn('w:shd'), {qn('w:fill'): 'D9E2F3', qn('w:val'): 'clear'})
        shading.append(elm)
    
    # Data rows
    for item in items:
        row = table.add_row()
        vals = [item[0], item[1], item[2], item[3], item[4], str(item[5]), item[6]]
        for i, val in enumerate(vals):
            cell = row.cells[i]
            cell.width = col_widths[i]
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(8)
            if i == 5:  # center align count
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtotal
    sub = sum(d[5] for d in items)
    doc.add_paragraph(f"  Subtotal: {len(items)} kịch bản, {sub} TC ước tính", style='List Bullet')

# Summary
doc.add_heading("Tổng Kết", level=2)
summary_data = []
for type_key, type_label in type_order:
    items = [d for d in DATA if d[3] == type_key]
    if items:
        summary_data.append((type_label.split("(")[0].strip(), len(items), sum(d[5] for d in items)))

st = doc.add_table(rows=1+len(summary_data)+1, cols=3)
st.style = 'Table Grid'
for i, h in enumerate(["Nhóm", "Số kịch bản", "Số TC ước tính"]):
    c = st.rows[0].cells[i]
    c.paragraphs[0].add_run(h).bold = True
    c.paragraphs[0].runs[0].font.size = Pt(9)

for idx, (name, sc, tc) in enumerate(summary_data, 1):
    st.rows[idx].cells[0].paragraphs[0].add_run(name).font.size = Pt(9)
    st.rows[idx].cells[1].paragraphs[0].add_run(str(sc)).font.size = Pt(9)
    st.rows[idx].cells[2].paragraphs[0].add_run(str(tc)).font.size = Pt(9)

# Total row
total_sc = sum(s[1] for s in summary_data)
total_tc = sum(s[2] for s in summary_data)
last = st.rows[-1]
last.cells[0].paragraphs[0].add_run("TỔNG CỘNG").bold = True
last.cells[1].paragraphs[0].add_run(str(total_sc)).bold = True
last.cells[2].paragraphs[0].add_run(str(total_tc)).bold = True

out = os.path.join(os.path.dirname(os.path.dirname(__file__)), "US06_Part_C_TestCase_Coverage.docx")
doc.save(out)
print(f"✅ Part C saved: {out}")
print(f"📊 Tổng: {total_sc} kịch bản, {total_tc} TC ước tính")
