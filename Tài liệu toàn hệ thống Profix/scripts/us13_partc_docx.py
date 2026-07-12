# -*- coding: utf-8 -*-
"""US13 Part C — Script sinh file .docx tổng hợp (Phần A update + Phần C Test Coverage)"""
import sys
sys.path.insert(0, '.')
from us13_partc_data import DATA, EXCLUSIONS, PENDING

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUTPUT = "US13_PartC_TestCoverage.docx"

doc = Document()

# Page setup: Landscape, Narrow margins
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

# Style
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(9)

# ==================== PHẦN C ====================
doc.add_heading("Phần C: Bảng Tổng Hợp Test Case Đề Xuất (Test Case Coverage)", level=1)
p = doc.add_paragraph()
p.add_run("Feature: ").bold = True
p.add_run("Khai báo CTƯĐ không xác định sẵn danh sách khách hàng áp dụng")
p = doc.add_paragraph()
p.add_run("Phiên bản: ").bold = True
p.add_run("v2.0 | Ngày phân tích: 26/05/2026")
p = doc.add_paragraph()
p.add_run(f"Tổng số Test Case: {len(DATA)} SC").bold = True

# Table
headers = ["Mã SC", "Feature", "Module", "Loại Test Case", "Tên Test Case / Kịch bản", "Trích dẫn tài liệu (Traceability)"]
col_widths = [Cm(1.5), Cm(3.5), Cm(3.5), Cm(2.8), Cm(8.5), Cm(7.0)]

table = doc.add_table(rows=1, cols=len(headers))
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
hdr = table.rows[0]
for i, h in enumerate(headers):
    cell = hdr.cells[i]
    cell.width = col_widths[i]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Blue background
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): '2E75B6'
    })
    shading.append(shading_elm)

# Data rows
for row_data in DATA:
    row = table.add_row()
    for i, val in enumerate(row_data):
        cell = row.cells[i]
        cell.width = col_widths[i]
        p = cell.paragraphs[0]
        run = p.add_run(str(val))
        run.font.size = Pt(8)

# ==================== DANH SÁCH SC LOẠI TRỪ ====================
doc.add_heading("Danh sách SC loại trừ theo phản hồi BA", level=2)
if EXCLUSIONS:
    ex_table = doc.add_table(rows=1, cols=2)
    ex_table.style = 'Table Grid'
    hdr = ex_table.rows[0]
    for i, h in enumerate(["QA ID", "Lý do loại trừ"]):
        p = hdr.cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(8)
        shading = hdr.cells[i]._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): 'FFC000'
        })
        shading.append(shading_elm)
    for qa_id, reason in EXCLUSIONS:
        row = ex_table.add_row()
        row.cells[0].paragraphs[0].add_run(qa_id).font.size = Pt(8)
        row.cells[1].paragraphs[0].add_run(reason).font.size = Pt(8)
else:
    doc.add_paragraph("Không có SC nào bị loại trừ.")

# ==================== DANH SÁCH QA PENDING ====================
doc.add_heading("Danh sách QA chờ BA cập nhật (Pending)", level=2)
if PENDING:
    pd_table = doc.add_table(rows=1, cols=2)
    pd_table.style = 'Table Grid'
    hdr = pd_table.rows[0]
    for i, h in enumerate(["QA ID", "Mô tả & Lý do Pending"]):
        p = hdr.cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(8)
        shading = hdr.cells[i]._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): 'FF6B6B'
        })
        shading.append(shading_elm)
    for qa_id, desc in PENDING:
        row = pd_table.add_row()
        row.cells[0].paragraphs[0].add_run(qa_id).font.size = Pt(8)
        row.cells[1].paragraphs[0].add_run(desc).font.size = Pt(8)
else:
    doc.add_paragraph("Không có QA nào đang Pending.")

doc.save(OUTPUT)
print(f"✅ Đã tạo file: {OUTPUT} ({len(DATA)} TC)")
