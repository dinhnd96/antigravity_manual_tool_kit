"""US32 Part B - Q&A Generator"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# Page setup
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

title = doc.add_heading('US32 — PHẦN B: DANH SÁCH CẢNH BÁO & Q&A', level=1)
doc.add_paragraph('Feature: Dashboard — Báo cáo tổng quan hoạt động thu phí')
doc.add_paragraph('Phiên bản: v1.0 | Ngày: 14/05/2026 | Người phân tích: AI QA Lead')
doc.add_paragraph('─' * 80)

# Table setup
headers = ['ID', 'Trích xuất (Reference)', 'Câu hỏi / Sự cố', 'Phân loại', 'Đề xuất từ QA', 'Trả lời của BA']
col_widths = [Cm(2.0), Cm(4.5), Cm(7.0), Cm(2.5), Cm(6.0), Cm(5.0)]

# Q&A Data
qa_data = [
    # === HẠ MỤC 1: Nghiệp vụ / Luồng xử lý ===
    (
        'US32-QA-01.1',
        'Mục "Yêu cầu nghiệp vụ", đoạn "Khi vào các màn hình báo cáo tổng quan, hệ thống dựa trên thông tin khối của người dùng..."',
        'Phân quyền dữ liệu QTC-10 áp dụng cho Dashboard: User thuộc Khối KHCN → Dashboard chỉ hiển thị dữ liệu KHCN hay hiển thị tất cả nhưng filter mặc định theo Khối? Tài liệu không nói rõ Dashboard lọc ở tầng nào (lọc toàn bộ biểu đồ hay chỉ lọc bảng danh sách).',
        'Nghiệp vụ',
        'Đề xuất: QTC-10 áp dụng cho TOÀN BỘ dữ liệu trên Dashboard (cả biểu đồ lẫn bảng). User Khối KHCN chỉ thấy dữ liệu giao dịch thu phí của KH thuộc Khối KHCN trên mọi biểu đồ.',
        ''
    ),
    (
        'US32-QA-01.2',
        'Mục "Yêu cầu nghiệp vụ", Phân vùng chỉ số tổng quan, đoạn "Phần trăm thay đổi..."',
        'Công thức % thay đổi: (Tháng này − Tháng trước) / Tháng trước × 100%. Nếu Tháng trước = 0 (không có giao dịch nào), phép chia cho 0 xảy ra. Hệ thống xử lý như thế nào?',
        'Nghiệp vụ',
        'Đề xuất: Nếu mẫu số = 0, hiển thị "N/A" hoặc "—" thay vì % và không áp dụng logic màu xanh/đỏ.',
        ''
    ),
    (
        'US32-QA-01.3',
        'Mục "Yêu cầu nghiệp vụ", Phân vùng chỉ số tổng quan, đoạn "Nếu kết quả >=100% thì hiển thị màu xanh; Nếu kết quả <100% thì hiển thị màu đỏ"',
        'Logic màu: >=100% = xanh, <100% = đỏ. Tuy nhiên công thức % thay đổi có thể cho ra giá trị ÂM (VD: tháng này = 100, tháng trước = 200 → % = -50%). Giá trị -50% < 100% → hiển thị đỏ. Vậy giá trị 0% hoặc giá trị âm có cần icon/biểu tượng đặc biệt nào không (mũi tên xuống, mũi tên lên)?',
        'Nghiệp vụ',
        'Đề xuất: Nếu % ≥ 0 → mũi tên lên xanh; Nếu % < 0 → mũi tên xuống đỏ. Giá trị = 0% → trung tính (không icon). Tham khảo Mockup: icon mũi tên đã có trên mockup.',
        ''
    ),
    (
        'US32-QA-01.4',
        'Mục "Dashboard theo Chi nhánh", biểu đồ "Nợ phí và truy thu theo chi nhánh", đoạn "hệ thống mặc định hiển thị 7 chi nhánh (theo thứ tự alphabet)"',
        'Biểu đồ Nợ phí và truy thu theo Chi nhánh mặc định hiển thị 7 chi nhánh theo alphabet, trong khi các biểu đồ khác (cùng Module 1) mặc định top 5. Tại sao con số 7 và tiêu chí alphabet, không phải top theo giá trị? Đây là thiết kế có chủ đích hay lỗi copy?',
        'Nghiệp vụ',
        'Đề xuất: Xác nhận 7 là con số thiết kế. Nếu user muốn thấy chi nhánh có nợ phí cao nhất thì cần tự chọn lại qua combobox.',
        ''
    ),
    (
        'US32-QA-01.5',
        'Mục "Dashboard theo Khối", bảng "Danh sách các khách hàng giao dịch nhiều nhất trong ngày"',
        'Tài liệu nêu "Chọn 1 hoặc nhiều Khối từ danh sách Khối KHCN, Khối KHDNL, Khối KHDN theo nguyên tắc phân quyền dữ liệu". Nếu user thuộc Khối KHCN, theo QTC-10 user chỉ thấy dữ liệu KHCN. Vậy dropdown "Chọn Khối" có bị disable/chỉ hiển thị 1 giá trị duy nhất cho user thuộc Khối KHCN/KHDN?',
        'Nghiệp vụ',
        'Đề xuất: Áp dụng QTC-10 — user thuộc Khối KHCN → dropdown "Chọn Khối" mặc định = KHCN, không cho sửa. User không thuộc KHCN/KHDN → dropdown cho chọn tự do.',
        ''
    ),
    (
        'US32-QA-01.6',
        'Mục "Dashboard theo Sản phẩm", biểu đồ "Số giao dịch thu phí theo sản phẩm", đoạn "Sắp xếp chi nhánh có số giao dịch thu phí cao nhất lên top"',
        'Tài liệu viết "Sắp xếp CHI NHÁNH có số giao dịch cao nhất" nhưng biểu đồ này là về SẢN PHẨM (SPDV cấp 1). Đây có phải lỗi copy-paste từ Module 1? Logic đúng phải là sắp xếp SẢN PHẨM có số giao dịch cao nhất.',
        'Nghiệp vụ',
        'Đề xuất: Đây là lỗi copy-paste. Sửa thành "Sắp xếp SẢN PHẨM (SPDV cấp 1) có số giao dịch thu phí cao nhất lên top".',
        ''
    ),
    (
        'US32-QA-01.7',
        'Mục "Dashboard theo Khối", biểu đồ "Nợ phí và truy thu", trục X ghi "Khối KHCN, Khối KHDNL, Khối KHDN"',
        'Trục X cố định 3 khối. Nếu user thuộc Khối KHCN (theo QTC-10 chỉ thấy dữ liệu KHCN), biểu đồ có hiển thị đủ 3 khối nhưng 2 khối kia = 0? Hay chỉ hiển thị 1 khối mà user có quyền?',
        'Nghiệp vụ',
        'Đề xuất: Chỉ hiển thị các Khối mà user có quyền xem dữ liệu. User KHCN → biểu đồ chỉ hiển thị 1 điểm (Khối KHCN). User không thuộc KHCN/KHDN → hiển thị đủ 3 khối.',
        ''
    ),

    # === HẠ MỤC 2: Giới hạn hệ thống & Exception Handling ===
    (
        'US32-QA-02.1',
        'Mục "Yêu cầu nghiệp vụ", đoạn "Các số tiền trên dashboard đều có đơn vị Triệu VND"',
        'Đơn vị "Triệu VND" nhưng không nêu rõ quy tắc làm tròn. Ví dụ: Tổng doanh thu = 12,231,456 VND → hiển thị 12.23 hay 12.2 hay 12 Triệu VND? Số thập phân sau dấu chấm là bao nhiêu?',
        'Giới hạn',
        'Đề xuất: Hiển thị 3 chữ số sau dấu phẩy (VD: 12,231 = 12 tỷ 231 triệu) hoặc 2 chữ số thập phân (12.23 triệu). Cần BA xác nhận.',
        ''
    ),
    (
        'US32-QA-02.2',
        'Mục "Dashboard theo Chi nhánh", biểu đồ "Nợ phí và truy thu", đoạn "được chọn tối đa 10 chi nhánh"',
        'Giới hạn tối đa 10 chi nhánh khi chọn. Nếu user cố chọn chi nhánh thứ 11 thì hệ thống xử lý thế nào? FE disable checkbox/hiển thị cảnh báo?',
        'Giới hạn',
        'Đề xuất: Khi đã chọn đủ 10 → FE disable các checkbox còn lại (không cho tick thêm). Hiển thị label "Chọn tối đa 10 Chi nhánh" phía trên danh sách.',
        ''
    ),
    (
        'US32-QA-02.3',
        'Mục "Dashboard theo Sản phẩm", biểu đồ "Doanh thu phí theo SPDV", đoạn "Cho phép chọn tối đa 10 SPDV cùng cấp"',
        'Tương tự giới hạn 10 chi nhánh. Nếu user chọn quá 10 SPDV thì hệ thống xử lý thế nào?',
        'Giới hạn',
        'Đề xuất: FE disable checkbox khi đã đủ 10. Hiển thị label "Chọn tối đa 10 Sản phẩm dịch vụ".',
        ''
    ),
    (
        'US32-QA-02.4',
        'Toàn bộ Dashboard',
        'Tài liệu không đề cập hành vi khi không có DỮ LIỆU (tháng mới, chưa có giao dịch nào). Tất cả biểu đồ + KPI card hiển thị gì? Biểu đồ có hiển thị empty state hay trống hoàn toàn?',
        'Giới hạn',
        'Đề xuất: KPI card hiển thị giá trị 0, % thay đổi = "N/A". Biểu đồ hiển thị trạng thái empty state (VD: "Chưa có dữ liệu").',
        ''
    ),

    # === HẠ MỤC 3: Toàn vẹn dữ liệu & Ràng buộc ===
    (
        'US32-QA-03.1',
        'Mục "Yêu cầu nghiệp vụ", đoạn "không tính các giao dịch đã bị reverse"',
        'Quy tắc "không tính giao dịch reverse" được lặp đi lặp lại ở nhiều biểu đồ. Tuy nhiên, nếu 1 giao dịch được reverse SAU khi Dashboard đã load, Dashboard có tự động refresh/cập nhật hay phải reload trang?',
        'Toàn vẹn dữ liệu',
        'Đề xuất: Dashboard hiển thị dữ liệu tại thời điểm load. Nếu muốn thấy dữ liệu mới nhất (bao gồm reverse), user cần refresh trang hoặc chuyển tab rồi quay lại.',
        ''
    ),
    (
        'US32-QA-03.2',
        'Mục "Dashboard theo Chi nhánh", đoạn "tổng doanh thu phí giao dịch đã thu được có số tiền phí thực thu VND > 0"',
        'Điều kiện lọc "phí thực thu VND > 0" xuất hiện ở nhiều biểu đồ. Câu hỏi: Nếu giao dịch có phí thực thu bằng ngoại tệ (USD, EUR) nhưng quy đổi VND > 0, có được tính không? Hay chỉ tính giao dịch phí gốc bằng VND?',
        'Toàn vẹn dữ liệu',
        'Đề xuất: "Phí thực thu VND" nghĩa là số tiền phí đã quy đổi sang VND. Mọi giao dịch thu phí (bất kể ngoại tệ gốc) nếu có quy đổi VND > 0 đều được tính.',
        ''
    ),
    (
        'US32-QA-03.3',
        'Mục "Dashboard theo Khối", đoạn trục X "Khối KHCN, Khối KHDNL, Khối KHDN"',
        'Tên Khối trên biểu đồ Nợ phí theo Khối ghi "Khối KHDNL", trong khi Mockup image3 ghi "KHDN". Ở QTC-10, danh sách Khối chỉ có KHCN và KHDN. Vậy "Khối KHDNL" là gì? Có phải Khối KHDN Lớn (= KHTC theo QTC-10)?',
        'Toàn vẹn dữ liệu',
        'Đề xuất: Xác nhận mapping: KHDNL = Khối khách hàng doanh nghiệp lớn (tương đương Khối KHDN trong QTC-10). Và xác nhận "Khối KHDN" trên trục X chính xác là Khối nào trong ma trận QTC-10.',
        ''
    ),
    (
        'US32-QA-03.4',
        'Mục "Dashboard theo Khối", bảng danh sách KH giao dịch nhiều nhất, đoạn "trong ngày hệ thống"',
        'Bảng Top KH giao dịch nhiều nhất lấy dữ liệu "trong ngày hệ thống". Các biểu đồ khác lấy dữ liệu "tháng hiện tại" hoặc "từ đầu tháng đến ngày hệ thống". Xác nhận: bảng Top KH chỉ tính giao dịch trong ĐÚNG 1 NGÀY (ngày hệ thống), không phải MTD?',
        'Toàn vẹn dữ liệu',
        'Đề xuất: Xác nhận đúng 1 ngày (ngày hệ thống hiện tại), khác biệt so với các biểu đồ MTD.',
        ''
    ),

    # === HẠ MỤC 4: UI/UX & Giao diện ===
    (
        'US32-QA-04.1',
        'Mục "Giao diện", Mockup image2 — Dashboard theo Sản phẩm',
        'Mockup image2 có 1 vùng ghi "Tính năng đang phát triển" (phần bên phải biểu đồ Line chart "Số giao dịch thu phí theo sản phẩm"). Tài liệu text không đề cập vùng này. Vùng "Tính năng đang phát triển" là placeholder cho biểu đồ nào? Hay là vùng sẽ bổ sung sau?',
        'UI-UX',
        'Đề xuất: Xác nhận vùng này là placeholder tạm thời và sẽ ẩn/không hiển thị trong phiên bản Release. Hoặc nếu là biểu đồ khác sẽ bổ sung → cần update US.',
        ''
    ),
    (
        'US32-QA-04.2',
        'Mục "Giao diện", Mockup image3 — Dashboard theo Khối',
        'Tương tự image2, Mockup image3 cũng có vùng "Tính năng đang phát triển" (bên phải biểu đồ Nợ phí và truy thu theo Khối). Xác nhận đây cũng là placeholder?',
        'UI-UX',
        'Đề xuất: Tương tự US32-QA-04.1, xác nhận vùng placeholder.',
        ''
    ),
    (
        'US32-QA-04.3',
        'Mục "Giao diện", Mockup image1 — Dashboard theo Chi nhánh, phía dưới bên trái có hiện avatar + tên "Trần Thành Đạt, Công nghệ thông tin"',
        'Mockup hiển thị thông tin user (tên + phòng ban) ở góc dưới trái màn hình biểu đồ "Nợ phí và truy thu theo chi nhánh". Tài liệu text không đề cập đến phần hiển thị thông tin user này. Đây là phần cố định của layout Dashboard hay chỉ là chi tiết mockup?',
        'UI-UX',
        'Đề xuất: Đây có thể là phần sidebar/user profile cố định của layout tổng thể ứng dụng (không thuộc phạm vi US32). Xác nhận để tránh nhầm lẫn khi test.',
        ''
    ),
    (
        'US32-QA-04.4',
        'Mục "Dashboard theo Sản phẩm", biểu đồ "Doanh thu phí theo SPDV", đoạn "Sau khi thoát khỏi dropdown list..."',
        'Tài liệu nói "Sau khi thoát khỏi dropdown list Chọn SPDV" nhưng không nói cách thoát. Click ngoài vùng dropdown? Hay có nút "Áp dụng"/"Đóng"? Trên Mockup image2, dropdown SPDV mở ra nhưng không thấy nút "Áp dụng".',
        'UI-UX',
        'Đề xuất: Thoát dropdown bằng cách click ra ngoài vùng dropdown (auto-apply sau khi đóng). Không cần nút "Áp dụng" riêng.',
        ''
    ),
    (
        'US32-QA-04.5',
        'Mục "Yêu cầu nghiệp vụ", đoạn "Các số tiền trên dashboard đều có đơn vị Triệu VND"',
        'Mockup hiển thị label "Đơn vị: Triệu VND" ở góc phải trên. Nhưng bảng "Danh sách KH giao dịch nhiều nhất" (Module 3.2) có cột "Doanh thu". Cột "Doanh thu" trong bảng này cũng hiển thị đơn vị Triệu VND hay hiển thị số VND gốc?',
        'UI-UX',
        'Đề xuất: Cột "Doanh thu" trong bảng danh sách KH cũng hiển thị đơn vị Triệu VND (thống nhất với toàn Dashboard). Tuy nhiên cần BA xác nhận vì bảng danh sách có thể cần chính xác hơn biểu đồ.',
        ''
    ),
    (
        'US32-QA-04.6',
        'Mục "Dashboard theo Chi nhánh", biểu đồ Donut chart, Mockup image1',
        'Mockup Donut chart hiển thị danh sách chi nhánh ở bên trái (Hà Nội 34,200; TP HCM 34,200; ...) kèm "Còn lại 34,200" ở bên phải. Nhưng con số 34,200 giống nhau cho tất cả chi nhánh → có vẻ là dữ liệu mẫu. Tài liệu text không mô tả format hiển thị legend bên trái donut. Xác nhận legend gồm: Tên chi nhánh + Giá trị doanh thu?',
        'UI-UX',
        'Đề xuất: Legend Donut: Tên chi nhánh + Giá trị doanh thu (Triệu VND). Phần "Còn lại" hiển thị riêng. Giá trị "Tổng" hiển thị ở trung tâm donut.',
        ''
    ),
]

# Create table
table = doc.add_table(rows=1, cols=6)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Set column widths
for i, width in enumerate(col_widths):
    table.columns[i].width = width

# Header row
hdr = table.rows[0]
for i, text in enumerate(headers):
    cell = hdr.cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(255, 255, 255)
    from docx.oxml import OxmlElement
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), '003366')
    cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    cell._element.get_or_add_tcPr().append(shading)

# Data rows
for row_data in qa_data:
    row = table.add_row()
    for i, text in enumerate(row_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(8.5)
        p.paragraph_format.space_after = Pt(2)

# Autofit
table.autofit = True

output = '/Users/mac/antigravity-testing-kit/Feature_02_SA_Tham_So_He_Thong/US32_PartB_QA.docx'
doc.save(output)
print(f'✅ Đã sinh file: {output}')
