# -*- coding: utf-8 -*-
import sys, os
sys.path.insert(0, os.path.dirname(__file__))
from us35_sc_data import FEATURE, SC_DATA
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUT = os.path.join(os.path.dirname(os.path.dirname(__file__)), "US35_PartC_TestCoverage.docx")

doc = Document()
# Page setup: Landscape, Narrow margins
for sec in doc.sections:
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    sec.top_margin = Cm(1.27)
    sec.bottom_margin = Cm(1.27)
    sec.left_margin = Cm(1.27)
    sec.right_margin = Cm(1.27)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
    return h

def set_cell(cell, text, bold=False, size=9, bg=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    if bg:
        shading = cell._element.get_or_add_tcPr()
        sh = shading.makeelement(qn('w:shd'), {qn('w:fill'): bg, qn('w:val'): 'clear'})
        shading.append(sh)

def set_col_width(table, widths):
    for row in table.rows:
        for i, w in enumerate(widths):
            if i < len(row.cells):
                row.cells[i].width = Cm(w)

# ============================================================
# PHẦN A (Cập nhật): Tóm Tắt Nghiệp Vụ
# ============================================================
add_heading("PHẦN A: TÓM TẮT NGHIỆP VỤ (Cập nhật sau phản hồi BA)", 1)

add_heading("A.1. Thông điệp cốt lõi", 2)
doc.add_paragraph(
    "US35 mô tả luồng backend tự động thu các loại phí định kỳ. "
    "Đầu ngày T, hệ thống ProfiX xác định các Job cần chạy, sinh danh sách khoản phí đến hạn, "
    "ghi vào Topic Kafka. Core T24 hạch toán thu phí và trả kết quả qua Topic kết quả. "
    "ProfiX đọc kết quả, cập nhật trạng thái + lịch sử. Không có UI (Giao diện = N/A)."
)

add_heading("A.2. Cấu trúc luồng nghiệp vụ & Module", 2)

# Module 1
doc.add_paragraph("Module 1: Sinh dữ liệu phí định kỳ", style='List Number')
doc.add_paragraph("Happy Path: Đầu ngày T → Xác định Job → Lấy DS code phí B → "
    "Xác định đối tượng tính phí (Customer/Account/Card) → Sinh DS khoản phí C.", style='List Bullet')
doc.add_paragraph("Nhánh ĐT = Customer: Bước 2 kiểm tra TK mặc định → nếu không đủ ĐK → tìm TK thay thế (ưu tiên số dư lớn nhất, cùng số dư → ngẫu nhiên). "
    "Nếu không có TK thay thế → vẫn dùng TK mặc định (T24 sẽ trả Chưa thanh toán).", style='List Bullet 2')
doc.add_paragraph("Nhánh ĐT = Account/Card: Bỏ qua bước 2, sang bước 3.", style='List Bullet 2')
doc.add_paragraph("TK thay thế CHỈ dùng cho phí đúng hạn, KHÔNG dùng cho truy thu/tận thu.", style='List Bullet 2')

# Module 2
doc.add_paragraph("Module 2: Tính phí", style='List Number')
doc.add_paragraph("Bước 3.1: Xác định Quy tắc tính phí (nếu có Nhóm KH).", style='List Bullet')
doc.add_paragraph("Bước 3.2: Tính phí theo Số cố định hoặc Công thức (ref US05). "
    "** Nếu Phí ≤ 0 → Phí thu được = 0, VAT = 0, không gửi T24 (BA đã cập nhật).", style='List Bullet')
doc.add_paragraph("Bước 3.3-3.4: Áp Min/Max Code phí (có quy đổi tỷ giá nếu khác loại tiền). "
    "Chỉ có Min → chỉ so Min; chỉ có Max → chỉ so Max (BA xác nhận). "
    "Tỷ giá: lấy bản ghi gần nhất từ Core (BA đã cập nhật).", style='List Bullet')

# Module 3
doc.add_paragraph("Module 3: Áp dụng CTƯĐ", style='List Number')
doc.add_paragraph("Bước 4: Xác định CTƯĐ áp dụng (2 trường hợp: không đánh giá ĐK / có đánh giá ĐK). "
    "Ngưỡng ƯĐ: theo số lần/số tiền GD (ref US12).", style='List Bullet')
doc.add_paragraph("Bước 5.1-5.2: Tính ƯĐ. Nhiều CTƯĐ → chọn ƯĐ lớn nhất → bằng nhau → thời gian khởi tạo xa nhất (BA đã sửa).", style='List Bullet')
doc.add_paragraph("Bước 5.3-5.4: Áp Min/Max CTƯĐ (quy đổi tỷ giá tương tự bước 3.3).", style='List Bullet')
doc.add_paragraph("Không có CTƯĐ → bỏ qua bước 5 (BA xác nhận).", style='List Bullet')
doc.add_paragraph("THỨ TỰ ÁP DỤNG: Bước 3.4 (Min/Max Code phí) TRƯỚC → Bước 5.4 (Min/Max CTƯĐ) SAU. Tuần tự, không xung đột (BA xác nhận).", style='List Bullet')

# Module 4
doc.add_paragraph("Module 4: Tính VAT", style='List Number')
doc.add_paragraph("VAT = '' → không có VAT. VAT ≠ '' → tính theo Có/Không bao gồm VAT.", style='List Bullet')
doc.add_paragraph("Làm tròn: VND/JPY → số nguyên; khác → 2 chữ số thập phân. Không làm tròn trung gian.", style='List Bullet')
doc.add_paragraph("'Số tiền phí thực thu' = 'Số tiền phí sau ưu đãi' tại thời điểm tính phí (BA đã thống nhất).", style='List Bullet')

# Module 5
doc.add_paragraph("Module 5: Ghi Kafka & Update kết quả", style='List Number')
doc.add_paragraph("Ghi danh sách vào Topic Kafka theo thứ tự ưu tiên nhóm code phí.", style='List Bullet')
doc.add_paragraph("Kết quả Thanh toán toàn bộ/một phần → ghi lịch sử + cập nhật trạng thái.", style='List Bullet')
doc.add_paragraph("Kết quả Chưa thanh toán → KHÔNG ghi lịch sử (BA xác nhận).", style='List Bullet')
doc.add_paragraph("Trạng thái: Thêm mới → Đang xử lý → TT toàn bộ/TT một phần/Chưa TT/Xóa nợ (Xóa nợ ở US36).", style='List Bullet')
doc.add_paragraph("Nợ phí: TT một phần + Chưa TT. Truy thu/tận thu → US36.", style='List Bullet')
doc.add_paragraph("Ngày thu = 31, tháng chỉ có 28/29/30 → chạy ngày cuối tháng (ref US02).", style='List Bullet')

add_heading("A.3. Bảng Điều Kiện Tiên Quyết", 2)
pre_tbl = doc.add_table(rows=6, cols=2)
pre_tbl.style = 'Table Grid'
hdrs = [("Điều kiện", "Mô tả")]
data_pre = [
    ("Dữ liệu đồng bộ T-1", "Customer, Account, Card, Tỷ giá từ Core T24 (US33)"),
    ("Code phí + Job", "Đã cài đặt Job với lịch chạy, nhóm code phí, đối tượng tính phí"),
    ("CTƯĐ (nếu có)", "Đã thiết lập CTƯĐ + bản ghi ƯĐ theo SPDV cấp cuối (US12)"),
    ("Tham số CA_PRODUCT", "DS sản phẩm TK được phép trích thu phí"),
    ("Kafka Topics", "Topic thu phí + Topic kết quả đã cấu hình"),
]
for i, (h1, h2) in enumerate(hdrs):
    set_cell(pre_tbl.rows[0].cells[0], h1, bold=True, bg="D9E2F3")
    set_cell(pre_tbl.rows[0].cells[1], h2, bold=True, bg="D9E2F3")
for i, (c1, c2) in enumerate(data_pre):
    set_cell(pre_tbl.rows[i+1].cells[0], c1, bold=True)
    set_cell(pre_tbl.rows[i+1].cells[1], c2)

add_heading("A.4. Quy Tắc Chung Áp Dụng", 2)
qtc_items = [
    "QTC-01.4 (Number): Phân cách hàng nghìn bằng dấu phẩy, thập phân bằng dấu chấm, 2 chữ số. VND/JPY không có thập phân.",
    "QTC-01.5 (Date): dd/mm/yyyy, Từ ngày 00:00:00.000 → Đến ngày 23:59:59.999.",
    "QTC-11 (Error Handling): FE-First, BE trả mã lỗi kỹ thuật nếu FE chưa chặn.",
    "QTC-14.7: Phân biệt Trạng thái phê duyệt vs Trạng thái hoạt động.",
]
for item in qtc_items:
    doc.add_paragraph(item, style='List Bullet')

# ============================================================
# PHẦN C: Bảng Tổng Hợp Test Case Đề Xuất
# ============================================================
doc.add_page_break()
add_heading("PHẦN C: BẢNG TỔNG HỢP TEST CASE ĐỀ XUẤT (Test Case Coverage)", 1)
doc.add_paragraph(f"Feature: {FEATURE}")
doc.add_paragraph(f"Tổng số SC: {len(SC_DATA)}")

# Group by type for summary
from collections import Counter
type_counts = Counter(sc[2] for sc in SC_DATA)
summary_p = doc.add_paragraph("Phân bổ: ")
for t, c in type_counts.items():
    summary_p.add_run(f"{t}: {c}  |  ")

headers = ["Mã SC", "Module", "Loại TC", "Tên Test Case / Kịch bản", "Trích dẫn tài liệu"]
tbl = doc.add_table(rows=1 + len(SC_DATA), cols=5)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
for i, h in enumerate(headers):
    set_cell(tbl.rows[0].cells[i], h, bold=True, size=9, bg="1F4E79")
    for run in tbl.rows[0].cells[i].paragraphs[0].runs:
        run.font.color.rgb = RGBColor(255, 255, 255)

# Data rows
TYPE_COLORS = {
    "Happy Path": "E2EFDA",
    "Negative Path": "FCE4EC",
    "Boundary Value": "FFF3E0",
    "Field Validation": "E3F2FD",
    "Business Logic": "F3E5F5",
    "Data Integrity": "FFF9C4",
    "NFR": "E0F7FA",
}
for idx, (sc_id, module, tc_type, title, ref) in enumerate(SC_DATA):
    row = tbl.rows[idx + 1]
    bg = TYPE_COLORS.get(tc_type, None)
    set_cell(row.cells[0], sc_id, bold=True, size=8, bg=bg)
    set_cell(row.cells[1], module, size=8, bg=bg)
    set_cell(row.cells[2], tc_type, size=8, bg=bg)
    set_cell(row.cells[3], title, size=8)
    set_cell(row.cells[4], ref, size=8)

# Column widths (landscape A4 ~27cm usable)
set_col_width(tbl, [2.0, 3.5, 3.0, 10.0, 8.5])

doc.save(OUT)
print(f"✅ Đã tạo: {OUT}")
print(f"📊 Tổng SC: {len(SC_DATA)}")
for t, c in sorted(type_counts.items()):
    print(f"   {t}: {c}")
