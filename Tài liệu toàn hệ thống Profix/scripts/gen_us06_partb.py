#!/usr/bin/env python3
"""Generate US06 Part B Q&A Report for BA."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# Page setup: Landscape, narrow margins
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

# Title
title = doc.add_heading("US06 – Phần B: Danh Sách Cảnh Báo & Q&A (Dành cho BA)", level=1)

doc.add_paragraph(
    "Tài liệu phân tích: US06 – Khai báo Biểu phí\n"
    "Người phân tích: QA Team\n"
    "Trạng thái: Chờ BA phản hồi"
)

# ============================================================
# Q&A DATA - 4 Categories
# ============================================================
qa_data = [
    # --- Category 1: Business Logic & Flow ---
    ("US06-QA-01.1", 
     'Mục "Thêm mới Biểu phí", đoạn mô tả gán Code phí', 
     "Tài liệu cho phép gán Code phí bằng 2 cách (tích chọn cây SPDV hoặc Upload). Nếu user tích chọn trước rồi upload sau, hệ thống thay thế toàn bộ. Nhưng ngược lại: user upload trước rồi tích chọn thêm thì sao? Gộp (merge) hay thay thế?", 
     "Nghiệp vụ",
     "Đề xuất: Upload trước + tích chọn sau → merge (thêm vào danh sách hiện tại). Chỉ upload sau mới thay thế toàn bộ."),

    ("US06-QA-01.2",
     'Mục "Chỉnh sửa Biểu phí", đoạn kiểm tra tác vụ Chờ duyệt',
     'Tài liệu ghi: "kiểm tra xem có đang tồn tại bản ghi Biểu phí với tác vụ Chỉnh sửa/Chỉnh sửa - Sửa Code phí ở trạng thái Chờ duyệt hay không". Câu hỏi: Tác vụ "Thêm mới" đang Chờ duyệt có chặn Chỉnh sửa không? Vì Biểu phí chưa chính thức tồn tại trên lưới nếu Thêm mới chưa duyệt.',
     "Nghiệp vụ",
     'Đề xuất: Nếu Biểu phí có tác vụ "Thêm mới" đang Chờ duyệt → nút Chỉnh sửa không hiển thị trên lưới (vì bản ghi chưa chính thức).'),

    ("US06-QA-01.3",
     'Mục "Chỉnh sửa Biểu phí", trạng thái Đang hiệu lực',
     'Khi Biểu phí ở trạng thái Đang hiệu lực, chỉ cho phép sửa Ngày hết hiệu lực. Tuy nhiên tài liệu không đề cập: user có được sửa Code phí (Điều kiện tính phí, Quy tắc tính phí) khi Biểu phí đang hiệu lực hay không?',
     "Nghiệp vụ",
     "Đề xuất: Khi Đang hiệu lực → chỉ sửa Ngày hết hiệu lực, KHÔNG cho sửa Code phí. Nếu cần sửa Code phí → phải tạo Biểu phí mới."),

    ("US06-QA-01.4",
     'Mục "Chỉnh sửa Biểu phí", trạng thái Đang hiệu lực, ràng buộc Ngày hết hiệu lực',
     'Tài liệu ghi: "Ngày hết hiệu lực mới phải > Ngày hệ thống và > Ngày hiệu lực". Nhưng khi Thêm mới thì ràng buộc là >=. Có mâu thuẫn giữa > (strict) khi sửa và >= khi thêm mới.',
     "Nghiệp vụ",
     'Đề xuất: Thống nhất dùng >= cho cả 2 luồng (Ngày hết hiệu lực >= Ngày hiệu lực). Riêng > Ngày hệ thống là hợp lý khi sửa vì phải gia hạn ra tương lai.'),

    ("US06-QA-01.5",
     'Mục "Thêm mới Biểu phí", đoạn về trạng thái sau khi Checker duyệt',
     'Tài liệu ghi: "Trường hợp nếu Ngày duyệt > Ngày hiệu lực, tại thời điểm duyệt hệ thống validate lại các ràng buộc như khi thêm mới". Nếu validate lại thất bại (VD: Ngày hiệu lực < Ngày hệ thống tại thời điểm duyệt), hệ thống xử lý thế nào? Tự động từ chối hay cho Checker quyết định?',
     "Nghiệp vụ",
     "Đề xuất: Hệ thống hiển thị cảnh báo cho Checker biết ràng buộc ngày đã vi phạm. Checker có thể chọn Từ chối và ghi lý do."),

    ("US06-QA-01.6",
     'Mục "Chỉnh sửa Biểu phí", đoạn về trạng thái Hết hiệu lực',
     'Khi Biểu phí Hết hiệu lực → chỉ cho phép "chuyển đổi code phí sang Biểu phí mới" (tham chiếu US09). Câu hỏi: User có thể sửa Ngày hết hiệu lực để gia hạn Biểu phí đã hết hiệu lực không? Hay bắt buộc phải tạo mới?',
     "Nghiệp vụ",
     "Đề xuất: Biểu phí Hết hiệu lực → KHÔNG cho gia hạn. Phải tạo Biểu phí mới hoặc chuyển đổi code phí qua US09."),

    ("US06-QA-01.7",
     'Mục "Thêm mới Biểu phí", đoạn mô tả luồng gán Code phí',
     'Tài liệu không nêu rõ: Biểu phí có bắt buộc phải gán ít nhất 1 Code phí trước khi Xác nhận hay không? Hay cho phép tạo Biểu phí rỗng (không có Code phí)?',
     "Nghiệp vụ",
     "Đề xuất: BẮT BUỘC phải gán ít nhất 1 Code phí. FE disable nút Xác nhận nếu lưới Thông tin chi tiết trống."),

    ("US06-QA-01.8",
     'Flowchart Thêm mới, Bước 5 (User chọn sửa thông tin chi tiết Code phí)',
     'Flowchart vẽ nhánh 5.1 "Sửa thông tin chi tiết Code phí" nhưng không có nhánh xử lý khi user sửa xong rồi nhấn Hủy/Đóng popup sửa Code phí. Dữ liệu đã sửa tại popup có bị discard không?',
     "Nghiệp vụ",
     "Đề xuất theo QTC-15: Nhấn Đóng popup sửa Code phí → discard thay đổi, không lưu."),

    # --- Category 2: System Limits & Exceptions ---
    ("US06-QA-02.1",
     'Mục "Thêm mới Biểu phí", trường Mã Biểu phí',
     'Tài liệu ghi Mã Biểu phí do user nhập và phải unique. Nhưng không nêu giới hạn ký tự, định dạng cho phép (chỉ chữ-số? cho phép ký tự đặc biệt? có phân biệt hoa/thường?). Áp dụng QTC-01.6 mặc định 50 ký tự, nhưng cần BA xác nhận format.',
     "Giới hạn",
     "Đề xuất: Mã Biểu phí chỉ cho phép chữ cái (A-Z, a-z), số (0-9), dấu gạch ngang (-) và gạch dưới (_). Không cho phép khoảng trắng và ký tự đặc biệt khác."),

    ("US06-QA-02.2",
     'Mục "Thêm mới Biểu phí", trường Sản phẩm/Dịch vụ/Code phí',
     'Tài liệu không nêu giới hạn số lượng Code phí tối đa được gán vào 1 Biểu phí. Nếu user tích chọn toàn bộ cây SPDV (hàng nghìn Code phí), hệ thống có xử lý được không?',
     "Giới hạn",
     "Đề xuất: BA xác nhận giới hạn tối đa Code phí / Biểu phí (nếu có) để đảm bảo hiệu năng hệ thống."),

    ("US06-QA-02.3",
     'Bảng mô tả trường, trường "Ngày ban hành"',
     'Ràng buộc "Không chọn ngày tương lai" nghĩa là Ngày ban hành <= Ngày hệ thống. Nhưng không rõ: có cho phép chọn ngày quá khứ rất xa (VD: 10 năm trước) hay có giới hạn khoảng thời gian?',
     "Giới hạn",
     "Đề xuất: Cho phép chọn bất kỳ ngày quá khứ nào, không giới hạn khoảng thời gian."),

    ("US06-QA-02.4",
     'Mục "Chỉnh sửa Biểu phí", đoạn chỉnh sửa khi Chưa hiệu lực',
     'Khi sửa Biểu phí Chưa hiệu lực: user được phép thêm/bỏ Code phí. Nếu user bỏ hết toàn bộ Code phí khỏi Biểu phí rồi nhấn Xác nhận → hệ thống có chặn không?',
     "Giới hạn",
     "Đề xuất: Tương tự thêm mới, BẮT BUỘC ít nhất 1 Code phí. FE chặn nếu lưới trống."),

    # --- Category 3: Data Integrity & Constraints ---
    ("US06-QA-03.1",
     'Mục "Thêm mới Biểu phí", đoạn gán Code phí qua cây SPDV',
     'Tài liệu ghi: "hiển thị ra cây SPDV với các cấp khai báo có trạng thái = Đang hoạt động". Câu hỏi: Nếu 1 SPDV cha ở trạng thái Hoạt động nhưng SPDV con ở trạng thái Ngừng hoạt động → SPDV con có hiển thị trên cây hay không?',
     "Toàn vẹn dữ liệu",
     "Đề xuất: Chỉ hiển thị SPDV có trạng thái = Hoạt động. SPDV con Ngừng hoạt động → ẩn khỏi cây, kể cả Code phí trực thuộc."),

    ("US06-QA-03.2",
     'Mục "Thêm mới Biểu phí", đoạn về Code phí',
     '1 Code phí có thể được gán vào nhiều Biểu phí đồng thời không? Hay 1 Code phí chỉ thuộc duy nhất 1 Biểu phí tại 1 thời điểm?',
     "Toàn vẹn dữ liệu",
     "Đề xuất: BA xác nhận ràng buộc. Nếu 1 Code phí chỉ thuộc 1 Biểu phí → cây SPDV phải ẩn/disable các Code phí đã được gán ở Biểu phí khác."),

    ("US06-QA-03.3",
     'Mục "Chỉnh sửa Biểu phí", đoạn sửa Code phí trong Biểu phí',
     'Khi user sửa Điều kiện tính phí / Quy tắc tính phí của Code phí tại Biểu phí: thay đổi này chỉ áp dụng trong phạm vi Biểu phí hiện tại hay ảnh hưởng đến Code phí gốc tại Danh mục SPDV?',
     "Toàn vẹn dữ liệu",
     "Đề xuất: Thay đổi chỉ áp dụng trong scope Biểu phí (override). Code phí gốc tại Danh mục SPDV không bị ảnh hưởng."),

    ("US06-QA-03.4",
     'Mục "Thêm mới Biểu phí", đoạn xác định trạng thái Biểu phí',
     'Trạng thái Biểu phí phụ thuộc Ngày hệ thống so với Ngày hiệu lực/Ngày hết hiệu lực. Khi Batch Job chạy qua ngày → trạng thái tự động chuyển. Câu hỏi: Batch Job chạy lúc mấy giờ? Nếu Ngày hết hiệu lực = Ngày hệ thống, trạng thái là "Đang hiệu lực" (theo công thức <=). Sang ngày hôm sau mới chuyển "Hết hiệu lực" → timing chính xác?',
     "Toàn vẹn dữ liệu",
     "Đề xuất: BA xác nhận thời điểm cắt trạng thái (00:00:00 đầu ngày hay real-time). Và xác nhận Batch Job interval."),

    ("US06-QA-03.5",
     'Mục "Thêm mới Biểu phí", đoạn hành động chờ duyệt',
     'Tài liệu phân biệt 2 hành động: "Thêm mới" vs "Thêm mới – Sửa Code phí". Khi Checker từ chối tác vụ "Thêm mới – Sửa Code phí", Maker chỉnh sửa lại rồi submit → hành động mới là gì? Vẫn giữ "Thêm mới – Sửa Code phí" hay đổi thành "Chỉnh sửa"?',
     "Toàn vẹn dữ liệu",
     'Đề xuất: Giữ nguyên hành động gốc "Thêm mới – Sửa Code phí" cho đến khi được duyệt hoặc xóa tác vụ.'),

    # --- Category 4: UI/UX & Interface ---
    ("US06-QA-04.1",
     'Mockup màn hình Thêm mới Biểu phí (image3)',
     'Mockup Thêm mới không hiển thị trường "Trạng thái". Nhưng Mockup Chỉnh sửa (image4) lại hiển thị trường "Trạng thái" = "Hoạt động". Xác nhận: Trạng thái chỉ hiển thị ở màn hình Chỉnh sửa, không hiển thị ở Thêm mới?',
     "UI-UX",
     "Đề xuất: Đúng, Thêm mới không cần hiển thị Trạng thái (vì chưa có). Chỉnh sửa hiển thị Trạng thái ở dạng readonly."),

    ("US06-QA-04.2",
     'Mockup màn hình Thêm mới (image3), phân vùng Thông tin chi tiết',
     'Mockup hiển thị cột "Loại tính phí" trong lưới Thông tin chi tiết, nhưng Bảng mô tả trường (Table 1) không có trường "Loại tính phí". Text mô tả tại P33 có nhắc đến "Loại tính phí". Cần xác nhận trường này có tồn tại trong lưới hay không.',
     "UI-UX",
     'Đề xuất: BA xác nhận "Loại tính phí" có phải là trường chính thức trên lưới hay bị thiếu trong bảng mô tả.'),

    ("US06-QA-04.3",
     'Mockup màn hình Chỉnh sửa (image4)',
     'Mockup Chỉnh sửa có nút "Chuyển đổi biểu phí 2 code phí" ở dưới cùng, nhưng không có nút "Xác nhận" và "Đóng" như màn hình Thêm mới. Khi Biểu phí ở trạng thái Chưa hiệu lực (cho phép sửa), nút Xác nhận/Đóng hiển thị ở đâu?',
     "UI-UX",
     "Đề xuất: Mockup chỉnh sửa cần bổ sung nút Xác nhận + Đóng giống Thêm mới. Nút Chuyển đổi chỉ hiển thị khi trạng thái = Hết hiệu lực."),

    ("US06-QA-04.4",
     'Mockup Chỉnh sửa (image4), checkbox ở lưới Thông tin chi tiết',
     'Mockup Chỉnh sửa hiển thị checkbox ở mỗi dòng Code phí trong lưới, nhưng Mockup Thêm mới không có checkbox. Tài liệu text không đề cập đến checkbox này. Mục đích checkbox là gì? Chọn nhiều Code phí để xóa hàng loạt?',
     "UI-UX",
     "Đề xuất: BA xác nhận mục đích checkbox. Nếu để xóa hàng loạt → cần bổ sung nút Xóa nhiều."),

    ("US06-QA-04.5",
     'Bảng mô tả trường (Table 1), trường "Sản phẩm/Dịch vụ/Code phí"',
     'Trường này có Bắt buộc = "－" (không bắt buộc), nhưng nếu theo đề xuất QA-01.7 Biểu phí phải có ít nhất 1 Code phí, thì trường này nên là bắt buộc (★). Cần BA xác nhận.',
     "UI-UX",
     'Đề xuất: Đổi dấu bắt buộc thành ★, hoặc validation ở cấp form (FE chặn khi lưới chi tiết trống).'),

    ("US06-QA-04.6",
     'Bảng mô tả trường (Table 1), R28 (dòng trống)',
     'Bảng mô tả trường có 1 dòng trống (R28) giữa "Loại khách hàng" và "Khai báo theo nhóm khách hàng". Đây có phải là trường "Loại tính phí" bị thiếu tên hay là lỗi format tài liệu?',
     "UI-UX",
     'Đề xuất: BA xác nhận. Khả năng cao đây là trường "Loại tính phí" bị mất tên do lỗi copy-paste tài liệu.'),

    ("US06-QA-04.7",
     'Bảng mô tả trường (Table 1), nút "Chỉnh sửa" tại Thông tin chi tiết',
     'Ràng buộc nút Chỉnh sửa ghi: "Chỉ cho phép chỉnh sửa Trạng thái Code phí". Nhưng text nghiệp vụ mô tả cho phép sửa Điều kiện tính phí (Mô tả, Giá trị), Quy tắc tính phí (Giá trị số, Tối thiểu, Tối đa). Hai nơi mâu thuẫn nhau.',
     "UI-UX",
     'Đề xuất: Ràng buộc tại bảng bị sai/cũ. Nên sửa thành: "Cho phép chỉnh sửa Điều kiện tính phí và Quy tắc tính phí" theo đúng text nghiệp vụ.'),
]

# ============================================================
# Build table
# ============================================================
headers = ["ID", "Trích xuất", "Câu hỏi / Sự cố", "Phân loại", "Đề xuất từ QA", "Trả lời của BA"]
col_widths = [Cm(2.5), Cm(4.5), Cm(7), Cm(2), Cm(6.5), Cm(4.5)]

# Category headers
cat_headers = {
    "Nghiệp vụ": "🔶 Hạng mục 1: Vấn đề Nghiệp vụ / Luồng xử lý (Business Logic & Flow Issues)",
    "Giới hạn": "🔴 Hạng mục 2: Giới hạn hệ thống & Exception Handling (System Limits & Exceptions)", 
    "Toàn vẹn dữ liệu": "🟠 Hạng mục 3: Toàn vẹn dữ liệu & Ràng buộc (Data Integrity & Constraints)",
    "UI-UX": "🔵 Hạng mục 4: UI/UX & Giao diện (UI/UX & Interface Issues)",
}

cat_order = ["Nghiệp vụ", "Giới hạn", "Toàn vẹn dữ liệu", "UI-UX"]

for cat in cat_order:
    doc.add_heading(cat_headers[cat], level=2)
    
    items = [q for q in qa_data if q[3] == cat]
    
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # Header row
    hdr = table.rows[0]
    for i, (h, w) in enumerate(zip(headers, col_widths)):
        cell = hdr.cells[i]
        cell.width = w
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        # Gray background
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): 'D9E2F3',
            qn('w:val'): 'clear'
        })
        shading.append(shading_elm)
    
    # Data rows
    for item in items:
        row = table.add_row()
        for i, val in enumerate(item):
            cell = row.cells[i]
            cell.width = col_widths[i]
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(8.5)
        # Last column empty (BA answer)
        row.cells[5].paragraphs[0].text = ""
    
    doc.add_paragraph("")  # spacer

# Summary
doc.add_heading("Tổng kết", level=2)
summary = doc.add_paragraph()
summary.add_run(f"Tổng số câu hỏi: {len(qa_data)}\n").bold = True
summary.add_run(f"  • Nghiệp vụ (01.x): {len([q for q in qa_data if q[3]=='Nghiệp vụ'])}\n")
summary.add_run(f"  • Giới hạn (02.x): {len([q for q in qa_data if q[3]=='Giới hạn'])}\n")
summary.add_run(f"  • Toàn vẹn dữ liệu (03.x): {len([q for q in qa_data if q[3]=='Toàn vẹn dữ liệu'])}\n")
summary.add_run(f"  • UI-UX (04.x): {len([q for q in qa_data if q[3]=='UI-UX'])}\n")

doc.add_paragraph("⚠️ DỪNG LẠI: Vui lòng gửi file này cho BA phản hồi cột 'Trả lời của BA' trước khi tiến hành Phase 2 (Phần C).")

# Save
out_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "US06_Part_B_QA_Report.docx")
doc.save(out_path)
print(f"✅ Đã tạo file: {out_path}")
print(f"📊 Tổng: {len(qa_data)} câu hỏi Q&A")
