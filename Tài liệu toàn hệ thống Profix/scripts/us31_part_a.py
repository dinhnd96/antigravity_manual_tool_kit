"""US31 - Part A v2: Tóm tắt Nghiệp vụ (cập nhật theo v2 + BA feedback)"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

title = doc.add_heading("US31 – Báo Cáo Tổng Doanh Thu Phí (v2)", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_heading("PHẦN A: TÓM TẮT NGHIỆP VỤ (Requirements Breakdown)", level=2)

# A.1
doc.add_heading("A.1. Thông Điệp Cốt Lõi (Core Business Value)", level=3)
p = doc.add_paragraph()
p.add_run("Mục đích: ").bold = True
p.add_run("US31 cung cấp chức năng tra cứu và xuất báo cáo tổng hợp doanh thu phí dịch vụ "
          "theo khoảng thời gian, giúp người dùng nắm bắt tổng doanh thu phí đã thu từ khách hàng. "
          "Dữ liệu được group by theo tổ hợp (Chi nhánh + Loại tiền + Loại tính phí + Biểu phí + Code phí).")
p2 = doc.add_paragraph()
p2.add_run("Người dùng cuối: ").bold = True
p2.add_run("Nhân viên ngân hàng, Quản lý nghiệp vụ phí — phân quyền theo Khối (KHCN, KHDN, KHDNL, hoặc Khối khác).")
p3 = doc.add_paragraph()
p3.add_run("Navigation: ").bold = True
p3.add_run("Báo cáo >> Báo cáo tổng doanh thu phí")
p4 = doc.add_paragraph()
p4.add_run("⚠️ Thay đổi v2: ").bold = True
p4.add_run("(1) Đã xóa nút Tra cứu CIF khỏi US. (2) Sửa lỗi copy-paste 'dự thu phí'. "
           "(3) Thống nhất 'Ngày giao dịch'. (4) Bổ sung rõ group by. (5) Tách Tổng cộng vào 'Thông tin khác'. "
           "(6) Xác nhận KHDNL là Khối riêng biệt.")

# A.2
doc.add_heading("A.2. Cấu Trúc Luồng Nghiệp Vụ & Phân Bổ Module", level=3)

doc.add_heading("Module 1: Tra cứu Báo cáo tổng doanh thu phí", level=4)
doc.add_paragraph("Luồng chính (Happy Path):", style="List Bullet")
for s in [
    "Bước 1: Người dùng chọn menu \"Báo cáo tổng doanh thu phí\".",
    "Bước 2: FE hiển thị màn hình gồm: vùng Điều kiện tìm kiếm (7 trường), lưới rỗng (QTC-04), Tổng cộng, nút chức năng.",
    "Bước 3.1: Người dùng nhập/chọn điều kiện (Khối, Mã Chi nhánh, Biểu phí, Code phí, Loại tính phí, Từ ngày ★, Đến ngày ★).",
    "Bước 4.2: Nhấn \"Tra cứu\" → FE gửi BE (Bước 6-7).",
    "Bước 8: FE hiển thị lưới danh sách tổng doanh thu phí (group by tổ hợp), Tổng cộng cập nhật.",
]:
    doc.add_paragraph(s, style="List Bullet 2")

doc.add_paragraph("Các Luồng Rẽ Nhánh / Ngoại lệ:", style="List Bullet")
for e in [
    "Xóa tra cứu (4.1→5): Xóa điều kiện, lưới về mặc định (QTC-04).",
    "Bỏ trống Từ ngày/Đến ngày: FE validate chặn (QTC-14.5).",
    "Từ ngày > Đến ngày: FE không cho phép chọn.",
    "Không có kết quả: Lưới rỗng, Tổng cộng = 0.",
    "Phân quyền Khối: KHCN/KHDN/KHDNL → Combobox disabled. Khối khác → chọn tự do (QTC-10).",
]:
    doc.add_paragraph(e, style="List Bullet 2")

doc.add_heading("Module 2: Tải xuống báo cáo (Export)", level=4)
doc.add_paragraph("Luồng chính (Happy Path):", style="List Bullet")
for s in [
    "Bước 3.2: Nhấn \"Tải xuống\" → FE hiển thị dropdown Excel/PDF.",
    "Bước 10: Chọn Excel (10.1) hoặc PDF (10.2) → Bước 11: Tải file theo QTC-05.",
    "Tải xuống luôn được cho phép dù lưới có dữ liệu hay không (QTC-05).",
]:
    doc.add_paragraph(s, style="List Bullet 2")

# A.3
doc.add_heading("A.3. Bảng Điều Kiện Tiên Quyết & Cấu Hình", level=3)
t = doc.add_table(rows=5, cols=2, style="Table Grid")
t.rows[0].cells[0].text = "Điều kiện"
t.rows[0].cells[1].text = "Chi tiết"
for i, (k, v) in enumerate([
    ("Phân quyền", "User có quyền truy cập menu Báo cáo >> Báo cáo tổng doanh thu phí"),
    ("Dữ liệu mồi", "Hệ thống có dữ liệu giao dịch thu phí đã ghi nhận (Biểu phí, Code phí, Chi nhánh)"),
    ("Khối user", "Xác định Khối: KHCN / KHDN / KHDNL / Khác"),
    ("Master data", "Danh sách Biểu phí, Code phí, Chi nhánh, Loại tính phí sẵn sàng"),
]):
    t.rows[i+1].cells[0].text = k
    t.rows[i+1].cells[1].text = v

# A.4
doc.add_heading("A.4. Quy Tắc Chung Áp Dụng", level=3)
qt = doc.add_table(rows=10, cols=2, style="Table Grid")
qt.rows[0].cells[0].text = "QTC"
qt.rows[0].cells[1].text = "Nội dung áp dụng cho US31"
for i, (c, d) in enumerate([
    ("QTC-01.1", "Combobox: Khối, Mã Chi nhánh, Biểu phí, Code phí"),
    ("QTC-01.2", "Dropdown List: Loại tính phí — chỉ 1 giá trị"),
    ("QTC-01.4", "Number: Doanh thu phí, Tổng cộng — format phân cách hàng nghìn. VND/JPY: không thập phân"),
    ("QTC-01.5", "Date: Từ ngày, Đến ngày — dd/mm/yyyy"),
    ("QTC-01.7", "Trường rỗng → blank"),
    ("QTC-04", "Tra cứu: Mặc định không hiển thị. Xóa tra cứu → về mặc định"),
    ("QTC-05", "Tải xuống: Excel/PDF, tên file theo format. Luôn cho phép tải"),
    ("QTC-06", "Phân trang: 50 bản ghi/trang. Sort mặc định giảm dần"),
    ("QTC-10", "Phân quyền Khối: KHCN, KHDN, KHDNL (mới xác nhận). Lưu ý: KHDNL chưa có trong QTC-10 cũ"),
]):
    qt.rows[i+1].cells[0].text = c
    qt.rows[i+1].cells[1].text = d

# Bold headers
for table in doc.tables:
    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True

out_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "US31_PartA_Summary.docx")
doc.save(out_path)
print(f"✅ Đã cập nhật: {out_path}")
