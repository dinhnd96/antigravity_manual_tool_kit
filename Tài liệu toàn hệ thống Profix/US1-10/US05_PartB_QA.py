"""
US05 - Part B: Danh Sách Cảnh Báo & Q&A (Dành cho BA)
Quy tắc tính phí cho Code phí
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# Page setup: Landscape, Narrow margins
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(9)

# Title
title = doc.add_heading('US05 — PHẦN B: DANH SÁCH CẢNH BÁO & Q&A (Dành cho BA) — Merged with VA Review', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')

# Q&A Data - 4 hạng mục
qa_data = [
    # Hạng mục 1: Nghiệp vụ / Luồng xử lý
    ("US05-QA-01.1", "Nghiệp vụ",
     'Mục "Yêu cầu nghiệp vụ", phần mô tả Loại quy tắc "Thỏa thuận"',
     'Trường "Tối thiểu" và "Tối đa" đều bắt buộc (★) và >= 0. Tuy nhiên tài liệu KHÔNG đề cập ràng buộc Tối thiểu <= Tối đa. '
     'Nếu người dùng nhập Tối thiểu = 500,000 và Tối đa = 100,000 → hệ thống xử lý như thế nào?',
     'Đề xuất: FE validate Tối thiểu <= Tối đa trước khi cho phép lưu. Nếu vi phạm, hiển thị inline error.'),

    ("US05-QA-01.2", "Nghiệp vụ",
     'Mục "Yêu cầu nghiệp vụ", phần Công thức — Tối thiểu/Tối đa chung tại Phân vùng Quy tắc tính phí',
     'Tài liệu ghi "Tối thiểu/Tối đa của cấu phần là độc lập với Tối thiểu/Tối đa chung tại Phân vùng Quy tắc tính phí". '
     'Câu hỏi: Nếu kết quả tính công thức (sau khi áp min/max từng cấu phần) vượt ngoài min/max CHUNG → hành vi cụ thể là gì? '
     'Hệ thống tự động cap về min/max chung, hay chặn lỗi, hay chỉ cảnh báo?',
     'Đề xuất: Cần làm rõ logic áp dụng min/max 2 tầng (tầng cấu phần + tầng công thức chung) để Tester viết TC chính xác.'),

    ("US05-QA-01.3", "Nghiệp vụ",
     'Mục "Yêu cầu nghiệp vụ", phần Công thức — Nguồn dữ liệu',
     'Tài liệu mô tả 3 ràng buộc Nguồn dữ liệu (ETL Tài khoản / ETL Thẻ / ETL Khách hàng) nhưng KHÔNG nêu rõ THỜI ĐIỂM validate: '
     'Validate khi nhấn "Xác nhận" ở Thiết lập cấu phần? Hay khi "Xác nhận" ở Thiết lập công thức? Hay khi "Xác nhận" tại màn hình Code phí (US02)?',
     'Đề xuất: Nên validate ngay tại Thiết lập cấu phần (gần nhất với hành động vi phạm) để giảm chi phí re-work cho người dùng.'),

    ("US05-QA-01.4", "Nghiệp vụ",
     'Mục "Yêu cầu nghiệp vụ", phần Công thức — toán tử giữa cấu phần',
     'Tài liệu ghi "toán tử giữa các cấu phần được phép chọn là toán tử + hoặc −". Mockup (image7 — Thiết lập công thức) '
     'hiển thị dropdown cho toán tử giữa cấu phần (có icon + và v). Câu hỏi: Giá trị mặc định của toán tử là gì (+ hay −)? '
     'Có validate trường hợp kết quả công thức tổng hợp < 0 không?',
     'Đề xuất: Mặc định = "+". Nếu kết quả < 0 → FE cảnh báo hoặc chặn.'),

    ("US05-QA-01.5", "Nghiệp vụ",
     'Mục "Yêu cầu nghiệp vụ", phần Công thức — Luồng chỉnh sửa quy tắc',
     'Tài liệu ghi "Để thực hiện chỉnh sửa Quy tắc tính phí, Maker thực hiện thông qua luồng chỉnh sửa code phí (US02) hoặc bằng cách upload (US07)". '
     'Câu hỏi: Khi Maker chỉnh sửa Code phí đã duyệt và thay đổi Loại quy tắc (VD: từ Số cố định → Công thức), '
     'dữ liệu quy tắc cũ (Số cố định) có bị XÓA hoàn toàn hay vẫn lưu lịch sử?',
     'Đề xuất: Khi đổi loại quy tắc → dữ liệu quy tắc cũ bị xóa (reset), chỉ giữ loại mới. Cần BA xác nhận.'),

    ("US05-QA-01.6", "Nghiệp vụ",
     'Mục "Yêu cầu nghiệp vụ", phần "Khai báo theo nhóm khách hàng" = Có — Sao chép',
     'Tài liệu ghi "Sao chép bản ghi Quy tắc tính phí từ một nhóm KH này để khai báo nhanh cho Nhóm KH khác" nhưng KHÔNG rõ: '
     'Sau khi sao chép, dropdown "Nhóm KH" của bản ghi mới có tự động gán nhóm KH chưa hay để trống? '
     'Bản ghi được sao chép có bao gồm cả Tối thiểu/Tối đa và toàn bộ công thức chi tiết (cấu phần + định dạng) không?',
     'Đề xuất: Sao chép toàn bộ dữ liệu trừ "Nhóm KH" (để trống cho người dùng chọn). Cần BA xác nhận.'),

    ("US05-QA-01.7", "Nghiệp vụ",
     'Mục "Yêu cầu nghiệp vụ", phần Công thức — Thiết lập cấu phần',
     'Tài liệu ghi "Không được phép bỏ trống các cấu phần đã thêm vào màn hình" và "không được phép bỏ trống các đầu vào đã thêm". '
     'Câu hỏi: Nếu người dùng nhấn "Thêm cấu phần" (hoặc "Thêm định dạng") rồi bỏ trống → validate THỜI ĐIỂM nào? '
     'Validate ngay khi blur field? Hay khi nhấn "Xác nhận"?',
     'Đề xuất: Validate khi nhấn "Xác nhận" — highlight các trường bỏ trống, chặn không cho lưu.'),

    ("US05-QA-01.8", "Nghiệp vụ",
     'Mục "Yêu cầu nghiệp vụ", phần Công thức — Nguồn dữ liệu ETL và trường điều kiện',
     'Hệ thống sẽ chặn việc chọn trường điều kiện không phù hợp (VD: chọn trường theo Thẻ khi Nguồn dữ liệu = ETL Tài khoản) '
     'bằng cách nào? Ẩn luôn các trường không phù hợp khỏi dropdown, hay vẫn cho chọn và chỉ báo lỗi khi ấn "Xác nhận"?',
     'Đề xuất: FE lọc và ẩn luôn các trường điều kiện không phù hợp ngay trong Dropdown list — giảm lỗi nhập liệu.'),

    ("US05-QA-01.9", "Nghiệp vụ",
     'Mục "Yêu cầu nghiệp vụ", phần Công thức — Thiết lập cấu phần, định dạng "Tỷ lệ"',
     'Trong định dạng Tỷ lệ, tài liệu yêu cầu nhập 3 trường: Tỷ lệ, Giá trị, Tên trường. '
     'Vậy phép tính chính xác của định dạng này là "Tỷ lệ * Giá trị" (Giá trị do người dùng nhập) '
     'hay "Tỷ lệ * Tên trường" (giá trị thực tế lấy từ Nguồn dữ liệu)? '
     'Hoặc "Tỷ lệ * Giá trị * Tên trường" (tổ hợp cả 3)?',
     'Đề xuất: BA xác nhận rõ công thức. Dựa trên mockup, khả năng cao là "Tỷ lệ * Tên trường" (Tên trường = biến lấy giá trị thực từ Nguồn dữ liệu).'),

    # Hạng mục 2: Giới hạn hệ thống
    ("US05-QA-02.1", "Giới hạn",
     'Mục "Yêu cầu nghiệp vụ", phần "Khai báo theo nhóm khách hàng" = Có',
     'Tài liệu không đề cập giới hạn số lượng nhóm KH có thể khai báo cho 1 Code phí. '
     'Nếu Danh mục Nhóm KH có 50 nhóm → người dùng có thể thêm cả 50 bản ghi quy tắc? '
     'Có giới hạn tối đa số nhóm KH/Code phí không?',
     'Đề xuất: Giới hạn = số nhóm KH Hoạt động hiện có. Không được chọn trùng nhóm KH trong cùng Code phí.'),

    ("US05-QA-02.2", "Giới hạn",
     'Mục "Yêu cầu nghiệp vụ", phần Công thức — cấu phần',
     'Tài liệu cho phép "Thêm cấu phần" không giới hạn. '
     'Câu hỏi: Có giới hạn tối đa số cấu phần trong 1 công thức không? '
     'Tương tự, có giới hạn tối đa số định dạng trong 1 cấu phần không?',
     'Đề xuất: Cần có giới hạn rõ ràng để tránh overflow / performance issue khi công thức quá phức tạp.'),

    ("US05-QA-02.3", "Giới hạn",
     'Bảng mô tả trường Quy tắc tính phí = Công thức, trường "Tỷ lệ"',
     'Bảng mô tả ghi ràng buộc Tỷ lệ: "số > 0 và <= 1". Câu hỏi: Giá trị Tỷ lệ = 1 (100%) có được chấp nhận hay chỉ < 1? '
     'Tài liệu ghi "<= 1" nhưng cần xác nhận lại ý định nghiệp vụ — Tỷ lệ 100% có ý nghĩa gì trong ngữ cảnh tính phí?',
     'Đề xuất: Giữ nguyên ràng buộc <= 1 theo tài liệu. Cần BA confirm 1 = 100% là hợp lệ.'),

    ("US05-QA-02.4", "Giới hạn",
     'Bảng mô tả trường Quy tắc tính phí = Số cố định, Thỏa thuận — trường "Số cố định", "Tối thiểu", "Tối đa"',
     'Ràng buộc ghi ">= 0" cho Số cố định, Tối thiểu, Tối đa. Câu hỏi: Giá trị = 0 có ý nghĩa gì? '
     'Phí = 0 đồng → miễn phí hoàn toàn? Trường hợp Tối thiểu = 0 VÀ Tối đa = 0 cùng lúc có được phép không?',
     'Đề xuất: Cho phép giá trị 0 (miễn phí). Tuy nhiên cần xác nhận trường hợp Tối thiểu = Tối đa = 0.'),

    # Hạng mục 3: Toàn vẹn dữ liệu
    ("US05-QA-03.1", "Toàn vẹn dữ liệu",
     'Mục "Yêu cầu nghiệp vụ", phần Công thức — trường "Nguồn dữ liệu" và "Tên trường"',
     'Trường "Tên trường" (Combobox) lọc từ Danh mục Điều kiện tính phí theo Nguồn dữ liệu đã khai báo. '
     'Câu hỏi: Nếu người dùng đã thiết lập công thức với Nguồn dữ liệu = ETL Tài khoản, sau đó QUAY LẠI thay đổi Nguồn dữ liệu → ETL Thẻ, '
     'thì các trường điều kiện đã chọn trước đó (theo ETL Tài khoản) có bị RESET / xóa tự động không?',
     'Đề xuất: Khi đổi Nguồn dữ liệu → reset toàn bộ công thức đã thiết lập + cảnh báo "Thay đổi nguồn dữ liệu sẽ xóa công thức hiện tại".'),

    ("US05-QA-03.2", "Toàn vẹn dữ liệu",
     'Mục "Yêu cầu nghiệp vụ", phần "Khai báo theo nhóm KH" — trường "Nhóm khách hàng"',
     'Dropdown "Nhóm KH" lấy từ Danh mục Nhóm KH trạng thái Hoạt động. '
     'Câu hỏi: Nếu 1 Nhóm KH đã được chọn tại bản ghi quy tắc thứ 1, bản ghi thứ 2 có được phép chọn TRÙNG nhóm KH đó không? '
     'Tài liệu không đề cập ràng buộc unique nhóm KH trong cùng Code phí.',
     'Đề xuất: Không cho phép chọn trùng nhóm KH trong cùng Code phí (logic unique). FE disable hoặc ẩn nhóm KH đã chọn ở dropdown.'),

    ("US05-QA-03.3", "Toàn vẹn dữ liệu",
     'Mục "Yêu cầu nghiệp vụ", phần "Khai báo theo nhóm KH" — trường "Nhóm khách hàng"',
     'Sau khi Code phí đã được duyệt, nếu Nhóm KH bị chuyển trạng thái → Không hoạt động tại US27, '
     'thì bản ghi Quy tắc tính phí đã khai báo với nhóm KH đó có bị ảnh hưởng gì không? '
     'Có bị vô hiệu hóa / ẩn khỏi lưới không?',
     'Đề xuất: Bản ghi đã duyệt vẫn giữ nguyên hiệu lực. Nhóm KH bị Không hoạt động chỉ ảnh hưởng khi Thêm mới/Chỉnh sửa mới.'),

    ("US05-QA-03.4", "Toàn vẹn dữ liệu",
     'Bảng mô tả trường Thiết lập cấu phần — "Ngày bắt đầu" và "Ngày kết thúc"',
     'Bảng mô tả ghi ràng buộc "Chọn 1 Trường điều kiện với định dạng là Number và trạng thái = Hoạt động" cho cả Ngày bắt đầu và Ngày kết thúc. '
     'Tuy nhiên, đây là trường NGÀY → lẽ ra phải lọc điều kiện có định dạng là Date, không phải Number. '
     'Phần text nghiệp vụ lại ghi đúng "định dạng là Date". Đây là MÂU THUẪN giữa bảng mô tả trường và phần yêu cầu nghiệp vụ.',
     'Đề xuất: Sửa bảng mô tả trường → "định dạng là Date" thay vì "Number" cho Ngày bắt đầu và Ngày kết thúc. BA cần xác nhận.'),

    ("US05-QA-03.5", "Toàn vẹn dữ liệu",
     'Bảng mô tả trường Thiết lập cấu phần — "Ngày bắt đầu" và "Ngày kết thúc"',
     'Tài liệu ghi ràng buộc "Ngày bắt đầu phải khác Ngày kết thúc" nhưng KHÔNG đề cập ràng buộc thứ tự: '
     'Có bắt buộc Ngày bắt đầu < Ngày kết thúc không? Nếu người dùng chọn Ngày bắt đầu nằm SAU Ngày kết thúc '
     '→ kết quả phép trừ ra số âm → ảnh hưởng đến công thức tính phí.',
     'Đề xuất: Bổ sung ràng buộc Ngày bắt đầu < Ngày kết thúc. FE chặn ngay khi chọn sai thứ tự.'),

    ("US05-QA-03.6", "Toàn vẹn dữ liệu",
     'Mục "Yêu cầu nghiệp vụ", phần "Khai báo theo nhóm KH" — Sao chép quy tắc khi Pending',
     'Nếu Quy tắc tính phí của Nhóm KH A đang ở trạng thái "Chờ duyệt" (Pending), '
     'thì Maker có được phép sao chép quy tắc từ Nhóm KH B đè lên Nhóm KH A không? '
     'Sao chép vào nhóm đang Pending có thể gây xung đột dữ liệu với bản ghi đang chờ Checker duyệt.',
     'Đề xuất: Khóa (Disabled) chức năng Sao chép đích đến đối với các Nhóm KH đang có trạng thái Pending.'),

    ("US05-QA-03.7", "Toàn vẹn dữ liệu",
     'Mục "Yêu cầu nghiệp vụ", phần Công thức — Loại tiền tối thiểu/tối đa',
     'Tài liệu cho phép Loại tiền tối thiểu/tối đa KHÁC với Loại tiền tệ của Code phí. '
     'Câu hỏi: Khi tính phí, nếu Loại tiền Tối thiểu/Tối đa ≠ Loại tiền Code phí → hệ thống có tự động quy đổi tỷ giá không? '
     'Nếu có, tỷ giá lấy từ đâu và tại thời điểm nào?',
     'Đề xuất: Cần BA làm rõ cơ chế quy đổi tiền tệ khi Loại tiền min/max ≠ Loại tiền Code phí.'),

    # Hạng mục 4: UI/UX
    ("US05-QA-04.1", "UI-UX",
     'Mockup "Quy tắc tính phí là Số cố định" (image3) — Phân vùng Quy tắc tính phí',
     'Mockup hiển thị radio button cho 3 loại quy tắc (Thỏa thuận / Số cố định / Công thức) ở phía trên phân vùng. '
     'Tuy nhiên tài liệu text KHÔNG mô tả trường "Loại quy tắc" (radio button) trong bảng mô tả trường cho BẤT KỲ loại nào. '
     'Trường này thiếu trong bảng mô tả → cần bổ sung: tên trường, bắt buộc hay không, giá trị mặc định.',
     'Đề xuất: Bổ sung trường "Loại quy tắc" (Radio button, ★ bắt buộc) vào bảng mô tả. Giá trị mặc định cần BA xác nhận.'),

    ("US05-QA-04.2", "UI-UX",
     'Mockup "Quy tắc tính phí là Thỏa thuận" (image6) so với bảng mô tả trường',
     'Mockup hiển thị trường "Nhóm khách hàng" với dropdown kèm danh sách nhóm KH (KH thông thường, Private, Diamond, Platinum, Gold, G-Diamond, G-Platinum, G-Gold). '
     'Bảng mô tả trường phần Thỏa thuận GHI nút "Xóa" có ràng buộc "Chỉ hiển thị khi tích chọn trường «Khai báo theo nhóm KH» tại phân vùng Thông tin chung" — '
     'KHÔNG THỐNG NHẤT VỚI bảng phần Số cố định, nơi ràng buộc ghi "Chỉ hiển thị khi tham số «Khai báo theo nhóm KH» của Code phí = Có". '
     'Hai cách diễn đạt khác nhau cho cùng 1 logic.',
     'Đề xuất: Thống nhất wording ràng buộc: "Chỉ hiển thị khi tham số «Khai báo theo nhóm KH» = Có".'),

    ("US05-QA-04.3", "UI-UX",
     'Mockup "Thiết lập cấu phần" (image8) — trường "Ngày kết thúc" và "Ngày bắt đầu"',
     'Mockup (image8) hiển thị thứ tự: Ngày kết thúc TRƯỚC, Ngày bắt đầu SAU (công thức = Ngày kết thúc - Ngày bắt đầu). '
     'Tuy nhiên bảng mô tả trường liệt kê thứ tự: Ngày bắt đầu trước, Ngày kết thúc sau. '
     'Cần thống nhất: Mockup lấy Ngày kết thúc - Ngày bắt đầu (phép trừ) → thứ tự hiển thị nào đúng?',
     'Đề xuất: Mockup là chuẩn (Ngày kết thúc - Ngày bắt đầu = số ngày dương). Sửa bảng mô tả trường cho khớp.'),

    ("US05-QA-04.4", "UI-UX",
     'Mockup "Thiết lập cấu phần" (image8)',
     'Mockup chỉ hiển thị nút "Xác nhận" ở Thiết lập cấu phần, KHÔNG thấy nút "Đóng". '
     'Nhưng bảng mô tả trường liệt kê CẢ nút "Xác nhận" VÀ nút "Đóng" (★ Button). '
     'Cần BA xác nhận: Thiết lập cấu phần có nút "Đóng" hay chỉ có nút "Xác nhận" + nút X (close popup)?',
     'Đề xuất: Bổ sung nút "Đóng" lên mockup Thiết lập cấu phần cho thống nhất với bảng mô tả.'),

    ("US05-QA-04.5", "UI-UX",
     'Flowchart (image1) — Bước 6',
     'Flowchart vẽ Bước 6 = "Người dùng khai báo các định dạng và ấn Xác nhận" → rồi mới đến Diamond "User chọn?" (Bước 7). '
     'Nhưng logic thực tế: người dùng khai báo xong → CHỌN "Xác nhận" HOẶC "Đóng" → đây mới đúng là Diamond ở Bước 7. '
     'Vậy Bước 6 lẽ ra chỉ là "Khai báo các định dạng" (chưa ấn gì). Flowchart đang gom hành động "ấn Xác nhận" vào Bước 6 nhưng lại tách ra nhánh ở Bước 7.',
     'Đề xuất: Sửa Flowchart Bước 6 → "Người dùng khai báo các định dạng" (bỏ "ấn Xác nhận"). Hành động chọn nút nằm ở Diamond Bước 7.'),

    ("US05-QA-04.6", "UI-UX",
     'Bảng mô tả trường Thiết lập công thức — các nút chức năng',
     'Bảng mô tả trường liệt kê nút "Thêm cấu phần" (icon +) và nút "Xóa cấu phần" (icon ×) nhưng KHÔNG ghi rõ tên nút. '
     'Cả 2 nút chỉ có icon, thiếu label text. Tương tự cho "Thêm định dạng" và "Xóa định dạng" ở Thiết lập cấu phần. '
     'Cần confirm: Có tooltip khi hover không? Label text là gì?',
     'Đề xuất: Bổ sung tooltip cho icon buttons. Label = "Thêm cấu phần" / "Xóa cấu phần" / "Thêm định dạng" / "Xóa định dạng".'),

    ("US05-QA-04.7", "UI-UX",
     'Bảng mô tả trường Thiết lập cấu phần — trường "Tỷ lệ"',
     'Ràng buộc ghi "giá trị > 0 và <= 1". Trên UI, người dùng có được nhập dạng phần trăm (%) '
     'ví dụ 5% không, hay bắt buộc nhập dạng số thập phân (0.05)? '
     'Nếu hỗ trợ nhập %, FE có tự động convert sang dạng thập phân trước khi gửi BE không?',
     'Đề xuất: UI cho phép nhập kèm ký hiệu % (VD: 5%) và tự động convert thành 0.05 dưới BE để thân thiện người dùng.'),
]

# Build table
doc.add_heading('Bảng Q&A — Phân tích US05', level=2)
table = doc.add_table(rows=1, cols=6, style='Light List Accent 1')
table.autofit = True

# Header
headers = ['ID', 'Trích xuất (Reference)', 'Câu hỏi / Sự cố', 'Phân loại', 'Đề xuất từ QA', 'Trả lời của BA']
for i, h_text in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h_text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(8)

# Set column widths
from docx.shared import Emu
col_widths = [Cm(2.2), Cm(4.5), Cm(8), Cm(2), Cm(6.5), Cm(3.5)]

for row_data in qa_data:
    row = table.add_row()
    values = [row_data[0], row_data[2], row_data[3], row_data[1], row_data[4], '']
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = val
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(8)

# Set column widths after all rows added
for row in table.rows:
    for i, width in enumerate(col_widths):
        row.cells[i].width = width

# Stats summary
doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run(f'Tổng số câu hỏi: {len(qa_data)}').bold = True
doc.add_paragraph(f'  • 🔶 Nghiệp vụ: {sum(1 for q in qa_data if q[1] == "Nghiệp vụ")}')
doc.add_paragraph(f'  • 🔴 Giới hạn: {sum(1 for q in qa_data if q[1] == "Giới hạn")}')
doc.add_paragraph(f'  • 🟠 Toàn vẹn dữ liệu: {sum(1 for q in qa_data if q[1] == "Toàn vẹn dữ liệu")}')
doc.add_paragraph(f'  • 🔵 UI-UX: {sum(1 for q in qa_data if q[1] == "UI-UX")}')

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('⚠️ LƯU Ý: ').bold = True
p.add_run('Vui lòng điền câu trả lời vào cột "Trả lời của BA" và gửi lại file này. '
           'Tuyệt đối không sinh Part C khi chưa có câu trả lời của BA.')

output_path = '/Users/mac/antigravity-testing-kit/Feature_02_SA_Tham_So_He_Thong/US05_PartB_QA.docx'
doc.save(output_path)
print(f'✅ Part B saved: {output_path}')
